"""
Cover letter formatter that generates DOCX files matching the exact template specification.
Every formatting value comes from styles.py.
"""

from docx import Document
from docx.shared import Pt
from . import styles
from .utils import validate_contact_info


class CoverLetterFormatter:
    """Generates formatted cover letter documents."""

    def __init__(self):
        """Initialize document with default settings."""
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """Configure page settings and margins."""
        sections = self.doc.sections
        for section in sections:
            section.page_height = styles.PAGE_HEIGHT
            section.page_width = styles.PAGE_WIDTH
            section.top_margin = styles.CL_MARGIN_ALL
            section.bottom_margin = styles.CL_MARGIN_ALL
            section.left_margin = styles.CL_MARGIN_ALL
            section.right_margin = styles.CL_MARGIN_ALL

    def add_header(self, name, tagline, contact_info):
        """
        Add cover letter header.

        Args:
            name (str): Full name (will be uppercased)
            tagline (str): Role/specialization
            contact_info (dict): {phone, email, linkedin, location}
        """
        contact_info = validate_contact_info(contact_info)

        # Name - left-aligned, bold, 14pt
        name_para = self.doc.add_paragraph()
        name_para.alignment = styles.ALIGN_LEFT
        name_run = name_para.add_run(name.upper())
        name_run.font.size = styles.FONT_SIZE_CL_NAME
        name_run.font.bold = True
        name_run.font.name = styles.FONT_FAMILY
        name_run.font.color.rgb = styles.COLOR_BLACK
        name_para.paragraph_format.space_after = styles.CL_SPACING_AFTER_NAME

        # Tagline - left-aligned, 11pt, gray
        tagline_para = self.doc.add_paragraph()
        tagline_para.alignment = styles.ALIGN_LEFT
        tagline_run = tagline_para.add_run(tagline)
        tagline_run.font.size = styles.FONT_SIZE_CL_TAGLINE
        tagline_run.font.name = styles.FONT_FAMILY
        tagline_run.font.color.rgb = styles.COLOR_DARK_GRAY
        tagline_para.paragraph_format.space_after = styles.CL_SPACING_AFTER_TAGLINE

        # Contact info - left-aligned, 10pt, space-separated (not bullets)
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = styles.ALIGN_LEFT
        contact_parts = [v for v in [
            contact_info['phone'],
            contact_info['email'],
            contact_info['linkedin'],
            contact_info['location']
        ] if v]
        contact_text = "  |  ".join(contact_parts)
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = styles.FONT_SIZE_CL_CONTACT
        contact_run.font.name = styles.FONT_FAMILY
        contact_run.font.color.rgb = styles.COLOR_BLACK
        contact_para.paragraph_format.space_after = styles.CL_SPACING_AFTER_CONTACT

    def add_section_label(self):
        """Add 'COVER LETTER' label."""
        label_para = self.doc.add_paragraph()
        label_para.alignment = styles.ALIGN_LEFT
        label_run = label_para.add_run("COVER LETTER")
        label_run.font.size = styles.FONT_SIZE_CL_BODY
        label_run.font.bold = True
        label_run.font.name = styles.FONT_FAMILY
        label_run.font.color.rgb = styles.COLOR_BLACK
        label_para.paragraph_format.space_before = styles.CL_SPACING_BEFORE_LABEL
        label_para.paragraph_format.space_after = styles.CL_SPACING_AFTER_LABEL

    def add_salutation(self, recipient_name=None):
        """
        Add salutation.

        Args:
            recipient_name (str): Hiring manager name (optional)
        """
        salutation_para = self.doc.add_paragraph()
        salutation_para.alignment = styles.ALIGN_LEFT

        if recipient_name:
            salutation_text = f"Dear {recipient_name},"
        else:
            salutation_text = "Dear Hiring Manager,"

        salutation_run = salutation_para.add_run(salutation_text)
        salutation_run.font.size = styles.FONT_SIZE_CL_BODY
        salutation_run.font.name = styles.FONT_FAMILY
        salutation_run.font.color.rgb = styles.COLOR_BLACK
        salutation_para.paragraph_format.space_before = styles.CL_SPACING_BEFORE_SALUTATION
        salutation_para.paragraph_format.space_after = styles.CL_SPACING_AFTER_SALUTATION

    def add_body_paragraph(self, text):
        """
        Add body paragraph.

        Args:
            text (str): Paragraph text
        """
        body_para = self.doc.add_paragraph()
        body_para.alignment = styles.ALIGN_LEFT
        body_run = body_para.add_run(text)
        body_run.font.size = styles.FONT_SIZE_CL_BODY
        body_run.font.name = styles.FONT_FAMILY
        body_run.font.color.rgb = styles.COLOR_BLACK

        # Slightly more open line spacing (1.15)
        body_para.paragraph_format.line_spacing = styles.LINE_SPACING_CL_BODY
        body_para.paragraph_format.space_after = styles.CL_SPACING_AFTER_BODY_PARA

    def add_signature(self, name):
        """
        Add closing and signature.

        Args:
            name (str): Full name
        """
        # Closing ("Sincerely,")
        closing_para = self.doc.add_paragraph()
        closing_para.alignment = styles.ALIGN_LEFT
        closing_run = closing_para.add_run("Sincerely,")
        closing_run.font.size = styles.FONT_SIZE_CL_BODY
        closing_run.font.name = styles.FONT_FAMILY
        closing_run.font.color.rgb = styles.COLOR_BLACK
        closing_para.paragraph_format.space_before = styles.CL_SPACING_BEFORE_CLOSING
        closing_para.paragraph_format.space_after = styles.CL_SPACING_AFTER_CLOSING

        # Name (not bold, not uppercase)
        name_para = self.doc.add_paragraph()
        name_para.alignment = styles.ALIGN_LEFT
        name_run = name_para.add_run(name)
        name_run.font.size = styles.FONT_SIZE_CL_BODY
        name_run.font.name = styles.FONT_FAMILY
        name_run.font.color.rgb = styles.COLOR_BLACK

    def save(self, filepath):
        """
        Save document to file.

        Args:
            filepath (str): Output path
        """
        self.doc.save(filepath)
        print(f"✓ Cover letter saved to: {filepath}")

    def get_document(self):
        """Return the document object for further processing."""
        return self.doc
