"""
Resume formatter that generates DOCX files matching the exact template specification.
Every formatting value comes from styles.py.

Also generates HTML preview and plain text from the same tracked content,
ensuring preview === download (single source of truth).
"""

import html as html_module

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from . import styles
from .utils import add_paragraph_border, format_date_range, validate_contact_info


class ResumeFormatter:
    """Generates formatted resume documents, HTML previews, and plain text."""

    def __init__(self):
        """Initialize document with default settings."""
        self.doc = Document()
        self._setup_document()
        # Content tracking for to_plain_text() and to_html()
        self._content = {
            "name": "",
            "tagline": "",
            "contact_parts": [],
            "sections": [],       # ordered list of (section_type, section_data)
        }

    def _setup_document(self):
        """Configure page settings and margins."""
        sections = self.doc.sections
        for section in sections:
            section.page_height = styles.PAGE_HEIGHT
            section.page_width = styles.PAGE_WIDTH
            section.top_margin = styles.RESUME_MARGIN_TOP
            section.bottom_margin = styles.RESUME_MARGIN_BOTTOM
            section.left_margin = styles.RESUME_MARGIN_LEFT
            section.right_margin = styles.RESUME_MARGIN_RIGHT

    def add_header(self, name, tagline, contact_info):
        """
        Add resume header with name, tagline, and contact info.

        Args:
            name (str): Full name (will be uppercased)
            tagline (str): Role/specialization
            contact_info (dict): {phone, email, linkedin, location}
        """
        contact_info = validate_contact_info(contact_info)

        # Track content
        self._content["name"] = name
        self._content["tagline"] = tagline
        self._content["contact_parts"] = [v for v in [
            contact_info['phone'],
            contact_info['email'],
            contact_info['linkedin'],
            contact_info['location']
        ] if v]

        # Name - centered, bold, 18pt
        name_para = self.doc.add_paragraph()
        name_para.alignment = styles.ALIGN_CENTER
        name_run = name_para.add_run(name.upper())
        name_run.font.size = styles.FONT_SIZE_RESUME_NAME
        name_run.font.bold = True
        name_run.font.name = styles.FONT_FAMILY
        name_run.font.color.rgb = styles.COLOR_BLACK
        name_para.paragraph_format.space_after = styles.SPACING_AFTER_NAME

        # Tagline - centered, 11pt, dark gray
        tagline_para = self.doc.add_paragraph()
        tagline_para.alignment = styles.ALIGN_CENTER
        tagline_run = tagline_para.add_run(tagline)
        tagline_run.font.size = styles.FONT_SIZE_RESUME_TAGLINE
        tagline_run.font.name = styles.FONT_FAMILY
        tagline_run.font.color.rgb = styles.COLOR_DARK_GRAY
        tagline_para.paragraph_format.space_after = styles.SPACING_AFTER_TAGLINE

        # Contact info - centered, 9pt, bullet-separated
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = styles.ALIGN_CENTER
        contact_text = " \u2022 ".join(self._content["contact_parts"])
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = styles.FONT_SIZE_RESUME_CONTACT
        contact_run.font.name = styles.FONT_FAMILY
        contact_run.font.color.rgb = styles.COLOR_BLACK
        contact_para.paragraph_format.space_after = styles.SPACING_AFTER_CONTACT

        # Add bottom border (1pt solid line)
        add_paragraph_border(contact_para, bottom=True, bottom_width=styles.BORDER_THIN)

    def add_section_header(self, title):
        """
        Add section header (SUMMARY, EXPERIENCE, etc.) with borders.

        Args:
            title (str): Section name (will be uppercased)
        """
        # Track: we record the section header. The actual content gets
        # attached by the next add_* call (summary, competencies, etc.)
        self._content["sections"].append(("header", title.upper()))

        header_para = self.doc.add_paragraph()
        header_para.alignment = styles.ALIGN_LEFT
        header_run = header_para.add_run(title.upper())
        header_run.font.size = styles.FONT_SIZE_SECTION_HEADER
        header_run.font.bold = True
        header_run.font.name = styles.FONT_FAMILY
        header_run.font.color.rgb = styles.COLOR_BLACK

        # Add top (2pt) and bottom (1pt) borders
        add_paragraph_border(
            header_para,
            top=True,
            bottom=True,
            top_width=styles.BORDER_THICK,
            bottom_width=styles.BORDER_THIN
        )

        header_para.paragraph_format.space_before = styles.SPACING_BEFORE_SECTION_HEADER
        header_para.paragraph_format.space_after = styles.SPACING_AFTER_SECTION_HEADER

    def add_summary(self, text):
        """
        Add summary paragraph (justified alignment).

        Args:
            text (str): Summary text
        """
        self._content["sections"].append(("summary", text))

        summary_para = self.doc.add_paragraph()
        summary_para.alignment = styles.ALIGN_JUSTIFY
        summary_run = summary_para.add_run(text)
        summary_run.font.size = styles.FONT_SIZE_DEFAULT
        summary_run.font.name = styles.FONT_FAMILY
        summary_run.font.color.rgb = styles.COLOR_BLACK
        summary_para.paragraph_format.space_after = styles.SPACING_AFTER_SUMMARY

    def add_core_competencies(self, competencies):
        """
        Add core competencies as 3-column layout with checkmarks.

        Args:
            competencies (list): List of competency strings
        """
        self._content["sections"].append(("competencies", list(competencies)))

        # Calculate rows needed for 3-column layout
        rows_needed = (len(competencies) + 2) // 3

        # Create table
        table = self.doc.add_table(rows=rows_needed, cols=3)

        # Remove all borders and set cell properties
        for row in table.rows:
            for cell in row.cells:
                # Remove borders
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        # Fill cells with competencies
        for i, comp in enumerate(competencies):
            row_idx = i // 3
            col_idx = i % 3
            cell = table.rows[row_idx].cells[col_idx]

            cell_para = cell.paragraphs[0]
            # Add hanging indent so text wraps under text, not under checkmark
            cell_para.paragraph_format.left_indent = Pt(12)  # Indent for the checkmark
            cell_para.paragraph_format.first_line_indent = Pt(-12)  # Negative indent for first line (checkmark)

            # Add checkmark
            check_run = cell_para.add_run("\u2713 ")
            check_run.font.size = styles.FONT_SIZE_DEFAULT
            check_run.font.name = styles.FONT_FAMILY

            # Add competency text
            comp_run = cell_para.add_run(comp)
            comp_run.font.size = styles.FONT_SIZE_DEFAULT
            comp_run.font.name = styles.FONT_FAMILY
            comp_run.font.color.rgb = styles.COLOR_BLACK

    def add_experience_entry(self, company, title, location, dates, overview=None, bullets=None):
        """
        Add a job entry in Experience section.

        Args:
            company (str): Company name
            title (str): Job title
            location (str): City, STATE
            dates (str): YYYY - YYYY
            overview (str): Company overview text (optional)
            bullets (list): List of bullet point strings
        """
        if bullets is None:
            bullets = []

        self._content["sections"].append(("experience_entry", {
            "company": company,
            "title": title,
            "location": location,
            "dates": dates,
            "overview": overview,
            "bullets": list(bullets),
        }))

        # Create 2-row table for proper layout:
        # Row 1: Company (left) | Location (right)
        # Row 2: Job Title (left) | Dates (right)
        company_table = self.doc.add_table(rows=2, cols=2)
        company_table.autofit = False
        company_table.allow_autofit = False

        # Remove borders from all cells
        for row in company_table.rows:
            for cell in row.cells:
                tc = cell._element
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'nil')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

        # Row 1, Left cell: Company name (bold, 11pt)
        company_cell = company_table.rows[0].cells[0]
        company_para = company_cell.paragraphs[0]
        company_para.paragraph_format.space_after = Pt(0)
        company_run = company_para.add_run(company)
        company_run.font.size = styles.FONT_SIZE_COMPANY
        company_run.font.bold = True
        company_run.font.name = styles.FONT_FAMILY
        company_run.font.color.rgb = styles.COLOR_BLACK

        # Row 1, Right cell: Location (right-aligned, gray)
        location_cell = company_table.rows[0].cells[1]
        location_para = location_cell.paragraphs[0]
        location_para.alignment = styles.ALIGN_RIGHT
        if location:
            location_run = location_para.add_run(location)
            location_run.font.size = styles.FONT_SIZE_DEFAULT
            location_run.font.name = styles.FONT_FAMILY
            location_run.font.color.rgb = styles.COLOR_DARK_GRAY

        # Row 2, Left cell: Job title (bold, 10pt)
        title_cell = company_table.rows[1].cells[0]
        title_para = title_cell.paragraphs[0]
        title_para.paragraph_format.space_before = Pt(0)
        title_run = title_para.add_run(title)
        title_run.font.size = styles.FONT_SIZE_JOB_TITLE
        title_run.font.bold = True
        title_run.font.name = styles.FONT_FAMILY
        title_run.font.color.rgb = styles.COLOR_BLACK

        # Row 2, Right cell: Dates (right-aligned, gray)
        dates_cell = company_table.rows[1].cells[1]
        dates_para = dates_cell.paragraphs[0]
        dates_para.alignment = styles.ALIGN_RIGHT
        dates_para.paragraph_format.space_before = Pt(0)
        dates_run = dates_para.add_run(dates)
        dates_run.font.size = styles.FONT_SIZE_DEFAULT
        dates_run.font.name = styles.FONT_FAMILY
        dates_run.font.color.rgb = styles.COLOR_DARK_GRAY

        # Company overview (italic, gray, same size as other details)
        if overview:
            overview_para = self.doc.add_paragraph()
            overview_para.alignment = styles.ALIGN_LEFT
            overview_run = overview_para.add_run(f"Company Overview: {overview}")
            overview_run.font.size = styles.FONT_SIZE_DEFAULT
            overview_run.font.italic = True
            overview_run.font.name = styles.FONT_FAMILY
            overview_run.font.color.rgb = styles.COLOR_DARK_GRAY
            overview_para.paragraph_format.space_before = Pt(0)
            overview_para.paragraph_format.space_after = Pt(6)

        # Bullets (left-aligned with hanging indent)
        for i, bullet in enumerate(bullets):
            bullet_para = self.doc.add_paragraph()
            bullet_para.alignment = styles.ALIGN_LEFT
            bullet_para.paragraph_format.left_indent = Inches(0.25)
            bullet_para.paragraph_format.first_line_indent = Inches(-0.25)
            bullet_run = bullet_para.add_run("\u2022\t" + bullet)
            bullet_run.font.size = styles.FONT_SIZE_DEFAULT
            bullet_run.font.name = styles.FONT_FAMILY
            bullet_run.font.color.rgb = styles.COLOR_BLACK
            bullet_para.paragraph_format.space_after = styles.SPACING_BETWEEN_BULLETS
            if i == len(bullets) - 1:
                bullet_para.paragraph_format.space_after = styles.SPACING_BETWEEN_JOBS

    def add_skills(self, skills_dict):
        """
        Add skills section with category headers.

        Args:
            skills_dict (dict): {"Category": ["skill1", "skill2", ...]}
        """
        self._content["sections"].append(("skills", dict(skills_dict)))

        for category, skills_list in skills_dict.items():
            skills_para = self.doc.add_paragraph()

            # Category header (bold)
            category_run = skills_para.add_run(f"{category}: ")
            category_run.font.size = styles.FONT_SIZE_DEFAULT
            category_run.font.bold = True
            category_run.font.name = styles.FONT_FAMILY
            category_run.font.color.rgb = styles.COLOR_BLACK

            # Skills (bullet-separated)
            skills_text = " \u2022 ".join(skills_list)
            skills_run = skills_para.add_run(skills_text)
            skills_run.font.size = styles.FONT_SIZE_DEFAULT
            skills_run.font.name = styles.FONT_FAMILY
            skills_run.font.color.rgb = styles.COLOR_BLACK
            skills_para.paragraph_format.space_after = Pt(6)

    def add_education(self, school, degree, details=None):
        """
        Add education entry.

        Args:
            school (str): School name (institution)
            degree (str): Degree type and major
            details (list): Optional list of detail strings (e.g., concentration, honors)
        """
        self._content["sections"].append(("education", {
            "school": school,
            "degree": degree,
            "details": details,
        }))

        # Line 1: Degree type and major (regular, 10pt)
        if degree and degree.strip():
            degree_para = self.doc.add_paragraph()
            degree_para.alignment = styles.ALIGN_LEFT
            degree_run = degree_para.add_run(degree)
            degree_run.font.size = styles.FONT_SIZE_DEFAULT
            degree_run.font.name = styles.FONT_FAMILY
            degree_run.font.color.rgb = styles.COLOR_BLACK
            degree_para.paragraph_format.space_after = Pt(0)

        # Line 2: Institution/School name (bold, 10pt)
        if school and school.strip():
            school_para = self.doc.add_paragraph()
            school_para.alignment = styles.ALIGN_LEFT
            school_run = school_para.add_run(school)
            school_run.font.size = styles.FONT_SIZE_DEFAULT
            school_run.font.bold = True
            school_run.font.name = styles.FONT_FAMILY
            school_run.font.color.rgb = styles.COLOR_BLACK
            school_para.paragraph_format.space_after = Pt(0)

        # Line 3+: Details like concentration (regular, 10pt)
        if details:
            if isinstance(details, str):
                details_list = [details]
            elif isinstance(details, list):
                details_list = details
            else:
                details_list = [str(details)]

            for detail in details_list:
                if detail and str(detail).strip():
                    detail_para = self.doc.add_paragraph()
                    detail_para.alignment = styles.ALIGN_LEFT
                    detail_run = detail_para.add_run(str(detail))
                    detail_run.font.size = styles.FONT_SIZE_DEFAULT
                    detail_run.font.name = styles.FONT_FAMILY
                    detail_run.font.color.rgb = styles.COLOR_BLACK
                    detail_para.paragraph_format.space_after = Pt(0)

    def save(self, filepath):
        """
        Save document to file.

        Args:
            filepath (str): Output path
        """
        self.doc.save(filepath)
        print(f"\u2713 Resume saved to: {filepath}")

    def get_document(self):
        """Return the document object for further processing."""
        return self.doc

    # =========================================================================
    # PLAIN TEXT OUTPUT (for canonical hash + fallback preview)
    # =========================================================================

    def to_plain_text(self):
        """
        Generate plain text representation of resume from tracked content.
        Used by canonical_document.py for hash computation and text preview.
        """
        lines = []
        esc = lambda s: str(s) if s else ""

        # Header
        lines.append(esc(self._content["name"]).upper())
        if self._content["tagline"]:
            lines.append(esc(self._content["tagline"]))
        if self._content["contact_parts"]:
            lines.append(" | ".join(self._content["contact_parts"]))
        lines.append("")
        lines.append("\u2500" * 60)
        lines.append("")

        for section_type, data in self._content["sections"]:
            if section_type == "header":
                lines.append(data)

            elif section_type == "summary":
                lines.append(esc(data))
                lines.append("")

            elif section_type == "competencies":
                lines.append(" \u2022 ".join(data))
                lines.append("")

            elif section_type == "experience_entry":
                company_line = esc(data["company"])
                if data["location"]:
                    company_line += f", {data['location']}"
                if data["dates"]:
                    company_line += f"  |  {data['dates']}"
                lines.append(company_line)
                if data["title"]:
                    lines.append(esc(data["title"]))
                if data["overview"]:
                    lines.append(f"Company Overview: {data['overview']}")
                for bullet in data["bullets"]:
                    if bullet:
                        bt = esc(bullet)
                        lines.append(f"\u2022 {bt}" if not bt.startswith("\u2022") else bt)
                lines.append("")

            elif section_type == "skills":
                if isinstance(data, dict):
                    for category, skill_list in data.items():
                        cat = category.title() if category else ""
                        if isinstance(skill_list, list):
                            lines.append(f"{cat}: {', '.join(skill_list)}")
                        else:
                            lines.append(f"{cat}: {skill_list}")
                elif isinstance(data, list):
                    lines.append(" \u2022 ".join(data))
                else:
                    lines.append(str(data))
                lines.append("")

            elif section_type == "education":
                if data.get("degree") and str(data["degree"]).strip():
                    lines.append(esc(data["degree"]))
                if data.get("school") and str(data["school"]).strip():
                    lines.append(esc(data["school"]))
                details = data.get("details")
                if details:
                    if isinstance(details, list):
                        for d in details:
                            if d and str(d).strip():
                                lines.append(str(d))
                    elif str(details).strip():
                        lines.append(str(details))
                lines.append("")

        return "\n".join(lines).strip()

    # =========================================================================
    # HTML OUTPUT (WYSIWYG preview matching DOCX layout)
    # =========================================================================

    def to_html(self):
        """
        Generate styled HTML that visually mirrors the DOCX output.
        Uses inline styles derived from styles.py constants for exact parity.
        """
        e = html_module.escape
        parts = []

        # Outer page container — mimics 8.5" page with proper margins
        parts.append(
            '<div class="resume-page" style="'
            'max-width: 816px; '  # 8.5" at 96dpi
            'margin: 0 auto; '
            'padding: 48px 62px; '  # ~0.5" top/bottom, ~0.65" left/right
            'font-family: Arial, Calibri, Helvetica, sans-serif; '
            'color: #000; '
            'line-height: 1.3; '
            'background: #fff; '
            '">'
        )

        # ── Header ──
        name = e(self._content["name"].upper()) if self._content["name"] else ""
        parts.append(
            f'<div class="resume-name" style="'
            f'text-align: center; '
            f'font-size: 18pt; '
            f'font-weight: bold; '
            f'margin-bottom: 2px; '
            f'">{name}</div>'
        )

        if self._content["tagline"]:
            parts.append(
                f'<div class="resume-tagline" style="'
                f'text-align: center; '
                f'font-size: 11pt; '
                f'color: #666666; '
                f'margin-bottom: 4px; '
                f'">{e(self._content["tagline"])}</div>'
            )

        if self._content["contact_parts"]:
            contact_html = " &bull; ".join(e(p) for p in self._content["contact_parts"])
            parts.append(
                f'<div class="resume-contact" style="'
                f'text-align: center; '
                f'font-size: 9pt; '
                f'padding-bottom: 6px; '
                f'border-bottom: 1px solid #000; '
                f'margin-bottom: 8px; '
                f'">{contact_html}</div>'
            )

        # ── Sections ──
        for section_type, data in self._content["sections"]:
            if section_type == "header":
                parts.append(
                    f'<div class="resume-section-header" style="'
                    f'font-size: 12pt; '
                    f'font-weight: bold; '
                    f'text-transform: uppercase; '
                    f'border-top: 2px solid #000; '
                    f'border-bottom: 1px solid #000; '
                    f'padding: 3px 0; '
                    f'margin-top: 8px; '
                    f'margin-bottom: 4px; '
                    f'">{e(data)}</div>'
                )

            elif section_type == "summary":
                parts.append(
                    f'<div class="resume-summary" style="'
                    f'font-size: 10pt; '
                    f'text-align: justify; '
                    f'margin-bottom: 8px; '
                    f'">{e(data)}</div>'
                )

            elif section_type == "competencies":
                parts.append(
                    '<div class="resume-competencies" style="'
                    'display: grid; '
                    'grid-template-columns: 1fr 1fr 1fr; '
                    'gap: 2px 12px; '
                    'font-size: 10pt; '
                    'margin-bottom: 8px; '
                    '">'
                )
                for comp in data:
                    parts.append(
                        f'<div style="padding: 1px 0;">'
                        f'<span style="color: #000;">\u2713</span> {e(comp)}'
                        f'</div>'
                    )
                parts.append('</div>')

            elif section_type == "experience_entry":
                parts.append('<div class="resume-experience-entry" style="margin-bottom: 8px;">')

                # Row 1: Company | Location
                parts.append(
                    '<div style="display: flex; justify-content: space-between; align-items: baseline;">'
                )
                parts.append(
                    f'<span style="font-size: 11pt; font-weight: bold;">'
                    f'{e(data["company"])}</span>'
                )
                if data["location"]:
                    parts.append(
                        f'<span style="font-size: 10pt; color: #666666;">'
                        f'{e(data["location"])}</span>'
                    )
                parts.append('</div>')

                # Row 2: Title | Dates
                parts.append(
                    '<div style="display: flex; justify-content: space-between; align-items: baseline;">'
                )
                parts.append(
                    f'<span style="font-size: 10pt; font-weight: bold;">'
                    f'{e(data["title"])}</span>'
                )
                if data["dates"]:
                    parts.append(
                        f'<span style="font-size: 10pt; color: #666666;">'
                        f'{e(data["dates"])}</span>'
                    )
                parts.append('</div>')

                # Overview
                if data.get("overview"):
                    parts.append(
                        f'<div style="font-size: 10pt; font-style: italic; color: #666666; '
                        f'margin-top: 2px; margin-bottom: 4px;">'
                        f'Company Overview: {e(data["overview"])}</div>'
                    )

                # Bullets
                if data["bullets"]:
                    parts.append(
                        '<ul style="'
                        'list-style: none; '
                        'padding-left: 18px; '
                        'margin: 3px 0 0 0; '
                        '">'
                    )
                    for bullet in data["bullets"]:
                        if bullet:
                            bt = e(bullet.lstrip("\u2022 ").strip())
                            parts.append(
                                f'<li style="'
                                f'font-size: 10pt; '
                                f'text-indent: -14px; '
                                f'padding-left: 14px; '
                                f'margin-bottom: 2px; '
                                f'">&bull;&ensp;{bt}</li>'
                            )
                    parts.append('</ul>')

                parts.append('</div>')

            elif section_type == "skills":
                parts.append('<div class="resume-skills" style="margin-bottom: 6px;">')
                if isinstance(data, dict):
                    for category, skill_list in data.items():
                        cat = e(category)
                        if isinstance(skill_list, list):
                            skills_str = " &bull; ".join(e(s) for s in skill_list)
                        else:
                            skills_str = e(str(skill_list))
                        parts.append(
                            f'<div style="font-size: 10pt; margin-bottom: 4px;">'
                            f'<strong>{cat}:</strong> {skills_str}'
                            f'</div>'
                        )
                parts.append('</div>')

            elif section_type == "education":
                parts.append('<div class="resume-education" style="font-size: 10pt;">')
                if data.get("degree") and str(data["degree"]).strip():
                    parts.append(f'<div>{e(str(data["degree"]))}</div>')
                if data.get("school") and str(data["school"]).strip():
                    parts.append(
                        f'<div style="font-weight: bold;">{e(str(data["school"]))}</div>'
                    )
                details = data.get("details")
                if details:
                    if isinstance(details, list):
                        for d in details:
                            if d and str(d).strip():
                                parts.append(f'<div>{e(str(d))}</div>')
                    elif str(details).strip():
                        parts.append(f'<div>{e(str(details))}</div>')
                parts.append('</div>')

        parts.append('</div>')  # close .resume-page
        return "\n".join(parts)
