"""
Document Generator Module for HenryAI

Professional DOCX formatters for resumes and cover letters.
Designed to be ATS-safe with clean, consistent formatting.

FORMATTING SPECIFICATION (NON-NEGOTIABLE):
See docs/guides/resume-builder/RESUME_COVERLETTER_FORMATTING_SPEC.md

Resume:
- Margins: 0.5 inches all sides
- Font: Calibri/Arial/Helvetica only
- Name: 16-18pt, Section headers: 11-12pt, Body: 10.5-11pt
- Line spacing: 1.0-1.15
- Single column, left-aligned, round bullets only
- Max 2 pages
- NO graduation years in education

Cover Letter:
- Margins: 1 inch all sides
- Font size: 11pt
- Single-spaced, one blank line between paragraphs
- Max 1 page, 3-4 paragraphs
- Executive tone, no enthusiasm statements or cliches
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List, Optional
import re


# =============================================================================
# FORMATTING CONSTANTS (FROM SPEC - NON-NEGOTIABLE)
# =============================================================================

# Resume formatting constants
RESUME_MARGINS_INCHES = 0.5
RESUME_NAME_SIZE_PT = 17  # 16-18 pt range
RESUME_SECTION_HEADER_SIZE_PT = 11.5  # 11-12 pt range
RESUME_BODY_SIZE_PT = 10.5  # 10.5-11 pt range
RESUME_LINE_SPACING = 1.08  # 1.0-1.15 range
RESUME_SPACE_AFTER_HEADER_PT = 7  # 6-8 pt range
RESUME_SPACE_BETWEEN_BULLETS_PT = 5  # 4-6 pt range
RESUME_MAX_PAGES = 2

# Cover letter formatting constants
COVER_LETTER_MARGINS_INCHES = 1.0
COVER_LETTER_FONT_SIZE_PT = 11
COVER_LETTER_MAX_PAGES = 1
COVER_LETTER_MAX_PARAGRAPHS = 4

# Allowed fonts (in order of preference)
ALLOWED_FONTS = ["Calibri", "Arial", "Helvetica"]
DEFAULT_FONT = ALLOWED_FONTS[0]

# Forbidden cover letter patterns -> replacements
FORBIDDEN_COVER_LETTER_PATTERNS = [
    (r"I'm excited to", "I am prepared to"),
    (r"I am excited to", "I am prepared to"),
    (r"I would be thrilled", "I am prepared"),
    (r"passionate about", "experienced in"),
    (r"team player", "collaborative professional"),
    (r"hard worker", "dedicated professional"),
    (r"I believe I would be", "I am"),
    (r"I feel that I", "I"),
]


def strip_graduation_year(education_text: str) -> str:
    """
    Remove graduation years from education entries.

    Per spec: Education must include ONLY degree, field, and institution.
    NO graduation years.

    Examples:
        "Stanford University, 2015" -> "Stanford University"
        "BS Computer Science (2018)" -> "BS Computer Science"
        "MBA, Harvard Business School, Class of 2020" -> "MBA, Harvard Business School"
    """
    if not education_text:
        return ""

    # Remove year patterns (order matters - more specific patterns first)
    patterns = [
        r',?\s*Class of \d{4}',  # "Class of 2020"
        r',?\s*Graduated:?\s*\d{4}',  # "Graduated 2015"
        r'\s*\(\d{4}\s*[-–]\s*\d{4}\)\s*',  # "(2012-2016)"
        r'\s*\(\d{4}\)\s*',  # "(2015)"
        r',?\s*\d{4}\s*[-–]\s*\d{4}',  # ", 2012-2016"
        r',?\s*\d{4}\s*[-–]\s*[Pp]resent',  # "2012-Present"
        r',?\s*\d{4}\s*$',  # ", 2015" at end
    ]

    result = education_text
    for pattern in patterns:
        result = re.sub(pattern, '', result, flags=re.IGNORECASE)

    # Clean up any trailing punctuation and whitespace
    result = result.strip()
    result = re.sub(r',\s*$', '', result)  # Remove trailing comma
    result = re.sub(r'\s+', ' ', result)  # Normalize whitespace

    return result.strip()


def filter_cover_letter_content(text: str) -> str:
    """
    Remove forbidden phrases from cover letter content.

    Per spec: Cover letters must use executive-to-executive tone.
    NO enthusiasm statements, NO autobiographical storytelling, NO cliches.
    """
    if not text:
        return ""

    result = text
    for pattern, replacement in FORBIDDEN_COVER_LETTER_PATTERNS:
        result = re.sub(pattern, replacement, result, flags=re.IGNORECASE)
    return result


class BaseFormatter:
    """Base class for document formatters with common utilities."""

    # Subclasses override these
    MARGIN_INCHES = 0.75  # Default, overridden by subclass

    def __init__(self):
        self.doc = Document()
        self._setup_document()

    def _setup_document(self):
        """Set up ATS-safe document defaults with spec-compliant margins."""
        for section in self.doc.sections:
            section.top_margin = Inches(self.MARGIN_INCHES)
            section.bottom_margin = Inches(self.MARGIN_INCHES)
            section.left_margin = Inches(self.MARGIN_INCHES)
            section.right_margin = Inches(self.MARGIN_INCHES)

    def _set_font(self, run, font_name=None, font_size=11, bold=False, italic=False):
        """Apply consistent font settings to a run. Uses spec-compliant fonts."""
        if font_name is None:
            font_name = DEFAULT_FONT
        # Ensure font is in allowed list
        if font_name not in ALLOWED_FONTS:
            font_name = DEFAULT_FONT

        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic

        # Ensure font is set for complex scripts
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rPr.insert(0, rFonts)

    def _add_paragraph(self, text: str, font_size: int = 11, bold: bool = False,
                       alignment=None, space_before: int = 0, space_after: int = 0):
        """Add a paragraph with standard formatting."""
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15

        if alignment:
            p.alignment = alignment

        run = p.add_run(text)
        self._set_font(run, font_size=font_size, bold=bold)

        return p

    def get_document(self):
        """Return the Document object."""
        return self.doc


class ResumeFormatter(BaseFormatter):
    """
    Professional resume formatter with ATS-safe formatting.

    SPEC COMPLIANCE (NON-NEGOTIABLE):
    - Margins: 0.5 inches all sides
    - Font: Calibri/Arial/Helvetica only
    - Name: 16-18pt, Section headers: 11-12pt, Body: 10.5-11pt
    - Line spacing: 1.0-1.15
    - Single column, left-aligned, round bullets only
    - Max 2 pages
    - NO graduation years in education

    Tracks all content added so it can generate both DOCX and plain text
    from the same source (single source of truth for preview === download).
    """

    # Resume uses 0.5" margins per spec
    MARGIN_INCHES = RESUME_MARGINS_INCHES

    def __init__(self):
        super().__init__()
        # Track content for plain text generation
        self._content = {
            "name": "",
            "tagline": "",
            "contact_parts": [],
            "summary": "",
            "competencies": [],
            "experience": [],
            "skills": None,
            "education": [],
        }

    def add_header(self, name: str, tagline: str = "", contact_info: Dict[str, Any] = None):
        """Add resume header with name, tagline, and contact info."""
        # Track for plain text
        self._content["name"] = name
        self._content["tagline"] = tagline

        # Name - 16-18pt per spec, bold, centered
        p = self._add_paragraph(name, font_size=RESUME_NAME_SIZE_PT, bold=True,
                               alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Tagline - section header size per spec
        if tagline:
            self._add_paragraph(tagline, font_size=RESUME_SECTION_HEADER_SIZE_PT,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              space_after=4)

        # Contact info - single line, centered, body size
        if contact_info:
            contact_parts = []
            if contact_info.get('phone'):
                contact_parts.append(contact_info['phone'])
            if contact_info.get('email'):
                contact_parts.append(contact_info['email'])
            if contact_info.get('linkedin'):
                contact_parts.append(contact_info['linkedin'])
            if contact_info.get('location'):
                contact_parts.append(contact_info['location'])

            # Track for plain text
            self._content["contact_parts"] = contact_parts

            if contact_parts:
                contact_text = " | ".join(contact_parts)
                self._add_paragraph(contact_text, font_size=RESUME_BODY_SIZE_PT,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  space_after=12)

    def add_section_header(self, title: str):
        """Add a section header (e.g., 'Experience', 'Education')."""
        # Add a little space before section
        self._add_paragraph("", space_before=8)

        # Section header - 11-12pt per spec, bold, with underline effect via border
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(RESUME_SPACE_AFTER_HEADER_PT)  # 6-8pt per spec

        run = p.add_run(title.upper())
        self._set_font(run, font_size=RESUME_SECTION_HEADER_SIZE_PT, bold=True)

        # Add bottom border for visual separation
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    def add_summary(self, summary: str):
        """Add professional summary section."""
        # Track for plain text
        self._content["summary"] = summary
        self._add_paragraph(summary, font_size=RESUME_BODY_SIZE_PT, space_after=6)

    def add_core_competencies(self, competencies: List[str]):
        """Add core competencies as a formatted list or grid."""
        if not competencies:
            return

        # Track for plain text
        self._content["competencies"] = competencies

        # Format as a comma-separated list with bullets
        competency_text = " • ".join(competencies)
        self._add_paragraph(competency_text, font_size=RESUME_BODY_SIZE_PT, space_after=6)

    def add_experience_entry(self, company: str, title: str, location: str = "",
                            dates: str = "", overview: str = None, bullets: List[str] = None):
        """Add a single experience entry with spec-compliant formatting."""
        # Track for plain text
        self._content["experience"].append({
            "company": company,
            "title": title,
            "location": location,
            "dates": dates,
            "overview": overview,
            "bullets": bullets or [],
        })

        # Company and dates on same line - NO EXTRA SPACING between roles per spec
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

        # Company name - bold, body size per spec
        company_run = p.add_run(company)
        self._set_font(company_run, font_size=RESUME_BODY_SIZE_PT, bold=True)

        # Location if provided
        if location:
            loc_run = p.add_run(f", {location}")
            self._set_font(loc_run, font_size=RESUME_BODY_SIZE_PT)

        # Dates - ATS-safe inline format
        if dates:
            dates_run = p.add_run(f"  |  {dates}")
            self._set_font(dates_run, font_size=RESUME_BODY_SIZE_PT)

        # Title - italic, body size per spec
        if title:
            title_p = self.doc.add_paragraph()
            title_p.style = self.doc.styles['Normal']
            title_p.paragraph_format.space_before = Pt(0)
            title_p.paragraph_format.space_after = Pt(4)

            title_run = title_p.add_run(title)
            self._set_font(title_run, font_size=RESUME_BODY_SIZE_PT, italic=True)

        # Overview paragraph if provided
        if overview:
            self._add_paragraph(f"Company Overview: {overview}", font_size=RESUME_BODY_SIZE_PT, space_after=4)

        # Bullet points - 4-6pt spacing between bullets per spec
        if bullets:
            for bullet in bullets:
                bullet_text = f"• {bullet}" if not bullet.startswith("•") else bullet
                self._add_paragraph(bullet_text, font_size=RESUME_BODY_SIZE_PT,
                                  space_after=RESUME_SPACE_BETWEEN_BULLETS_PT)

    def add_skills(self, skills: Dict[str, Any]):
        """Add skills section - handles various formats."""
        # Track for plain text
        self._content["skills"] = skills

        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                # Title case the category name
                category_title = category.title() if category else ""
                if isinstance(skill_list, list):
                    skill_text = f"{category_title}: {', '.join(skill_list)}"
                else:
                    skill_text = f"{category_title}: {skill_list}"
                self._add_paragraph(skill_text, font_size=RESUME_BODY_SIZE_PT, space_after=2)
        elif isinstance(skills, list):
            skills_text = " • ".join(skills)
            self._add_paragraph(skills_text, font_size=RESUME_BODY_SIZE_PT, space_after=4)
        elif isinstance(skills, str):
            self._add_paragraph(skills, font_size=RESUME_BODY_SIZE_PT, space_after=4)

    def add_education(self, school: str, degree: str = "", details=None):
        """
        Add education entry.

        SPEC COMPLIANCE: NO GRADUATION YEARS
        Per spec, education must include ONLY degree, field, and institution.
        Years are automatically stripped.
        """
        # Strip graduation years from all education fields
        school = strip_graduation_year(school)
        degree = strip_graduation_year(degree)
        if details:
            if isinstance(details, list):
                details = ', '.join(str(d) for d in details if d)
            details = strip_graduation_year(str(details))

        # Track for plain text
        self._content["education"].append({
            "school": school,
            "degree": degree,
            "details": details,
        })

        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(2)

        # School name - bold, body size per spec
        school_run = p.add_run(school)
        self._set_font(school_run, font_size=RESUME_BODY_SIZE_PT, bold=True)

        # Degree - body size per spec
        if degree:
            degree_p = self.doc.add_paragraph()
            degree_p.style = self.doc.styles['Normal']
            degree_p.paragraph_format.space_before = Pt(0)
            degree_p.paragraph_format.space_after = Pt(2)

            degree_run = degree_p.add_run(degree)
            self._set_font(degree_run, font_size=RESUME_BODY_SIZE_PT)

        # Additional details (already stripped of years)
        if details:
            self._add_paragraph(str(details), font_size=RESUME_BODY_SIZE_PT, space_after=4)

    def to_plain_text(self) -> str:
        """
        Generate plain text representation of resume.
        This is the CANONICAL text that preview displays.
        Matches exactly what goes into the DOCX.
        """
        lines = []

        # Header
        lines.append(self._content["name"].upper())
        if self._content["tagline"]:
            lines.append(self._content["tagline"])
        if self._content["contact_parts"]:
            lines.append(" | ".join(self._content["contact_parts"]))
        lines.append("")
        lines.append("─" * 60)
        lines.append("")

        # Summary
        if self._content["summary"]:
            lines.append("SUMMARY")
            lines.append(self._content["summary"])
            lines.append("")

        # Core Competencies
        if self._content["competencies"]:
            lines.append("CORE COMPETENCIES")
            lines.append(" • ".join(self._content["competencies"]))
            lines.append("")

        # Experience
        if self._content["experience"]:
            lines.append("EXPERIENCE")
            lines.append("")
            for exp in self._content["experience"]:
                # Company line with location and dates
                company_line = exp["company"]
                if exp["location"]:
                    company_line += f", {exp['location']}"
                if exp["dates"]:
                    company_line += f"  |  {exp['dates']}"
                lines.append(company_line)

                # Title
                if exp["title"]:
                    lines.append(exp["title"])

                # Overview
                if exp["overview"]:
                    lines.append(f"Company Overview: {exp['overview']}")

                # Bullets
                for bullet in exp["bullets"]:
                    if bullet:
                        bullet_text = f"• {bullet}" if not bullet.startswith("•") else bullet
                        lines.append(bullet_text)
                lines.append("")

        # Skills
        if self._content["skills"]:
            lines.append("SKILLS")
            skills = self._content["skills"]
            if isinstance(skills, dict):
                for category, skill_list in skills.items():
                    category_title = category.title() if category else ""
                    if isinstance(skill_list, list):
                        lines.append(f"{category_title}: {', '.join(skill_list)}")
                    else:
                        lines.append(f"{category_title}: {skill_list}")
            elif isinstance(skills, list):
                lines.append(" • ".join(skills))
            else:
                lines.append(str(skills))
            lines.append("")

        # Education
        if self._content["education"]:
            lines.append("EDUCATION")
            for edu in self._content["education"]:
                if edu["school"]:
                    lines.append(edu["school"])
                if edu["degree"]:
                    lines.append(edu["degree"])
                if edu["details"]:
                    details = edu["details"]
                    if isinstance(details, list):
                        details = ", ".join(str(d) for d in details if d)
                    if details:
                        lines.append(str(details))
            lines.append("")

        return "\n".join(lines).strip()


class CoverLetterFormatter(BaseFormatter):
    """
    Professional cover letter formatter with ATS-safe formatting.

    SPEC COMPLIANCE (NON-NEGOTIABLE):
    - Margins: 1 inch all sides
    - Font size: 11pt
    - Single-spaced, one blank line between paragraphs
    - Max 1 page, 3-4 paragraphs
    - Executive tone, no enthusiasm statements or cliches
    """

    # Cover letter uses 1" margins per spec
    MARGIN_INCHES = COVER_LETTER_MARGINS_INCHES

    def __init__(self):
        super().__init__()
        self._paragraph_count = 0

    def add_header(self, name: str, tagline: str = "", contact_info: Dict[str, Any] = None):
        """Add cover letter header matching resume style."""
        # Name - large and bold, centered (same as resume)
        self._add_paragraph(name, font_size=RESUME_NAME_SIZE_PT, bold=True,
                          alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Tagline
        if tagline:
            self._add_paragraph(tagline, font_size=RESUME_SECTION_HEADER_SIZE_PT,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER,
                              space_after=4)

        # Contact info
        if contact_info:
            contact_parts = []
            if contact_info.get('email'):
                contact_parts.append(contact_info['email'])
            if contact_info.get('phone'):
                contact_parts.append(contact_info['phone'])
            if contact_info.get('location'):
                contact_parts.append(contact_info['location'])

            if contact_parts:
                contact_text = " | ".join(contact_parts)
                self._add_paragraph(contact_text, font_size=COVER_LETTER_FONT_SIZE_PT,
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER,
                                  space_after=12)

    def add_section_label(self, label: str = "COVER LETTER"):
        """Add section label similar to resume sections."""
        # Add space before
        self._add_paragraph("", space_before=12)

        # Section label with border
        p = self.doc.add_paragraph()
        p.style = self.doc.styles['Normal']
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(12)

        run = p.add_run(label)
        self._set_font(run, font_size=COVER_LETTER_FONT_SIZE_PT, bold=True)

        # Add bottom border
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)

    def add_salutation(self, recipient_name: Optional[str] = None):
        """Add greeting/salutation."""
        if recipient_name:
            salutation = f"Dear {recipient_name},"
        else:
            salutation = "Dear Hiring Manager,"

        self._add_paragraph(salutation, font_size=COVER_LETTER_FONT_SIZE_PT, space_after=12)

    def add_body_paragraph(self, text: str):
        """
        Add a body paragraph with spec-compliant formatting.

        SPEC COMPLIANCE:
        - 11pt font size
        - Single-spaced, one blank line between paragraphs
        - Executive tone: forbidden phrases are automatically filtered
        - Max 4 paragraphs
        """
        # Filter out forbidden phrases (enthusiasm, cliches)
        filtered_text = filter_cover_letter_content(text)

        # Track paragraph count for compliance
        self._paragraph_count += 1
        if self._paragraph_count > COVER_LETTER_MAX_PARAGRAPHS:
            # Log warning but still add (let the compliance check catch it)
            pass

        self._add_paragraph(filtered_text, font_size=COVER_LETTER_FONT_SIZE_PT, space_after=12)

    def add_signature(self, name: str):
        """Add closing and signature."""
        # Closing
        self._add_paragraph("Sincerely,", font_size=COVER_LETTER_FONT_SIZE_PT, space_before=12)

        # Space for signature
        self._add_paragraph("", space_before=24)

        # Name
        self._add_paragraph(name, font_size=COVER_LETTER_FONT_SIZE_PT)

    def get_paragraph_count(self) -> int:
        """Return the number of body paragraphs added."""
        return self._paragraph_count
