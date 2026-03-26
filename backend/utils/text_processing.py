"""Text processing utilities for document extraction and JSON cleaning"""

from fastapi import HTTPException


def clean_claude_json(text: str) -> str:
    """Aggressively clean Claude's response to extract valid JSON"""
    # Remove code fences
    if text.strip().startswith("```"):
        parts = text.strip().split("```")
        text = parts[1] if len(parts) > 1 else text
    # Remove "json" prefix inside code blocks
    text = text.strip()
    if text.startswith("json"):
        text = text[4:].strip()
    # Remove leading non-json content
    first_brace = text.find("{")
    if first_brace != -1:
        text = text[first_brace:]
    # Remove trailing junk after last closing brace
    last_brace = text.rfind("}")
    if last_brace != -1:
        text = text[:last_brace + 1]
    return text


def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from PDF file using PyMuPDF (fitz)"""
    try:
        import fitz  # PyMuPDF

        # Open PDF from bytes
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()

        # Clean the text
        text = text.replace('\x00', '')  # Remove null bytes

        if not text.strip():
            raise ValueError("No text extracted from PDF")

        return text.strip()
    except ImportError:
        # Fallback to PyPDF2 if fitz not available
        try:
            import PyPDF2
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
            return text
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF extraction failed: {str(e)}")
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF extraction failed: {str(e)}")


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from DOCX file using python-docx.

    Handles paragraphs, tables, text boxes, content controls,
    headers, footers, and shapes so that resume templates using
    any of these layouts are parsed correctly.
    """
    try:
        from docx import Document
        from io import BytesIO
        from lxml import etree

        print("Attempting to parse DOCX file...")
        doc = Document(BytesIO(file_bytes))

        # --- Method 1: standard paragraphs ---
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        print(f"Found {len(paragraphs)} paragraphs")

        # --- Method 2: tables (resumes often use tables for layout) ---
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    table_text.append(' | '.join(row_text))
        print(f"Found {len(table_text)} table rows")

        # --- Method 3: headers and footers ---
        header_footer_text = []
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header and header.is_linked_to_previous is False:
                    for p in header.paragraphs:
                        if p.text.strip():
                            header_footer_text.append(p.text.strip())
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer and footer.is_linked_to_previous is False:
                    for p in footer.paragraphs:
                        if p.text.strip():
                            header_footer_text.append(p.text.strip())
        print(f"Found {len(header_footer_text)} header/footer lines")

        # --- Method 4: text boxes, shapes, and content controls via raw XML ---
        xml_extra = []
        ns = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
            'v': 'urn:schemas-microsoft-com:vml',
        }
        body = doc.element.body
        # Content controls (structured document tags)
        for sdt in body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt'):
            texts = sdt.itertext()
            content = ' '.join(t.strip() for t in texts if t.strip())
            if content:
                xml_extra.append(content)
        # Text boxes inside shapes (wps:txbx and v:textbox)
        for txbx in body.iter('{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx'):
            texts = txbx.itertext()
            content = ' '.join(t.strip() for t in texts if t.strip())
            if content:
                xml_extra.append(content)
        for vtextbox in body.iter('{urn:schemas-microsoft-com:vml}textbox'):
            texts = vtextbox.itertext()
            content = ' '.join(t.strip() for t in texts if t.strip())
            if content:
                xml_extra.append(content)
        print(f"Found {len(xml_extra)} text box / content control blocks")

        # --- Combine all sources ---
        all_text = paragraphs + table_text + header_footer_text + xml_extra
        text = '\n'.join(all_text)

        # Clean the text
        text = text.replace('\x00', '')  # Remove null bytes

        # --- Fallback: brute-force all w:t elements if still empty ---
        if not text.strip():
            print("Standard extraction empty. Falling back to raw XML w:t extraction...")
            all_wt = []
            for wt in body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                if wt.text and wt.text.strip():
                    all_wt.append(wt.text)
            text = ' '.join(all_wt)
            print(f"Raw XML fallback extracted {len(text)} characters")

        if not text.strip():
            raise ValueError("No text extracted from DOCX")

        print(f"Successfully extracted {len(text)} characters from DOCX")
        return text.strip()
    except Exception as e:
        print(f"DOCX EXTRACTION ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"DOCX extraction failed: {str(e)}")
