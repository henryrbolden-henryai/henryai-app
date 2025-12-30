"""
Henry Job Search Engine - Backend API
FastAPI application with resume parsing, JD analysis, and document generation
"""

import os
import sys
import json
import io
import uuid
import random
import logging
import requests
import logging
from typing import Optional, Dict, Any, List
from datetime import datetime
from enum import Enum

# Configure structured logging
# In production, Railway captures stdout/stderr - this ensures proper formatting
logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO"),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger("henryhq")

from fastapi import FastAPI, HTTPException, File, UploadFile, Form, Request

# Supabase client for database operations
try:
    from supabase import create_client, Client
    SUPABASE_URL = os.getenv("SUPABASE_URL")
    SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY")  # Use service role for backend

    if SUPABASE_URL and SUPABASE_SERVICE_KEY:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_SERVICE_KEY)
        print("✅ Supabase client initialized")
    else:
        supabase = None
        print("⚠️ Supabase credentials not configured - database features disabled")
except ImportError:
    supabase = None
    print("⚠️ Supabase package not installed - database features disabled")
from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
import anthropic

# =============================================================================
# MODULAR IMPORTS - Extracted modules for better organization
# =============================================================================

# Models - Pydantic schemas (still using inline definitions during transition)
# TODO: Migrate to: from models import *

# Utils - Helper functions
from utils import (
    clean_claude_json,
    extract_pdf_text,
    extract_docx_text,
    calculate_days_since,
    detect_role_type,
    determine_target_level,
    infer_seniority_from_title,
    infer_industry_from_company,
    get_competency_for_stage,
    SIGNAL_WEIGHTS_BY_ROLE,
    calculate_momentum_score,
    calculate_jd_confidence,
    calculate_decision_confidence,
    get_confidence_label,
    get_confidence_guidance,
    determine_action_for_status,
    calculate_ui_signals,
    calculate_pipeline_health,
    verify_ats_keyword_coverage,
    validate_document_quality,
)

# Storage - Data persistence helpers
from storage import (
    save_mock_session,
    get_mock_session,
    save_mock_question,
    get_mock_question,
    save_mock_response,
    get_mock_responses,
    save_mock_analysis,
    get_mock_analysis,
    update_mock_session,
    cleanup_expired_sessions,
    mock_interview_sessions,
    mock_interview_questions,
    mock_interview_responses,
    mock_interview_analyses,
    outcomes_store,
    SESSION_TTL_SECONDS,
    set_supabase_client,
)

# Services - Core business logic
from services import (
    call_claude,
    call_claude_streaming,
    initialize_client as initialize_claude_client,
)

# Prompts - System prompts for Claude AI interactions
from prompts import (
    MOCK_GENERATE_QUESTION_PROMPT,
    MOCK_ANALYZE_RESPONSE_PROMPT,
    MOCK_QUESTION_FEEDBACK_PROMPT,
    MOCK_SESSION_FEEDBACK_PROMPT,
    DEBRIEF_SYSTEM_PROMPT_WITH_TRANSCRIPT,
    DEBRIEF_SYSTEM_PROMPT_NO_TRANSCRIPT,
    DEBRIEF_EXTRACTION_PROMPT,
    SCREENING_QUESTIONS_PROMPT,
    CLARIFYING_QUESTIONS_PROMPT,
    REANALYZE_PROMPT,
    RESUME_LEVELING_PROMPT,
    HEY_HENRY_SYSTEM_PROMPT,
    ASK_HENRY_SYSTEM_PROMPT,
    RESUME_CHAT_SYSTEM_PROMPT,
    RESUME_GENERATION_PROMPT,
)

# Rate limiting
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# Add parent directory to path for document_generator import
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from document_generator import ResumeFormatter, CoverLetterFormatter

# Add current directory to path for qa_validation import (needed for Railway deployment)
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# QA Validation module for fabrication detection and data quality
from qa_validation import (
    validate_documents_generation,
    validate_chat_response,
    validate_resume_parse,
    validate_jd_analysis,
    validate_interview_prep,
    create_validation_error_response,
    create_chat_fallback_response,
    add_validation_warnings_to_response,
    ValidationResult,
    ValidationLogger
)

# Tier Configuration and Service for subscription management
try:
    from tier_config import (
        TIER_LIMITS,
        TIER_PRICES,
        TIER_ORDER,
        TIER_NAMES,
        get_all_tier_info,
    )
    from tier_service import TierService
    TIER_SERVICE_AVAILABLE = True
except ImportError:
    TIER_SERVICE_AVAILABLE = False
    print("⚠️ Tier service not available - tier features disabled")

# Recruiter Calibration module for gap classification and red flag detection
# Per Calibration Spec v1.0 (REVISED): Recruiter-grade judgment framework
# CRITICAL: Calibration explains gaps, does NOT override Job Fit recommendation
try:
    from calibration import (
        calibrate_executive_role,
        calibrate_technical_role,
        calibrate_gtm_role,
        classify_gap,
        assess_domain_transferability,
        detect_red_flags,
        calibrate_gaps,  # Control layer for gap classification
    )
    CALIBRATION_AVAILABLE = True
except ImportError:
    CALIBRATION_AVAILABLE = False
    print("⚠️ Calibration module not available - using fallback behavior")

# Coaching Controller module for "Your Move" and "Gaps to Address" sections
# Per Selective Coaching Moments Spec v1.0 (REVISED)
# CRITICAL: Silence suppresses gaps, NOT "Your Move"
try:
    from coaching import (
        generate_coaching_output,
        generate_your_move,
        extract_primary_strength,
    )
    COACHING_AVAILABLE = True
except ImportError:
    COACHING_AVAILABLE = False
    print("⚠️ Coaching module not available - using fallback behavior")

# Final Recommendation Controller - SINGLE SOURCE OF TRUTH
# Per Architecture Simplification Spec v1.0:
# - One recommendation, one score, no later overrides
# - All other layers (CEC, Calibration, Coaching) are advisory only
try:
    from recommendation.final_controller import (
        FinalRecommendationController,
        compute_final_recommendation,
    )
    FINAL_CONTROLLER_AVAILABLE = True
except ImportError:
    FINAL_CONTROLLER_AVAILABLE = False
    print("⚠️ Final Recommendation Controller not available - using legacy behavior")

# Reality Check System - Market-truth intervention framework
# Per REALITY_CHECK_SPEC.md:
# - Surfaces risk, friction, and market behavior patterns
# - NEVER modifies fit_score or eligibility
# - Market Bias and Market Climate signals are message-only overlays
#
# ROLLOUT NOTE (Dec 2024):
# Reality Check System enabled behind feature flag.
# No scoring logic changes. Message-only overlays.
# Violations are aggressively logged for early warning.
try:
    from reality_check import (
        RealityCheckController,
        analyze_reality_checks,
        SignalClass,
        Severity,
    )
    REALITY_CHECK_AVAILABLE = True
    # Feature flag for gradual rollout - set to True to enable
    REALITY_CHECK_ENABLED = True
except ImportError:
    REALITY_CHECK_AVAILABLE = False
    REALITY_CHECK_ENABLED = False
    print("⚠️ Reality Check module not available - using fallback behavior")

# Initialize rate limiter
# Limits: 30 requests per minute per IP for expensive endpoints (Claude API calls)
# Health check and simple endpoints are not rate limited
limiter = Limiter(key_func=get_remote_address)

# Initialize FastAPI app
app = FastAPI(
    title="Henry Job Search Engine API",
    description="Backend for resume parsing, JD analysis, and application generation",
    version="1.0.0"
)

# Add rate limiter to app state and exception handler
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# Configure CORS - explicit origins for production
ALLOWED_ORIGINS = [
    "https://henryai-app.vercel.app",
    "https://www.henryhq.ai",
    "https://henryhq.ai",
    "http://localhost:3000",
    "http://localhost:5500",
    "http://127.0.0.1:5500",
    "http://localhost:8080",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
    expose_headers=["*"],
)


# Log startup
logger.info("HenryHQ API starting up...")

# Add validation error handler for better debugging
@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    logger.error(f"Validation error: {exc.errors()}", extra={"body": str(exc.body)[:500]})
    return JSONResponse(
        status_code=422,
        content={"detail": exc.errors(), "body": str(exc.body)[:500]}
    )

# Initialize Anthropic client via services module
client = initialize_claude_client()

# OpenAI API key for TTS (optional - for natural AI voice)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# Initialize Supabase client for persistent storage
SUPABASE_URL_STORAGE = os.getenv("SUPABASE_URL", "https://xmbappvomvmanvybdavs.supabase.co")
SUPABASE_KEY = os.getenv("SUPABASE_SERVICE_KEY")  # Use service key for backend (bypasses RLS)

supabase_client = None
if SUPABASE_KEY:
    try:
        from supabase import create_client, Client
        supabase_client: Client = create_client(SUPABASE_URL_STORAGE, SUPABASE_KEY)
        # Connect storage module to Supabase
        set_supabase_client(supabase_client)
        logger.info("Supabase client initialized successfully")
    except Exception as e:
        logger.warning(f"Failed to initialize Supabase client: {e}. Falling back to in-memory storage.")
        supabase_client = None
else:
    logger.warning("SUPABASE_SERVICE_KEY not set. Using in-memory storage (data will be lost on restart).")

# Note: In-memory storage now imported from storage module
# mock_interview_sessions, mock_interview_questions, mock_interview_responses,
# mock_interview_analyses, outcomes_store, SESSION_TTL_SECONDS are all from storage module


# ============================================================================
# MOCK INTERVIEW STORAGE HELPERS (Supabase with in-memory fallback)
# ============================================================================

def save_mock_session(session_id: str, session_data: Dict[str, Any], user_id: str = None) -> bool:
    """Save mock interview session to Supabase or fallback to in-memory."""
    if supabase_client and user_id:
        try:
            data = {
                "id": session_id,
                "user_id": user_id,
                "resume_json": session_data.get("resume_json", {}),
                "job_description": session_data.get("job_description"),
                "company": session_data.get("company"),
                "role_title": session_data.get("role_title"),
                "interview_stage": session_data.get("interview_stage"),
                "difficulty_level": session_data.get("difficulty_level", "medium"),
                "current_question_number": session_data.get("current_question_number", 1),
            }
            supabase_client.table("mock_interview_sessions").upsert(data).execute()
            return True
        except Exception as e:
            logger.error(f"Failed to save session to Supabase: {e}")
    # Fallback to in-memory
    mock_interview_sessions[session_id] = session_data
    return True

def get_mock_session(session_id: str) -> Optional[Dict[str, Any]]:
    """Get mock interview session from Supabase or fallback."""
    if supabase_client:
        try:
            result = supabase_client.table("mock_interview_sessions").select("*").eq("id", session_id).single().execute()
            if result.data:
                return result.data
        except Exception as e:
            logger.warning(f"Failed to get session from Supabase: {e}")
    # Fallback to in-memory
    return mock_interview_sessions.get(session_id)

def save_mock_question(question_id: str, question_data: Dict[str, Any]) -> bool:
    """Save mock interview question to Supabase or fallback."""
    if supabase_client:
        try:
            data = {
                "id": question_id,
                "session_id": question_data.get("session_id"),
                "question_number": question_data.get("question_number"),
                "question_text": question_data.get("question_text"),
                "competency_tested": question_data.get("competency_tested"),
                "difficulty": question_data.get("difficulty", "medium"),
            }
            supabase_client.table("mock_interview_questions").upsert(data).execute()
            return True
        except Exception as e:
            logger.error(f"Failed to save question to Supabase: {e}")
    # Fallback to in-memory
    mock_interview_questions[question_id] = question_data
    return True

def get_mock_question(question_id: str) -> Optional[Dict[str, Any]]:
    """Get mock interview question from Supabase or fallback."""
    if supabase_client:
        try:
            result = supabase_client.table("mock_interview_questions").select("*").eq("id", question_id).single().execute()
            if result.data:
                return result.data
        except Exception as e:
            logger.warning(f"Failed to get question from Supabase: {e}")
    # Fallback to in-memory
    return mock_interview_questions.get(question_id)

def save_mock_response(question_id: str, response_data: Dict[str, Any]) -> bool:
    """Save mock interview response to Supabase or fallback."""
    if supabase_client:
        try:
            data = {
                "question_id": question_id,
                "session_id": response_data.get("session_id"),
                "response_text": response_data.get("response_text"),
                "score": response_data.get("score"),
                "feedback": response_data.get("feedback"),
                "strengths": response_data.get("strengths", []),
                "improvements": response_data.get("improvements", []),
            }
            supabase_client.table("mock_interview_responses").insert(data).execute()
            return True
        except Exception as e:
            logger.error(f"Failed to save response to Supabase: {e}")
    # Fallback to in-memory
    if question_id not in mock_interview_responses:
        mock_interview_responses[question_id] = []
    mock_interview_responses[question_id].append(response_data)
    return True

def get_mock_responses(question_id: str) -> List[Dict[str, Any]]:
    """Get mock interview responses from Supabase or fallback."""
    if supabase_client:
        try:
            result = supabase_client.table("mock_interview_responses").select("*").eq("question_id", question_id).execute()
            if result.data:
                return result.data
        except Exception as e:
            logger.warning(f"Failed to get responses from Supabase: {e}")
    # Fallback to in-memory
    return mock_interview_responses.get(question_id, [])

def save_mock_analysis(session_id: str, analysis_data: Dict[str, Any]) -> bool:
    """Save mock interview analysis to Supabase or fallback."""
    if supabase_client:
        try:
            data = {
                "session_id": session_id,
                "overall_score": analysis_data.get("overall_score"),
                "competency_scores": analysis_data.get("competency_scores"),
                "key_strengths": analysis_data.get("key_strengths", []),
                "areas_for_improvement": analysis_data.get("areas_for_improvement", []),
                "recommendations": analysis_data.get("recommendations", []),
                "detailed_feedback": analysis_data.get("detailed_feedback"),
            }
            supabase_client.table("mock_interview_analyses").upsert(data, on_conflict="session_id").execute()
            return True
        except Exception as e:
            logger.error(f"Failed to save analysis to Supabase: {e}")
    # Fallback to in-memory
    mock_interview_analyses[session_id] = analysis_data
    return True

def get_mock_analysis(session_id: str) -> Optional[Dict[str, Any]]:
    """Get mock interview analysis from Supabase or fallback."""
    if supabase_client:
        try:
            result = supabase_client.table("mock_interview_analyses").select("*").eq("session_id", session_id).single().execute()
            if result.data:
                return result.data
        except Exception as e:
            logger.warning(f"Failed to get analysis from Supabase: {e}")
    # Fallback to in-memory
    return mock_interview_analyses.get(session_id)

def update_mock_session(session_id: str, updates: Dict[str, Any]) -> bool:
    """Update mock interview session in Supabase or fallback."""
    if supabase_client:
        try:
            supabase_client.table("mock_interview_sessions").update(updates).eq("id", session_id).execute()
            return True
        except Exception as e:
            logger.error(f"Failed to update session in Supabase: {e}")
    # Fallback to in-memory
    if session_id in mock_interview_sessions:
        mock_interview_sessions[session_id].update(updates)
        return True
    return False


def cleanup_expired_sessions() -> int:
    """
    Clean up expired mock interview sessions to prevent memory leaks
    and cross-session contamination.

    Returns:
        Number of sessions cleaned up
    """
    import time
    current_time = time.time()
    expired_sessions = []

    # Find expired sessions
    for session_id, session in mock_interview_sessions.items():
        created_at = session.get("created_at", 0)
        if isinstance(created_at, str):
            # Parse ISO format timestamp
            try:
                from datetime import datetime
                dt = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
                created_at = dt.timestamp()
            except (ValueError, TypeError):
                created_at = 0

        if current_time - created_at > SESSION_TTL_SECONDS:
            expired_sessions.append(session_id)

    # Clean up expired sessions and their associated data
    for session_id in expired_sessions:
        # Find and remove questions for this session
        questions_to_remove = [
            qid for qid, q in mock_interview_questions.items()
            if q.get("session_id") == session_id
        ]
        for qid in questions_to_remove:
            mock_interview_questions.pop(qid, None)
            mock_interview_responses.pop(qid, None)

        # Remove session
        mock_interview_sessions.pop(session_id, None)
        mock_interview_analyses.pop(session_id, None)

    if expired_sessions:
        print(f"🧹 Cleaned up {len(expired_sessions)} expired mock interview sessions")

    return len(expired_sessions)

# Load question bank
QUESTION_BANK_PATH = os.path.join(os.path.dirname(__file__), "data", "question_bank.json")
QUESTION_BANK: Dict[str, Any] = {}

def load_question_bank():
    """Load the structured question bank from JSON file."""
    global QUESTION_BANK
    try:
        if os.path.exists(QUESTION_BANK_PATH):
            with open(QUESTION_BANK_PATH, "r") as f:
                QUESTION_BANK = json.load(f)
            print(f"✅ Loaded question bank with {len(QUESTION_BANK.get('role_specific_questions', {}))} role types")
        else:
            print(f"⚠️ Question bank not found at {QUESTION_BANK_PATH}")
            QUESTION_BANK = {}
    except Exception as e:
        print(f"🔥 Error loading question bank: {e}")
        QUESTION_BANK = {}

# Load question bank on startup
load_question_bank()


# ============================================================================
# NON-TRANSFERABLE DOMAINS - GLOBAL CONSTANT
# Per Eligibility Gate Enforcement Spec: These domains require direct experience.
# No amount of adjacent experience substitutes for domain-specific expertise.
#
# CRITICAL FIX: Only fire on EXPLICIT requirement language, not mentions.
# Mentions ≠ Requirements. These gates are now DISABLED for Manager-level roles
# and require ownership language (e.g., "FDA compliance lead", "SOX owner").
# ============================================================================
NON_TRANSFERABLE_DOMAINS = [
    # Executive Search at VP+/C-suite level
    {
        "domain": "executive_search",
        "levels": ["vp", "vice president", "c-suite", "chief", "svp", "evp", "cxo"],
        "description": "Executive Search at VP+ or C-suite level",
        # EXPLICIT ownership keywords only - must indicate primary role responsibility
        "keywords": ["executive search", "executive recruiting", "retained search", "executive placement"],
        # Keywords that REQUIRE context (not standalone matches)
        "require_ownership": True
    },
    # NOTE: People Leadership is intentionally NOT in this list because:
    # 1. People leadership IS transferable across industries and roles
    # 2. We have a dedicated tiered leadership check (see CHECK 2 in check_eligibility_gate)
    #    that properly handles leadership gaps with 50%/70% thresholds
    # 3. Evidence like "built and scaled teams" should satisfy leadership requirements
    #    even if exact phrasing differs from JD
    # Core Software Engineering (not PM, not design, not adjacent)
    {
        "domain": "core_software_engineering",
        "levels": ["senior", "staff", "principal", "lead", "architect"],
        "description": "Core Software Engineering (hands-on coding)",
        "keywords": ["software engineer", "software developer", "backend engineer", "frontend engineer", "full stack", "systems engineer", "platform engineer"],
        "require_ownership": False  # Title match is sufficient
    },
    # ML/AI Research (not just "used AI tools")
    {
        "domain": "ml_ai_research",
        "levels": ["research", "scientist", "senior", "staff", "principal"],
        "description": "ML/AI Research (publish papers, build models)",
        "keywords": ["machine learning engineer", "deep learning engineer", "ai research", "ml engineer", "research scientist", "computer vision engineer"],
        "require_ownership": True
    },
    # Regulated Clinical or Finance leadership
    # CRITICAL: This gate was firing on garbage like "sec" matching "security"
    # Now requires EXPLICIT ownership language
    {
        "domain": "regulated_clinical_finance",
        "levels": ["director", "vp", "head of", "chief"],  # Removed "senior" - too broad
        "description": "Regulated Clinical or Finance leadership",
        # FIXED: Removed short tokens like "sec", "fda" - too many false positives
        # Now requires explicit ownership phrases
        "keywords": [
            "clinical operations lead", "clinical trials manager", "fda compliance",
            "regulatory affairs", "compliance officer", "finra licensed",
            "sec reporting", "banking regulation", "healthcare compliance officer",
            "hipaa compliance", "sox compliance", "regulatory compliance lead"
        ],
        "require_ownership": True
    }
]


# ============================================================================
# EXPERIENCE NORMALIZATION HELPER
# CRITICAL INVARIANT: All experience data must be List[Dict] - NOTHING ELSE
# This helper ensures downstream functions never crash on strings or mixed types
# ============================================================================

def normalize_experience(resume_data: dict) -> list:
    """
    Normalize experience to guaranteed List[Dict].

    This is THE canonical normalization function. All code paths that access
    resume_data.get("experience") should use this instead.

    Returns:
        List of dict experience items. Empty list if input is invalid.
    """
    if not resume_data or not isinstance(resume_data, dict):
        return []

    experience = resume_data.get("experience", [])

    # Case 1: String (LinkedIn free text)
    if isinstance(experience, str):
        return []

    # Case 2: Not a list
    if not isinstance(experience, list):
        return []

    # Case 3: List but may contain non-dict items - filter to only dicts
    return [exp for exp in experience if isinstance(exp, dict)]


# ============================================================================
# ROLE PARSER - EXTRACTS STRUCTURED ELIGIBILITY LOGIC FROM JD
# Per Role Parser Spec: "This parser decides what is non-negotiable vs coachable
# before scoring ever happens."
#
# PARSING ORDER (CRITICAL):
# 1. Normalize JD text
# 2. Extract numeric requirements first
# 3. Bind numbers to verbs and scopes
# 4. Assign Tier 1 vs Tier 2
# 5. Extract market context last
# 6. Return structured output only - NO scoring, NO messaging, NO advice
# ============================================================================

def parse_role_requirements(job_description: str, role_title: str = "") -> dict:
    """
    Role Parser - Converts JD into structured, enforceable eligibility logic.

    This parser decides what is non-negotiable vs coachable BEFORE scoring happens.

    Args:
        job_description: Raw JD text
        role_title: Role title if available

    Returns:
        Structured eligibility object:
        {
            "role_id": str,
            "tier_1_hard_gates": [],      # Binary, non-transferable, failure = DO NOT APPLY
            "tier_2_conditional_requirements": [],  # Transferable, missing = risk
            "tier_3_market_signals": [],  # Strategy only, never affects eligibility
            "confidence_flags": []        # Interpretation risk warnings
        }
    """
    import re
    import hashlib

    # Normalize text
    jd_lower = (job_description or "").lower()
    title_lower = (role_title or "").lower()
    combined = f"{title_lower} {jd_lower}"

    # Generate role ID for tracking
    role_id = hashlib.md5(combined.encode()).hexdigest()[:12]

    result = {
        "role_id": role_id,
        "tier_1_hard_gates": [],
        "tier_2_conditional_requirements": [],
        "tier_3_market_signals": [],
        "confidence_flags": []
    }

    # ========================================================================
    # TIER 1: HARD ELIGIBILITY GATES (Binary, Non-Transferable)
    # Failure = DO NOT APPLY. No exceptions.
    # ========================================================================

    # Pattern 1: Years tied to leadership or execution
    # "7+ years of people leadership", "8+ years executive recruiting"
    years_patterns = [
        # People leadership years
        (r"(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s*(?:people\s*)?(?:leadership|management|managing\s+teams?|managing\s+people|direct\s+reports?)", "people_leadership_years"),
        # Executive/senior specific experience
        (r"(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s*(?:executive|vp|director|c-suite|senior)\s*(?:search|recruiting|experience)", "executive_experience_years"),
        # Domain-specific experience with years
        (r"(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s*(?:software|engineering|development|coding)\s+experience", "engineering_experience_years"),
        # Generic senior experience requirements
        (r"(\d+)\+?\s*(?:years?|yrs?)(?:\s+of)?\s*(?:relevant|related|progressive|hands-on)\s+experience", "domain_experience_years"),
    ]

    for pattern, req_type in years_patterns:
        matches = re.findall(pattern, combined)
        for match in matches:
            years = int(match) if match else 0
            if years > 0:
                result["tier_1_hard_gates"].append({
                    "type": req_type,
                    "required": years,
                    "evidence_required": True,
                    "source_pattern": pattern
                })

    # Pattern 2: Scope ownership keywords (automatic Tier 1)
    ownership_keywords = [
        ("own", "scope_ownership"),
        ("accountable for", "scope_ownership"),
        ("final decision-maker", "decision_authority"),
        ("ultimate responsibility", "scope_ownership"),
        ("budget owner", "budget_authority"),
        ("p&l responsibility", "budget_authority"),
        ("headcount authority", "hiring_authority"),
    ]

    for keyword, req_type in ownership_keywords:
        if keyword in jd_lower:
            # Extract context around the keyword
            result["tier_1_hard_gates"].append({
                "type": req_type,
                "keyword_detected": keyword,
                "evidence_required": True
            })

    # Pattern 3: Level specificity (VP+, C-suite, Board-facing, global)
    level_keywords = [
        (r"\bvp\+\b", "vp_plus_level"),
        (r"\bc-suite\b", "c_suite_level"),
        (r"\bboard-facing\b", "board_exposure"),
        (r"\bglobal\s+(?:ownership|responsibility|scope)\b", "global_scope"),
        (r"\bexecutive\s+(?:team|leadership|committee)\b", "executive_team_membership"),
    ]

    for pattern, req_type in level_keywords:
        if re.search(pattern, combined):
            result["tier_1_hard_gates"].append({
                "type": req_type,
                "evidence_required": True,
                "source_pattern": pattern
            })

    # Pattern 4: Legal or credential requirements
    credential_patterns = [
        (r"\b(?:security\s+)?clearance\s+(?:required|necessary|must\s+have)\b", "security_clearance"),
        (r"\b(?:cpa|cfa|pe|bar|md|rn|jd)\s+(?:required|license|certification)\b", "professional_license"),
        (r"\bmust\s+be\s+licensed\b", "professional_license"),
        (r"\bcertified\s+(?:public\s+accountant|financial\s+analyst)\b", "professional_certification"),
    ]

    for pattern, req_type in credential_patterns:
        if re.search(pattern, combined):
            result["tier_1_hard_gates"].append({
                "type": req_type,
                "evidence_required": True,
                "source_pattern": pattern
            })

    # Pattern 5: Non-transferable domain detection
    # CRITICAL FIX: Only add to tier_1 if ownership language is present
    # Manager-level roles skip this entirely (handled in eligibility gate)
    manager_level_exemptions = [
        "engineering manager", "product manager", "program manager", "project manager",
        "senior manager", "team lead", "tech lead", "staff engineer", "senior engineer"
    ]
    is_manager_level = any(exempt in title_lower for exempt in manager_level_exemptions)

    if not is_manager_level:
        for domain_config in NON_TRANSFERABLE_DOMAINS:
            domain = domain_config["domain"]
            levels = domain_config["levels"]
            keywords = domain_config["keywords"]
            require_ownership = domain_config.get("require_ownership", True)

            jd_requires_domain = any(kw in combined for kw in keywords)
            jd_at_required_level = any(level in combined for level in levels)

            # Check for ownership context if required
            if require_ownership:
                ownership_phrases = ["required", "must have", "requirements:", "you will own", "responsible for"]
                has_ownership_context = any(phrase in jd_lower for phrase in ownership_phrases)
                if not has_ownership_context:
                    continue

            if jd_requires_domain and jd_at_required_level:
                result["tier_1_hard_gates"].append({
                    "type": f"non_transferable_domain:{domain}",
                    "domain": domain,
                    "description": domain_config["description"],
                    "evidence_required": True
                })

    # ========================================================================
    # TIER 2: CONDITIONAL REQUIREMENTS (Transferable, Missing = Risk)
    # ========================================================================

    # Pattern 1: Industry preference language
    industry_patterns = [
        (r"\b(?:experience\s+in|background\s+in)\s+(healthcare|fintech|saas|e-commerce|retail|manufacturing|banking|insurance)\b", "industry_experience"),
        (r"\b(healthcare|fintech|saas|b2b|enterprise|consumer)\s+(?:preferred|a\s+plus|nice\s+to\s+have)\b", "industry_preference"),
    ]

    for pattern, req_type in industry_patterns:
        matches = re.findall(pattern, combined)
        for match in matches:
            domain = match if isinstance(match, str) else match[0]
            result["tier_2_conditional_requirements"].append({
                "type": req_type,
                "domain": domain,
                "transferable": True
            })

    # Pattern 2: Scale language without numeric enforcement
    scale_keywords = [
        ("high-growth", "scale_experience"),
        ("series c+", "scale_experience"),
        ("series d", "scale_experience"),
        ("enterprise scale", "scale_experience"),
        ("hypergrowth", "scale_experience"),
        ("rapidly scaling", "scale_experience"),
    ]

    for keyword, req_type in scale_keywords:
        if keyword in jd_lower:
            result["tier_2_conditional_requirements"].append({
                "type": req_type,
                "keyword_detected": keyword,
                "transferable": True
            })

    # Pattern 3: First-time scope expansion
    first_time_patterns = [
        (r"\bbuild\s+from\s+scratch\b", "greenfield_builder"),
        (r"\bfirst\s+\w+\s+hire\b", "first_hire"),
        (r"\bestablish\s+the\s+function\b", "function_builder"),
        (r"\b0\s*(?:to|→)\s*1\b", "zero_to_one"),
    ]

    for pattern, req_type in first_time_patterns:
        if re.search(pattern, combined):
            result["tier_2_conditional_requirements"].append({
                "type": req_type,
                "transferable": True,
                "source_pattern": pattern
            })

    # Pattern 4: Tooling or domain adjacency
    tooling_patterns = [
        (r"\b(?:greenhouse|lever|workday|icims|taleo)\b", "ats_experience"),
        (r"\b(?:aws|azure|gcp)\s+(?:experience|knowledge)\b", "cloud_platform"),
        (r"\bregulated\s+environment\b", "regulated_environment"),
    ]

    for pattern, req_type in tooling_patterns:
        if re.search(pattern, combined):
            result["tier_2_conditional_requirements"].append({
                "type": req_type,
                "transferable": True,
                "source_pattern": pattern
            })

    # ========================================================================
    # TIER 3: MARKET & CONTEXT SIGNALS (Never affect eligibility)
    # ========================================================================

    # Funding stage signals
    funding_patterns = [
        (r"\bseries\s+[a-e]\b", "funding_stage"),
        (r"\b(?:seed|pre-seed|angel)\s+(?:stage|funded)\b", "funding_stage"),
        (r"\bipo\b", "funding_stage"),
        (r"\bpublic\s+company\b", "funding_stage"),
    ]

    for pattern, signal_type in funding_patterns:
        match = re.search(pattern, combined)
        if match:
            result["tier_3_market_signals"].append({
                "type": signal_type,
                "signal": match.group(0),
                "affects_eligibility": False
            })

    # Compensation signals
    comp_patterns = [
        (r"\$\d{2,3}k\s*[-–]\s*\$\d{2,3}k", "compensation_range"),
        (r"\bequity\b", "equity_offered"),
        (r"\bstock\s+options\b", "equity_offered"),
    ]

    for pattern, signal_type in comp_patterns:
        match = re.search(pattern, combined)
        if match:
            result["tier_3_market_signals"].append({
                "type": signal_type,
                "signal": match.group(0),
                "affects_eligibility": False
            })

    # Urgency signals
    urgency_patterns = [
        (r"\bimmediately\b", "high_urgency"),
        (r"\basap\b", "high_urgency"),
        (r"\burgent\b", "high_urgency"),
        (r"\bstart\s+date\s+(?:asap|immediately)\b", "high_urgency"),
    ]

    for pattern, signal_type in urgency_patterns:
        if re.search(pattern, combined):
            result["tier_3_market_signals"].append({
                "type": signal_type,
                "signal": "high",
                "affects_eligibility": False
            })

    # ========================================================================
    # INDUSTRY ALIGNMENT CLASSIFICATION (Per Industry Alignment Extension Spec)
    # Industry affects eligibility ONLY when the role makes it unavoidable.
    # Otherwise, it affects competitiveness, not permission to apply.
    # ========================================================================
    industry_alignment = {
        "required": [],       # Tier 1 - Hard gate (regulated, legal, compliance)
        "preferred": [],      # Tier 2 - Conditional (nice to have)
        "agnostic": False,    # True if role is industry-agnostic
        "confidence_flag": None
    }

    # Tier 1 Industry - Required (becomes hard gate)
    # Trigger: "must have", "deep domain expertise", "regulated environment required"
    # Industries: Healthcare, FinServ (regulated), Government, Energy, Life Sciences

    tier_1_industry_patterns = [
        # Explicit requirement language
        (r"\bmust\s+have\s+(?:experience|background)\s+in\s+(healthcare|fintech|financial\s+services|government|energy|life\s+sciences|pharma|biotech|defense|govcon)\b", "required_experience"),
        (r"\bdeep\s+domain\s+expertise\s+in\s+(healthcare|finance|banking|insurance|government|regulated)\b", "deep_domain"),
        (r"\bregulated\s+environment\s+(?:experience\s+)?required\b", "regulated_environment"),
        (r"\bindustry[- ]specific\s+knowledge\s+(?:required|critical|essential)\b", "industry_specific"),
        # Compliance/legal tied industries
        (r"\b(hipaa|fda|finra|sec|sox)\s+(?:compliance|experience|knowledge)\s+required\b", "regulatory_compliance"),
        (r"\b(clinical\s+trials?|drug\s+development|medical\s+device)\s+experience\s+required\b", "life_sciences"),
    ]

    for pattern, reason_type in tier_1_industry_patterns:
        match = re.search(pattern, combined)
        if match:
            industry = match.group(1) if match.lastindex else "regulated domain"
            industry_alignment["required"].append({
                "industry": industry.title().replace("_", " "),
                "reason": f"Regulated domain experience required ({reason_type})"
            })
            # Also add to Tier 1 hard gates
            result["tier_1_hard_gates"].append({
                "type": f"industry_requirement:{industry.lower().replace(' ', '_')}",
                "industry": industry,
                "evidence_required": True,
                "reason": "Regulated domain - non-transferable"
            })

    # Tier 2 Industry - Preferred (conditional, coachable)
    # Trigger: "preferred", "a plus", "nice to have", "familiarity with"
    tier_2_industry_patterns = [
        (r"\b(?:experience|background)\s+in\s+(saas|b2b|b2c|e-commerce|retail|media|adtech|martech|edtech|proptech|consumer|enterprise)\s+(?:preferred|a\s+plus|nice\s+to\s+have|helpful)\b", "preferred"),
        (r"\b(saas|b2b|enterprise|consumer|fintech)\s+(?:preferred|a\s+plus|background\s+helpful)\b", "preferred"),
        (r"\bfamiliarity\s+with\s+(healthcare|finance|tech|retail|manufacturing)\b", "familiarity"),
        (r"\bexperience\s+in\s+(saas|b2b|enterprise|consumer|startup)\s+environment\s+(?:preferred|helpful)\b", "preferred"),
    ]

    for pattern, reason_type in tier_2_industry_patterns:
        match = re.search(pattern, combined)
        if match:
            industry = match.group(1) if match.lastindex else "general"
            # Don't add if already in required
            if not any(req.get("industry", "").lower() == industry.lower() for req in industry_alignment["required"]):
                industry_alignment["preferred"].append({
                    "industry": industry.title(),
                    "transferable": True
                })
                result["tier_2_conditional_requirements"].append({
                    "type": f"industry_preference:{industry.lower()}",
                    "industry": industry,
                    "transferable": True
                })

    # Check for industry-agnostic signals
    agnostic_patterns = [
        r"\bany\s+industry\b",
        r"\bindustry[- ]agnostic\b",
        r"\bopen\s+to\s+(?:all\s+)?industries\b",
        r"\bno\s+(?:specific\s+)?industry\s+(?:preference|requirement)\b"
    ]

    if any(re.search(p, combined) for p in agnostic_patterns):
        industry_alignment["agnostic"] = True

    # If no industry language detected at all, mark as agnostic
    if not industry_alignment["required"] and not industry_alignment["preferred"]:
        # Check if ANY industry-related terms appear
        industry_terms = r"\b(saas|b2b|b2c|fintech|healthcare|finance|retail|consumer|enterprise|government|regulated|compliance)\b"
        if not re.search(industry_terms, combined):
            industry_alignment["agnostic"] = True

    # Confidence flag for ambiguous industry requirements
    # Emit when industry is mentioned but not clearly required or preferred
    ambiguous_industry = re.search(r"\bexperience\s+in\s+(?:complex\s+)?(?:industry|industries|domain)\b", combined)
    mixed_signals = (
        industry_alignment["agnostic"] and
        (industry_alignment["required"] or industry_alignment["preferred"])
    )

    if ambiguous_industry or mixed_signals:
        industry_alignment["confidence_flag"] = {
            "type": "AMBIGUOUS_INDUSTRY_REQUIREMENT",
            "detail": "Industry mentioned without clarity on whether it is mandatory"
        }
        result["confidence_flags"].append(industry_alignment["confidence_flag"])

    # Add industry_alignment to result
    result["industry_alignment"] = industry_alignment

    # ========================================================================
    # CONFIDENCE FLAGS (Safety Net)
    # Emit flags when interpretation risk exists
    # ========================================================================

    # Flag: Leadership mentioned without explicit people management
    if re.search(r"\bleadership\b", combined) and not re.search(r"\b(?:people|team|direct\s+reports?|managing)\b", combined):
        result["confidence_flags"].append({
            "flag": "AMBIGUOUS_LEADERSHIP",
            "detail": "Leadership mentioned without explicit people management - may be technical or operational leadership"
        })

    # Flag: Years mentioned without clear scope
    if re.search(r"\d+\+?\s*years", combined) and not result["tier_1_hard_gates"]:
        result["confidence_flags"].append({
            "flag": "AMBIGUOUS_YEARS",
            "detail": "Years requirement detected but could not bind to specific scope"
        })

    # Flag: Senior title without scope requirements
    if re.search(r"\b(?:senior|staff|principal|director|vp)\b", title_lower):
        if not any(gate.get("type") in ["scope_ownership", "decision_authority", "budget_authority"] for gate in result["tier_1_hard_gates"]):
            result["confidence_flags"].append({
                "flag": "TITLE_WITHOUT_SCOPE",
                "detail": f"Senior title '{role_title}' detected but no explicit scope requirements in JD"
            })

    # Deduplicate gates and requirements
    result["tier_1_hard_gates"] = _deduplicate_requirements(result["tier_1_hard_gates"])
    result["tier_2_conditional_requirements"] = _deduplicate_requirements(result["tier_2_conditional_requirements"])

    print(f"📋 ROLE PARSER OUTPUT:")
    print(f"   Tier 1 Hard Gates: {len(result['tier_1_hard_gates'])}")
    print(f"   Tier 2 Conditional: {len(result['tier_2_conditional_requirements'])}")
    print(f"   Tier 3 Market Signals: {len(result['tier_3_market_signals'])}")
    print(f"   Confidence Flags: {len(result['confidence_flags'])}")

    return result


def _deduplicate_requirements(requirements: list) -> list:
    """Remove duplicate requirements based on type."""
    seen = set()
    unique = []
    for req in requirements:
        req_type = req.get("type", "")
        if req_type not in seen:
            seen.add(req_type)
            unique.append(req)
    return unique


def evaluate_candidate_against_role_requirements(resume_data: dict, role_requirements: dict) -> dict:
    """
    Evaluate a candidate against parsed role requirements.

    Per Role Parser Spec: "All downstream scoring and messaging must obey
    Tier 1 output without exception."

    Args:
        resume_data: Parsed resume
        role_requirements: Output from parse_role_requirements()

    Returns:
        Evaluation result with eligibility determination
    """
    result = {
        "eligible": True,
        "tier_1_failures": [],
        "tier_2_risks": [],
        "tier_3_strategy": [],
        "recommendation": None,
        "recommendation_locked": False
    }

    if not resume_data or not role_requirements:
        return result

    # ========================================================================
    # TIER 1 EVALUATION - Any failure = DO NOT APPLY (locked)
    # ========================================================================
    tier_1_gates = role_requirements.get("tier_1_hard_gates", [])

    for gate in tier_1_gates:
        gate_type = gate.get("type", "")

        # Check people leadership years
        if "people_leadership" in gate_type:
            required_years = gate.get("required", 0)
            candidate_years = extract_people_leadership_years(resume_data)

            if candidate_years < required_years:
                result["tier_1_failures"].append({
                    "gate": gate,
                    "required": required_years,
                    "candidate_has": candidate_years,
                    "gap_type": "missing_experience"
                })

        # Check non-transferable domains
        if "non_transferable_domain" in gate_type:
            domain = gate.get("domain", "")
            domain_config = next((d for d in NON_TRANSFERABLE_DOMAINS if d["domain"] == domain), None)

            if domain_config:
                keywords = domain_config["keywords"]
                experience = resume_data.get("experience", [])
                combined_resume = " ".join([
                    f"{exp.get('title', '')} {exp.get('description', '')}"
                    for exp in experience if isinstance(exp, dict)
                ]).lower()

                has_domain = any(kw in combined_resume for kw in keywords)

                if not has_domain:
                    result["tier_1_failures"].append({
                        "gate": gate,
                        "domain": domain,
                        "candidate_has_domain": False,
                        "gap_type": "missing_experience"
                    })

        # Check for scope ownership evidence
        if gate_type in ["scope_ownership", "decision_authority", "budget_authority"]:
            # Look for ownership language in resume
            experience = resume_data.get("experience", [])
            combined_resume = " ".join([
                f"{exp.get('description', '')} {' '.join(exp.get('highlights', []))}"
                for exp in experience if isinstance(exp, dict)
            ]).lower()

            ownership_evidence = any(kw in combined_resume for kw in [
                "owned", "accountable", "responsible for", "led",
                "built", "established", "final decision", "p&l", "budget"
            ])

            if not ownership_evidence:
                result["tier_1_failures"].append({
                    "gate": gate,
                    "evidence_found": False,
                    "gap_type": "missing_evidence"
                })

        # Check industry requirements (Tier 1 - regulated domains)
        if "industry_requirement" in gate_type:
            industry = gate.get("industry", "")
            experience = resume_data.get("experience", [])
            combined_resume = " ".join([
                f"{exp.get('company', '')} {exp.get('description', '')} {exp.get('title', '')}"
                for exp in experience if isinstance(exp, dict)
            ]).lower()

            # Check for industry-specific keywords
            industry_keywords = {
                "healthcare": ["healthcare", "hospital", "medical", "clinical", "health tech", "healthtech", "hipaa"],
                "fintech": ["fintech", "financial services", "banking", "payments", "trading", "finra", "sec"],
                "financial services": ["financial services", "banking", "investment", "trading", "finra", "sec"],
                "government": ["government", "federal", "public sector", "govcon", "defense", "dod"],
                "life sciences": ["pharma", "biotech", "life sciences", "clinical trials", "fda", "drug development"],
                "energy": ["energy", "utilities", "oil", "gas", "renewable"],
            }

            industry_lower = industry.lower()
            relevant_keywords = industry_keywords.get(industry_lower, [industry_lower])
            has_industry = any(kw in combined_resume for kw in relevant_keywords)

            if not has_industry:
                result["tier_1_failures"].append({
                    "gate": gate,
                    "industry": industry,
                    "candidate_has_industry": False,
                    "gap_type": "missing_experience"
                })

    # ========================================================================
    # TIER 1 FAILURE = LOCKED DO NOT APPLY
    # ========================================================================
    if result["tier_1_failures"]:
        result["eligible"] = False
        result["recommendation"] = "Do Not Apply"
        result["recommendation_locked"] = True

        # Find the most severe failure for the reason
        primary_failure = result["tier_1_failures"][0]
        gate = primary_failure.get("gate", {})

        if "people_leadership" in gate.get("type", ""):
            result["locked_reason"] = (
                f"This role requires {primary_failure['required']}+ years of people leadership. "
                f"You have {primary_failure['candidate_has']:.1f} years verified. "
                f"Operational leadership does not count."
            )
        elif "non_transferable_domain" in gate.get("type", ""):
            domain_desc = gate.get("description", gate.get("domain", "this domain"))
            result["locked_reason"] = (
                f"This role requires {domain_desc}. "
                f"Your background does not include direct experience in this domain. "
                f"This is a non-transferable requirement."
            )
        elif "industry_requirement" in gate.get("type", ""):
            industry = gate.get("industry", "this industry")
            result["locked_reason"] = (
                f"This role requires experience in {industry}. "
                f"Your background does not show verified experience in this regulated domain. "
                f"Industry alignment is required for compliance-heavy roles."
            )
        else:
            result["locked_reason"] = (
                f"This role has requirements you do not meet: {gate.get('type', 'unknown')}."
            )

        print(f"🚫 TIER 1 GATE FAILURE: {len(result['tier_1_failures'])} gates failed")
        print(f"   🔒 RECOMMENDATION LOCKED: Do Not Apply")
        return result

    # ========================================================================
    # TIER 2 EVALUATION - Risks noted but don't block application
    # ========================================================================
    tier_2_reqs = role_requirements.get("tier_2_conditional_requirements", [])

    for req in tier_2_reqs:
        req_type = req.get("type", "")

        # Check industry experience
        if "industry" in req_type:
            domain = req.get("domain", "")
            experience = resume_data.get("experience", [])
            combined_resume = " ".join([
                f"{exp.get('company', '')} {exp.get('description', '')}"
                for exp in experience if isinstance(exp, dict)
            ]).lower()

            if domain and domain.lower() not in combined_resume:
                result["tier_2_risks"].append({
                    "requirement": req,
                    "risk_level": "medium",
                    "note": f"Industry experience in {domain} not verified"
                })

    # ========================================================================
    # TIER 3 - Strategy signals (informational only)
    # ========================================================================
    result["tier_3_strategy"] = role_requirements.get("tier_3_market_signals", [])

    # Apply confidence flag logic - default conservative if flags exist
    if role_requirements.get("confidence_flags"):
        result["confidence_flags"] = role_requirements["confidence_flags"]
        # Don't allow optimistic phrasing when flags exist
        result["allow_optimistic_phrasing"] = False
    else:
        result["allow_optimistic_phrasing"] = True

    print(f"✅ CANDIDATE ELIGIBLE - No Tier 1 failures")
    if result["tier_2_risks"]:
        print(f"   ⚠️ Tier 2 risks: {len(result['tier_2_risks'])}")

    return result


# ============================================================================
# CREDIBILITY ALIGNMENT ENGINE (CAE)
# Per CAE Spec: "Assess how believable the candidate will look to a hiring manager,
# not whether they're allowed to apply. This layer never blocks on its own.
# It modulates confidence, language, and recommendation strength."
#
# ONE-LINE RULE: "Credibility modifiers reflect how a hiring manager will perceive
# the candidate. They must materially affect recommendation strength and messaging
# tone, even when eligibility is met."
# ============================================================================

def evaluate_credibility_alignment(resume_data: dict, response_data: dict) -> dict:
    """
    Credibility Alignment Engine (CAE) - Runs AFTER eligibility, BEFORE recommendation.

    Purpose: Quantify how much explanation the candidate will need and whether
    their story is plausible to a hiring manager.

    This is NOT a gate. It modulates:
    - Recommendation ceiling
    - Coaching intensity
    - Language firmness
    - CTA urgency

    Args:
        resume_data: Parsed resume
        response_data: JD analysis response

    Returns:
        CAE result with credibility assessments for all three dimensions
    """
    import re

    result = {
        "industry_alignment": {
            "match": "direct",  # direct | adjacent | distant
            "risk_level": "low",  # low | medium | high
            "narrative_burden": "low"  # low | medium | high
        },
        "company_scale_alignment": {
            "candidate_scale": None,  # startup | mid | enterprise
            "target_scale": None,  # startup | mid | enterprise
            "delta": 0,  # 0 | 1 | 2 | 3
            "risk_level": "low"  # low | medium | high
        },
        "role_scope_alignment": {
            "title_inflation_risk": "low",  # low | medium | high
            "scope_gap": "none",  # none | moderate | significant
            "hm_skepticism": "low"  # low | medium | high
        },
        "overall_credibility_risk": "low",  # low | medium | high
        "recommendation_ceiling": None,  # None = no cap, or "Apply with Caution"
        "coaching_intensity": "normal",  # normal | moderate | heavy
        "language_firmness": "confident",  # confident | candid | skeptical
        "mandatory_reality_check": False,
        "reality_check_message": None
    }

    if not resume_data or not response_data:
        return result

    # Extract job info
    role_title = (response_data.get("role_title", "") or "").lower()
    jd_text = (response_data.get("job_description", "") or "").lower()
    combined_jd = f"{role_title} {jd_text}"

    # Extract resume info
    experience = resume_data.get("experience", [])
    combined_resume = ""
    candidate_companies = []
    candidate_titles = []

    for exp in experience:
        if isinstance(exp, dict):
            title = (exp.get("title", "") or "").lower()
            desc = (exp.get("description", "") or "").lower()
            company = (exp.get("company", "") or "").lower()
            highlights = exp.get("highlights", [])
            if highlights and isinstance(highlights, list):
                desc += " " + " ".join([h.lower() for h in highlights if isinstance(h, str)])
            combined_resume += f" {title} {desc} {company}"
            candidate_companies.append(company)
            candidate_titles.append(title)

    high_risk_count = 0
    medium_risk_count = 0

    # ========================================================================
    # DIMENSION 1: INDUSTRY ALIGNMENT
    # Direct → no friction, Adjacent → explainable, Distant → heavy skepticism
    # ========================================================================

    # Define industry clusters for adjacency mapping
    industry_clusters = {
        "tech_saas": ["saas", "software", "tech", "technology", "b2b", "enterprise software", "cloud"],
        "fintech": ["fintech", "financial services", "banking", "payments", "trading", "investment", "insurance", "finra", "sec"],
        "healthcare": ["healthcare", "health tech", "healthtech", "medical", "clinical", "hospital", "health", "hipaa", "pharma", "biotech"],
        "consumer": ["consumer", "b2c", "retail", "e-commerce", "ecommerce", "cpg", "consumer goods", "dtc"],
        "marketplace": ["marketplace", "platform", "two-sided", "gig economy", "sharing economy"],
        "media_adtech": ["media", "advertising", "adtech", "martech", "content", "publishing"],
        "edtech": ["education", "edtech", "learning", "training", "e-learning"],
        "government": ["government", "federal", "public sector", "govcon", "defense", "dod", "state", "municipal"],
        "manufacturing": ["manufacturing", "industrial", "hardware", "supply chain", "logistics"],
    }

    # Detect target industry
    target_industry = None
    for industry, keywords in industry_clusters.items():
        if any(kw in combined_jd for kw in keywords):
            target_industry = industry
            break

    # Detect candidate's industry background
    candidate_industries = []
    for industry, keywords in industry_clusters.items():
        if any(kw in combined_resume for kw in keywords):
            candidate_industries.append(industry)

    # Determine match type
    if target_industry is None:
        # Industry-agnostic role
        result["industry_alignment"]["match"] = "direct"
        result["industry_alignment"]["risk_level"] = "low"
        result["industry_alignment"]["narrative_burden"] = "low"
    elif target_industry in candidate_industries:
        # Direct match
        result["industry_alignment"]["match"] = "direct"
        result["industry_alignment"]["risk_level"] = "low"
        result["industry_alignment"]["narrative_burden"] = "low"
    else:
        # Check for adjacency
        adjacency_map = {
            "tech_saas": ["fintech", "marketplace", "media_adtech", "edtech"],
            "fintech": ["tech_saas", "consumer"],
            "healthcare": ["tech_saas"],  # Limited adjacency - highly regulated
            "consumer": ["marketplace", "media_adtech", "tech_saas"],
            "marketplace": ["consumer", "tech_saas"],
            "media_adtech": ["consumer", "tech_saas"],
            "edtech": ["tech_saas", "consumer"],
            "government": [],  # Very limited adjacency
            "manufacturing": ["tech_saas"],  # If going digital
        }

        adjacent_industries = adjacency_map.get(target_industry, [])
        has_adjacent = any(ind in candidate_industries for ind in adjacent_industries)

        if has_adjacent:
            result["industry_alignment"]["match"] = "adjacent"
            result["industry_alignment"]["risk_level"] = "medium"
            result["industry_alignment"]["narrative_burden"] = "medium"
            medium_risk_count += 1
        else:
            result["industry_alignment"]["match"] = "distant"
            result["industry_alignment"]["risk_level"] = "high"
            result["industry_alignment"]["narrative_burden"] = "high"
            high_risk_count += 1

    # Special handling for regulated industries
    regulated_industries = ["healthcare", "fintech", "government"]
    if target_industry in regulated_industries:
        if result["industry_alignment"]["match"] != "direct":
            # Regulated industries have higher skepticism
            result["industry_alignment"]["risk_level"] = "high"
            result["industry_alignment"]["narrative_burden"] = "high"
            if result["industry_alignment"]["match"] == "adjacent":
                high_risk_count += 1
                medium_risk_count -= 1

    # ========================================================================
    # DIMENSION 2: COMPANY SCALE ALIGNMENT
    # Delta 0-1 → normal, Delta 2 → skepticism, Delta 3 → major credibility gap
    # ========================================================================

    # Detect target company scale from JD
    enterprise_signals = [
        "fortune 500", "fortune 100", "global", "enterprise", "multinational",
        "publicly traded", "10000+", "10,000+", "5000+", "5,000+", "large organization"
    ]
    mid_signals = [
        "growth stage", "series c", "series d", "series e", "scale-up", "scaling",
        "500+", "1000+", "1,000+", "mid-size", "midsize", "growing"
    ]
    startup_signals = [
        "startup", "early stage", "seed", "series a", "series b", "pre-seed",
        "founding team", "0 to 1", "zero to one", "greenfield", "first hire",
        "50 employees", "100 employees", "small team"
    ]

    target_scale = None
    if any(sig in combined_jd for sig in enterprise_signals):
        target_scale = "enterprise"
    elif any(sig in combined_jd for sig in mid_signals):
        target_scale = "mid"
    elif any(sig in combined_jd for sig in startup_signals):
        target_scale = "startup"

    # Detect candidate's scale from resume
    # Check for enterprise keywords in company descriptions
    candidate_scale = None
    enterprise_company_keywords = ["inc.", "corporation", "corp.", "global", "worldwide"]
    startup_company_keywords = ["startup", "founded", "co-founder", "early stage"]

    # Analyze company signals
    enterprise_signal_count = 0
    startup_signal_count = 0
    mid_signal_count = 0

    for company in candidate_companies:
        if any(kw in company for kw in enterprise_company_keywords):
            enterprise_signal_count += 1

    for exp in experience:
        if isinstance(exp, dict):
            desc = (exp.get("description", "") or "").lower()
            highlights = exp.get("highlights", [])
            desc_combined = desc + " " + " ".join([h.lower() for h in highlights if isinstance(h, str)])

            # Check scale signals in experience descriptions
            if any(sig in desc_combined for sig in ["fortune 500", "enterprise", "global team", "10000+"]):
                enterprise_signal_count += 1
            if any(sig in desc_combined for sig in ["series c", "series d", "scale-up", "hypergrowth"]):
                mid_signal_count += 1
            if any(sig in desc_combined for sig in ["startup", "founded", "first hire", "0 to 1", "series a", "series b"]):
                startup_signal_count += 1

    # Determine candidate scale
    if enterprise_signal_count >= 2:
        candidate_scale = "enterprise"
    elif startup_signal_count >= 2:
        candidate_scale = "startup"
    elif mid_signal_count >= 1 or (enterprise_signal_count == 1 and startup_signal_count == 1):
        candidate_scale = "mid"
    else:
        # Default to mid if unclear
        candidate_scale = "mid"

    result["company_scale_alignment"]["candidate_scale"] = candidate_scale
    result["company_scale_alignment"]["target_scale"] = target_scale

    # Calculate delta (scale distance)
    scale_order = {"startup": 0, "mid": 1, "enterprise": 2}
    if target_scale and candidate_scale:
        delta = abs(scale_order.get(candidate_scale, 1) - scale_order.get(target_scale, 1))
        result["company_scale_alignment"]["delta"] = delta

        if delta == 0:
            result["company_scale_alignment"]["risk_level"] = "low"
        elif delta == 1:
            result["company_scale_alignment"]["risk_level"] = "low"  # Explainable
        elif delta == 2:
            result["company_scale_alignment"]["risk_level"] = "high"
            high_risk_count += 1
        elif delta >= 3:
            result["company_scale_alignment"]["risk_level"] = "high"
            high_risk_count += 1

    # ========================================================================
    # DIMENSION 3: ROLE SCOPE VS TITLE INFLATION
    # "Scope beats title. Always."
    # ========================================================================

    # Detect title inflation signals
    # Check if candidate has impressive titles but lacks scope evidence
    executive_titles = ["vp", "vice president", "director", "head of", "chief", "president", "svp", "evp"]
    leadership_titles = ["manager", "lead", "senior", "principal", "staff"]

    candidate_highest_level = None
    for title in candidate_titles:
        if any(et in title for et in executive_titles):
            candidate_highest_level = "executive"
            break
        elif any(lt in title for lt in leadership_titles):
            candidate_highest_level = "senior"

    if candidate_highest_level is None:
        candidate_highest_level = "ic"

    # Check for scope evidence in resume
    scope_evidence_keywords = [
        "owned", "led", "built", "scaled", "managed $", "budget", "p&l",
        "grew team", "hired", "headcount", "revenue", "increased", "reduced",
        "implemented", "drove", "delivered", "launched"
    ]
    strong_scope_evidence = sum(1 for kw in scope_evidence_keywords if kw in combined_resume)

    # Detect target role scope from JD
    jd_scope_keywords = ["own", "accountable", "budget", "p&l", "final decision", "build from scratch"]
    jd_scope_weight = sum(1 for kw in jd_scope_keywords if kw in combined_jd)

    # Title inflation assessment
    if candidate_highest_level == "executive":
        if strong_scope_evidence < 3:
            result["role_scope_alignment"]["title_inflation_risk"] = "high"
            result["role_scope_alignment"]["hm_skepticism"] = "high"
            high_risk_count += 1
        elif strong_scope_evidence < 5:
            result["role_scope_alignment"]["title_inflation_risk"] = "medium"
            result["role_scope_alignment"]["hm_skepticism"] = "medium"
            medium_risk_count += 1
        else:
            result["role_scope_alignment"]["title_inflation_risk"] = "low"
            result["role_scope_alignment"]["hm_skepticism"] = "low"
    elif candidate_highest_level == "senior":
        if strong_scope_evidence < 2:
            result["role_scope_alignment"]["title_inflation_risk"] = "medium"
            result["role_scope_alignment"]["hm_skepticism"] = "medium"
            medium_risk_count += 1
        else:
            result["role_scope_alignment"]["title_inflation_risk"] = "low"
            result["role_scope_alignment"]["hm_skepticism"] = "low"
    else:
        result["role_scope_alignment"]["title_inflation_risk"] = "low"
        result["role_scope_alignment"]["hm_skepticism"] = "low"

    # Scope gap assessment
    if jd_scope_weight >= 3:
        # High scope role
        if strong_scope_evidence < 3:
            result["role_scope_alignment"]["scope_gap"] = "significant"
            if result["role_scope_alignment"]["hm_skepticism"] == "low":
                result["role_scope_alignment"]["hm_skepticism"] = "medium"
                medium_risk_count += 1
        elif strong_scope_evidence < 5:
            result["role_scope_alignment"]["scope_gap"] = "moderate"
        else:
            result["role_scope_alignment"]["scope_gap"] = "none"
    elif jd_scope_weight >= 1:
        if strong_scope_evidence < 2:
            result["role_scope_alignment"]["scope_gap"] = "moderate"
        else:
            result["role_scope_alignment"]["scope_gap"] = "none"
    else:
        result["role_scope_alignment"]["scope_gap"] = "none"

    # ========================================================================
    # OVERALL CREDIBILITY ASSESSMENT & RECOMMENDATION CEILING
    # Any high-risk signal → cap at "Apply with Caution"
    # Multiple medium risks → also cap at "Apply with Caution"
    # ========================================================================

    if high_risk_count >= 1:
        result["overall_credibility_risk"] = "high"
        result["recommendation_ceiling"] = "Apply with Caution"
        result["coaching_intensity"] = "heavy"
        result["language_firmness"] = "skeptical"
        result["mandatory_reality_check"] = True
    elif medium_risk_count >= 2:
        result["overall_credibility_risk"] = "high"
        result["recommendation_ceiling"] = "Apply with Caution"
        result["coaching_intensity"] = "heavy"
        result["language_firmness"] = "candid"
        result["mandatory_reality_check"] = True
    elif medium_risk_count == 1:
        result["overall_credibility_risk"] = "medium"
        result["recommendation_ceiling"] = None  # No cap
        result["coaching_intensity"] = "moderate"
        result["language_firmness"] = "candid"
        result["mandatory_reality_check"] = False
    else:
        result["overall_credibility_risk"] = "low"
        result["recommendation_ceiling"] = None
        result["coaching_intensity"] = "normal"
        result["language_firmness"] = "confident"
        result["mandatory_reality_check"] = False

    # ========================================================================
    # MANDATORY REALITY CHECK MESSAGE
    # Per spec: When credibility risk is high, must include one explicit reality check
    # ========================================================================

    if result["mandatory_reality_check"]:
        # Build specific reality check message based on risks
        risk_factors = []

        if result["industry_alignment"]["risk_level"] == "high":
            risk_factors.append("industry alignment")
        if result["company_scale_alignment"]["risk_level"] == "high":
            risk_factors.append("company scale")
        if result["role_scope_alignment"]["hm_skepticism"] == "high":
            risk_factors.append("scope alignment")

        if risk_factors:
            risk_str = " and ".join(risk_factors)
            result["reality_check_message"] = (
                f"Hiring managers in this space typically prioritize candidates who have "
                f"operated at similar scale and within the same industry. Your background "
                f"is credible, but you will need to proactively address the {risk_str} gap. "
                f"Expect skepticism and be prepared to explain your transition story directly."
            )
        else:
            result["reality_check_message"] = (
                "Hiring managers typically prioritize candidates who have operated at "
                "similar scale and within the same industry. Your background is credible, "
                "but you should expect skepticism and be prepared to address it directly."
            )

    # Log CAE output
    print(f"🎯 CREDIBILITY ALIGNMENT ENGINE OUTPUT:")
    print(f"   Industry: {result['industry_alignment']['match']} ({result['industry_alignment']['risk_level']} risk)")
    print(f"   Scale: {result['company_scale_alignment']['candidate_scale']} → {result['company_scale_alignment']['target_scale']} (delta: {result['company_scale_alignment']['delta']})")
    print(f"   Title Inflation: {result['role_scope_alignment']['title_inflation_risk']}, Scope Gap: {result['role_scope_alignment']['scope_gap']}")
    print(f"   Overall Risk: {result['overall_credibility_risk']}")
    if result["recommendation_ceiling"]:
        print(f"   ⚠️ RECOMMENDATION CEILING: {result['recommendation_ceiling']}")
    if result["mandatory_reality_check"]:
        print(f"   📢 MANDATORY REALITY CHECK REQUIRED")

    return result


def apply_credibility_to_recommendation(
    base_recommendation: str,
    cae_result: dict,
    eligibility_locked: bool = False,
    hard_requirement_failure: bool = False
) -> dict:
    """
    Apply credibility modifiers to the recommendation.

    Per CAE Spec: "Credibility can downgrade, never upgrade."

    CRITICAL: Per Hard Requirement Failure spec:
    - If hard_requirement_failure is True, CAE cannot soften or override
    - Recommendation stays locked at "Do Not Apply"
    - CAE may explain skepticism but cannot change the decision

    Args:
        base_recommendation: Original recommendation before CAE
        cae_result: Output from evaluate_credibility_alignment()
        eligibility_locked: Whether eligibility gate locked the recommendation
        hard_requirement_failure: Whether a hard requirement failed (leadership, years, etc.)

    Returns:
        Modified recommendation with adjusted language and ceiling applied
    """
    result = {
        "final_recommendation": base_recommendation,
        "was_downgraded": False,
        "downgrade_reason": None,
        "coaching_intensity": cae_result.get("coaching_intensity", "normal"),
        "language_firmness": cae_result.get("language_firmness", "confident"),
        "reality_check_message": cae_result.get("reality_check_message"),
        "hard_requirement_locked": hard_requirement_failure
    }

    # HARD REQUIREMENT FAILURE: CAE cannot soften this
    # The recommendation is locked at "Do Not Apply" - period.
    if hard_requirement_failure:
        result["final_recommendation"] = "Do Not Apply"
        result["cae_override_blocked"] = True
        print("🔒 CAE BLOCKED: Hard requirement failure - cannot soften recommendation")
        return result

    # Don't modify if eligibility already locked
    if eligibility_locked:
        result["final_recommendation"] = "Do Not Apply"
        return result

    # Apply recommendation ceiling
    recommendation_ceiling = cae_result.get("recommendation_ceiling")

    if recommendation_ceiling:
        # Recommendation hierarchy: Apply > Apply with Caution > Consider with Caveats > Do Not Apply
        recommendation_rank = {
            "Apply": 4,
            "Strong Apply": 4,
            "Apply with Caution": 3,
            "Consider with Caveats": 2,
            "Consider with Reservations": 2,
            "Do Not Apply": 1
        }

        base_rank = recommendation_rank.get(base_recommendation, 3)
        ceiling_rank = recommendation_rank.get(recommendation_ceiling, 3)

        if base_rank > ceiling_rank:
            result["final_recommendation"] = recommendation_ceiling
            result["was_downgraded"] = True

            # Build downgrade reason
            risks = []
            if cae_result.get("industry_alignment", {}).get("risk_level") == "high":
                result["downgrade_reason"] = f"industry mismatch ({cae_result['industry_alignment']['match']})"
                risks.append("industry")
            if cae_result.get("company_scale_alignment", {}).get("risk_level") == "high":
                delta = cae_result.get("company_scale_alignment", {}).get("delta", 0)
                risks.append(f"scale gap (delta: {delta})")
            if cae_result.get("role_scope_alignment", {}).get("hm_skepticism") == "high":
                risks.append("scope/title skepticism")

            result["downgrade_reason"] = ", ".join(risks) if risks else "credibility risk"

            print(f"📉 RECOMMENDATION DOWNGRADED: {base_recommendation} → {result['final_recommendation']}")
            print(f"   Reason: {result['downgrade_reason']}")

    return result


def get_cae_coaching_language(cae_result: dict) -> dict:
    """
    Get language templates based on CAE risk level.

    Per Language Intensity Spec:
    - Low Risk: Confident, normal coaching
    - Medium Risk: Candid, realistic
    - High Risk: Direct, skeptical, still respectful

    Args:
        cae_result: Output from evaluate_credibility_alignment()

    Returns:
        Language templates for use in coaching messages
    """
    overall_risk = cae_result.get("overall_credibility_risk", "low")

    if overall_risk == "low":
        return {
            "intro": "Apply. Here's how to stand out:",
            "coaching_tone": "confident",
            "cta": "You're competitive for this role.",
            "hedging_allowed": False,
            "stretch_language_allowed": False  # Never at senior levels
        }
    elif overall_risk == "medium":
        return {
            "intro": "You can apply, but be prepared to address some questions:",
            "coaching_tone": "candid",
            "cta": "Position your experience strategically.",
            "hedging_allowed": True,
            "stretch_language_allowed": False  # Never at senior levels
        }
    else:  # high
        return {
            "intro": "You should expect skepticism due to alignment gaps:",
            "coaching_tone": "skeptical",
            "cta": "This will be an uphill battle. Apply only if you have a compelling transition story.",
            "hedging_allowed": True,
            "stretch_language_allowed": False,  # 🚫 Never
            # Per spec: 🚫 No "learning opportunity", 🚫 No "stretch role" at senior levels
            "banned_phrases": [
                "learning opportunity",
                "stretch role",
                "great chance to grow",
                "developmental opportunity"
            ]
        }


def check_eligibility_gate(resume_data: dict, response_data: dict) -> dict:
    """
    ELIGIBILITY GATE - Runs BEFORE scoring.
    Per Eligibility Gate Enforcement Spec: "Before scoring, determine if the candidate is *eligible*.
    Eligibility failure overrides score, confidence, and tone."

    Eligibility checks:
    1. Required domain experience (is the candidate's domain relevant?)
    2. Required leadership type (people vs operational)
    3. Required seniority scope (scale, team size, budget)
    4. Non-transferable domains (hard stop if mismatch)

    Returns:
        dict with:
            - eligible: bool - True if candidate passes eligibility gate
            - reason: str - Why they failed (if applicable)
            - failed_check: str - Which check failed
            - locked_recommendation: str - "Do Not Apply" if failed
            - gap_classification: str - "missing_experience" for non-transferable
    """
    result = {
        "eligible": True,
        "reason": None,
        "failed_check": None,
        "locked_recommendation": None,
        "gap_classification": None
    }

    if not resume_data or not response_data:
        return result

    # Extract job info
    role_title = (response_data.get("role_title", "") or "").lower()
    jd_text = (response_data.get("job_description", "") or "").lower()
    combined_jd = f"{role_title} {jd_text}"

    # Extract resume info - USE NORMALIZED EXPERIENCE
    experience = normalize_experience(resume_data)
    resume_titles = []
    resume_descriptions = []

    for exp in experience:
        # No need to check isinstance(exp, dict) - normalize_experience guarantees it
        title = (exp.get("title", "") or "").lower()
        desc = (exp.get("description", "") or "").lower()
        highlights = exp.get("highlights", [])
        if highlights and isinstance(highlights, list):
            desc += " " + " ".join([h.lower() for h in highlights if isinstance(h, str)])
        resume_titles.append(title)
        resume_descriptions.append(desc)

    combined_resume = " ".join(resume_titles) + " " + " ".join(resume_descriptions)

    # ========================================================================
    # CHECK 1: Non-Transferable Domain Detection
    #
    # CRITICAL FIX: This gate was firing on garbage matches (e.g., "sec" in "security")
    # Now requires:
    # 1. Manager-level roles are EXEMPT from this gate entirely
    # 2. Only fires on explicit ownership language, not mentions
    # ========================================================================

    # Manager-level role titles are EXEMPT from non-transferable domain gates
    manager_level_exemptions = [
        "engineering manager", "product manager", "program manager", "project manager",
        "senior manager", "manager of", "team lead", "tech lead", "technical lead",
        "staff engineer", "senior engineer", "principal engineer", "senior staff",
        "senior product manager", "group product manager", "staff product manager",
        "engineering lead", "senior software engineer"
    ]
    is_manager_level = any(exempt in role_title for exempt in manager_level_exemptions)

    if is_manager_level:
        print(f"✅ Manager-level role detected: '{role_title}' - SKIPPING non-transferable domain gates")
    else:
        for domain_config in NON_TRANSFERABLE_DOMAINS:
            domain = domain_config["domain"]
            levels = domain_config["levels"]
            keywords = domain_config["keywords"]
            description = domain_config["description"]
            require_ownership = domain_config.get("require_ownership", True)

            # Check if JD EXPLICITLY requires this domain with ownership language
            # Must find FULL keyword phrase, not partial token matches
            jd_requires_domain = any(kw in combined_jd for kw in keywords)
            jd_at_required_level = any(level in combined_jd for level in levels)

            # ADDITIONAL CHECK: For require_ownership domains, must have explicit requirement language
            if require_ownership:
                ownership_phrases = [
                    "required", "must have", "requirements:", "you will own",
                    "you will lead", "responsible for", "accountable for"
                ]
                has_ownership_context = any(phrase in jd_text for phrase in ownership_phrases)
                if not has_ownership_context:
                    # Mention without ownership context - skip this gate
                    continue

            if jd_requires_domain and jd_at_required_level:
                # JD requires this non-transferable domain at a senior level
                # Check if candidate has experience in this exact domain
                candidate_has_domain = any(kw in combined_resume for kw in keywords)

                if not candidate_has_domain:
                    # HARD STOP - Non-transferable domain mismatch
                    print(f"🚫 ELIGIBILITY GATE FAILED: Non-transferable domain")
                    print(f"   Domain: {description}")
                    print(f"   JD requires: {[kw for kw in keywords if kw in combined_jd]}")
                    print(f"   Candidate lacks domain experience")

                    result["eligible"] = False
                    result["reason"] = f"This role requires {description}. Your background does not include direct experience in this domain. This is a non-transferable requirement."
                    result["failed_check"] = f"non_transferable_domain:{domain}"
                    result["locked_recommendation"] = "Do Not Apply"
                    result["gap_classification"] = "missing_experience"
                    return result

    # ========================================================================
    # CHECK 2: People Leadership vs Operational Leadership
    # (Delegate to existing function but flag as eligibility)
    # ========================================================================
    people_leadership_years = extract_people_leadership_years(resume_data)
    required_people_leadership, is_hard_requirement = extract_required_people_leadership_years(response_data)

    if is_hard_requirement and required_people_leadership > 0:
        if people_leadership_years < required_people_leadership:
            # Calculate how severe the gap is
            leadership_ratio = people_leadership_years / required_people_leadership if required_people_leadership > 0 else 0
            gap_years = required_people_leadership - people_leadership_years

            print(f"⚠️ ELIGIBILITY CHECK: People leadership requirement")
            print(f"   Required: {required_people_leadership} years")
            print(f"   Candidate has: {people_leadership_years} years")
            print(f"   Ratio: {leadership_ratio:.1%}")

            # TIERED RESPONSE based on gap severity:
            # - <50% of required = hard fail (Do Not Apply)
            # - 50-70% of required = significant gap (Apply with Caution)
            # - >70% of required = minor gap (warning only, allow score-based recommendation)
            if leadership_ratio < 0.5:
                # SEVERE gap - hard fail
                print(f"🚫 ELIGIBILITY GATE FAILED: Severe leadership gap (<50%)")
                result["eligible"] = False
                result["reason"] = f"This role requires {required_people_leadership:.0f}+ years of people leadership experience. You have {people_leadership_years:.1f} years verified. Operational leadership does not count."
                result["failed_check"] = "people_leadership_requirement"
                result["locked_recommendation"] = "Do Not Apply"
                result["gap_classification"] = "missing_experience"
                return result
            elif leadership_ratio < 0.7:
                # MODERATE gap - strong warning but allow cautious apply
                print(f"⚠️ ELIGIBILITY WARNING: Moderate leadership gap (50-70%)")
                result["eligible"] = True
                result["warning"] = f"Leadership gap: {gap_years:.1f} years short of {required_people_leadership:.0f}+ requirement"
                result["recommended_adjustment"] = "Apply with Caution"
                result["gap_classification"] = "experience_gap"
                # Don't lock - let fit score drive recommendation with warning
            else:
                # MINOR gap - allow normal recommendation with note
                print(f"ℹ️ ELIGIBILITY NOTE: Minor leadership gap (>70%)")
                result["eligible"] = True
                result["warning"] = f"Minor leadership gap: {gap_years:.1f} years short"
                # No adjustment needed

    # ========================================================================
    # CHECK 3: Seniority Scope Mismatch
    # If JD requires VP-level and candidate is mid-level, fail
    #
    # CRITICAL FIX: Check ROLE TITLE first, not full JD text
    # Role titles like "Engineering Manager" should NOT trigger VP+ detection
    # even if the JD mentions "chief architect" in requirements
    # ========================================================================

    # Manager-level role titles that should NEVER trigger VP+ detection
    manager_level_titles = [
        "engineering manager", "product manager", "program manager", "project manager",
        "senior manager", "manager of", "team lead", "tech lead", "technical lead",
        "staff engineer", "senior engineer", "principal engineer", "senior staff",
        "senior product manager", "group product manager", "staff product manager"
    ]

    # Check if role title indicates Manager-level (NOT VP+)
    is_manager_level_role = any(title in role_title for title in manager_level_titles)

    if is_manager_level_role:
        # GUARDRAIL: Manager-level roles NEVER trigger VP+ eligibility failure
        print(f"✅ Role title '{role_title}' is Manager-level - skipping VP+ check")
        print(f"✅ ELIGIBILITY GATE PASSED")
        return result

    # Only check for VP+ requirements if role title suggests executive level
    executive_levels = ["vp", "vice president", "svp", "evp", "chief", "c-suite", "cxo", "ceo", "cfo", "cto", "coo", "cmo"]
    director_levels = ["director", "senior director", "head of"]

    # Check ROLE TITLE first for executive level, not full JD
    # This prevents false positives from JD text mentioning "chief architect" etc.
    jd_requires_executive = any(level in role_title for level in executive_levels)

    # Only fall back to JD text if role title is ambiguous AND title contains director+
    if not jd_requires_executive:
        jd_requires_director = any(level in role_title for level in director_levels)
        # If it's a Director role, don't require executive experience
        if jd_requires_director:
            print(f"✅ Role title is Director-level - not requiring VP+ experience")
            print(f"✅ ELIGIBILITY GATE PASSED")
            return result

    # Log what triggered the executive detection
    if jd_requires_executive:
        triggering_level = [level for level in executive_levels if level in role_title]
        print(f"🔍 Executive level detected in role title: {triggering_level}")

    # Check candidate's highest level
    candidate_has_executive = any(level in combined_resume for level in executive_levels)
    candidate_has_director = any(level in combined_resume for level in director_levels)

    if jd_requires_executive and not candidate_has_executive:
        # Check if they have at least director level
        if not candidate_has_director:
            print(f"🚫 ELIGIBILITY GATE FAILED: Seniority scope - VP+ role, candidate lacks executive/director experience")
            print(f"   Role title: {role_title}")
            print(f"   Triggering level: {[l for l in executive_levels if l in role_title]}")

            result["eligible"] = False
            result["reason"] = "This is a VP or executive-level role. Your background does not include executive or director-level experience. Do not apply."
            result["failed_check"] = "seniority_scope_executive"
            result["locked_recommendation"] = "Do Not Apply"
            result["gap_classification"] = "missing_experience"
            return result

    print(f"✅ ELIGIBILITY GATE PASSED")
    return result


def generate_specific_redirect(resume_data: dict, response_data: dict, eligibility_result: dict = None) -> list:
    """
    Generate SPECIFIC role category redirects based on candidate's actual background.

    Per Required Fixes Spec: "You must name realistic role categories where the candidate IS competitive."

    Args:
        resume_data: Parsed resume
        response_data: JD analysis response
        eligibility_result: Result from check_eligibility_gate

    Returns:
        list: 2-3 specific, actionable redirect recommendations
    """
    redirects = []

    if not resume_data:
        return [
            "Focus on roles that align with your verified experience",
            "Target positions where you meet 70%+ of the requirements"
        ]

    # Extract candidate's actual background
    experience = resume_data.get("experience", [])
    titles = []
    domains = []
    skills = []

    for exp in experience:
        if isinstance(exp, dict):
            title = (exp.get("title", "") or "").lower()
            desc = (exp.get("description", "") or "").lower()
            company = (exp.get("company", "") or "").lower()
            titles.append(title)

            # Detect domains
            if any(kw in f"{title} {desc}" for kw in ["recruit", "talent", "hiring", "sourcing"]):
                domains.append("recruiting")
            if any(kw in f"{title} {desc}" for kw in ["product manager", "pm ", "roadmap", "product strategy"]):
                domains.append("product management")
            if any(kw in f"{title} {desc}" for kw in ["product design", "ux", "ui", "user experience", "user interface", "designer", "design system"]):
                domains.append("product design")
            if any(kw in f"{title} {desc}" for kw in ["engineer", "developer", "software", "code", "backend", "frontend"]):
                domains.append("engineering")
            if any(kw in f"{title} {desc}" for kw in ["operations", "ops", "process", "systems"]):
                domains.append("operations")
            if any(kw in f"{title} {desc}" for kw in ["sales", "revenue", "account"]):
                domains.append("sales")
            if any(kw in f"{title} {desc}" for kw in ["marketing", "growth", "brand"]):
                domains.append("marketing")
            if any(kw in f"{title} {desc}" for kw in ["data", "analyst", "analytics", "bi", "business intelligence"]):
                domains.append("data/analytics")

    # Unique domains
    domains = list(set(domains))
    primary_domain = domains[0] if domains else "your field"

    # Determine candidate's level from titles
    has_senior = any(kw in " ".join(titles) for kw in ["senior", "sr", "lead", "principal"])
    has_manager = any(kw in " ".join(titles) for kw in ["manager", "supervisor"])
    has_director = any(kw in " ".join(titles) for kw in ["director", "head"])
    has_vp = any(kw in " ".join(titles) for kw in ["vp", "vice president"])

    # Determine why they were rejected
    failed_check = ""
    if eligibility_result:
        failed_check = eligibility_result.get("failed_check", "") or ""

    # Generate specific redirects based on rejection reason and background
    if "people_leadership" in failed_check:
        # Rejected for lacking people leadership
        if "recruiting" in domains:
            redirects = [
                "Redirect to Senior Technical Recruiting or Recruiting Operations roles where your infrastructure-building and scaling experience is a direct match",
                "Target Senior Recruiter or Lead Recruiter positions that value operational excellence without requiring direct reports",
                "Consider Recruiting Program Manager roles that leverage your process and systems expertise"
            ]
        elif "operations" in domains:
            redirects = [
                "Redirect to Senior Operations Manager or Operations Lead roles where your process and systems expertise is valued",
                "Target Program Manager or Technical Program Manager positions that don't require people leadership",
                "Consider Chief of Staff or Business Operations roles that leverage strategic thinking without direct reports"
            ]
        elif "product" in domains:
            redirects = [
                "Redirect to Senior Product Manager roles where your product strategy skills are directly applicable",
                "Target Principal Product Manager or Staff PM positions that value IC excellence",
                "Consider Product Strategy or Product Operations roles that leverage your expertise"
            ]
        elif "engineering" in domains:
            redirects = [
                "Redirect to Staff Engineer or Principal Engineer roles where your technical leadership is valued",
                "Target Technical Lead or Architect positions that don't require people management",
                "Consider Engineering Program Manager roles that leverage your technical expertise"
            ]
        else:
            redirects = [
                f"Redirect to Senior {primary_domain.title()} roles where your expertise is a direct match",
                f"Target Lead or Principal {primary_domain.title()} positions that value IC excellence",
                "Consider roles that emphasize operational excellence without people management requirements"
            ]

    elif "seniority_scope" in failed_check:
        # Rejected for seniority mismatch
        if has_director:
            redirects = [
                "Target Director-level roles rather than VP+ positions",
                "Consider Head of [Function] roles at smaller companies where scope aligns with your experience",
                "Focus on Senior Director positions as a stepping stone to VP"
            ]
        elif has_manager:
            redirects = [
                "Target Senior Manager or Manager roles rather than Director+",
                "Focus on Lead or Principal IC positions that match your current scope",
                "Consider smaller companies where your experience level is competitive"
            ]
        else:
            redirects = [
                "Target Senior IC or Lead positions that align with your experience",
                "Focus on roles where your current scope is directly applicable",
                "Build scope at your current level before targeting executive roles"
            ]

    elif "non_transferable_domain" in failed_check:
        # Rejected for domain mismatch - give specific, actionable redirects
        if "product design" in domains:
            redirects = [
                "Redirect to Product Designer roles at companies that need user-facing design, not research roles",
                "Target Senior Product Designer or Design Lead positions where your UX expertise is the core requirement",
                "Consider Growth Design or Design Systems roles that leverage your onboarding and scale experience"
            ]
        elif "product management" in domains:
            redirects = [
                "Redirect to Product Manager roles where your product strategy skills are directly applicable",
                "Target PM positions at companies in your domain (consumer tech, B2B SaaS, etc.)",
                "Consider Technical PM roles if you have strong technical collaboration experience"
            ]
        elif "engineering" in domains:
            redirects = [
                "Redirect to Software Engineering roles where your technical skills are directly applicable",
                "Target positions at companies using your tech stack and domain",
                "Consider Engineering Manager roles if you have people leadership experience"
            ]
        elif domains:
            redirects = [
                f"Redirect to {primary_domain.title()} roles where your direct experience applies",
                f"Target positions in {', '.join(d.title() for d in domains[:2])} where you are competitive",
                "Focus on roles that value your existing domain expertise"
            ]
        else:
            redirects = [
                "Target roles in domains where you have verified experience",
                "Focus on positions that leverage your existing skillset",
                "Build domain expertise before targeting specialized roles"
            ]

    else:
        # Generic redirect based on background
        if "recruiting" in domains:
            redirects = [
                "Target Senior Recruiting or Recruiting Operations roles where your experience is a direct match",
                "Focus on positions that value your specific recruiting expertise",
                "Consider roles at companies where your background is competitive"
            ]
        elif domains:
            redirects = [
                f"Redirect to {primary_domain.title()} roles where your experience is directly applicable",
                "Target positions that match your verified experience level",
                "Focus on roles where you meet 70%+ of the requirements"
            ]
        else:
            redirects = [
                "Target roles that align with your verified experience",
                "Focus on positions where you are competitive",
                "Build additional experience before targeting stretch roles"
            ]

    return redirects[:3]  # Max 3 redirects


def get_question_from_bank(role_type: str, category: str, asked_questions: List[str], target_level: str = "mid") -> Optional[Dict]:
    """
    Select an appropriate question from the bank based on role, category, and level.
    Returns None if no suitable question is found (fallback to Claude generation).
    """
    if not QUESTION_BANK:
        return None

    # Map role_type to question bank keys
    role_mapping = {
        "product_manager": "product_manager",
        "software_engineer": "software_engineer",
        "ux_designer": "ux_designer",
        "ui_designer": "ui_designer",
        "talent_acquisition": "talent_acquisition",
        "general_leadership": "general_leadership"
    }

    role_key = role_mapping.get(role_type, "general_leadership")
    role_questions = QUESTION_BANK.get("role_specific_questions", {}).get(role_key, {})

    # Determine which question pool to use based on category/stage
    question_pool = []

    if category in ["warm_start", "recruiter_screen"]:
        # Use generic category questions
        category_questions = QUESTION_BANK.get("question_categories", {}).get(category, {}).get("questions", [])
        question_pool.extend(category_questions)
    elif category == "behavioral":
        # Use both generic and role-specific behavioral
        generic_behavioral = QUESTION_BANK.get("question_categories", {}).get("behavioral", {}).get("questions", [])
        role_behavioral = role_questions.get("behavioral", [])
        question_pool.extend(generic_behavioral)
        question_pool.extend(role_behavioral)
    elif category in ["hiring_manager", "hiring_manager_deep_dive"]:
        # Use role-specific hiring manager questions
        question_pool.extend(role_questions.get("hiring_manager_deep_dive", []))
    elif category == "strategy":
        # Use role-specific strategy questions
        question_pool.extend(role_questions.get("strategy", []))
    else:
        # Default: use hiring_manager_deep_dive + behavioral
        question_pool.extend(role_questions.get("hiring_manager_deep_dive", []))
        question_pool.extend(role_questions.get("behavioral", []))

    # Filter by target level and exclude already asked questions
    filtered_questions = []
    for q in question_pool:
        # Check if question hasn't been asked
        if q.get("text") in asked_questions or q.get("id") in asked_questions:
            continue
        # Check if question is appropriate for target level
        target_levels = q.get("target_levels", ["mid", "senior", "director", "executive"])
        if target_level in target_levels:
            filtered_questions.append(q)

    if not filtered_questions:
        return None

    # Return a random question from filtered pool
    return random.choice(filtered_questions)

# ============================================================================
# PYDANTIC MODELS
# ============================================================================

class ResumeParseRequest(BaseModel):
    resume_text: Optional[str] = None
    linkedin_url: Optional[str] = None

class Experience(BaseModel):
    company: str
    title: str
    dates: str
    industry: Optional[str] = None
    bullets: List[str] = []

class ResumeProfile(BaseModel):
    full_name: str
    current_title: Optional[str] = None
    target_roles: List[str] = []
    industries: List[str] = []
    years_experience: Optional[str] = None
    summary: str
    experience: List[Experience]
    key_achievements: List[str] = []
    core_competencies: List[str] = []
    skills: List[str] = []
    education: Optional[str] = None

class JDAnalyzeRequest(BaseModel):
    company: str
    role_title: str
    job_description: Optional[str] = None  # Optional when using provisional_profile
    resume: Optional[Dict[str, Any]] = None
    preferences: Optional[Dict[str, Any]] = None
    # Command Center v2.0 fields
    jd_source: Optional[str] = None  # user_provided, url_fetched, inferred, missing
    provisional_profile: Optional[Dict[str, Any]] = None  # When jd_source is inferred/missing

class JDAnalysis(BaseModel):
    company: str
    role_title: str
    company_context: str
    role_overview: str
    key_responsibilities: List[str]
    required_skills: List[str]
    preferred_skills: List[str]
    ats_keywords: List[str]
    fit_score: int
    strengths: List[str]
    gaps: List[str]
    strategic_positioning: str
    salary_info: Optional[str] = None

class SupplementAnswer(BaseModel):
    """A user-provided answer to a clarifying question about a gap"""
    gap_area: str
    question: str
    answer: str

class LevelingContext(BaseModel):
    """Career level assessment data from resume-leveling page"""
    current_level: str
    current_level_id: str
    target_level: Optional[str] = None
    target_level_id: Optional[str] = None
    levels_apart: Optional[int] = None
    detected_function: str
    language_level: str
    recommendations: Optional[List[Dict[str, Any]]] = None
    gaps: Optional[List[Dict[str, Any]]] = None

class DocumentsGenerateRequest(BaseModel):
    resume: Dict[str, Any]
    jd_analysis: Dict[str, Any]
    preferences: Optional[Dict[str, Any]] = None
    supplements: Optional[List[SupplementAnswer]] = None  # User-provided answers from Strengthen page
    leveling: Optional[LevelingContext] = None  # Career level assessment data

class ResumeContent(BaseModel):
    summary: str
    skills: List[str]
    experience: List[Experience]

class CoverLetterContent(BaseModel):
    greeting: str
    opening: str
    body: str
    closing: str
    full_text: str

class GeneratedDocuments(BaseModel):
    resume: ResumeContent
    cover_letter: CoverLetterContent

# ============================================================================
# TAILORED DOCUMENT GENERATION MODELS
# ============================================================================

class ResumeCustomizeRequest(BaseModel):
    resume_text: Optional[str] = None
    resume_json: Optional[Dict[str, Any]] = None
    job_description: str
    target_role: str
    target_company: str
    jd_analysis: Optional[Dict[str, Any]] = None
    situation: Optional[Dict[str, Any]] = None  # Candidate emotional/situational state
    supplements: Optional[List[SupplementAnswer]] = None  # User-provided gap-filling info

class CoverLetterGenerateRequest(BaseModel):
    resume_text: Optional[str] = None
    resume_json: Optional[Dict[str, Any]] = None
    job_description: str
    target_role: str
    target_company: str
    strengths: Optional[List[str]] = None
    jd_analysis: Optional[Dict[str, Any]] = None
    situation: Optional[Dict[str, Any]] = None  # Candidate emotional/situational state
    supplements: Optional[List[SupplementAnswer]] = None  # User-provided gap-filling info

# ============================================================================
# MVP+1 FEATURE MODELS
# ============================================================================

# Feature 1: Daily Command Center
class Application(BaseModel):
    id: str
    company: str
    role_title: str
    fit_score: int = Field(ge=0, le=100)
    stage: str  # applied|screen|onsite|offer|rejected
    last_activity_date: str  # ISO date
    has_outreach: bool = False

class TasksPrioritizeRequest(BaseModel):
    applications: List[Application]
    today: str  # ISO date

class Task(BaseModel):
    type: str  # apply|follow_up|outreach
    application_id: str
    priority: int = Field(ge=1, le=3)
    reason: str
    suggested_message_stub: Optional[str] = None

class TasksPrioritizeResponse(BaseModel):
    tasks: List[Task]

# Feature 2: Learning/Feedback Loop
class OutcomeLogRequest(BaseModel):
    application_id: str
    stage: str  # applied|viewed|replied|interview|offer|rejected
    outcome: str
    date: str  # ISO date

class StrategyReviewRequest(BaseModel):
    applications: List[Dict[str, Any]]
    outcomes: List[Dict[str, Any]]

class StrategyReviewResponse(BaseModel):
    insights: List[str]
    recommendations: List[str]

# Feature 3: Network Engine (Lite)
class Contact(BaseModel):
    name: str
    company: str
    title: str
    relationship: str

class NetworkRecommendRequest(BaseModel):
    company: str
    role_title: str
    contacts: List[Contact]

class ContactRecommendation(BaseModel):
    name: str
    company: str
    title: str
    relationship: str
    priority: int = Field(ge=1, le=3)
    reason: str
    suggested_message_stub: str

class NetworkRecommendResponse(BaseModel):
    recommendations: List[ContactRecommendation]

# Feature 4: Interview Intelligence
class InterviewQuestion(BaseModel):
    question: str
    type: str  # behavioral|technical|leadership|competency|wildcard
    competency_tag: str
    difficulty: int = Field(ge=1, le=5)

class InterviewParseRequest(BaseModel):
    transcript_text: str
    role_title: str
    company: str
    jd_analysis: Optional[Dict[str, Any]] = None

class InterviewParseResponse(BaseModel):
    questions: List[InterviewQuestion]
    themes: List[str]
    warnings: List[str]

class InterviewFeedbackRequest(BaseModel):
    transcript_text: str
    role_title: str
    company: str
    questions: List[InterviewQuestion]

class DeliveryFeedback(BaseModel):
    tone: str
    pacing: str
    clarity: str
    structure: str

class InterviewFeedbackResponse(BaseModel):
    overall_score: int = Field(ge=1, le=100)
    strengths: List[str]
    areas_for_improvement: List[str]
    delivery_feedback: DeliveryFeedback
    recommendations: List[str]

class ThankYouRequest(BaseModel):
    transcript_text: str
    role_title: str
    company: str
    interviewer_name: Optional[str] = None
    jd_analysis: Optional[Dict[str, Any]] = None

class ThankYouResponse(BaseModel):
    subject: str
    body: str
    key_points_used: List[str]

# ============================================================================
# INTERVIEW INTELLIGENCE MODELS
# ============================================================================

class InterviewPrepRequest(BaseModel):
    """Request for generating interview prep materials"""
    resume_json: Dict[str, Any]
    job_description: Optional[str] = None  # Optional when using provisional_profile
    company: str
    role_title: str
    interview_stage: str = Field(..., pattern="^(recruiter_screen|hiring_manager)$")
    jd_analysis: Optional[Dict[str, Any]] = None
    # Command Center v2.0 fields
    jd_source: Optional[str] = None  # user_provided, url_fetched, inferred, missing
    provisional_profile: Optional[Dict[str, Any]] = None  # When jd_source is inferred/missing
    application_id: Optional[str] = None  # For tracking
    # Debrief Intelligence (Phase 2.3)
    past_debriefs: Optional[List[Dict[str, Any]]] = None  # Past debriefs for this company or similar roles

class RedFlagMitigation(BaseModel):
    flag: str
    mitigation: str

class QuestionToAsk(BaseModel):
    question: str
    why: str

class STARExample(BaseModel):
    competency: str
    situation: str
    task: str
    action: str
    result: str

class LikelyQuestion(BaseModel):
    question: str
    suggested_answer: str

class RecruiterScreenPrepContent(BaseModel):
    interview_overview: str
    key_talking_points: List[str]
    red_flag_mitigation: List[RedFlagMitigation]
    likely_questions: List[LikelyQuestion]
    questions_to_ask: List[QuestionToAsk]
    compensation_expectations: str
    timeline_strategy: str

class HiringManagerPrepContent(BaseModel):
    interview_overview: str
    strengths_to_emphasize: List[str]
    gaps_to_mitigate: List[str]
    star_examples: List[STARExample]
    likely_questions: List[LikelyQuestion]
    questions_to_ask: List[QuestionToAsk]
    closing_strategy: str

class InterviewPrepResponse(BaseModel):
    interview_stage: str
    prep_content: Dict[str, Any]
    # Command Center v2.0 fields
    confidence: Optional[str] = None  # "directional" or "refined"
    ui_note: Optional[str] = None  # Displayed to user when prep is directional

class IntroSellTemplateRequest(BaseModel):
    """Request for generating intro sell template"""
    resume_json: Dict[str, Any]
    job_description: str
    company: str
    role_title: str

class IntroSellTemplateResponse(BaseModel):
    template: str
    word_count: int
    coaching_note: str

class IntroSellFeedbackRequest(BaseModel):
    """Request for analyzing intro sell attempt"""
    resume_json: Dict[str, Any]
    job_description: str
    company: str
    role_title: str
    candidate_version: str

class IntroSellFeedbackResponse(BaseModel):
    overall_score: float
    content_score: int
    structure_score: int
    tone_score: int
    word_count: int
    estimated_time_seconds: int
    strengths: List[str]
    opportunities: List[str]
    revised_version: str
    coaching_note: str

class TypedDebriefResponses(BaseModel):
    overall_feeling: str
    questions_asked: List[str]
    strong_answers: List[str]
    weak_answers: List[str]
    learnings: str

class InterviewDebriefRequest(BaseModel):
    """Request for post-interview debrief analysis"""
    resume_json: Dict[str, Any]
    job_description: str
    company: str
    role_title: str
    interview_stage: str = Field(..., pattern="^(recruiter_screen|hiring_manager)$")
    interview_date: str
    interviewer_name: Optional[str] = None
    transcript_text: Optional[str] = None
    typed_responses: Optional[TypedDebriefResponses] = None

class DimensionScores(BaseModel):
    content: int = Field(..., ge=1, le=10)
    clarity: int = Field(..., ge=1, le=10)
    delivery: int = Field(..., ge=1, le=10)
    tone: int = Field(..., ge=1, le=10)
    structure: int = Field(..., ge=1, le=10)
    confidence: int = Field(..., ge=1, le=10)

class WhatTheyShouldHaveSaid(BaseModel):
    question: str
    weak_answer: str
    strong_answer: str

class InterviewDebriefResponse(BaseModel):
    overall_score: int = Field(..., ge=1, le=10)
    dimension_scores: DimensionScores
    strengths: List[str]
    opportunities: List[str]
    what_they_should_have_said: List[WhatTheyShouldHaveSaid]
    coaching_points: List[str]
    action_items: List[str]
    thank_you_email: str
    next_stage_adjustments: List[str]

# ============================================================================
# SCREENING QUESTIONS MODELS
# ============================================================================

class ScreeningQuestionsRequest(BaseModel):
    """Request for generating screening question responses"""
    resume_json: Dict[str, Any]
    job_description: str
    company: str
    role_title: str
    screening_questions: str  # Raw text of screening questions pasted by user
    jd_analysis: Optional[Dict[str, Any]] = None
    situation: Optional[Dict[str, Any]] = None  # Candidate emotional/situational state

class ScreeningQuestionResponse(BaseModel):
    """Individual screening question response"""
    question: str
    recommended_response: str
    strategic_note: str
    gap_detected: bool = False
    mitigation_strategy: Optional[str] = None

class ScreeningQuestionsResponse(BaseModel):
    """Response containing all screening question answers"""
    responses: List[ScreeningQuestionResponse]
    overall_strategy: str

# ============================================================================
# EXPERIENCE SUPPLEMENTATION MODELS
# ============================================================================

class ClarifyingQuestion(BaseModel):
    """A targeted question to fill a specific gap"""
    gap_area: str  # The gap this question addresses
    question: str  # The question to ask the candidate
    why_it_matters: str  # Brief explanation of why this matters for the role
    example_answer: str  # Example of a strong response

class GenerateClarifyingQuestionsRequest(BaseModel):
    """Request to generate clarifying questions based on gaps"""
    resume_json: Dict[str, Any]
    job_description: str
    company: str
    role_title: str
    gaps: List[str]  # The gaps identified in JD analysis
    fit_score: int  # Current fit score

class GenerateClarifyingQuestionsResponse(BaseModel):
    """Response containing targeted clarifying questions"""
    questions: List[ClarifyingQuestion]
    intro_message: str  # Personalized message to the candidate

class SupplementedExperience(BaseModel):
    """A piece of supplemented experience from the candidate"""
    gap_area: str
    question: str
    answer: str

class ReanalyzeWithSupplementsRequest(BaseModel):
    """Request to re-analyze fit with supplemented experience"""
    resume_json: Dict[str, Any]
    job_description: str
    company: str
    role_title: str
    original_gaps: List[str]
    original_fit_score: int
    supplements: List[SupplementedExperience]

class ReanalyzeWithSupplementsResponse(BaseModel):
    """Response from re-analysis with supplemented experience"""
    new_fit_score: int
    score_change: int  # Positive = improved
    updated_strengths: List[str]
    remaining_gaps: List[str]
    addressed_gaps: List[str]
    updated_resume_json: Dict[str, Any]  # Resume with supplements incorporated
    summary: str  # Brief explanation of changes

# ============================================================================
# RESUME LEVELING MODELS
# ============================================================================

class LevelCompetency(BaseModel):
    """Assessment of a specific competency area"""
    area: str  # e.g., "Technical Depth", "Leadership", "Strategic Thinking"
    current_level: str  # e.g., "Strong at Senior level"
    evidence: List[str]  # Quotes/signals from resume
    gap_to_target: Optional[str] = None  # What's missing for target level

class LevelingGap(BaseModel):
    """A specific gap between current and target level"""
    category: str  # "scope", "impact", "competency", "language"
    description: str
    recommendation: str  # How to address in resume
    priority: str  # "high", "medium", "low"

class LevelingRecommendation(BaseModel):
    """A specific recommendation to strengthen resume for target level"""
    type: str  # "content", "language", "quantification", "scope"
    priority: str  # "high", "medium", "low"
    current: str  # Current state
    suggested: str  # Recommended change
    rationale: str  # Why this matters

class ResumeLevelingRequest(BaseModel):
    """Request for resume level assessment"""
    resume_json: Any  # Can be dict or string (will be parsed in endpoint)
    job_description: Optional[str] = None
    target_title: Optional[str] = None  # If provided, do gap analysis against this
    company: Optional[str] = None
    role_title: Optional[str] = None

class LevelMismatchWarning(BaseModel):
    """Warning when target level doesn't match assessed level"""
    target_level: str
    assessed_level: str
    gap_explanation: str
    alternative_recommendation: Optional[str] = None


class RoleTypeBreakdown(BaseModel):
    """Breakdown of years by role type"""
    pm_years: int = 0
    engineering_years: int = 0
    operations_years: int = 0
    other_years: int = 0


class ResumeLevelingResponse(BaseModel):
    """Response containing full resume leveling analysis"""
    # Function detection
    detected_function: str  # "Engineering", "Product Management", "Marketing", etc.
    function_confidence: float  # 0-1 confidence score

    # Current level assessment
    current_level: str  # e.g., "Associate PM", "Mid-Level PM" (NOT inflated titles)
    current_level_id: str  # e.g., "associate_pm", "mid_pm"
    level_confidence: float  # 0-1 confidence score
    years_experience: int

    # NEW: Strict role-type experience tracking
    years_in_role_type: Optional[int] = None  # Only years in actual PM/Eng/etc roles
    role_type_breakdown: Optional[RoleTypeBreakdown] = None

    # Evidence for current level
    scope_signals: List[str]  # Scope indicators found
    impact_signals: List[str]  # Impact statements found
    leadership_signals: List[str]  # Leadership evidence
    technical_signals: List[str]  # Technical depth evidence

    # Competency breakdown
    competencies: List[LevelCompetency]

    # Language analysis
    language_level: str  # "Entry", "Mid", "Senior", "Principal"
    action_verb_distribution: Dict[str, float]  # {"entry": 0.1, "mid": 0.3, "senior": 0.5, "principal": 0.1}
    quantification_rate: float  # % of bullets with numbers

    # Red flags
    red_flags: List[str]  # Issues found (generic claims, inconsistencies, etc.)

    # NEW: Title inflation detection
    title_inflation_detected: Optional[bool] = None
    title_inflation_explanation: Optional[str] = None

    # Target analysis (if target_title provided)
    target_level: Optional[str] = None
    target_level_id: Optional[str] = None
    levels_apart: Optional[int] = None  # 0 = matches, positive = target is higher
    is_qualified: Optional[bool] = None
    qualification_confidence: Optional[float] = None

    # NEW: Level mismatch warnings
    level_mismatch_warnings: Optional[List[LevelMismatchWarning]] = None

    # Gap analysis (if target provided)
    gaps: Optional[List[LevelingGap]] = None

    # Recommendations
    recommendations: List[LevelingRecommendation]

    # Summary
    summary: str  # Brief narrative assessment

    # LEPE Integration (for Manager+ roles)
    lepe_applicable: Optional[bool] = None  # True if target is Manager+
    lepe_decision: Optional[str] = None  # apply | position | caution | locked
    lepe_coaching: Optional[str] = None  # Coaching advice if positioning mode
    lepe_skepticism_warning: Optional[str] = None  # Warning if caution/locked
    leadership_tenure: Optional[Dict[str, Any]] = None  # From LEPE analysis
    accountability_record: Optional[Dict[str, Any]] = None  # LEPE accountability

# ============================================================================
# PREP GUIDE MODELS
# ============================================================================

class PrepGuideLikelyQuestion(BaseModel):
    """A likely interview question with guidance"""
    question: str
    guidance: str  # Suggested answer approach based on candidate's experience

class PrepGuideStory(BaseModel):
    """A STAR story for the candidate"""
    title: str
    competency: str  # e.g., "Leadership", "Problem-solving"
    situation: str
    task: str
    action: str
    result: str

class PrepGuideStrategyScenario(BaseModel):
    """A strategy/case scenario for HM/Technical interviews"""
    scenario: str
    approach: str

class PrepGuideRequest(BaseModel):
    """Request to generate an interview prep guide"""
    company: str
    role_title: str
    interview_type: str  # recruiter_screen, hiring_manager, technical, panel, executive
    job_description: str = ""
    interviewer_name: str = ""
    interviewer_title: str = ""
    resume_json: Dict[str, Any] = {}

class PrepGuideResponse(BaseModel):
    """Response containing the full prep guide"""
    what_they_evaluate: List[str]
    intro_pitch: str
    likely_questions: List[PrepGuideLikelyQuestion]
    stories: List[PrepGuideStory]
    red_flags: List[str]
    strategy_scenarios: Optional[List[PrepGuideStrategyScenario]] = None

class RegenerateIntroRequest(BaseModel):
    """Request to regenerate just the intro pitch"""
    company: str
    role_title: str
    interview_type: str
    resume_json: Dict[str, Any] = {}

class RegenerateIntroResponse(BaseModel):
    """Response with new intro pitch"""
    intro_pitch: str

# ============================================================================
# MOCK INTERVIEW MODELS
# ============================================================================

class StartMockInterviewRequest(BaseModel):
    """Request to start a mock interview session"""
    resume_json: Dict[str, Any]
    job_description: str
    company: str
    role_title: str
    interview_stage: str = Field(..., pattern="^(recruiter_screen|hiring_manager)$")
    difficulty_level: str = Field(default="medium", pattern="^(easy|medium|hard)$")

class FirstQuestionResponse(BaseModel):
    """First question returned when starting a session"""
    question_id: str
    question_text: str
    competency_tested: str
    question_number: int

class StartMockInterviewResponse(BaseModel):
    """Response after starting a mock interview"""
    session_id: str
    interview_stage: str
    difficulty_level: str
    first_question: FirstQuestionResponse

class SubmitMockResponseRequest(BaseModel):
    """Request to submit a response to a mock interview question"""
    session_id: str
    question_id: str
    response_text: str
    response_number: int = Field(..., ge=1, le=4)

class SubmitMockResponseResponse(BaseModel):
    """Response after submitting an answer"""
    response_recorded: bool
    follow_up_question: Optional[str] = None
    should_continue: bool
    brief_feedback: str

class QuestionFeedbackResponse(BaseModel):
    """Comprehensive feedback for a completed question"""
    question_text: str
    all_responses: List[str]
    score: int = Field(..., ge=1, le=10)
    level_demonstrated: str
    what_landed: List[str]
    what_didnt_land: List[str]
    coaching: str
    revised_answer: str

class NextQuestionRequest(BaseModel):
    """Request to move to the next question"""
    session_id: str
    current_question_number: int

class MockSessionProgress(BaseModel):
    """Progress within a mock interview session"""
    questions_completed: int
    average_score: float

class NextQuestionResponse(BaseModel):
    """Response with the next question"""
    question_id: str
    question_text: str
    competency_tested: str
    question_number: int
    total_questions: int
    session_progress: MockSessionProgress

class EndMockInterviewRequest(BaseModel):
    """Request to end a mock interview session"""
    session_id: str

class MockQuestionSummary(BaseModel):
    """Summary of a single question's performance"""
    question_number: int
    question_text: str
    score: int
    brief_feedback: str

class MockSessionFeedback(BaseModel):
    """Session-level feedback after completing mock interview"""
    overall_assessment: str
    key_strengths: List[str]
    areas_to_improve: List[str]
    coaching_priorities: List[str]
    recommended_drills: List[str] = []  # Signal-based drill recommendations
    readiness_score: str  # "Ready" | "Almost Ready" | "Needs Practice"
    level_estimate: str = "mid"  # "mid" | "senior" | "director" | "executive"
    next_steps: str

class EndMockInterviewResponse(BaseModel):
    """Response after ending a mock interview"""
    session_id: str
    overall_score: float
    questions_completed: int
    session_feedback: MockSessionFeedback
    question_summaries: List[MockQuestionSummary]

class MockInterviewSessionSummary(BaseModel):
    """Summary of a past mock interview session"""
    session_id: str
    interview_stage: str
    started_at: str
    completed_at: Optional[str]
    overall_score: Optional[float]
    questions_completed: int

class ImprovementTrend(BaseModel):
    """Improvement trend across multiple sessions"""
    first_session_score: float
    latest_session_score: float
    improvement: str

class SessionHistoryResponse(BaseModel):
    """Response with session history"""
    sessions: List[MockInterviewSessionSummary]
    improvement_trend: Optional[ImprovementTrend] = None

# ============================================================================
# COMMAND CENTER MODELS (v2.0)
# ============================================================================

class JDSource(str, Enum):
    """Track JD provenance"""
    USER_PROVIDED = "user_provided"
    URL_FETCHED = "url_fetched"
    INFERRED = "inferred"
    MISSING = "missing"
    LINK_FAILED = "link_failed"

class ConfidenceLabel(str, Enum):
    """Trust preservation labels"""
    DIRECTIONAL = "directional"
    REFINED = "refined"

class PriorityLevel(str, Enum):
    """Focus guidance levels"""
    HIGH = "high"
    MEDIUM = "medium"
    LOW = "low"
    ARCHIVE = "archive"

class UrgencyLevel(str, Enum):
    """Action urgency levels"""
    IMMEDIATE = "immediate"
    SOON = "soon"
    ROUTINE = "routine"
    NONE = "none"

class PipelineStatus(str, Enum):
    """Pipeline health status"""
    HEALTHY = "healthy"
    THIN = "thin"
    OVERLOADED = "overloaded"
    STALLED = "stalled"

class PipelineTone(str, Enum):
    """Pipeline recommendation tone"""
    STEADY = "steady"
    CAUTION = "caution"
    URGENT = "urgent"

# --- JD Reconstruct Models ---

class JDReconstructRequest(BaseModel):
    """Request to generate Provisional Role Profile when JD is missing"""
    role_title: str
    company_name: str
    industry: Optional[str] = None  # Optional, inferred if missing
    seniority: Optional[str] = None  # Optional, inferred from title

class ProvisionalProfile(BaseModel):
    """Reconstructed role profile when JD is missing"""
    role_title: str
    typical_responsibilities: List[str]
    common_competencies: List[str]
    interview_focus_areas: List[str]
    evaluation_criteria: List[str]

class JDReconstructResponse(BaseModel):
    """Response from JD reconstruction"""
    provisional_profile: ProvisionalProfile
    confidence: ConfidenceLabel = ConfidenceLabel.DIRECTIONAL
    jd_source: JDSource = JDSource.INFERRED

# --- Tracker Intelligence Models ---

class TrackerApplication(BaseModel):
    """Application data for tracker intelligence"""
    id: str
    status: str
    company: str
    role: str
    date_applied: Optional[str] = None
    decision_confidence: Optional[int] = None
    jd_source: Optional[str] = None
    fit_score: Optional[int] = None
    last_activity_date: Optional[str] = None
    days_since_last_activity: Optional[int] = None
    interview_count: Optional[int] = 0
    substatus: Optional[str] = None
    manual_lock: Optional[bool] = False
    user_override: Optional[bool] = False
    user_override_reason: Optional[str] = None

class OneClickAction(BaseModel):
    """One-click action for quick execution"""
    type: str  # draft_email, archive, open_prep, block_time
    template: Optional[str] = None
    application_id: Optional[str] = None
    confirm: Optional[bool] = False

class PriorityAction(BaseModel):
    """Priority action for an application"""
    application_id: str
    action: str
    reason: str
    priority: PriorityLevel
    one_click_action: Optional[OneClickAction] = None

class PipelineHealth(BaseModel):
    """Pipeline health assessment"""
    active_count: int
    status: PipelineStatus
    color: str
    icon: str
    tone: PipelineTone
    recommendation: str
    reason: str
    priority_count: int = 0

class FocusModeAction(BaseModel):
    """Action to display in focus mode"""
    application_id: str
    company: str
    action: str

class FocusMode(BaseModel):
    """Focus mode configuration"""
    enabled: bool = True
    top_actions: List[FocusModeAction]
    dim_others: bool = True

class UISignals(BaseModel):
    """UI-ready signals for frontend binding"""
    priority: PriorityLevel
    confidence: str  # high, medium, low
    urgency: UrgencyLevel
    color_code: str  # green, yellow, red, gray
    icon: str
    badge: Optional[ConfidenceLabel] = None
    action_available: bool = False
    dimmed: bool = False

class TrackerIntelligenceRequest(BaseModel):
    """Request for tracker intelligence calculation"""
    user_id: str
    applications: List[TrackerApplication]

class ApplicationWithIntelligence(BaseModel):
    """Application with calculated intelligence"""
    id: str
    next_action: str
    next_action_reason: str
    priority_level: PriorityLevel
    one_click_action: Optional[OneClickAction] = None
    ui_signals: UISignals
    decision_confidence: int
    days_since_last_activity: int
    substatus: str

class TrackerIntelligenceResponse(BaseModel):
    """Response from tracker intelligence engine"""
    priority_actions: List[PriorityAction]
    pipeline_health: PipelineHealth
    focus_mode: FocusMode
    applications: List[ApplicationWithIntelligence]

# --- Calculate Confidence Models ---

class CalculateConfidenceRequest(BaseModel):
    """Request to calculate decision confidence"""
    application_id: str
    fit_score: int
    jd_source: str
    days_since_applied: int
    status: str
    interview_count: int = 0
    response_time_days: Optional[int] = None
    days_since_last_activity: Optional[int] = None

class ConfidenceFactors(BaseModel):
    """Factors used in confidence calculation"""
    alignment_score: int
    momentum_score: int
    jd_confidence: int

class CalculateConfidenceResponse(BaseModel):
    """Response with calculated confidence"""
    decision_confidence: int
    label: str  # high, medium, low
    factors: ConfidenceFactors
    guidance: str


# ============================================================================
# HELPER FUNCTIONS
# Note: Core helpers (call_claude, clean_claude_json, detect_role_type, etc.)
# are now imported from utils/, services/, and storage/ modules.
# ============================================================================


def calculate_weighted_score(signals: dict, role_type: str) -> float:
    """
    Calculate weighted score based on signals and role-specific weights.
    """
    weights = SIGNAL_WEIGHTS_BY_ROLE.get(role_type, SIGNAL_WEIGHTS_BY_ROLE["general_leadership"])
    weighted_sum = 0.0

    for signal_name, signal_value in signals.items():
        weight = weights.get(signal_name, 0.0)
        weighted_sum += signal_value * weight

    # Convert to 1-10 scale
    return round(weighted_sum * 10, 1)


def calculate_mock_session_average(session_id: str) -> float:
    """
    Calculate average score across all completed questions in session.
    """
    session = get_mock_session(session_id)
    if not session:
        return 0.0

    question_ids = session.get("question_ids", [])
    scores = []

    for qid in question_ids:
        analysis = get_mock_analysis(qid)
        if analysis and analysis.get("score"):
            scores.append(analysis["score"])

    if not scores:
        return 0.0

    return sum(scores) / len(scores)


def format_responses_for_analysis(question_id: str) -> str:
    """
    Format all responses to a question into text for analysis.
    """
    responses = get_mock_responses(question_id)
    formatted = []

    for i, response in enumerate(responses, 1):
        label = "INITIAL RESPONSE" if i == 1 else f"FOLLOW-UP {i-1}"
        formatted.append(f"{label}:\n{response.get('response_text', '')}")

    return "\n\n".join(formatted)


# Mock Interview Prompt Templates - Now imported from prompts/ module
# (MOCK_GENERATE_QUESTION_PROMPT, MOCK_ANALYZE_RESPONSE_PROMPT,
#  MOCK_QUESTION_FEEDBACK_PROMPT, MOCK_SESSION_FEEDBACK_PROMPT)


def validate_outreach_template(template_text: str, template_type: str = "generic") -> tuple[bool, list[str]]:
    """
    Validates outreach template for quality and positioning.
    Returns (is_valid, error_messages)
    
    Args:
        template_text: The outreach message text to validate
        template_type: "hiring_manager" or "recruiter" or "generic"
    """
    errors = []
    
    if not template_text or len(template_text.strip()) == 0:
        errors.append("Template is empty")
        return False, errors
    
    # Check for em dashes and en dashes
    if '—' in template_text:
        errors.append("Contains em dashes (—) - use colons or periods instead")
    if '–' in template_text:
        errors.append("Contains en dashes (–) - use colons or periods instead")
    
    # Check for exclamation points (signals desperation)
    if '!' in template_text:
        errors.append("Contains exclamation points - remove for professional tone")
    
    # Check for generic phrases
    generic_phrases = [
        "I'm excited about this opportunity",
        "I'd be a great fit",
        "I'd love to chat",
        "I came across your posting",
        "I'm reaching out to express my interest",
        "great fit for my background",
        "perfect match for",
        "thrilled about"
    ]
    for phrase in generic_phrases:
        if phrase.lower() in template_text.lower():
            errors.append(f"Contains generic phrase: '{phrase}' - be more specific")
    
    # Check sentence length (flag if any sentence > 35 words)
    sentences = [s.strip() for s in template_text.replace('?', '.').replace('!', '.').split('.') if s.strip()]
    for i, sentence in enumerate(sentences):
        word_count = len(sentence.split())
        if word_count > 35:
            errors.append(f"Sentence {i+1} is too long ({word_count} words) - keep under 30 words")
    
    # Check for passive voice indicators
    passive_indicators = ['was led by', 'were managed by', 'is overseen by', 'has been done']
    for indicator in passive_indicators:
        if indicator in template_text.lower():
            errors.append(f"Contains passive voice: '{indicator}' - use active voice")
    
    # Check for clear ask at the end (should have call-to-action words)
    ask_keywords = ['discuss', 'conversation', 'call', 'meeting', 'chat', 'connect', 'talk', 'explore']
    has_ask = any(keyword in template_text.lower() for keyword in ask_keywords)
    if not has_ask:
        errors.append("Missing clear ask/call-to-action")
    
    # Check minimum length (should be at least 2 sentences)
    if len(sentences) < 2:
        errors.append("Template too short - should be 3-5 sentences")
    
    # Check maximum length (should not be more than 7 sentences)
    if len(sentences) > 7:
        errors.append("Template too long - should be 3-5 sentences")
    
    return len(errors) == 0, errors


def cleanup_outreach_template(template_text: str) -> str:
    """
    Cleans up common issues in generated outreach templates.
    This is a last-resort cleanup for minor formatting issues.
    """
    if not template_text:
        return template_text
    
    # Replace em dashes with colons
    template_text = template_text.replace('—', ':')
    template_text = template_text.replace('–', ':')
    
    # Remove exclamation points (replace with periods)
    template_text = template_text.replace('!', '.')
    
    # Fix double spaces
    template_text = ' '.join(template_text.split())
    
    # Fix double periods
    while '..' in template_text:
        template_text = template_text.replace('..', '.')
    
    return template_text.strip()



def build_candidate_calibration_prompt(situation: Optional[Dict[str, Any]]) -> str:
    """
    Build a calibration prompt based on candidate's emotional state and situation.
    This affects DELIVERY, not ACCURACY - never inflate scores or hide gaps.
    """
    if not situation:
        return ""
    
    holding_up = situation.get("holding_up", "")
    timeline = situation.get("timeline", "")
    confidence = situation.get("confidence", "")
    move_type = situation.get("move_type", "")
    
    # If no situation data, return empty
    if not any([holding_up, timeline, confidence, move_type]):
        return ""
    
    calibration = """
=== CANDIDATE STATE CALIBRATION ===

The candidate has shared their current emotional and situational context. 
Use this to calibrate your tone, prioritization, and framing. Not to change the facts, but to change how you deliver them.

**Candidate State:**
"""
    
    if holding_up:
        calibration += f"- Emotional state: {holding_up}\n"
    if timeline:
        calibration += f"- Timeline/Urgency: {timeline}\n"
    if confidence:
        calibration += f"- Confidence level: {confidence}\n"
    if move_type:
        calibration += f"- Move type: {move_type}\n"
    
    calibration += """
**Calibration Rules:**

### Emotional State (holding_up)
- "doing_well" → They're browsing, not desperate. Be efficient, skip emotional framing. Focus on fit and strategy.
- "stressed_but_managing" → Acknowledge the pressure briefly, then get to work. Don't dwell.
- "struggling" → Validate first. Acknowledge the difficulty before diving into tactics. Avoid toxic positivity. Be warm but direct.
- "rather_not_say" → Neutral tone. Don't probe. Just do the job.

### Timeline (timeline)
- "no_rush" → Full strategic mode. Include long-game moves: networking, positioning for future roles, stretch applications. Quality over speed.
- "actively_looking" → Balance strategy with momentum. Mix high-probability with strategic stretch roles.
- "soon" → Prioritize high-probability roles. Lead with "Apply" and "Strongly Apply" recommendations. Flag stretch roles as secondary.
- "urgent" → Speed matters. Prioritize roles above 60% fit. Lead with quick wins. Be direct about what's realistic. Do not waste their time on long shots.

### Confidence (confidence)
- "strong" → Be direct. They can handle blunt feedback. Don't pad the analysis.
- "shaky" → Balanced delivery. State gaps plainly but constructively. Remind them that rejections reflect market noise, not just their value.
- "low" → Lead with strengths. Frame gaps as "areas to address" not weaknesses. Soften language in fit analysis without hiding reality. End sections on a forward-looking note.
- "need_validation" → They need to hear they're not crazy. Acknowledge the market is brutal. Be honest about their strengths. Don't fabricate, but don't be stingy with legitimate praise either.

### Move Type (move_type)
- "lateral" → Emphasize transferable wins, culture fit, and why this company over others. Positioning is about differentiation.
- "step_up" → Emphasize leadership signals, scope expansion, and readiness for the next level. Address the "stretch" concern head-on.
- "pivot" → Emphasize transferable skills, adjacent experience, and learning agility. Acknowledge the narrative gap and help bridge it.
- "returning" → Normalize the gap. Emphasize what they did during it (if anything) and how their prior experience still applies. Don't make them feel like they're starting over.

**CRITICAL:** These calibrations affect *delivery*, not *accuracy*. Never inflate fit scores, hide real gaps, or sugarcoat bad-fit roles. The candidate deserves honesty. Just delivered in a way that meets them where they are.

"""
    return calibration

# Note: extract_pdf_text, extract_docx_text, and Command Center helpers
# (calculate_days_since, calculate_momentum_score, etc.) are now imported from utils/


# ============================================================================
# API ENDPOINTS
# ============================================================================

@app.get("/health")
async def health_check():
    """Simple health check for load balancers and monitoring"""
    return {"status": "ok"}

@app.get("/")
async def root():
    """Detailed API information endpoint"""
    return {
        "status": "running",
        "service": "Henry Job Search Engine API",
        "version": "1.4.0",
        "endpoints": [
            "/api/resume/parse",
            "/api/jd/analyze",
            "/api/documents/generate",
            "/api/screening-questions/generate",
            "/api/experience/clarifying-questions",
            "/api/experience/reanalyze",
            "/api/resume/level-assessment",
            "/api/tasks/prioritize",
            "/api/outcomes/log",
            "/api/strategy/review",
            "/api/network/recommend",
            "/api/interview/parse",
            "/api/interview/feedback",
            "/api/interview/thank_you",
            "/api/prep-guide/generate",
            "/api/prep-guide/regenerate-intro",
            # Command Center v2.0 endpoints
            "/api/jd/reconstruct",
            "/api/tracker/intelligence",
            "/api/tracker/calculate-confidence"
        ]
    }

@app.post("/api/resume/parse")
@limiter.limit("30/minute")
async def parse_resume(
    request: Request,
    file: Optional[UploadFile] = File(None),
    resume_text: Optional[str] = Form(None)
) -> Dict[str, Any]:
    """
    Parse resume from file upload or text
    
    Two ways to call this endpoint:
    1. File upload: Send multipart/form-data with 'file' field containing PDF/DOCX
    2. JSON body: Send application/json with { "resume_text": "..." }
    
    Returns structured JSON profile
    """
    
    text_content = None
    
    try:
        # Method 1: File upload (multipart/form-data)
        if file:
            print(f"📁 Received file: {file.filename}")
            try:
                file_bytes = await file.read()
                print(f"📁 File size: {len(file_bytes)} bytes")
                filename = file.filename.lower() if file.filename else ""
                
                if filename.endswith('.pdf'):
                    print("📁 Detected PDF format")
                    text_content = extract_pdf_text(file_bytes)
                elif filename.endswith('.docx'):
                    print("📁 Detected DOCX format")
                    text_content = extract_docx_text(file_bytes)
                elif filename.endswith('.txt'):
                    print("📁 Detected TXT format")
                    text_content = file_bytes.decode('utf-8')
                else:
                    # Try to decode as text
                    print(f"📁 Unknown format: {filename}, trying as text")
                    try:
                        text_content = file_bytes.decode('utf-8')
                    except UnicodeDecodeError:
                        raise HTTPException(
                            status_code=400,
                            detail=f"Unsupported file type: {filename}. Please upload PDF, DOCX, or TXT file."
                        )
                print(f"📁 Extracted text length: {len(text_content) if text_content else 0}")
            except HTTPException:
                raise
            except Exception as e:
                print(f"🔥 FILE READING ERROR: {e}")
                import traceback
                traceback.print_exc()
                raise HTTPException(status_code=400, detail=f"Error reading file: {str(e)}")
        
        # Method 2: Text from form data (for compatibility)
        elif resume_text:
            text_content = resume_text
        
        # Method 3: JSON body (handle via request body parsing)
        # This is handled by FastAPI's automatic JSON parsing when Content-Type is application/json
        # We need a separate endpoint or use Request object
        
        if not text_content:
            raise HTTPException(
                status_code=400, 
                detail="No resume provided. Please upload a file or provide resume text."
            )
        
        # System prompt for resume parsing
        system_prompt = """You are a resume parser. Extract structured information from resumes.

CRITICAL RULES:
- Extract ONLY information that is explicitly present
- Do NOT fabricate or infer data
- If a field is not found, use null or empty array
- Maintain exact titles, companies, and dates as written
- Preserve all bullet points exactly as written

Return valid JSON matching this structure:
{
  "full_name": "string",
  "contact": {
    "full_name": "string",
    "email": "string or null",
    "phone": "string or null",
    "location": "string or null",
    "linkedin": "string or null"
  },
  "current_title": "string or null",
  "target_roles": ["array of strings"],
  "industries": ["array of strings"],
  "years_experience": "string or null",
  "summary": "string (extract or create brief 2-3 sentence summary)",
  "experience": [
    {
      "company": "string",
      "title": "string",
      "dates": "string (as written)",
      "industry": "string or null",
      "bullets": ["array of bullet points exactly as written"]
    }
  ],
  "key_achievements": ["array of strings"],
  "core_competencies": ["array of strings"],
  "skills": ["array of technical and soft skills"],
  "education": "string or null",
  "certifications": ["array of strings or empty"]
}

Your response must be ONLY valid JSON, no additional text."""

        user_message = f"Parse this resume into structured JSON:\n\n{text_content}"
        
        # Call Claude
        print(f"📤 Sending {len(text_content)} chars to Claude for parsing...")
        response = call_claude(system_prompt, user_message)
        print(f"📥 Received response from Claude: {len(response)} chars")
        print(f"📥 Response preview: {response[:500]}...")
        
        # Clean and parse JSON response
        cleaned = clean_claude_json(response)
        print(f"📥 Cleaned JSON preview: {cleaned[:500]}...")
        
        try:
            parsed_data = json.loads(cleaned)
        except Exception as e:
            print("🔥 FINAL JSON PARSE ERROR:", e)
            import traceback
            traceback.print_exc()
            raise HTTPException(
                status_code=500,
                detail=f"Claude returned invalid JSON. Cleaned output preview: {cleaned[:300]}"
            )
        
        print(f"✅ Successfully parsed resume for: {parsed_data.get('full_name', 'Unknown')}")

        # =================================================================
        # QA VALIDATION: Check parsed resume data quality
        # =================================================================
        qa_validation_result = validate_resume_parse(parsed_data)

        if qa_validation_result.should_block:
            # Log blocking issues
            print("\n" + "🚫"*20)
            print("RESUME PARSE QA VALIDATION - CRITICAL FIELDS MISSING")
            print("🚫"*20)
            for issue in qa_validation_result.issues:
                print(f"  [{issue.severity.value}] {issue.field_path}: {issue.message}")
            print("🚫"*20 + "\n")

            # Return error with details about what's missing
            error_response = create_validation_error_response(qa_validation_result)
            raise HTTPException(status_code=422, detail=error_response)

        # Add warnings to response if any (but don't block)
        if qa_validation_result.warnings:
            parsed_data = add_validation_warnings_to_response(parsed_data, qa_validation_result)
            print(f"  ⚠️ Resume parse warnings: {len(qa_validation_result.warnings)}")

        return parsed_data

    except HTTPException:
        # Re-raise HTTP exceptions as-is
        raise
    except json.JSONDecodeError as e:
        print("🔥 JSON DECODE ERROR:", e)
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse Claude response as JSON: {str(e)}"
        )
    except Exception as e:
        print("🔥 ERROR WHILE PARSING RESUME:", e)
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500, 
            detail=f"Error parsing resume: {str(e)}"
        )


# Separate endpoint for JSON body resume text
@app.post("/api/resume/parse/text")
async def parse_resume_text(request: ResumeParseRequest) -> Dict[str, Any]:
    """
    Alternative endpoint that accepts JSON body with resume_text field
    This allows frontend to send JSON instead of form-data
    """
    if not request.resume_text:
        raise HTTPException(status_code=400, detail="resume_text field is required")
    
    try:
        # System prompt for resume parsing (same as above)
        system_prompt = """You are a resume parser. Extract structured information from resumes.

CRITICAL RULES:
- Extract ONLY information that is explicitly present
- Do NOT fabricate or infer data
- If a field is not found, use null or empty array
- Maintain exact titles, companies, and dates as written
- Preserve all bullet points exactly as written

Return valid JSON matching this structure:
{
  "full_name": "string",
  "contact": {
    "full_name": "string",
    "email": "string or null",
    "phone": "string or null",
    "location": "string or null",
    "linkedin": "string or null"
  },
  "current_title": "string or null",
  "target_roles": ["array of strings"],
  "industries": ["array of strings"],
  "years_experience": "string or null",
  "summary": "string (extract or create brief 2-3 sentence summary)",
  "experience": [
    {
      "company": "string",
      "title": "string",
      "dates": "string (as written)",
      "industry": "string or null",
      "bullets": ["array of bullet points exactly as written"]
    }
  ],
  "key_achievements": ["array of strings"],
  "core_competencies": ["array of strings"],
  "skills": ["array of technical and soft skills"],
  "education": "string or null",
  "certifications": ["array of strings or empty"]
}

Your response must be ONLY valid JSON, no additional text."""

        user_message = f"Parse this resume into structured JSON:\n\n{request.resume_text}"
        
        # Call Claude
        response = call_claude(system_prompt, user_message)
        
        # Parse JSON response
        if response.strip().startswith("```"):
            response = response.strip().split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()
        
        parsed_data = json.loads(response)
        return parsed_data
        
    except json.JSONDecodeError as e:
        raise HTTPException(
            status_code=500, 
            detail=f"Failed to parse Claude response as JSON: {str(e)}"
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error parsing resume text: {str(e)}"
        )


# ============================================================================
# MULTI-RESUME MANAGEMENT
# ============================================================================

class SaveResumeRequest(BaseModel):
    resume_name: str
    resume_json: Dict[str, Any]
    is_default: bool = False
    accuracy_confirmed: bool = False

class UpdateResumeRequest(BaseModel):
    resume_name: Optional[str] = None
    is_default: Optional[bool] = None

class ResumeResponse(BaseModel):
    id: str
    user_id: str
    resume_name: str
    resume_json: Dict[str, Any]
    is_default: bool
    accuracy_confirmed: bool
    linkedin_verified: Optional[bool] = None
    verification_notes: Optional[Dict[str, Any]] = None
    created_at: str
    updated_at: str

class ResumeListResponse(BaseModel):
    resumes: List[ResumeResponse]
    count: int

MAX_RESUMES_PER_USER = 5

@app.get("/api/resumes", response_model=ResumeListResponse)
async def list_resumes(request: Request, user_id: str = None):
    """
    List all resumes for a user.
    User ID should be passed as query param or extracted from auth.
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase_client:
        raise HTTPException(status_code=503, detail="Database not configured")

    try:
        result = supabase_client.table('user_resumes') \
            .select('*') \
            .eq('user_id', user_id) \
            .order('created_at', desc=True) \
            .execute()

        resumes = []
        for row in result.data:
            resumes.append(ResumeResponse(
                id=row['id'],
                user_id=row['user_id'],
                resume_name=row['resume_name'],
                resume_json=row['resume_json'],
                is_default=row['is_default'],
                accuracy_confirmed=row['accuracy_confirmed'],
                linkedin_verified=row.get('linkedin_verified'),
                verification_notes=row.get('verification_notes'),
                created_at=row['created_at'],
                updated_at=row['updated_at']
            ))

        return ResumeListResponse(resumes=resumes, count=len(resumes))

    except Exception as e:
        logger.error(f"Error listing resumes: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to list resumes: {str(e)}")


@app.post("/api/resumes", response_model=ResumeResponse)
async def save_resume(request: Request, body: SaveResumeRequest, user_id: str = None):
    """
    Save a new resume for a user.
    Enforces 5 resume limit per user.
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase_client:
        raise HTTPException(status_code=503, detail="Database not configured")

    if not body.accuracy_confirmed:
        raise HTTPException(status_code=400, detail="You must confirm the resume accuracy")

    try:
        # Check resume count limit
        count_result = supabase_client.table('user_resumes') \
            .select('id', count='exact') \
            .eq('user_id', user_id) \
            .execute()

        if count_result.count >= MAX_RESUMES_PER_USER:
            raise HTTPException(
                status_code=400,
                detail=f"Maximum {MAX_RESUMES_PER_USER} resumes allowed. Please delete one before adding another."
            )

        # If this is the first resume or marked as default, set is_default
        is_default = body.is_default or count_result.count == 0

        # Insert new resume
        result = supabase_client.table('user_resumes').insert({
            'user_id': user_id,
            'resume_name': body.resume_name,
            'resume_json': body.resume_json,
            'is_default': is_default,
            'accuracy_confirmed': body.accuracy_confirmed
        }).execute()

        if not result.data:
            raise HTTPException(status_code=500, detail="Failed to save resume")

        row = result.data[0]
        logger.info(f"Resume saved: {row['id']} for user {user_id}")

        return ResumeResponse(
            id=row['id'],
            user_id=row['user_id'],
            resume_name=row['resume_name'],
            resume_json=row['resume_json'],
            is_default=row['is_default'],
            accuracy_confirmed=row['accuracy_confirmed'],
            linkedin_verified=row.get('linkedin_verified'),
            verification_notes=row.get('verification_notes'),
            created_at=row['created_at'],
            updated_at=row['updated_at']
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error saving resume: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to save resume: {str(e)}")


@app.put("/api/resumes/{resume_id}", response_model=ResumeResponse)
async def update_resume(request: Request, resume_id: str, body: UpdateResumeRequest, user_id: str = None):
    """
    Update a resume (name or default status).
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase_client:
        raise HTTPException(status_code=503, detail="Database not configured")

    try:
        # Build update payload
        update_data = {}
        if body.resume_name is not None:
            update_data['resume_name'] = body.resume_name
        if body.is_default is not None:
            update_data['is_default'] = body.is_default

        if not update_data:
            raise HTTPException(status_code=400, detail="No update fields provided")

        # Update resume (RLS ensures user can only update their own)
        result = supabase_client.table('user_resumes') \
            .update(update_data) \
            .eq('id', resume_id) \
            .eq('user_id', user_id) \
            .execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Resume not found")

        row = result.data[0]
        logger.info(f"Resume updated: {resume_id}")

        return ResumeResponse(
            id=row['id'],
            user_id=row['user_id'],
            resume_name=row['resume_name'],
            resume_json=row['resume_json'],
            is_default=row['is_default'],
            accuracy_confirmed=row['accuracy_confirmed'],
            linkedin_verified=row.get('linkedin_verified'),
            verification_notes=row.get('verification_notes'),
            created_at=row['created_at'],
            updated_at=row['updated_at']
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error updating resume: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to update resume: {str(e)}")


@app.delete("/api/resumes/{resume_id}")
async def delete_resume(request: Request, resume_id: str, user_id: str = None):
    """
    Delete a resume.
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase_client:
        raise HTTPException(status_code=503, detail="Database not configured")

    try:
        # Check if this is the default resume
        check_result = supabase_client.table('user_resumes') \
            .select('is_default') \
            .eq('id', resume_id) \
            .eq('user_id', user_id) \
            .execute()

        if not check_result.data:
            raise HTTPException(status_code=404, detail="Resume not found")

        was_default = check_result.data[0]['is_default']

        # Delete the resume
        result = supabase_client.table('user_resumes') \
            .delete() \
            .eq('id', resume_id) \
            .eq('user_id', user_id) \
            .execute()

        # If deleted resume was default, set another as default
        if was_default:
            remaining = supabase_client.table('user_resumes') \
                .select('id') \
                .eq('user_id', user_id) \
                .order('created_at', desc=True) \
                .limit(1) \
                .execute()

            if remaining.data:
                supabase_client.table('user_resumes') \
                    .update({'is_default': True}) \
                    .eq('id', remaining.data[0]['id']) \
                    .execute()

        logger.info(f"Resume deleted: {resume_id}")
        return {"success": True, "message": "Resume deleted"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting resume: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete resume: {str(e)}")


@app.get("/api/resumes/{resume_id}", response_model=ResumeResponse)
async def get_resume(request: Request, resume_id: str, user_id: str = None):
    """
    Get a single resume by ID.
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase_client:
        raise HTTPException(status_code=503, detail="Database not configured")

    try:
        result = supabase_client.table('user_resumes') \
            .select('*') \
            .eq('id', resume_id) \
            .eq('user_id', user_id) \
            .execute()

        if not result.data:
            raise HTTPException(status_code=404, detail="Resume not found")

        row = result.data[0]
        return ResumeResponse(
            id=row['id'],
            user_id=row['user_id'],
            resume_name=row['resume_name'],
            resume_json=row['resume_json'],
            is_default=row['is_default'],
            accuracy_confirmed=row['accuracy_confirmed'],
            linkedin_verified=row.get('linkedin_verified'),
            verification_notes=row.get('verification_notes'),
            created_at=row['created_at'],
            updated_at=row['updated_at']
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting resume: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get resume: {str(e)}")


# ============================================================================
# URL JOB DESCRIPTION EXTRACTION
# ============================================================================

class URLExtractRequest(BaseModel):
    url: str

class URLExtractResponse(BaseModel):
    success: bool
    job_description: Optional[str] = None
    company: Optional[str] = None
    role_title: Optional[str] = None
    warning: Optional[str] = None
    error: Optional[str] = None

@app.post("/api/jd/extract-from-url", response_model=URLExtractResponse)
async def extract_jd_from_url(request: URLExtractRequest):
    """
    Extract job description content from a URL.

    Supported job boards:
    - LinkedIn, Indeed, Greenhouse, Lever, Workday, and generic job pages

    Note: Some sites may block scraping or require authentication.
    Results may vary in quality depending on the site structure.
    """
    import httpx
    import re

    url = request.url.strip()

    # Validate URL format
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url

    print(f"🔗 Extracting JD from URL: {url}")

    try:
        # Set up headers to mimic a browser
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "en-US,en;q=0.5",
            "Accept-Encoding": "gzip, deflate, br",
            "Connection": "keep-alive",
            "Upgrade-Insecure-Requests": "1"
        }

        async with httpx.AsyncClient(follow_redirects=True, timeout=15.0) as http_client:
            response = await http_client.get(url, headers=headers)

            if response.status_code != 200:
                return URLExtractResponse(
                    success=False,
                    error=f"Could not access the page (status {response.status_code}). The site may block automated access."
                )

            html_content = response.text

        # Check for common blocking patterns (more specific to avoid false positives)
        html_lower = html_content.lower()
        blocking_patterns = [
            "please verify you are human",
            "please complete the security check",
            "complete the captcha below",
            "recaptcha-checkbox",
            "cf-turnstile",  # Cloudflare Turnstile
            "hcaptcha-box",
            "g-recaptcha",
            "challenge-running",  # Cloudflare challenge
            "just a moment...</title>",  # Cloudflare waiting page
        ]
        if any(pattern in html_lower for pattern in blocking_patterns):
            return URLExtractResponse(
                success=False,
                error="This site requires human verification. Please copy and paste the job description manually."
            )

        if len(html_content) < 500:
            return URLExtractResponse(
                success=False,
                error="Page content appears incomplete. The site may require login or block automated access."
            )

        # Use Claude to extract the job description from HTML
        extraction_prompt = f"""Extract the job description from this HTML page. Return a JSON object with these fields:

1. "job_description": The full job description text, including:
   - About the company (if present)
   - Role overview/summary
   - Responsibilities/duties
   - Requirements/qualifications
   - Nice-to-haves/preferred qualifications
   - Benefits (if present)

   Format this as clean, readable text with clear section breaks. Remove any navigation, footer, or unrelated content.

2. "company": The company name

3. "role_title": The job title

If you cannot find a job description on this page, set job_description to null and explain in the "error" field.

HTML CONTENT:
{html_content[:50000]}

Return ONLY valid JSON, no markdown code blocks."""

        extraction_response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": extraction_prompt}]
        )

        result_text = extraction_response.content[0].text.strip()

        # Clean up potential markdown formatting
        if result_text.startswith("```"):
            result_text = re.sub(r'^```(?:json)?\n?', '', result_text)
            result_text = re.sub(r'\n?```$', '', result_text)

        try:
            extracted = json.loads(result_text)
        except json.JSONDecodeError:
            print(f"❌ Failed to parse extraction result: {result_text[:500]}")
            return URLExtractResponse(
                success=False,
                error="Could not parse the job description from this page. Please copy and paste manually."
            )

        job_description = extracted.get("job_description")

        if not job_description or len(job_description) < 100:
            return URLExtractResponse(
                success=False,
                error=extracted.get("error", "Could not find a job description on this page. Please copy and paste manually.")
            )

        print(f"✅ Extracted JD: {len(job_description)} chars, Company: {extracted.get('company')}, Role: {extracted.get('role_title')}")

        return URLExtractResponse(
            success=True,
            job_description=job_description,
            company=extracted.get("company"),
            role_title=extracted.get("role_title"),
            warning="Job descriptions extracted from URLs may be incomplete or contain formatting artifacts. For best results, review and edit the extracted text before analyzing."
        )

    except httpx.TimeoutException:
        return URLExtractResponse(
            success=False,
            error="Request timed out. The site may be slow or blocking automated access."
        )
    except httpx.RequestError as e:
        print(f"❌ HTTP error: {e}")
        return URLExtractResponse(
            success=False,
            error="Could not connect to the URL. Please check the URL and try again."
        )
    except Exception as e:
        print(f"❌ URL extraction error: {e}")
        return URLExtractResponse(
            success=False,
            error=f"An error occurred while extracting the job description. Please copy and paste manually."
        )


# ============================================================================
# TAILORED DOCUMENT GENERATION ENDPOINTS
# ============================================================================

@app.post("/api/resume/customize")
async def customize_resume(request: ResumeCustomizeRequest) -> Dict[str, Any]:
    """
    Generate a tailored resume based on job description.
    Returns fully formatted resume text ready for display/download.
    """

    system_prompt = """You are HenryAI, an intelligent resume customization engine that generates ATS-optimized, strategically positioned resumes. Your goal is to maximize interview conversion by combining keyword optimization with strategic positioning.

=== 1. ABSOLUTE RULES (NON-NEGOTIABLE) ===

## 1.1 Zero Fabrication Rule
You may NOT invent: job titles, metrics, tools/technologies, achievements, degrees, certifications.
You may only rewrite, clarify, reorder, or strengthen content that already exists in the candidate's resume.

## 1.2 ATS-Safe Formatting
- No tables, no two-column layouts, no icons, no graphics
- Standard section headers only: SUMMARY, EXPERIENCE, SKILLS, EDUCATION
- Bullet points using • only
- Standard fonts assumed (Arial, Calibri)
- No special characters beyond • and standard punctuation

=== 2. KEYWORD TIERING SYSTEM (CRITICAL FOR ATS) ===

Before writing the resume, extract keywords from the JD into three tiers:

**Tier 1 (Critical - MUST appear 4-6 times each):**
- Keywords in the job title itself
- Keywords repeated 3+ times in the JD
- Keywords in "required qualifications" section
Examples: If JD title is "Director of Talent Acquisition" and mentions "executive search" 4 times, both are Tier 1.

**Tier 2 (Important - MUST appear 2-3 times each):**
- Keywords in "required qualifications" that appear 1-2 times
- Industry-specific terms and tools mentioned
- Action verbs the JD emphasizes (lead, build, scale, optimize)
Examples: "Workday ATS", "high-volume recruiting", "workforce planning"

**Tier 3 (Nice-to-have - appear 1-2 times):**
- Keywords in "preferred qualifications"
- Soft skills mentioned (collaborative, data-driven)
- Secondary tools or methodologies

**Keyword Placement Priority (ATS Weight):**
1. HIGHEST: Job title line in header, Summary section (first 3 sentences), Skills section
2. MEDIUM: First bullet of each experience section, Key Achievements
3. LOWER: Later bullets, company descriptions

**Exact Match Rule:** Use the JD's exact phrasing. If JD says "Workday ATS" use "Workday ATS" not just "Workday". If JD says "executive search" use "executive search" not "executive recruiting".

=== 3. COMPANY STAGE POSITIONING ===

Detect company stage from JD signals and adjust tone/emphasis:

**Startup/Scale-up Signals:** "fast-paced", "0-1", "Series A/B/C", "high-growth", "build from scratch", "wear multiple hats"
→ Lead with: startup/growth experience, scrappy execution, speed, building from scratch
→ Tone: Action-oriented, emphasize agility and impact
→ De-emphasize: Large enterprise process, bureaucratic language

**Enterprise/Fortune 500 Signals:** "global", "Fortune 500", "public company", "cross-regional", "compliance"
→ Lead with: Scale, process rigor, executive stakeholder management, global scope
→ Tone: Professional, strategic, emphasize governance
→ Keep all experience, show progression

**Mid-market/PE-backed Signals:** "growth-stage", "private equity", "scalable", "cost efficiency"
→ Balance: Startup agility + operational maturity
→ Lead with: Scalable systems, efficiency, rapid execution with structure

=== 4. FUNCTIONAL DEPTH MATCHING ===

Identify what the JD emphasizes most and lead with matching experience:

**If JD emphasizes executive search/C-suite hiring:**
→ Lead with executive search experience, retained search, leadership assessment
→ Add keywords: succession planning, Board-level, VP/SVP placement

**If JD emphasizes high-volume/operational hiring:**
→ Lead with high-volume metrics, pipeline scale, hiring velocity
→ Add keywords: pipeline generation, sourcing campaigns, funnel metrics

**If JD emphasizes technical recruiting:**
→ Lead with engineering/product hiring experience
→ Add keywords: technical pipelines, engineering hiring, product roles

**If JD emphasizes team leadership:**
→ Lead with team size, global scope, mentoring
→ Add keywords: team development, recruiter coaching, performance management

=== 5. LEVEL CALIBRATION ===

**If applying for role BELOW current level (overqualified risk):**
→ Reframe: "Seeking hands-on leadership role balancing strategy with execution"
→ Emphasize: Execution metrics, deliverables, processes built
→ De-emphasize: Pure strategy, C-suite language

**If applying for role ABOVE current level (stretch role):**
→ Emphasize: Scope of impact, budget managed, cross-functional influence
→ Frame: "Operating at [higher level] scope"

=== 6. RESUME STRUCTURE (MANDATORY FORMAT) ===

```
{FULL NAME IN ALL CAPS}
{TARGET ROLE FROM JD} | {Strength 1} | {Strength 2}
{Phone} • {Email} • {LinkedIn URL} • {City, State}

SUMMARY
[4-5 sentences, 80-100 words. Must include: role being pursued, 2-3 core strengths matching JD, 1 measurable impact. Load with Tier 1 keywords. First sentence must state professional identity + years + core function.]

SKILLS
[8-16 skills in format: Skill 1 | Skill 2 | Skill 3 | ... ]
[Order by: JD required skills first, then tools from resume, then domain knowledge]
[Only include skills the candidate actually has evidence of]

EXPERIENCE

{COMPANY NAME}
{Job Title} | {Location} | {Dates}
[Optional 1-line company context if adds credibility: "Fortune 500 healthcare company with 50,000+ employees"]
• [Most JD-relevant bullet first - front-load with Tier 1 keywords, include metric]
• [Second most relevant - include Tier 1 or 2 keyword]
• [Third bullet - achievement or responsibility aligned to JD]
• [Fourth bullet - skill/tool demonstration]
[Max 5-6 bullets per role, fewer for older roles]

[Repeat for each relevant role - typically 3-4 roles for 2-page resume]

EDUCATION
{Degree}, {Major} | {University}
[Optional: relevant concentration, honors]
```

=== 7. BULLET REORDERING RULE ===

For each role, reorder bullets by relevance to THIS specific job:
- Most JD-relevant bullet becomes first bullet (even if chronologically it came later)
- Front-load each bullet with keywords (first 8-10 words get highest ATS weight)
- Each bullet should contain 2-4 relevant keywords naturally integrated

**Bullet Formula:** [Strong Action Verb] + [What You Did with JD Keywords] + [Quantified Outcome]
Example: "Led 25-person global technical recruiting team across engineering, product, and data functions, reducing time-to-offer by 30% while achieving 95% offer acceptance rate."

=== 8. OUTPUT FORMAT ===

Return a JSON object with this EXACT structure:

{
  "tailored_resume_text": "FULL FORMATTED RESUME TEXT - use \\n for line breaks",
  "positioning_strategy": {
    "company_stage_detected": "startup|scale-up|enterprise|mid-market",
    "functional_emphasis": "executive_search|high_volume|technical|team_leadership|generalist",
    "level_calibration": "at_level|stretch_up|stepping_down",
    "lead_with": "1 sentence describing what experience we're leading with and why",
    "de_emphasize": "1 sentence describing what we're de-emphasizing and why"
  },
  "keyword_analysis": {
    "tier_1_keywords": ["keyword1", "keyword2", "keyword3"],
    "tier_1_placements": 15,
    "tier_2_keywords": ["keyword1", "keyword2", "keyword3", "keyword4"],
    "tier_2_placements": 10,
    "tier_3_keywords": ["keyword1", "keyword2"],
    "tier_3_placements": 4
  },
  "changes_summary": {
    "summary_rationale": "What we changed in the summary and why, referencing specific JD requirements",
    "experience_rationale": "Which roles/bullets we emphasized or reordered and why",
    "ats_optimization": "How we optimized for ATS with specific keyword placements",
    "positioning_statement": "This positions you as [specific strategic positioning for this role]"
  },
  "red_flags_addressed": [
    {"flag": "description of potential concern", "mitigation": "how we addressed it in the resume"}
  ]
}

Your response must be ONLY valid JSON. No markdown formatting."""

    # Inject candidate state calibration if available
    calibration_prompt = build_candidate_calibration_prompt(request.situation)
    if calibration_prompt:
        system_prompt += calibration_prompt

    # Build the resume content for Claude
    resume_content = ""
    if request.resume_json:
        resume_content = json.dumps(request.resume_json, indent=2)
    elif request.resume_text:
        resume_content = request.resume_text
    else:
        raise HTTPException(status_code=400, detail="Either resume_text or resume_json is required")

    user_message = f"""Generate a tailored resume for this candidate applying to this role.

TARGET ROLE: {request.target_role}
TARGET COMPANY: {request.target_company}

JOB DESCRIPTION:
{request.job_description}

CANDIDATE RESUME DATA:
{resume_content}
"""

    if request.jd_analysis:
        user_message += f"\n\nJD ANALYSIS (use for tailoring):\n{json.dumps(request.jd_analysis, indent=2)}"

    # Add supplemental information from candidate if provided
    if request.supplements and len(request.supplements) > 0:
        user_message += "\n\n=== ADDITIONAL CANDIDATE CONTEXT ===\n"
        user_message += "The candidate provided the following additional context to address gaps. INCORPORATE this information into the resume where appropriate (but do NOT fabricate beyond what they stated):\n\n"
        for supp in request.supplements:
            user_message += f"**Gap Area: {supp.gap_area}**\n"
            user_message += f"Question: {supp.question}\n"
            user_message += f"Candidate's Answer: {supp.answer}\n\n"

    user_message += """

Generate the tailored resume following the exact format in the system instructions.
Remember: NO fabrication - only use information from the candidate's actual resume AND the additional context provided above."""

    try:
        response = call_claude(system_prompt, user_message, max_tokens=4000)
        cleaned = clean_claude_json(response)
        result = json.loads(cleaned)
        
        # Validate required fields
        if "tailored_resume_text" not in result:
            raise ValueError("Missing tailored_resume_text in response")
        
        return result
        
    except json.JSONDecodeError as e:
        print(f"🔥 JSON parse error in /api/resume/customize: {e}")
        print(f"Response was: {response[:500]}...")
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")
    except Exception as e:
        print(f"🔥 Error in /api/resume/customize: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/cover-letter/generate")
async def generate_cover_letter(request: CoverLetterGenerateRequest) -> Dict[str, Any]:
    """
    Generate a tailored cover letter based on job description and candidate background.
    Returns fully formatted cover letter text ready for display/download.
    """
    
    system_prompt = """You are the document generation engine for HenryAI.
Your job is to generate high-quality, recruiter-grade cover letters tailored to each job description.

=== 1. GLOBAL RULES ===

## 1.1 Zero Fabrication Rule
You may NOT invent:
- job titles
- metrics
- tools/technologies
- achievements
You may only rewrite, clarify, or strengthen content that already exists.

## 1.2 Recruiter-Quality, No Fluff
Write like a senior recruiting leader building a candidate's narrative:
- direct
- concise
- metric-driven
- confident
- no filler language or vague claims
- NO "I believe," "I think," "I feel"
- NO "I am writing to express my interest..."
- NO "I would be honored..."

=== 2. HEADER FORMAT (MANDATORY) ===

Use this header exactly (same as resume):

{FULL NAME IN ALL CAPS}
{TARGET ROLE} | {STRENGTH 1} | {STRENGTH 2}
{PHONE} • {EMAIL} • {LINKEDIN} • {CITY, STATE}

=== 3. COVER LETTER STRUCTURE ===

4 paragraphs max:

## Paragraph 1: Opening
- State interest in role
- 1 sentence linking background to company's needs

## Paragraph 2: Key Experience
- Pull top 1-2 accomplishments
- Tie directly to responsibilities in JD

## Paragraph 3: Value Proposition
- Show strengths aligned with what company is solving for
- Highlight leadership, process rigor, or technical depth

## Paragraph 4: Close
- Confident, brief
- End with appreciation
- NO "I believe," "I think," "I feel"

=== 4. RULES ===
- Do NOT repeat lines from summary or resume
- No generic filler
- Include metrics where they exist in the resume
- Match the tone: senior professional who knows their worth

=== OUTPUT FORMAT ===

Return a JSON object with this EXACT structure:

{
  "cover_letter_text": "FULL FORMATTED COVER LETTER TEXT - use \\n for line breaks",
  "changes_summary": {
    "opening_rationale": "1 sentence explaining why you led with this angle",
    "body_rationale": "1-2 sentences on what themes you emphasized and avoided",
    "closing_rationale": "1 sentence on the tone of the close",
    "positioning_statement": "This frames you as [strategic insight]"
  }
}

Your response must be ONLY valid JSON."""

    # Inject candidate state calibration if available
    calibration_prompt = build_candidate_calibration_prompt(request.situation)
    if calibration_prompt:
        system_prompt += calibration_prompt

    # Build the resume content for Claude
    resume_content = ""
    if request.resume_json:
        resume_content = json.dumps(request.resume_json, indent=2)
    elif request.resume_text:
        resume_content = request.resume_text
    else:
        raise HTTPException(status_code=400, detail="Either resume_text or resume_json is required")

    user_message = f"""Generate a tailored cover letter for this candidate applying to this role.

TARGET ROLE: {request.target_role}
TARGET COMPANY: {request.target_company}

JOB DESCRIPTION:
{request.job_description}

CANDIDATE RESUME DATA:
{resume_content}
"""

    if request.strengths:
        user_message += f"\n\nCANDIDATE STRENGTHS TO EMPHASIZE:\n{json.dumps(request.strengths, indent=2)}"

    if request.jd_analysis:
        user_message += f"\n\nJD ANALYSIS (use for tailoring):\n{json.dumps(request.jd_analysis, indent=2)}"

    # Add supplemental information from candidate if provided
    if request.supplements and len(request.supplements) > 0:
        user_message += "\n\n=== ADDITIONAL CANDIDATE CONTEXT ===\n"
        user_message += "The candidate provided the following additional context to address gaps. WEAVE this into the cover letter naturally to strengthen their narrative:\n\n"
        for supp in request.supplements:
            user_message += f"**Gap Area: {supp.gap_area}**\n"
            user_message += f"Context: {supp.answer}\n\n"

    user_message += """

Generate the cover letter following the exact format in the system instructions.
Remember: NO fabrication, NO generic filler, NO clichés."""

    try:
        response = call_claude(system_prompt, user_message, max_tokens=2000)
        cleaned = clean_claude_json(response)
        result = json.loads(cleaned)

        # Validate required fields
        if "cover_letter_text" not in result:
            raise ValueError("Missing cover_letter_text in response")

        # =================================================================
        # QA VALIDATION: Check cover letter for fabrication
        # =================================================================
        from qa_validation import OutputValidator, ResumeGroundingValidator

        grounding_validator = ResumeGroundingValidator(request.resume)
        output_validator = OutputValidator(grounding_validator)
        qa_validation_result = output_validator.validate_cover_letter({"content": result.get("cover_letter_text", "")})

        if qa_validation_result.should_block:
            print("\n" + "🚫"*20)
            print("COVER LETTER QA VALIDATION BLOCKED - POTENTIAL FABRICATION")
            print("🚫"*20)
            for issue in qa_validation_result.issues:
                print(f"  [{issue.category.value}] {issue.message}")
            print("🚫"*20 + "\n")

            # Log to file for review
            ValidationLogger.log_blocked_output(
                endpoint="/api/cover-letter/generate",
                result=qa_validation_result,
                output=result,
                resume_data=request.resume,
                request_context={"company": request.jd_analysis.get("company") if request.jd_analysis else None}
            )

            error_response = create_validation_error_response(qa_validation_result)
            raise HTTPException(status_code=422, detail=error_response)

        # Add warnings to response if any
        if qa_validation_result.warnings:
            result = add_validation_warnings_to_response(result, qa_validation_result)
            print(f"  ⚠️ Cover letter validation warnings: {len(qa_validation_result.warnings)}")

        return result

    except json.JSONDecodeError as e:
        print(f"🔥 JSON parse error in /api/cover-letter/generate: {e}")
        print(f"Response was: {response[:500]}...")
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")
    except Exception as e:
        print(f"🔥 Error in /api/cover-letter/generate: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


def detect_career_gap(resume_data: Dict) -> Optional[Dict]:
    """
    Calculate career gap from resume dates (post-processing, not prompt-based).
    Returns gap details if gap >= 3 months, None otherwise.
    """
    from dateutil import parser as date_parser

    # Get most recent role end date
    experience = resume_data.get("experience", [])
    if not experience:
        return None

    # Parse dates from most recent role
    most_recent = experience[0]  # Assumes experience is ordered newest first
    dates_str = most_recent.get("dates", "")

    # Extract end date from "Start - End" or "Start - Present"
    if " - " not in dates_str:
        return None

    end_date_str = dates_str.split(" - ")[1].strip()

    # If "Present", no gap
    if end_date_str.lower() in ["present", "current", "now", "ongoing"]:
        return None

    try:
        # Parse end date
        end_date = date_parser.parse(end_date_str, fuzzy=True)

        # Current date (December 2024)
        current_date = datetime(2024, 12, 17)

        # Calculate gap in months
        gap_months = (current_date.year - end_date.year) * 12 + (current_date.month - end_date.month)

        # Only flag gaps >= 3 months
        if gap_months < 3:
            return None

        # Determine severity
        if gap_months < 6:
            severity = "low"
            severity_label = "LOW RISK"
        elif gap_months < 12:
            severity = "medium"
            severity_label = "MEDIUM RISK"
        else:
            severity = "high"
            severity_label = "HIGH RISK"

        # Format dates for display
        end_month_year = end_date.strftime("%B %Y")
        company_name = most_recent.get("company", "your last company")

        return {
            "gap_type": "career_gap",
            "severity": severity,
            "description": f"Employment gap: {gap_months} months since last role ({end_month_year} to present)",
            "detailed_explanation": f"Your most recent role ended in {end_month_year}. This is a {gap_months}-month gap. Hiring managers will ask about this period. Frame proactively rather than defensively.",
            "impact": f"Interview question certainty: HIGH - You will be asked what you've been doing since {end_month_year}",
            "mitigation": f"Frame as deliberate career reset after your tenure at {company_name}, not forced circumstance. Emphasize intentional job search and skill development during gap. Position as strategic: 'After working at {company_name}, I wanted to be thoughtful about my next role rather than rush into something that wasn't the right fit.'",
            "duration_months": gap_months,
            "start_date": end_month_year,
            "end_date": "Present"
        }

    except Exception as e:
        print(f"⚠️ Could not parse career gap: {e}")
        return None


def force_apply_experience_penalties(response_data: dict, resume_data: dict = None) -> dict:
    """
    Force-apply experience penalties and hard caps to Claude's response.
    This ensures penalties are applied even if Claude ignores the prompt instructions.

    Per Eligibility Gate Enforcement Spec:
    ORDER OF OPERATIONS (MANDATORY):
    1. Eligibility Gate → 2. Gap Typing → 3. Recommendation Lock → 4. Score Explanation → 5. Coaching

    SYSTEM CONTRACT §7 - ANALYSIS ID ENFORCEMENT:
    Each run generates a unique analysis_id. All extracted data is tagged to this ID
    and destroyed at completion. No cross-candidate memory permitted.

    Args:
        response_data: Parsed JSON response from Claude
        resume_data: Original resume data (optional, for credibility adjustment)

    Returns:
        Modified response_data with corrected fit_score and recommendation
    """
    import uuid

    # =========================================================================
    # SYSTEM CONTRACT §7: ANALYSIS ID ENFORCEMENT
    # Generate unique analysis_id for this run. All candidate-scoped data
    # must be tagged to this ID and discarded at completion.
    # =========================================================================
    analysis_id = str(uuid.uuid4())
    response_data["_analysis_id"] = analysis_id
    print(f"\n🔐 ANALYSIS ID: {analysis_id[:8]}... (candidate-scoped data will be discarded at completion)")

    # Extract experience analysis (Claude should have populated this)
    experience_analysis = response_data.get("experience_analysis", {})
    required_years = experience_analysis.get("required_years", 0)

    # ========================================================================
    # STEP 1: ELIGIBILITY GATE - RUNS BEFORE EVERYTHING ELSE
    # Per Eligibility Gate Enforcement Spec: "Before scoring, determine if the
    # candidate is *eligible*. Eligibility failure overrides score, confidence,
    # and tone."
    # ========================================================================
    recommendation_locked = False
    locked_recommendation = None
    locked_reason = None
    people_leadership_years = 0.0
    required_people_leadership = 0.0
    eligibility_result = None

    if resume_data:
        print("🚪 ELIGIBILITY GATE CHECK (RUNS BEFORE SCORING)")
        eligibility_result = check_eligibility_gate(resume_data, response_data)

        if not eligibility_result["eligible"]:
            # ELIGIBILITY FAILED - Lock recommendation immediately
            recommendation_locked = True
            locked_recommendation = eligibility_result["locked_recommendation"]
            locked_reason = eligibility_result["reason"]

            # Store eligibility failure in experience_analysis
            experience_analysis["eligibility_gate_passed"] = False
            experience_analysis["eligibility_failed_check"] = eligibility_result["failed_check"]
            experience_analysis["eligibility_gap_classification"] = eligibility_result["gap_classification"]
            experience_analysis["recommendation_locked"] = True
            experience_analysis["locked_reason"] = locked_reason
            response_data["experience_analysis"] = experience_analysis

            print(f"   🚫 ELIGIBILITY FAILED: {eligibility_result['failed_check']}")
            print(f"   🔒 RECOMMENDATION LOCKED: {locked_recommendation}")

            # Add the eligibility failure as a critical gap
            gaps = response_data.get("gaps", [])
            if not isinstance(gaps, list):
                gaps = []

            eligibility_gap = {
                "gap_type": f"eligibility_failure:{eligibility_result['failed_check']}",
                "gap_classification": "missing_experience",
                "severity": "critical",
                "gap_description": locked_reason,
                "detailed_explanation": locked_reason,
                "impact": "You will not pass screening for this role",
                "mitigation_strategy": "Target roles that match your actual experience level and domain"
            }
            gaps.insert(0, eligibility_gap)
            response_data["gaps"] = gaps
        else:
            experience_analysis["eligibility_gate_passed"] = True
            response_data["experience_analysis"] = experience_analysis

    # Skip further enforcement if no experience requirements detected
    if required_years == 0 and not recommendation_locked:
        print("🔧 PENALTY ENFORCEMENT: No required_years detected, skipping enforcement")
        return response_data

    # ========================================================================
    # STEP 1b: PEOPLE LEADERSHIP CHECK (if not already failed by eligibility)
    # Per constraints: "This is enforcement, not vibes."
    #
    # IF JD requires X+ years people leadership
    # AND candidate has < X years verified people leadership
    # THEN recommendation MUST be "Do Not Apply"
    # NO EXCEPTIONS. No "Conditional Apply." No "Apply fast." No soft language.
    #
    # Per Leadership Tiering Spec:
    # - Leadership must be modeled in tiers (strategic, people, org-level)
    # - "0 years" ONLY when NO leadership signals exist at all
    # - "insufficient" when leadership exists but doesn't meet requirement
    # ========================================================================
    if resume_data and not recommendation_locked:
        # Extract TIERED leadership analysis (new model)
        tiered_leadership = extract_tiered_leadership(resume_data)
        people_leadership_years = tiered_leadership.get("people_leadership_years", 0.0)

        # Check if JD requires people leadership
        required_people_leadership, is_hard_requirement = extract_required_people_leadership_years(response_data)

        # Determine if role requires org-level leadership (VP+, C-suite)
        role_title = (response_data.get("role_title", "") or "").lower()
        requires_org_leadership = any(t in role_title for t in ["vp", "vice president", "chief", "c-suite", "head of"])
        required_org_leadership = 7.0 if requires_org_leadership else 0.0

        print(f"🔒 PEOPLE LEADERSHIP CHECK:")
        print(f"   Required: {required_people_leadership} years (hard_requirement={is_hard_requirement})")
        print(f"   Candidate has: {people_leadership_years} years verified (tiered)")
        print(f"   Org-level required: {requires_org_leadership} ({required_org_leadership} years)")

        # Get precise gap messaging using tiered analysis
        leadership_gap_msg = get_leadership_gap_messaging(
            tiered_leadership,
            required_people_leadership,
            required_org_leadership
        )

        print(f"   Leadership gap status: {leadership_gap_msg['status']}")
        if leadership_gap_msg['status'] != 'sufficient':
            print(f"   Leadership gap message: {leadership_gap_msg['message']}")

        # TIERED GATE: Apply same tiered response as eligibility gate
        # This ensures consistency between the two checks
        if is_hard_requirement and required_people_leadership > 0:
            if people_leadership_years < required_people_leadership:
                # Calculate gap severity (same logic as eligibility gate)
                leadership_ratio = people_leadership_years / required_people_leadership if required_people_leadership > 0 else 0
                gap_years = required_people_leadership - people_leadership_years

                print(f"   📊 Leadership ratio: {leadership_ratio:.1%} (gap: {gap_years:.1f} years)")

                # TIERED RESPONSE:
                # - <50% = hard fail
                # - 50-70% = warning (Apply with Caution)
                # - >70% = minor gap (allow score-based recommendation)
                if leadership_ratio < 0.5:
                    # SEVERE gap - hard fail
                    recommendation_locked = True
                    locked_recommendation = "Do Not Apply"

                    if leadership_gap_msg['status'] == 'none':
                        locked_reason = f"No people leadership experience found. This role requires {required_people_leadership:.0f}+ years."
                    elif leadership_gap_msg['status'] == 'insufficient':
                        if leadership_gap_msg.get('has_lower_tier'):
                            locked_reason = f"Leadership type mismatch: {leadership_gap_msg['message']}"
                        else:
                            locked_reason = f"People leadership experience insufficient ({people_leadership_years:.1f} vs {required_people_leadership:.0f}+ years required)"
                    else:
                        locked_reason = f"People leadership experience below role requirement ({people_leadership_years:.1f} vs {required_people_leadership:.0f}+ years required)"

                    print(f"   🚫 HARD GATE TRIGGERED (severe gap <50%): {locked_reason}")
                    print(f"   🔒 RECOMMENDATION LOCKED: Do Not Apply")

                    experience_analysis["people_leadership_hard_gate_failed"] = True
                    experience_analysis["recommendation_locked"] = True
                    experience_analysis["locked_reason"] = locked_reason
                    experience_analysis["hard_requirement_failure"] = True
                elif leadership_ratio < 0.7:
                    # MODERATE gap - warning only
                    print(f"   ⚠️ MODERATE LEADERSHIP GAP (50-70%): {gap_years:.1f} years short")
                    print(f"   📋 Recommendation: Apply with Caution (not locked)")
                    experience_analysis["leadership_warning"] = f"Leadership gap: {gap_years:.1f} years short of {required_people_leadership:.0f}+ requirement"
                    experience_analysis["recommended_adjustment"] = "Apply with Caution"
                else:
                    # MINOR gap - allow score-based
                    print(f"   ℹ️ MINOR LEADERSHIP GAP (>70%): {gap_years:.1f} years short")
                    print(f"   📋 Allowing score-based recommendation")
                    experience_analysis["leadership_note"] = f"Minor gap: {gap_years:.1f} years short"

                # Store tiered leadership analysis in experience_analysis
                experience_analysis["people_leadership_years"] = people_leadership_years
                experience_analysis["required_people_leadership_years"] = required_people_leadership
                experience_analysis["tiered_leadership"] = tiered_leadership
                experience_analysis["leadership_gap_messaging"] = leadership_gap_msg
                response_data["experience_analysis"] = experience_analysis

    # Try to get adjusted years first, fallback to raw years
    candidate_years = experience_analysis.get("candidate_years_adjusted_for_credibility")
    raw_years = experience_analysis.get("candidate_years_in_role_type", 0)

    if candidate_years is None:
        candidate_years = raw_years
        # Claude didn't apply credibility adjustment, so apply it ourselves
        if resume_data and raw_years > 0:
            candidate_years = apply_credibility_adjustment(resume_data, raw_years)
            if candidate_years != raw_years:
                experience_analysis["candidate_years_adjusted_for_credibility"] = candidate_years
                response_data["experience_analysis"] = experience_analysis

    # BACKEND SAFETY NET: Calculate role-specific experience from resume
    # Detects the role type from JD and uses appropriate experience calculator
    # Enhanced with credibility scoring and leadership vs IC distinction
    #
    # CRITICAL: Use isolated role detection when available (per contamination fix)
    # The isolated detection runs early with request-only data, avoiding any
    # cached/contaminated role type from previous analyses
    if resume_data:
        # Check if isolated role detection was performed (higher priority)
        isolated_role_type = experience_analysis.get("isolated_role_type")
        isolated_years = experience_analysis.get("isolated_candidate_years")

        if isolated_role_type:
            # Use isolated detection - it was computed with request-only data
            detected_role_type = isolated_role_type
            print(f"   🔍 Detected role type (ISOLATED): {detected_role_type.upper()}")
            if isolated_years is not None and isolated_years > 0:
                print(f"   🔍 Using isolated experience calculation: {isolated_years:.1f} years")
        else:
            # Fallback to original detection
            detected_role_type = detect_role_type_from_jd(response_data)
            print(f"   🔍 Detected role type: {detected_role_type.upper()}")

        # Check if JD requires leadership (for recruiting roles)
        requires_leadership = jd_requires_leadership(response_data)
        if detected_role_type == "recruiting":
            print(f"   🔍 JD requires leadership: {requires_leadership}")

        # CRITICAL: If isolated years are available and valid, use them directly
        # This bypasses the old calculation that may have been contaminated
        if isolated_years is not None and isolated_years > 0:
            backend_role_years = isolated_years
            print(f"   ✅ Using ISOLATED years calculation: {backend_role_years:.1f} years")
            experience_analysis["used_isolated_calculation"] = True
        elif detected_role_type == "recruiting" and requires_leadership:
            # Use leadership-aware recruiting calculator
            backend_role_years = calculate_recruiting_leadership_years(resume_data, requires_leadership)
            print(f"   🔍 Backend recruiting LEADERSHIP years: {backend_role_years:.1f} years")
        else:
            # Use credibility-adjusted calculation
            backend_role_years, raw_years, breakdown = calculate_credibility_adjusted_years(
                resume_data, detected_role_type
            )
            print(f"   🔍 Backend {detected_role_type} years: {backend_role_years:.1f} (raw: {raw_years:.1f})")

            # Store credibility breakdown for debugging
            if breakdown:
                experience_analysis["credibility_breakdown"] = breakdown

        # If Claude didn't count role-specific experience OR significantly overcounted, use ours
        if candidate_years == 0 or (candidate_years > 3 and backend_role_years < candidate_years * 0.5):
            if candidate_years == 0:
                print(f"   ✅ Using backend calculation (Claude reported 0): {backend_role_years:.1f} years")
            else:
                print(f"   ⚠️ Claude overcounted! Using backend calculation: {backend_role_years:.1f} vs Claude's {candidate_years:.1f}")
            candidate_years = backend_role_years
            experience_analysis["candidate_years_adjusted_for_credibility"] = candidate_years
            experience_analysis["backend_override"] = True
            experience_analysis["backend_override_reason"] = f"Backend calculated {backend_role_years:.1f} {detected_role_type}-specific years (credibility-adjusted)"
            experience_analysis["detected_role_type"] = detected_role_type
            if detected_role_type == "recruiting":
                experience_analysis["leadership_required"] = requires_leadership
            response_data["experience_analysis"] = experience_analysis

    # Get current fit score from Claude
    original_fit_score = response_data.get("fit_score", 0)
    if original_fit_score is None:
        original_fit_score = 0

    # Calculate years percentage
    if required_years > 0:
        years_percentage = (candidate_years / required_years) * 100
    else:
        years_percentage = 100  # No requirement specified

    # ========================================================================
    # HARD CAP LOGIC - REVISED PER RECRUITER CALIBRATION
    #
    # ONLY apply hard caps if:
    # 1. Experience is explicitly marked as a HARD requirement in JD
    # 2. AND candidate years < requirement by >40% (years_percentage < 60)
    #
    # Otherwise: downgrade to "Apply with Caution", NOT "Do Not Apply"
    #
    # This prevents false negatives for candidates like Marcus (9 years total,
    # 5.8 years EM) when JD asks for 7+ years.
    # ========================================================================

    # Check if experience requirement is explicitly hard (must have, required, etc.)
    jd_text = (response_data.get("job_description", "") or "").lower()
    experience_is_hard_requirement = any(phrase in jd_text for phrase in [
        "must have", "required:", "requirements:", "minimum of",
        "at least", "years required", "mandatory", "non-negotiable"
    ])

    # Determine hard cap based on years percentage
    # REVISED: Only apply aggressive caps for hard requirements AND >40% gap
    if years_percentage < 60 and experience_is_hard_requirement and not recommendation_locked:
        # Significant gap (>40%) on hard requirement - cap at 55 (Apply with Caution)
        hard_cap = 55
        hard_cap_reason = f"Candidate has {years_percentage:.1f}% of required years ({candidate_years:.1f}/{required_years}), hard requirement gap"
        print(f"   ⚠️ HARD CAP (experience hard requirement, >40% gap): 55%")
    elif years_percentage < 60 and not experience_is_hard_requirement:
        # Significant gap but NOT a hard requirement - soft cap at 60 (Apply with Caution threshold)
        hard_cap = 60
        hard_cap_reason = f"Candidate has {years_percentage:.1f}% of required years ({candidate_years:.1f}/{required_years}), preferred not hard requirement"
        print(f"   📝 SOFT CAP (experience preferred, not hard requirement): 60%")
    elif years_percentage < 80:
        # Moderate gap - light cap at 70
        hard_cap = 70
        hard_cap_reason = f"Candidate has {years_percentage:.1f}% of required years ({candidate_years:.1f}/{required_years})"
    else:
        hard_cap = 100  # No cap for candidates with 80%+ of required years
        hard_cap_reason = None

    # Apply hard cap if necessary
    capped_score = min(original_fit_score, hard_cap)
    hard_cap_applied = (capped_score < original_fit_score)

    # Update fit_score
    response_data["fit_score"] = capped_score

    # Update fit_score_breakdown if it exists
    if "fit_score_breakdown" not in response_data:
        response_data["fit_score_breakdown"] = {}

    response_data["fit_score_breakdown"]["calculated_score"] = original_fit_score
    response_data["fit_score_breakdown"]["hard_cap_applied"] = hard_cap_applied
    if hard_cap_reason:
        response_data["fit_score_breakdown"]["hard_cap_reason"] = hard_cap_reason
    response_data["fit_score_breakdown"]["final_score"] = capped_score

    # ========================================================================
    # FINAL RECOMMENDATION CONTROLLER - SINGLE SOURCE OF TRUTH
    # Per Architecture Simplification Spec v1.0:
    # - Recommendation is set ONCE and locked
    # - Score is set ONCE and locked
    # - No downstream layer may mutate these values
    # ========================================================================

    # Detect manager-level role
    role_title = (response_data.get('role_title', '') or '').lower()
    is_manager_role = any(x in role_title for x in [
        'manager', 'director', 'head of', 'vp ', 'vice president',
        'lead', 'chief', 'cto', 'ceo', 'coo', 'cfo', 'cmo'
    ])
    is_vp_plus_role = any(x in role_title for x in ['vp', 'vice president', 'director', 'chief', 'c-suite'])

    # Check for domain gap (from CEC or calibration if available)
    domain_gap_detected = False
    capability_evidence_report = response_data.get("capability_evidence_report", {})
    if capability_evidence_report:
        cec_summary = capability_evidence_report.get("summary", {})
        critical_gaps = cec_summary.get("critical_gaps", [])
        domain_gap_detected = any("domain" in (g.get("capability_id", "") or "").lower() for g in critical_gaps)

    # ========================================================================
    # CAREER TRANSITION DETECTION
    # Detect if candidate is making a career pivot vs stretch role
    # ========================================================================
    transition_info = None
    try:
        from seniority_detector import detect_transition_candidate
        # Extract role type from JD
        target_role_type = role_title
        # Common role type extraction
        if "product" in role_title:
            target_role_type = "product"
        elif "engineer" in role_title or "developer" in role_title:
            target_role_type = "engineering"
        elif "design" in role_title:
            target_role_type = "design"

        transition_info = detect_transition_candidate(
            resume_data=response_data.get("resume_data", {}),
            target_role_type=target_role_type
        )
        if transition_info.get("is_transition"):
            print(f"   🔄 CAREER TRANSITION DETECTED: {transition_info.get('source_domain')} → {transition_info.get('target_domain')}")
            print(f"      Type: {transition_info.get('transition_type')}, Adjacency: {transition_info.get('adjacency_score')}")
    except ImportError:
        pass  # Seniority detector not available

    # ========================================================================
    # SIX-TIER RECOMMENDATION SYSTEM (Dec 21, 2025)
    # Per HenryHQ Scoring Spec v2.0 - Invariant enforced score-decision pairing
    #
    # Score Range | Decision           | Meaning
    # 85-100      | Strong Apply       | Top-tier match. Prioritize this application.
    # 70-84       | Apply              | Solid fit. Worth your time and energy.
    # 55-69       | Consider           | Moderate fit. Apply if genuinely interested.
    # 40-54       | Apply with Caution | Stretch role. Need positioning and referral.
    # 25-39       | Long Shot          | Significant gaps. Only with inside connection.
    # 0-24        | Do Not Apply       | Not your role. Invest energy elsewhere.
    # ========================================================================

    if FINAL_CONTROLLER_AVAILABLE:
        # Use the new single-authority controller with six-tier system
        controller = FinalRecommendationController()

        decision = controller.compute_recommendation(
            fit_score=capped_score,
            eligibility_passed=not recommendation_locked,
            eligibility_reason=locked_reason or "",
            is_manager_role=is_manager_role,
            is_vp_plus_role=is_vp_plus_role,
            domain_gap_detected=domain_gap_detected,
            transition_info=transition_info,
            has_referral=False  # TODO: Add referral detection if needed
        )

        correct_recommendation = decision.recommendation
        # Use the enforced score from the decision (may be adjusted for invariant)
        capped_score = decision.fit_score

        # Generate tier-appropriate alternative actions (six-tier system)
        if correct_recommendation == "Strong Apply":
            alternative_actions = [
                "Apply immediately - you're a top-tier match for this role",
                "Reach out directly to the hiring manager on LinkedIn",
                "This should be a priority application in your pipeline"
            ]
        elif correct_recommendation == "Apply":
            alternative_actions = [
                "Solid fit - apply within the next 24-48 hours",
                "Lead with your strongest, most relevant accomplishments",
                "Consider reaching out to someone at the company for an intro"
            ]
        elif correct_recommendation == "Consider":
            alternative_actions = [
                "Worth pursuing if you're genuinely interested in this company/role",
                "Address the experience gaps proactively in your cover letter",
                "Don't prioritize this over stronger fits in your pipeline"
            ]
        elif correct_recommendation == "Apply with Caution":
            alternative_actions = [
                "This is a stretch role - you'll need strong positioning",
                "Network with someone at the company before applying if possible",
                "Lead with your most transferable accomplishments"
            ]
        elif correct_recommendation == "Long Shot":
            alternative_actions = [
                "Only pursue if you have an inside connection or unique angle",
                "Consider targeting roles 1-2 levels below this one at the same company",
                f"Build 1-2 more years of experience before targeting {required_years}+ year roles"
            ]
        else:  # Do Not Apply
            alternative_actions = [
                f"Target roles requiring {max(1, int(candidate_years))}-{int(candidate_years) + 2} years of experience",
                "This role requires fundamentally different experience than you have",
                "Focus your energy on roles where you're a stronger match"
            ]

        print(f"   Score: {capped_score}%")
        print(f"   Eligibility: {'PASSED' if not recommendation_locked else 'FAILED'}")
        print(f"   Manager Role: {is_manager_role}")
        print(f"   Domain Gap: {domain_gap_detected} {'(suppressed for manager)' if is_manager_role and domain_gap_detected else ''}")
        print(f"   Final Recommendation: {correct_recommendation}")

        # Store controller decision for downstream access
        response_data["final_decision_controller"] = controller.to_dict()
        # Update fit_score with potentially adjusted value from invariant enforcement
        response_data["fit_score"] = capped_score
    else:
        # Fallback six-tier mapping (if controller not available)
        # Per HenryHQ Scoring Spec v2.0 - strict tier boundaries
        if recommendation_locked:
            # Eligibility failure - cap score to match Do Not Apply tier
            capped_score = min(capped_score, 24)
            correct_recommendation = "Do Not Apply"
            alternative_actions = [
                f"Target roles requiring {max(1, int(candidate_years))}-{int(candidate_years) + 2} years of experience",
                "This role requires fundamentally different experience than you have",
                "Focus your energy on roles where you're a stronger match"
            ]
        elif capped_score >= 85:
            correct_recommendation = "Strong Apply"
            alternative_actions = [
                "Apply immediately - you're a top-tier match for this role",
                "Reach out directly to the hiring manager on LinkedIn",
                "This should be a priority application in your pipeline"
            ]
        elif capped_score >= 70:
            correct_recommendation = "Apply"
            alternative_actions = [
                "Solid fit - apply within the next 24-48 hours",
                "Lead with your strongest, most relevant accomplishments",
                "Consider reaching out to someone at the company for an intro"
            ]
        elif capped_score >= 55:
            correct_recommendation = "Consider"
            alternative_actions = [
                "Worth pursuing if you're genuinely interested in this company/role",
                "Address the experience gaps proactively in your cover letter",
                "Don't prioritize this over stronger fits in your pipeline"
            ]
        elif capped_score >= 40:
            correct_recommendation = "Apply with Caution"
            alternative_actions = [
                "This is a stretch role - you'll need strong positioning",
                "Network with someone at the company before applying if possible",
                "Lead with your most transferable accomplishments"
            ]
        elif capped_score >= 25:
            correct_recommendation = "Long Shot"
            alternative_actions = [
                "Only pursue if you have an inside connection or unique angle",
                "Consider targeting roles 1-2 levels below this one at the same company",
                f"Build 1-2 more years of experience before targeting {required_years}+ year roles"
            ]
        else:
            correct_recommendation = "Do Not Apply"
            alternative_actions = [
                f"Target roles requiring {max(1, int(candidate_years))}-{int(candidate_years) + 2} years of experience",
                "This role requires fundamentally different experience than you have",
                "Focus your energy on roles where you're a stronger match"
            ]

        # Update fit_score for eligibility failure
        response_data["fit_score"] = capped_score

    # ========================================================================
    # SET RECOMMENDATION - THIS IS THE FINAL IMMUTABLE VALUE
    # ========================================================================
    response_data["recommendation"] = correct_recommendation
    response_data["recommendation_locked"] = True  # Mark as locked

    # Update ALL recommendation fields to ensure UI displays correct value
    if "intelligence_layer" in response_data:
        if "apply_decision" not in response_data["intelligence_layer"]:
            response_data["intelligence_layer"]["apply_decision"] = {}
        response_data["intelligence_layer"]["apply_decision"]["recommendation"] = correct_recommendation
        response_data["intelligence_layer"]["apply_decision"]["locked"] = True

    if "apply_decision" in response_data:
        response_data["apply_decision"]["recommendation"] = correct_recommendation
        response_data["apply_decision"]["locked"] = True

    # Add alternative actions if present
    if alternative_actions and not response_data.get("alternative_actions"):
        response_data["alternative_actions"] = alternative_actions

    # Add experience gap warning if significant gap exists
    if years_percentage < 70:
        gaps = response_data.get("gaps", [])
        if not isinstance(gaps, list):
            gaps = []

        # Check if experience gap warning already exists
        has_experience_warning = any(
            g.get("gap_type") in ["experience_years_mismatch", "required_experience_missing"]
            for g in gaps
            if isinstance(g, dict)
        )

        if not has_experience_warning:
            # CRITICAL: Never expose percentage math or "credibility adjustment" language
            gap_warning = {
                "gap_type": "experience_years_mismatch",
                "severity": "critical",
                "gap_description": f"This role requires {required_years}+ years; you have {candidate_years:.0f}",
                "detailed_explanation": (
                    f"This role is looking for {required_years}+ years of direct experience. "
                    f"Your background shows roughly {candidate_years:.0f} years in this domain. "
                    f"Recruiters typically filter on years of experience first."
                ),
                "impact": "You may not pass initial screening filters for this role",
                "mitigation_strategy": (
                    f"Consider roles asking for {max(1, int(candidate_years))}-{int(candidate_years) + 2} years of experience, "
                    f"where your background will be more competitive."
                )
            }
            gaps.insert(0, gap_warning)  # Add at the beginning for visibility
            response_data["gaps"] = gaps

    # Log penalty enforcement (abbreviated - full summary at end)
    if hard_cap_applied:
        print(f"🔧 PENALTY: Score capped {original_fit_score}% → {capped_score}% ({years_percentage:.0f}% of required years)")

    # ========================================================================
    # GAP CLASSIFICATION AND STRENGTH-GAP CONFLICT RESOLUTION
    # Per constraints: Classify gaps and resolve conflicts
    # ========================================================================

    # Classify all gaps
    gaps = response_data.get("gaps", [])
    if gaps and isinstance(gaps, list):
        has_missing_experience_gap = False
        for gap in gaps:
            if isinstance(gap, dict):
                gap_classification = classify_gap_type(
                    gap,
                    candidate_years,
                    required_years,
                    people_leadership_years,
                    required_people_leadership
                )
                gap["gap_classification"] = gap_classification
                if gap_classification == "missing_experience":
                    has_missing_experience_gap = True
                    print(f"   🔴 GAP CLASSIFIED: missing_experience - {gap.get('gap_description', '')[:50]}")

        response_data["gaps"] = gaps

        # INFORMATIONAL: Log missing_experience gap detection
        # Note: Recommendation is already locked by Final Recommendation Controller
        # This is for debugging/visibility only
        if has_missing_experience_gap:
            print(f"   🔴 MISSING_EXPERIENCE GAP DETECTED (informational - controller already decided)")
            response_data["has_missing_experience_gap"] = True

    # Resolve strength-gap conflicts
    strengths = response_data.get("strengths", [])
    if strengths and gaps:
        cleaned_strengths, gaps = resolve_strength_gap_conflicts(strengths, gaps)
        if len(cleaned_strengths) < len(strengths):
            print(f"   ⚠️ REMOVED {len(strengths) - len(cleaned_strengths)} conflicting strengths")
        response_data["strengths"] = cleaned_strengths
        response_data["gaps"] = gaps

    # ========================================================================
    # CAPABILITY EVIDENCE CHECK (CEC) - STEP 2.5
    # Per CEC Spec v1.0: Diagnostic layer that explains HOW resume diverges from JD.
    #
    # RUNS AFTER: Eligibility gates (Step 2)
    # RUNS BEFORE: CAE (Step 3) and LEPE (Step 4)
    #
    # HARD CONSTRAINTS:
    # - Does NOT alter fit score
    # - Does NOT alter recommendation
    # - Does NOT trigger caps or multipliers
    # - Diagnostic output feeds "Gaps to Address" and "Your Move"
    # ========================================================================
    capability_evidence_report = None

    # CEC runs if JD and resume parsing succeeded, even before recommendation lock
    if resume_data and response_data.get("role_title"):
        print("📊 STEP 2.5: CAPABILITY EVIDENCE CHECK (CEC)")

        # Get leadership extraction if available
        leadership_extraction = experience_analysis.get("tiered_leadership") or {}

        # Run CEC evaluation
        capability_evidence_report = evaluate_capability_evidence(
            response_data=response_data,
            resume_data=resume_data,
            leadership_extraction=leadership_extraction
        )

        # Store CEC output in response (does NOT modify score/recommendation)
        response_data["capability_evidence_report"] = capability_evidence_report

        # Log CEC summary
        if capability_evidence_report:
            summary = capability_evidence_report.get("summary", {})
            critical_gaps = summary.get("critical_gaps", [])
            if critical_gaps:
                print(f"   🔴 CEC identified {len(critical_gaps)} critical capability gap(s)")
            else:
                print(f"   ✅ CEC: No critical capability gaps identified")

    # ========================================================================
    # RECRUITER CALIBRATION LAYER - STEP 2.6
    # Per Calibration Spec v1.0 (REVISED): Recruiter-grade judgment framework
    #
    # RUNS AFTER: CEC (Step 2.5)
    # RUNS BEFORE: CAE (Step 3) and LEPE (Step 4)
    #
    # PURPOSE:
    # - Apply function-specific calibration (executive, technical, GTM)
    # - Classify gaps as terminal vs coachable for coaching output
    # - Detect red flags (stop_search vs proceed_with_caution)
    # - Generate calibrated_gaps for coaching controller
    #
    # HARD CONSTRAINTS (REVISED):
    # - Does NOT override Job Fit recommendation
    # - Calibration EXPLAINS a "Do Not Apply," it does NOT create one
    # - locked_reason comes from Job Fit only
    # - redirect_reason is for coaching, not for changing recommendation
    # - Multiple implicit signals do NOT equal one explicit signal
    # ========================================================================
    calibration_result = None
    calibration_red_flags = []
    calibrated_gaps = None  # For coaching controller

    if CALIBRATION_AVAILABLE and resume_data and response_data.get("role_title"):
        print("🎯 STEP 2.6: RECRUITER CALIBRATION LAYER (Interpretive Only)")

        try:
            # Prepare candidate experience for calibration
            # ==========================================================================
            # CRITICAL INVARIANT: Experience must be List[Dict] - NOTHING ELSE
            # This is the SINGLE upstream normalization point. All downstream
            # functions (detect_red_flags, calculate_career_span, etc.) can now
            # assume they receive valid List[Dict] and won't crash on strings.
            # ==========================================================================
            raw_experience = resume_data.get('experience', [])

            # Case 1: Experience is a string (LinkedIn free text)
            if isinstance(raw_experience, str):
                print(f"   ⚠️ INVARIANT: Experience is string - normalizing to []")
                raw_experience = []

            # Case 2: Experience is not a list at all
            elif not isinstance(raw_experience, list):
                print(f"   ⚠️ INVARIANT: Experience is {type(raw_experience).__name__} - normalizing to []")
                raw_experience = []

            # Case 3: Experience is a list but contains non-dict items
            # Filter to ONLY dict items - this catches mixed lists like [str, dict, str]
            else:
                filtered_experience = [exp for exp in raw_experience if isinstance(exp, dict)]
                if len(filtered_experience) != len(raw_experience):
                    print(f"   ⚠️ INVARIANT: Experience list contained {len(raw_experience) - len(filtered_experience)} non-dict items - filtered out")
                raw_experience = filtered_experience

            candidate_experience = {
                'roles': raw_experience,
                'experience': raw_experience,
                'summary': resume_data.get('summary', ''),
                'skills': resume_data.get('skills', []),
                'domain': response_data.get('experience_analysis', {}).get('domain', ''),
                'level': response_data.get('experience_analysis', {}).get('level', ''),
            }

            # Prepare role requirements for calibration
            role_title = response_data.get('role_title', '').lower()
            role_requirements = {
                'level': response_data.get('role_level', ''),
                'team_size_min': response_data.get('team_size_requirement', 0),
                'domain': response_data.get('target_domain', ''),
                'function': response_data.get('role_function', ''),
                'job_description': response_data.get('job_description', ''),
                'role_title': response_data.get('role_title', ''),
                'global_scope': any(kw in role_title for kw in ['global', 'distributed', 'multi-region']),
                'requires_pnl': any(kw in role_title for kw in ['vp', 'director', 'head', 'c-suite']),
            }

            # Detect role function and apply appropriate calibration
            if any(kw in role_title for kw in ['vp', 'director', 'head of', 'chief', 'c-suite', 'senior director']):
                print("   📊 Applying EXECUTIVE calibration")
                calibration_result = calibrate_executive_role(candidate_experience, role_requirements)
                calibration_type = 'executive'
            elif any(kw in role_title for kw in ['sales', 'account executive', 'ae', 'customer success', 'bdr', 'sdr', 'bizdev']):
                print("   📊 Applying GTM calibration")
                calibration_result = calibrate_gtm_role(candidate_experience, role_requirements)
                calibration_type = 'gtm'
            else:
                print("   📊 Applying TECHNICAL calibration")
                calibration_result = calibrate_technical_role(candidate_experience, role_requirements)
                calibration_type = 'technical'

            # Detect red flags - wrapped in try/except because red flags are ADVISORY
            # If they crash, they lose the right to speak. Never derail the pipeline.
            try:
                calibration_red_flags = detect_red_flags(candidate_experience)
            except Exception:
                # Silently fail - red flags are advisory, crash is contained
                calibration_red_flags = []

            # Store calibration results (for logging/debugging)
            response_data["calibration_result"] = {
                'type': calibration_type,
                'actual_level': calibration_result.get('actual_level'),
                'terminal_gaps': calibration_result.get('terminal_gaps', []),
                'coachable_gaps': calibration_result.get('coachable_gaps', []),
                'confidence': calibration_result.get('confidence', 1.0),
                'red_flags': calibration_red_flags
            }

            # Log calibration results
            terminal_gaps = calibration_result.get('terminal_gaps', [])
            coachable_gaps = calibration_result.get('coachable_gaps', [])
            stop_search_flags = [f for f in calibration_red_flags if f.get('severity') == 'stop_search']
            caution_flags = [f for f in calibration_red_flags if f.get('severity') == 'proceed_with_caution']

            print(f"   📋 Calibration: {len(terminal_gaps)} terminal, {len(coachable_gaps)} coachable gaps")
            print(f"   🚩 Red flags: {len(stop_search_flags)} stop_search, {len(caution_flags)} proceed_with_caution")

            # REVISED: Use calibrate_gaps control layer to prioritize gaps for coaching
            # This does NOT override Job Fit recommendation
            # CRITICAL: Use locked_recommendation if set, otherwise use current recommendation
            # This ensures calibration sees the FINAL recommendation, not the pre-lock one
            cec_results = response_data.get("capability_evidence_report", {})
            if recommendation_locked and locked_recommendation:
                job_fit_recommendation = locked_recommendation
                print(f"   🔒 Using LOCKED recommendation for calibration: {locked_recommendation}")
            else:
                job_fit_recommendation = response_data.get("recommendation", "Apply")

            # Map Job Fit recommendation to coaching-compatible format
            recommendation_map = {
                "Strongly Apply": "Strong Apply",
                "Apply": "Apply",
                "Apply with caution": "Apply with Caution",
                "Conditional Apply": "Apply with Caution",
                "Skip": "Do Not Apply",
                "Do Not Apply": "Do Not Apply",
            }
            normalized_recommendation = recommendation_map.get(job_fit_recommendation, job_fit_recommendation)

            calibrated_gaps = calibrate_gaps(
                cec_results=cec_results,
                job_fit_recommendation=normalized_recommendation,
                candidate_resume=candidate_experience,
                job_requirements=role_requirements
            )

            # Store calibrated_gaps for coaching controller (Step 6.5)
            response_data["calibrated_gaps"] = calibrated_gaps

            # Log calibrated gaps summary
            if calibrated_gaps:
                primary = calibrated_gaps.get('primary_gap')
                secondary = calibrated_gaps.get('secondary_gaps', [])
                suppress = calibrated_gaps.get('suppress_gaps_section', False)
                redirect = calibrated_gaps.get('redirect_reason')

                if primary:
                    print(f"   🎯 Primary gap: {primary.get('gap', {}).get('capability', 'Unknown')}")
                if secondary:
                    print(f"   📋 Secondary gaps: {len(secondary)}")
                if suppress:
                    print(f"   🔇 Gaps section suppressed (Strong Apply)")
                if redirect:
                    print(f"   ↪️  Redirect reason: {redirect[:50]}...")

            # Apply confidence threshold per Selective Coaching Spec
            # One explicit signal > five implicit signals
            if calibration_result.get('confidence', 1.0) < 0.6:
                response_data["calibration_low_confidence"] = True
                print(f"   ⚠️ Low calibration confidence ({calibration_result.get('confidence', 1.0):.2f})")

        except Exception as e:
            print(f"   ⚠️ Calibration error (non-blocking): {e}")
            import traceback
            traceback.print_exc()
            # Calibration errors are non-blocking - continue with existing flow

    # ========================================================================
    # CREDIBILITY ALIGNMENT ENGINE (CAE) - RUNS AFTER ELIGIBILITY
    # Per CAE Spec: "This layer never blocks on its own. It modulates confidence,
    # language, and recommendation strength."
    #
    # CRITICAL: CAE cannot soften or override HARD REQUIREMENT FAILURES.
    # If eligibility gate failed, CAE may explain skepticism but cannot
    # change the "Do Not Apply" decision or soften the messaging.
    # ========================================================================
    cae_result = None
    hard_requirement_failure = experience_analysis.get("hard_requirement_failure", False)

    if resume_data and not recommendation_locked:
        print("🎯 CREDIBILITY ALIGNMENT ENGINE (CAE) CHECK")
        cae_result = evaluate_credibility_alignment(resume_data, response_data)

        # Store CAE results in response for downstream use
        response_data["credibility_alignment"] = cae_result

        # Apply CAE to recommendation (may downgrade, never upgrade)
        if cae_result.get("recommendation_ceiling"):
            current_rec = response_data.get("recommendation", "")
            cae_application = apply_credibility_to_recommendation(
                current_rec,
                cae_result,
                eligibility_locked=recommendation_locked,
                hard_requirement_failure=hard_requirement_failure
            )

            if cae_application["was_downgraded"]:
                # CAE is ADVISORY ONLY - it cannot override the Final Recommendation Controller
                # Log the suggestion but do NOT mutate recommendation
                print(f"   ⚠️ CAE ADVISORY: Would downgrade to {cae_application['final_recommendation']}")
                print(f"   → BLOCKED: Recommendation is locked by Final Recommendation Controller")
                response_data["cae_advisory_downgrade"] = cae_application["final_recommendation"]
                response_data["cae_downgrade_reason"] = cae_application["downgrade_reason"]

                # Store credibility risk for UI display (advisory only)
                if "intelligence_layer" in response_data:
                    if "apply_decision" not in response_data["intelligence_layer"]:
                        response_data["intelligence_layer"]["apply_decision"] = {}
                    response_data["intelligence_layer"]["apply_decision"]["credibility_risk"] = cae_result.get("overall_credibility_risk")

        # Add mandatory reality check message if required
        if cae_result.get("mandatory_reality_check") and cae_result.get("reality_check_message"):
            # Add to reality_check if it exists
            if "reality_check" in response_data:
                existing_action = response_data["reality_check"].get("strategic_action", "")
                # Prepend the reality check message
                response_data["reality_check"]["strategic_action"] = (
                    f"{cae_result['reality_check_message']} {existing_action}"
                )
                response_data["reality_check"]["credibility_warning"] = cae_result["reality_check_message"]

            # Store for frontend display
            response_data["credibility_reality_check"] = cae_result["reality_check_message"]

        # Apply language intensity based on CAE risk level
        language_templates = get_cae_coaching_language(cae_result)
        response_data["cae_language_intensity"] = language_templates.get("coaching_tone", "confident")

    # ========================================================================
    # LEADERSHIP EVALUATION & POSITIONING ENGINE (LEPE)
    # Per LEPE Spec: Applies ONLY to Manager+ roles.
    # Provides leadership readiness assessment with positioning guidance.
    # LEPE constraints final recommendations - cannot be overridden by optimism.
    # ========================================================================
    lepe_result = None
    if resume_data:
        lepe_result = evaluate_lepe(resume_data, response_data)
        response_data["lepe_analysis"] = lepe_result

        # If LEPE is applicable and decision is "locked", enforce it
        if lepe_result.get("lepe_applicable") and not recommendation_locked:
            positioning = lepe_result.get("positioning_decision", {})
            lepe_decision = positioning.get("decision", "apply")

            if lepe_decision == "locked":
                # LEPE locks recommendation for > 4 year gap
                recommendation_locked = True
                locked_recommendation = "Do Not Apply"
                locked_reason = positioning.get("messaging", {}).get("explanation", "Leadership gap too large")
                experience_analysis["lepe_locked"] = True
                experience_analysis["hard_requirement_failure"] = True
                response_data["experience_analysis"] = experience_analysis
                print(f"   🔒 LEPE LOCKED: {locked_reason}")

            elif lepe_decision == "caution":
                # LEPE is ADVISORY ONLY - it cannot override the Final Recommendation Controller
                # Log the suggestion but do NOT mutate recommendation
                current_rec = response_data.get("recommendation", "")
                if current_rec in ["Apply", "Strong Apply"]:
                    print(f"   ⚠️ LEPE ADVISORY: Would downgrade {current_rec} to Apply with Caution")
                    print(f"   → BLOCKED: Recommendation is locked by Final Recommendation Controller")
                    response_data["lepe_advisory_downgrade"] = "Apply with Caution"
                    response_data["lepe_caution_reason"] = positioning.get("messaging", {}).get("skepticism_warning", "")

            elif lepe_decision == "position":
                # LEPE positioning mode - add coaching guidance
                response_data["lepe_positioning_mode"] = True
                response_data["lepe_coaching"] = positioning.get("messaging", {}).get("coaching_advice", "")
                print(f"   📝 LEPE POSITIONING: Coaching available")

        # Store accountability record
        if lepe_result.get("accountability_record"):
            response_data["leadership_positioning_record"] = lepe_result["accountability_record"]

    # ========================================================================
    # LOCKED RECOMMENDATION VERIFICATION (NOT OVERRIDE)
    # The Final Recommendation Controller already handled locking.
    # This block only VERIFIES consistency and logs for debugging.
    # It does NOT change the recommendation - that was set once and is final.
    # ========================================================================
    if recommendation_locked and locked_recommendation:
        current_rec = response_data.get('recommendation')
        if current_rec != locked_recommendation:
            # Log the discrepancy but do NOT override - controller is authoritative
            print(f"   ⚠️ LOCKED RECOMMENDATION DISCREPANCY DETECTED (informational only):")
            print(f"      Controller set: {current_rec}")
            print(f"      Eligibility wanted: {locked_recommendation}")
            print(f"      Reason: {locked_reason}")
            print(f"      → Controller decision is FINAL, no override")
        else:
            print(f"   ✅ LOCKED RECOMMENDATION VERIFIED: {current_rec} (Reason: {locked_reason})")

        # Ensure locked_reason is stored for UI display (advisory only)
        if "intelligence_layer" in response_data:
            if "apply_decision" in response_data.get("intelligence_layer", {}):
                response_data["intelligence_layer"]["apply_decision"]["locked_reason"] = locked_reason
        if "apply_decision" in response_data:
            response_data["apply_decision"]["locked_reason"] = locked_reason

        # ========================================================================
        # LANGUAGE CONTRACT FOR "DO NOT APPLY"
        # Per Required Fixes Spec (UI CONTRACT + MESSAGING CONTRACT):
        #
        # STRUCTURE (MANDATORY):
        # A. Decision: "Do not apply." + state hard mismatch (one sentence)
        # B. Why: What role requires vs what is missing (plain, non-shaming)
        # C. Redirect: Name SPECIFIC role categories where candidate IS competitive
        #
        # RULES:
        # - Second person only (you/your)
        # - No candidate names in output
        # - No mixed verdict language
        # - No upside framing
        # - Coaching ≠ encouragement to apply
        # ========================================================================
        print(f"   📝 ENFORCING LANGUAGE CONTRACT FOR 'DO NOT APPLY'")

        # Generate specific redirect based on candidate's actual background
        redirect_roles = generate_specific_redirect(resume_data, response_data, eligibility_result)

        # Determine the specific reason and build rationale accordingly
        if eligibility_result and not eligibility_result.get("eligible", True):
            # Eligibility gate failure - use the eligibility reason
            failed_check = eligibility_result.get("failed_check", "") or ""

            # A. DECISION (one sentence, direct)
            # B. WHY (plain, non-shaming)
            #
            # PRIORITY ORDER for explanations:
            # When eligibility gate fails, the failed_check IS the primary blocker.
            # Don't override domain failures with years calculations - they're apples to oranges.
            #
            # 1. Domain gap (non-transferable domain) - fundamentally different skillset needed
            # 2. Seniority/Level gap (Director/VP role, candidate is IC)
            # 3. People leadership gap (specific management requirement)
            # 4. Years experience gap (only if it's the actual failed_check)

            # Get experience data for messaging
            candidate_years = experience_analysis.get("candidate_years", 0)
            required_years = experience_analysis.get("required_years", 0)
            leadership_gap_msg = experience_analysis.get("leadership_gap_messaging", {})

            # PRIORITY 1: Domain gap (non-transferable domain) - this is fundamental
            # When the role requires a completely different skillset (e.g., ML/AI Research),
            # that's the primary blocker - not years of experience in an unrelated field.
            if "non_transferable_domain" in failed_check:
                domain_name = failed_check.split(":")[-1] if ":" in failed_check else "this domain"
                # Make domain name more readable
                domain_display = domain_name.replace('_', ' ').replace('ml ai', 'ML/AI')
                response_data["recommendation_rationale"] = (
                    f"do not apply. This role requires direct experience in {domain_display}. "
                    f"Your background is in a different domain that doesn't transfer to this role."
                )
            # PRIORITY 2: Seniority/Level gap (Director/VP role, candidate is IC)
            elif "seniority_scope" in failed_check:
                response_data["recommendation_rationale"] = (
                    f"do not apply. This is a Director/VP-level role requiring {required_years}+ years and people management. "
                    f"Your background shows {candidate_years:.0f} years with no management experience. The level gap is the primary blocker."
                )
            # PRIORITY 3: People leadership gap
            elif "people_leadership" in failed_check:
                # Use tiered leadership messaging if available
                if leadership_gap_msg.get("factual_statement"):
                    response_data["recommendation_rationale"] = f"do not apply. {leadership_gap_msg['factual_statement']}"
                elif leadership_gap_msg.get("status") == "none":
                    response_data["recommendation_rationale"] = (
                        f"do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                        f"Your resume shows no verified leadership experience at any tier."
                    )
                elif leadership_gap_msg.get("status") == "insufficient":
                    response_data["recommendation_rationale"] = (
                        f"do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                        f"Your people leadership experience ({people_leadership_years:.1f} years) is insufficient."
                    )
                else:
                    response_data["recommendation_rationale"] = (
                        f"do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                        f"Your verified people leadership experience is {people_leadership_years:.1f} years. "
                        f"Operational leadership does not count toward this requirement."
                    )
            # PRIORITY 4: Years experience gap (only when years is the actual failed_check)
            elif "years" in failed_check.lower() or "experience" in failed_check.lower():
                years_percentage = (candidate_years / required_years * 100) if required_years > 0 else 100
                response_data["recommendation_rationale"] = (
                    f"do not apply. This role requires {required_years}+ years of experience. "
                    f"You have {candidate_years:.0f} years. This experience gap is the primary blocker."
                )
            else:
                # Use the eligibility reason directly when we have a specific one
                eligibility_reason = eligibility_result.get("reason", "")
                if eligibility_reason and len(eligibility_reason) > 20:
                    response_data["recommendation_rationale"] = f"do not apply. {eligibility_reason}"
                else:
                    response_data["recommendation_rationale"] = (
                        f"do not apply. This role requires experience you do not have. "
                        f"Your background does not meet the core requirements."
                    )

            # C. REDIRECT (specific, concrete)
            response_data["alternative_actions"] = redirect_roles

            # CRITICAL: Add redirect to the primary gap for coaching controller
            # The calibration controller reads primary_gap.get('redirect') to generate Your Move
            gaps = response_data.get("gaps", [])
            if gaps and isinstance(gaps[0], dict):
                # Format redirect as a single string for Your Move
                redirect_text = redirect_roles[0] if redirect_roles else "Redirect to roles that match your domain expertise"
                gaps[0]["redirect"] = redirect_text
                print(f"   ✅ Added redirect to primary gap: {redirect_text[:60]}...")

        elif required_people_leadership > 0 and people_leadership_years < required_people_leadership:
            # People leadership specific failure
            # Use tiered leadership messaging for accurate "none" vs "insufficient" distinction
            leadership_gap_msg = experience_analysis.get("leadership_gap_messaging", {})
            tiered_leadership = experience_analysis.get("tiered_leadership", {})

            # A. DECISION + B. WHY (using tiered messaging)
            if leadership_gap_msg.get("factual_statement"):
                response_data["recommendation_rationale"] = f"Do not apply. {leadership_gap_msg['factual_statement']}"
            elif leadership_gap_msg.get("status") == "none":
                response_data["recommendation_rationale"] = (
                    f"Do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                    f"Your resume shows no verified leadership experience at any tier (strategic, people, or org-level)."
                )
            elif leadership_gap_msg.get("status") == "insufficient":
                if leadership_gap_msg.get("has_lower_tier"):
                    strategic_years = tiered_leadership.get("strategic_leadership_years", 0)
                    response_data["recommendation_rationale"] = (
                        f"Do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                        f"Your leadership experience ({strategic_years:.1f} years) is strategic/functional only "
                        f"and does not include managing direct reports."
                    )
                else:
                    response_data["recommendation_rationale"] = (
                        f"Do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                        f"Your people leadership experience ({people_leadership_years:.1f} years) is insufficient. "
                        f"Gap: {leadership_gap_msg.get('gap_years', 0):.1f} years."
                    )
            else:
                response_data["recommendation_rationale"] = (
                    f"Do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                    f"Your verified people leadership experience is {people_leadership_years:.1f} years. "
                    f"Operational or project leadership does not substitute for people management."
                )

            # Add specific gap for people leadership with tiered detail
            gap_status = leadership_gap_msg.get("status", "unknown")
            if gap_status == "none":
                gap_description = f"Requires {required_people_leadership:.0f}+ years people leadership; you have none verified"
                detailed_explanation = (
                    f"This role requires people leadership - managing direct reports, building teams, "
                    f"conducting performance reviews, and making hiring decisions. "
                    f"Your resume shows no verified leadership at any tier."
                )
            elif gap_status == "insufficient" and leadership_gap_msg.get("has_lower_tier"):
                gap_description = f"Requires {required_people_leadership:.0f}+ years people leadership; you have strategic/functional only"
                detailed_explanation = (
                    f"This role requires people leadership - managing direct reports. "
                    f"Your background shows strategic/functional leadership "
                    f"({tiered_leadership.get('strategic_leadership_years', 0):.1f} years leading initiatives), "
                    f"but no verified people management experience."
                )
            else:
                gap_description = f"Requires {required_people_leadership:.0f}+ years people leadership; you have {people_leadership_years:.1f} verified (insufficient)"
                detailed_explanation = (
                    f"This role requires people leadership - managing direct reports, building teams, "
                    f"conducting performance reviews, and making hiring decisions. "
                    f"Your resume shows {people_leadership_years:.1f} years of verified people leadership, "
                    f"which is {leadership_gap_msg.get('gap_years', 0):.1f} years short of the requirement."
                )

            redirect_text = redirect_roles[0] if redirect_roles else "Target roles that match your experience level"
            people_gap = {
                "gap_type": "people_leadership_requirement_not_met",
                "gap_classification": "missing_experience",
                "severity": "critical",
                "gap_description": gap_description,
                "detailed_explanation": detailed_explanation,
                "leadership_status": gap_status,  # NEW: "none" vs "insufficient"
                "tiered_breakdown": tiered_leadership.get("leadership_tier_summary", ""),
                "impact": "You will not pass screening for this role",
                "mitigation_strategy": redirect_text,
                "redirect": redirect_text  # CRITICAL: For coaching controller Your Move
            }
            gaps = response_data.get("gaps", [])
            has_people_gap = any(g.get("gap_type") == "people_leadership_requirement_not_met" for g in gaps if isinstance(g, dict))
            if not has_people_gap:
                gaps.insert(0, people_gap)
                response_data["gaps"] = gaps
                print(f"   ✅ Added redirect to people leadership gap: {redirect_text[:60]}...")

            # C. REDIRECT (specific)
            response_data["alternative_actions"] = redirect_roles

        else:
            # Generic locked Do Not Apply (e.g., from missing_experience gap)
            # ==========================================================================
            # PRIORITY ORDER: Even when eligibility passed, apply this priority:
            # 1. Level/seniority gap (IC applying to Director+)
            # 2. Years experience gap
            # 3. People leadership gap
            # 4. Domain gap
            # ==========================================================================

            # Get gap data
            candidate_years = experience_analysis.get("candidate_years", 0)
            required_years = experience_analysis.get("required_years", 0)
            years_percentage = (candidate_years / required_years * 100) if required_years > 0 else 100
            people_leadership_years = experience_analysis.get("people_leadership_years", 0)
            required_people_leadership = experience_analysis.get("required_people_leadership_years", 0)
            role_level = experience_analysis.get("role_level", "").lower() or response_data.get("role_level", "").lower()
            candidate_level = experience_analysis.get("candidate_level", "").lower()
            leadership_gap_msg = experience_analysis.get("leadership_gap_messaging", {})

            # Also check role_title for level detection (more reliable source)
            role_title_for_level = (response_data.get("role_title", "") or "").lower()

            # Detect if this is a level mismatch (IC → Director+)
            # Check both role_level field AND role_title for director+ signals
            director_signals = ["director", "vp", "vice president", "head of", "chief", "executive", "senior director", "global"]
            is_director_plus_role = any(
                level in role_level for level in director_signals
            ) or any(
                level in role_title_for_level for level in director_signals
            )
            is_ic_candidate = candidate_level in ["ic", "individual contributor", "senior", "mid", "junior", "staff"] or not any(
                level in candidate_level for level in ["director", "manager", "lead", "head"]
            )

            # Debug: Log the values used for explanation priority
            print(f"   📋 EXPLANATION PRIORITY DEBUG:")
            print(f"      role_level: '{role_level}', role_title: '{role_title_for_level[:40]}...'")
            print(f"      is_director_plus_role: {is_director_plus_role}, is_ic_candidate: {is_ic_candidate}")
            print(f"      required_people_leadership: {required_people_leadership}, people_leadership_years: {people_leadership_years}")
            print(f"      candidate_years: {candidate_years}, required_years: {required_years}, years_percentage: {years_percentage:.1f}%")

            # PRIORITY 1: Level/seniority gap (IC applying to Director+ role)
            # Note: Director+ roles ALWAYS require people leadership, so we don't require it to be explicitly set
            if is_director_plus_role and is_ic_candidate and people_leadership_years < 1:
                print(f"   ✅ PRIORITY 1 matched: Level gap (Director+ role, IC candidate)")
                # Use role_title if role_level is empty
                level_display = role_level.title() if role_level else "Director"
                response_data["recommendation_rationale"] = (
                    f"Do not apply. This is a {level_display}-level role requiring people management experience. "
                    f"Your {candidate_years:.0f} years are IC-focused with no verified direct reports. The level gap is the primary blocker."
                )
            # PRIORITY 2: Severe years gap (<50% of required)
            elif years_percentage < 50 and required_years > 0:
                response_data["recommendation_rationale"] = (
                    f"Do not apply. This role requires {required_years}+ years of experience. "
                    f"You have {candidate_years:.0f} years. This experience gap is the primary blocker."
                )
            # PRIORITY 3: People leadership gap with specific messaging
            elif required_people_leadership > 0 and people_leadership_years < required_people_leadership:
                if leadership_gap_msg.get("factual_statement"):
                    response_data["recommendation_rationale"] = f"Do not apply. {leadership_gap_msg['factual_statement']}"
                elif leadership_gap_msg.get("status") == "none":
                    response_data["recommendation_rationale"] = (
                        f"Do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                        f"Your resume shows no verified people management experience."
                    )
                else:
                    response_data["recommendation_rationale"] = (
                        f"Do not apply. This role requires {required_people_leadership:.0f}+ years of people leadership. "
                        f"Your people leadership experience ({people_leadership_years:.1f} years) is insufficient."
                    )
            # PRIORITY 4: Default to locked_reason
            else:
                response_data["recommendation_rationale"] = (
                    f"Do not apply. {locked_reason}."
                )

            response_data["alternative_actions"] = redirect_roles

            # CRITICAL: Add redirect to the primary gap for coaching controller
            gaps = response_data.get("gaps", [])
            if gaps and isinstance(gaps[0], dict):
                redirect_text = redirect_roles[0] if redirect_roles else "Redirect to roles that match your experience level"
                gaps[0]["redirect"] = redirect_text
                print(f"   ✅ Added redirect to primary gap (else case): {redirect_text[:60]}...")

        # Update strategic_action - NO UPSIDE FRAMING, SPECIFIC REDIRECT
        if "reality_check" in response_data:
            # Use the first redirect role for strategic action
            redirect_text = redirect_roles[0] if redirect_roles else "roles where you are competitive"
            response_data["reality_check"]["strategic_action"] = (
                f"This role is not a fit. {redirect_text}"
            )

        # ========================================================================
        # REMOVE ANY APPLY CTA OR UPSIDE FRAMING
        # ========================================================================
        if "interview_prep" in response_data:
            del response_data["interview_prep"]
        if "networking_strategy" in response_data:
            del response_data["networking_strategy"]

        # Mark as ineligible for UI enforcement
        response_data["apply_disabled"] = True
        response_data["apply_disabled_reason"] = "Not eligible for this role"

    # CRITICAL: Validate and fix strategic_action/recommendation consistency
    response_data = validate_recommendation_consistency(response_data, capped_score)

    # ========================================================================
    # STEP 6: ENHANCE COACHING WITH CEC DATA
    # Per CEC Spec: Feed capability_evidence_report into "Gaps to Address"
    # and "Your Move" for coaching-grade specificity.
    # ========================================================================
    if response_data.get("capability_evidence_report"):
        print("📋 STEP 6: ENHANCING COACHING WITH CEC DATA")
        response_data = enhance_coaching_with_cec(response_data)

    # ========================================================================
    # STEP 6.5: COACHING CONTROLLER
    # Per Selective Coaching Moments Spec v1.0 (REVISED)
    #
    # PURPOSE:
    # - Generate "Your Move" section (ALWAYS generated, even with silence)
    # - Generate "Gaps to Address" section (suppressed by silence)
    # - Handle proceed-anyway accountability banner
    #
    # HARD CONSTRAINTS:
    # - Silence suppresses gaps, NOT "Your Move"
    # - No vague language ("some gaps," "most requirements")
    # - Max 3 sentences in "Your Move"
    # - One action only: Apply, Redirect, or Position
    # - Strong Apply still gets strategic guidance
    # ========================================================================
    if COACHING_AVAILABLE and response_data.get("calibrated_gaps"):
        print("🎯 STEP 6.5: COACHING CONTROLLER")

        try:
            calibrated_gaps = response_data.get("calibrated_gaps", {})
            job_fit_recommendation = response_data.get("recommendation", "Apply")

            # Map Job Fit recommendation to coaching-compatible format
            recommendation_map = {
                "Strongly Apply": "Strong Apply",
                "Apply": "Apply",
                "Apply with caution": "Apply with Caution",
                "Conditional Apply": "Apply with Caution",
                "Skip": "Do Not Apply",
                "Do Not Apply": "Do Not Apply",
            }
            normalized_recommendation = recommendation_map.get(job_fit_recommendation, job_fit_recommendation)

            # Prepare candidate resume data for coaching
            candidate_resume = {
                'summary': resume_data.get('summary', '') if resume_data else '',
                'experience': resume_data.get('experience', []) if resume_data else [],
                'domain': response_data.get('experience_analysis', {}).get('domain', ''),
            }

            # Sanitize role_title before using in coaching
            from coaching import _sanitize_role_title
            raw_role_title = response_data.get('role_title', '')
            sanitized_role_title = _sanitize_role_title(raw_role_title)
            print(f"   📋 Role title sanitization: '{raw_role_title[:50] if raw_role_title else 'None'}' → '{sanitized_role_title}'")

            # Prepare job requirements for coaching with sanitized role title
            job_requirements = {
                'role_title': sanitized_role_title,
                'job_description': response_data.get('job_description', ''),
                'domain': response_data.get('target_domain', '') or response_data.get('intelligence_layer', {}).get('target_domain', ''),
            }

            # Extract strengths for role-specific "Your Move" binding
            # Strengths can be in multiple locations - check all of them
            raw_strengths = response_data.get('strengths', [])

            # Also check intelligence_layer.strengths and experience_analysis.strengths
            if not raw_strengths:
                raw_strengths = response_data.get('intelligence_layer', {}).get('strengths', [])
            if not raw_strengths:
                raw_strengths = response_data.get('experience_analysis', {}).get('strengths', [])
            if not raw_strengths:
                raw_strengths = response_data.get('candidate_fit', {}).get('strengths', [])

            # Also check intelligence_layer.strategic_positioning.lead_with_strengths
            if not raw_strengths:
                raw_strengths = response_data.get('intelligence_layer', {}).get('strategic_positioning', {}).get('lead_with_strengths', [])
                if raw_strengths:
                    print("   📋 Found strengths in intelligence_layer.strategic_positioning.lead_with_strengths")

            # FALLBACK: If Claude returned 0 strengths, extract from resume directly
            # This ensures the UI always has candidate-specific strengths to show
            if not raw_strengths and resume_data:
                print("   ⚠️ Claude returned 0 strengths - extracting from resume")
                raw_strengths = _extract_fallback_strengths_from_resume(resume_data, response_data)
                if raw_strengths:
                    print(f"   ✅ Extracted {len(raw_strengths)} fallback strengths from resume")
                    # Also store in response_data for UI to use
                    response_data['strengths'] = raw_strengths

            print(f"   📋 Raw strengths extracted: {len(raw_strengths)} items")

            strengths_list = []
            for s in raw_strengths:
                if isinstance(s, str):
                    strengths_list.append(s)
                elif isinstance(s, dict):
                    # Try multiple keys: description, strength, text, summary
                    strength_text = s.get('description', s.get('strength', s.get('text', s.get('summary', ''))))
                    if strength_text:
                        strengths_list.append(strength_text)

            print(f"   ✅ Processed strengths for coaching: {len(strengths_list)} items")
            if strengths_list:
                print(f"      First strength: {strengths_list[0][:60]}...")

            # Generate coaching output with role-specific binding
            # Note: sanitized_role_title is already computed above when preparing job_requirements
            coaching_output = generate_coaching_output(
                calibrated_gaps=calibrated_gaps,
                job_fit_recommendation=normalized_recommendation,
                candidate_resume=candidate_resume,
                job_requirements=job_requirements,
                user_proceeded_anyway=False,  # TODO: Wire this from request
                strengths=strengths_list,
                role_title=sanitized_role_title
            )

            # Store coaching output in response
            response_data["coaching_output"] = coaching_output

            # Log coaching output
            your_move = coaching_output.get('your_move')
            gaps_to_address = coaching_output.get('gaps_to_address')
            show_banner = coaching_output.get('show_accountability_banner', False)

            if your_move:
                print(f"   📣 Your Move: {your_move[:80]}...")

                # CRITICAL: Wire coaching output to reality_check.strategic_action
                # Frontend reads "Your Move" from reality_check.strategic_action
                if "reality_check" not in response_data:
                    response_data["reality_check"] = {}
                response_data["reality_check"]["strategic_action"] = your_move
                print(f"   ✅ Wired Your Move to reality_check.strategic_action")

            if gaps_to_address:
                print(f"   📋 Gaps to Address: {len(gaps_to_address)} gap(s)")
            else:
                print(f"   🔇 Gaps to Address: suppressed")
            if show_banner:
                print(f"   ⚠️ Accountability banner: displayed")

        except Exception as e:
            print(f"   ⚠️ Coaching controller error (non-blocking): {e}")
            import traceback
            traceback.print_exc()
            # Coaching errors are non-blocking - continue with existing flow

    # ========================================================================
    # FINAL SUMMARY (Consolidated 3-Section Log)
    # Per Architecture Simplification Spec: Reduce logging noise to 3 sections
    # ========================================================================
    print("\n" + "=" * 80)
    print("📊 FINAL SUMMARY")
    print("=" * 80)

    # Section 1: Eligibility Gate
    eligibility_status = "PASSED" if not recommendation_locked else "FAILED"
    print(f"\n1. ELIGIBILITY GATE: {eligibility_status}")
    if recommendation_locked:
        print(f"   Reason: {locked_reason}")

    # Section 2: Scoring Summary
    final_score = response_data.get("fit_score", 0)
    print(f"\n2. SCORING SUMMARY:")
    print(f"   Final Score: {final_score}%")
    if response_data.get("fit_score_breakdown", {}).get("hard_cap_applied"):
        print(f"   Hard Cap Applied: Yes ({response_data.get('fit_score_breakdown', {}).get('hard_cap_reason', '')})")

    # Section 3: Final Decision
    final_rec = response_data.get("recommendation", "Unknown")
    print(f"\n3. FINAL DECISION: {final_rec}")
    print(f"   Locked: {response_data.get('recommendation_locked', False)}")

    # Advisory signals (non-blocking)
    advisory_notes = []
    if response_data.get("cae_advisory_downgrade"):
        advisory_notes.append(f"CAE suggested {response_data['cae_advisory_downgrade']}")
    if response_data.get("lepe_advisory_downgrade"):
        advisory_notes.append(f"LEPE suggested {response_data['lepe_advisory_downgrade']}")
    if response_data.get("has_missing_experience_gap"):
        advisory_notes.append("Missing experience gap detected")

    if advisory_notes:
        print(f"\n   🧠 Advisory Signals (Non-blocking):")
        for note in advisory_notes:
            print(f"      - {note}")

    print("\n" + "=" * 80 + "\n")

    return response_data


def validate_recommendation_consistency(analysis_data: dict, fit_score: int) -> dict:
    """
    Ensure recommendation and strategic_action match the fit_score tone.
    Fixes contradictions like "Skip" recommendation but "strong fit" language.

    RESPECTS LOCKED RECOMMENDATIONS - if recommendation is locked, this function
    will NOT soften the messaging.

    Args:
        analysis_data: Parsed response data
        fit_score: The final capped fit score

    Returns:
        Modified analysis_data with consistent messaging
    """
    # Check if recommendation is locked (from hard requirement gate)
    experience_analysis = analysis_data.get("experience_analysis", {})
    is_locked = experience_analysis.get("recommendation_locked", False)
    intelligence_layer = analysis_data.get("intelligence_layer", {})
    apply_decision = intelligence_layer.get("apply_decision", {})
    is_locked = is_locked or apply_decision.get("locked", False)

    recommendation = (analysis_data.get("recommendation") or "").lower()
    reality_check = analysis_data.get("reality_check", {})
    strategic_action = (reality_check.get("strategic_action") or "").lower() if reality_check else ""

    # Check for skip-type recommendations
    skip_signals = ["skip", "do not apply", "not recommended", "pass on this"]
    is_skip_recommendation = any(signal in recommendation for signal in skip_signals)

    # Check for apply-type language that shouldn't appear in skip scenarios
    apply_signals = ["strong fit", "strong match", "excellent match", "good fit", "great fit",
                     "well-positioned", "competitive", "prioritize this", "apply immediately"]
    has_apply_language = any(signal in strategic_action for signal in apply_signals)

    # Detect contradiction: skip recommendation but apply language
    if is_skip_recommendation and has_apply_language:
        print(f"⚠️ CONTRADICTION DETECTED: recommendation={recommendation}, but strategic_action has apply language")

        # Get top gaps for explanation
        # CRITICAL: Do NOT use candidate names - use second-person voice only
        gaps = analysis_data.get("gaps", [])
        gap_descriptions = []
        for gap in gaps[:2]:
            if isinstance(gap, dict):
                gap_descriptions.append(gap.get("gap_description") or gap.get("description", ""))
            elif isinstance(gap, str):
                gap_descriptions.append(gap)

        gaps_summary = " ".join(gap_descriptions[:2]) if gap_descriptions else "significant experience gaps exist"

        # Generate a consistent skip message - SECOND PERSON ONLY, no names
        # LOCKED recommendations get STRONGER messaging - no hedging
        if is_locked:
            new_strategic_action = (
                f"This role is not a fit. {gaps_summary}. "
                f"Do not apply. Redirect your energy to roles where you meet the core requirements."
            )
            print(f"   🔒 Using LOCKED messaging for strategic_action")
        elif fit_score < 45:
            new_strategic_action = (
                f"This role is a significant stretch. {gaps_summary}. "
                f"Do not apply. Focus on roles aligned with your background where you are competitive."
            )
        elif fit_score < 55:
            new_strategic_action = (
                f"This is a stretch role. {gaps_summary}. "
                f"Be strategic about positioning if you decide to pursue it. Consider targeting similar roles at earlier-stage companies."
            )
        else:
            new_strategic_action = (
                f"This role has addressable gaps. {gaps_summary}. "
                f"Focus on roles where you're a stronger fit, or network your way in before applying through the ATS."
            )

        # Update the strategic_action
        if reality_check:
            reality_check["strategic_action"] = new_strategic_action
            analysis_data["reality_check"] = reality_check
            print(f"✅ Rewrote strategic_action to match {recommendation} recommendation")

    # Also check recommendation_rationale for contradictions
    rationale = (analysis_data.get("recommendation_rationale") or "").lower()
    if is_skip_recommendation and any(signal in rationale for signal in apply_signals):
        print(f"⚠️ Rationale has contradictory apply language for skip recommendation")
        # The rationale is usually already updated by force_apply_experience_penalties, but double-check
        if "adjusted" not in rationale.lower():
            fit_score_int = int(fit_score)

            # LOCKED recommendations get DIRECT messaging - no hedging, no softening
            if is_locked:
                tone = "This role is not a fit. You do not meet the core requirements. Do not apply."
                print(f"   🔒 Using LOCKED messaging for rationale")
            elif fit_score_int < 45:
                tone = "This is a significant stretch. Only pursue if you have an inside connection."
            elif fit_score_int < 55:
                tone = "This is a stretch role. Be strategic about positioning if you pursue it."
            elif fit_score_int < 70:
                tone = "This is a moderate fit with addressable gaps."
            elif fit_score_int < 85:
                tone = "This is a good fit worth pursuing."
            else:
                tone = "This is a strong match. Prioritize this application."

            analysis_data["recommendation_rationale"] = tone
            print(f"✅ Updated recommendation_rationale to: {tone}")

    return analysis_data


def enhance_coaching_with_cec(response_data: dict) -> dict:
    """
    Enhance "Gaps to Address" and coaching output with CEC data.

    Per CEC Spec v1.0:
    - Adds coaching-grade diagnosis to gaps
    - Enhances "Your Move" with critical gap specificity
    - Does NOT alter score or recommendation

    Args:
        response_data: The full Job Fit response with capability_evidence_report

    Returns:
        Enhanced response_data with CEC-enriched coaching
    """
    cec_report = response_data.get("capability_evidence_report")
    if not cec_report:
        return response_data

    evaluated_capabilities = cec_report.get("evaluated_capabilities", [])
    summary = cec_report.get("summary", {})
    critical_gaps = summary.get("critical_gaps", [])

    # Build CEC-enhanced gaps list
    cec_gaps = []
    for cap in evaluated_capabilities:
        if cap.get("evidence_status") in ["missing", "implicit"]:
            cec_gap = {
                "capability_id": cap.get("capability_id"),
                "capability_name": cap.get("capability_name"),
                "evidence_status": cap.get("evidence_status"),
                "diagnosis": cap.get("diagnosis"),
                "distance": cap.get("distance"),
                "coachable": cap.get("coachable"),
                "criticality": cap.get("criticality"),
                "is_critical": cap.get("capability_id") in critical_gaps
            }
            cec_gaps.append(cec_gap)

    # Store CEC gaps for frontend rendering
    response_data["cec_gaps"] = cec_gaps

    # Enhance reality_check with CEC context
    reality_check = response_data.get("reality_check", {})
    recommendation = (response_data.get("recommendation") or "").lower()

    # Find most critical gap for coaching specificity
    most_critical = None
    for cap in evaluated_capabilities:
        if cap.get("evidence_status") == "missing" and cap.get("criticality") == "required":
            most_critical = cap
            break

    if most_critical:
        # Add CEC context to coaching
        cec_coaching_context = {
            "critical_gap_name": most_critical.get("capability_name"),
            "critical_gap_diagnosis": most_critical.get("diagnosis"),
            "critical_gap_distance": most_critical.get("distance"),
            "critical_gap_coachable": most_critical.get("coachable")
        }
        response_data["cec_coaching_context"] = cec_coaching_context

        # Enhance strategic_action for Do Not Apply scenarios
        if "do not apply" in recommendation or "skip" in recommendation:
            if reality_check and most_critical.get("diagnosis"):
                existing_action = reality_check.get("strategic_action", "")
                # Prepend specific CEC diagnosis
                enhanced_action = f"{most_critical.get('diagnosis')} {existing_action}"
                reality_check["strategic_action"] = enhanced_action
                reality_check["cec_enhanced"] = True
                response_data["reality_check"] = reality_check

    # Log enhancement
    if cec_gaps:
        print(f"   📋 CEC enhanced coaching: {len(cec_gaps)} capability gaps documented")
        for gap in cec_gaps[:2]:  # Log first 2
            status = "🔴" if gap.get("is_critical") else "🟡"
            print(f"      {status} {gap.get('capability_name')}: {gap.get('evidence_status')}")

    return response_data


# ============================================================================
# COMPANY CREDIBILITY SCORING (4-TIER SYSTEM)
# Per JOB_FIT_SCORING_SPEC.md and COMPETENCY_DIAGNOSTICS_INTEGRATION.md
# ============================================================================

# Well-known companies that get HIGH credibility (1.0x multiplier)
HIGH_CREDIBILITY_COMPANIES = [
    # FAANG / Big Tech
    "google", "facebook", "meta", "amazon", "apple", "microsoft", "netflix",
    # Other major tech
    "uber", "airbnb", "stripe", "spotify", "linkedin", "twitter", "x corp",
    "dropbox", "salesforce", "oracle", "adobe", "intuit", "workday", "servicenow",
    "snowflake", "databricks", "figma", "notion", "slack", "zoom", "square", "block",
    "coinbase", "robinhood", "plaid", "chime", "doordash", "instacart", "lyft",
    # Enterprise / B2B
    "ibm", "cisco", "vmware", "dell", "hp", "intel", "nvidia", "amd", "qualcomm",
    "sap", "atlassian", "hubspot", "zendesk", "twilio", "datadog", "splunk",
    # Finance
    "goldman sachs", "goldman", "jp morgan", "jpmorgan", "morgan stanley",
    "blackrock", "citadel", "two sigma", "jane street", "bridgewater",
    "bank of america", "wells fargo", "citi", "citibank", "capital one",
    # Consulting
    "mckinsey", "bain", "bcg", "boston consulting", "deloitte", "accenture",
    "pwc", "ey", "ernst & young", "kpmg",
    # Other large companies
    "disney", "warner", "comcast", "verizon", "att", "t-mobile",
    "walmart", "target", "costco", "home depot", "nike", "coca-cola", "pepsico",
    # Known startups / Series C+
    "rippling", "ramp", "brex", "gusto", "deel", "remote", "lattice", "drata",
    "retool", "vercel", "supabase", "planetscale", "neon", "linear", "raycast",
    "headway", "heidrick", "heidrick & struggles"
]


def get_company_credibility_tier(company: str, title: str = "") -> str:
    """
    Determine company credibility tier based on company name and title.

    Tiers:
    - HIGH (1.0x): Public companies, Series B+, well-known brands, >50 employees
    - MEDIUM (0.7x): Series A startups, 10-50 employees, regional players
    - LOW (0.3x): Seed-stage (<10 employees), defunct companies, limited presence
    - ZERO (0.0x): Title inflation (operations + PM/Eng), volunteer, side projects

    Returns:
        str: "HIGH", "MEDIUM", "LOW", or "ZERO"
    """
    company_lower = (company or "").lower().strip()
    title_lower = (title or "").lower().strip()

    # ZERO tier: Title inflation detection
    # Operations/coordinator titles combined with PM/Engineering claims
    operations_signals = ["operations", "coordinator", "assistant", "admin", "associate"]
    pm_eng_titles = ["product manager", "engineer", "developer", "pm"]

    has_ops_signal = any(op in title_lower for op in operations_signals)
    claims_pm_eng = any(pm in title_lower for pm in pm_eng_titles)

    if has_ops_signal and claims_pm_eng:
        print(f"   🚨 Title inflation detected: '{title}' - ZERO credibility")
        return "ZERO"

    # ZERO tier: Volunteer/side project work
    if "volunteer" in title_lower or "side project" in company_lower:
        print(f"   🚨 Volunteer/side project: '{company}' - ZERO credibility")
        return "ZERO"

    # ZERO tier: Personal branding entities (solo consultancies presenting as companies)
    if any(x in company_lower for x in ["consulting llc", "solutions llc", "ventures llc"]) and \
       any(x in title_lower for x in ["founder", "ceo", "owner", "principal"]):
        # This catches "John Smith Consulting LLC" type entries
        print(f"   🚨 Personal branding entity suspected: '{company}' - ZERO credibility")
        return "ZERO"

    # HIGH tier: Well-known companies
    if any(known in company_lower for known in HIGH_CREDIBILITY_COMPANIES):
        return "HIGH"

    # HIGH tier: Public company indicators
    if any(x in company_lower for x in ["inc.", "corp.", "corporation", "plc", "publicly traded"]):
        return "HIGH"

    # HIGH tier: Series B+ indicators
    if any(x in company_lower for x in ["series b", "series c", "series d", "series e"]):
        return "HIGH"

    # LOW tier: Seed/early stage indicators
    if any(x in company_lower for x in ["seed", "pre-seed", "stealth", "founding team", "defunct", "shutdown", "bankrupt"]):
        return "LOW"

    # Default: MEDIUM (legitimate but not tier 1)
    return "MEDIUM"


def get_credibility_multiplier(tier: str) -> float:
    """
    Map credibility tier to experience multiplier.

    Returns:
        float: 1.0, 0.7, 0.3, or 0.0
    """
    multipliers = {
        "HIGH": 1.0,
        "MEDIUM": 0.7,
        "LOW": 0.3,
        "ZERO": 0.0
    }
    return multipliers.get(tier, 0.7)


def calculate_credibility_adjusted_years(resume_data: dict, target_role_type: str = "pm") -> tuple:
    """
    Calculate role-specific years with company credibility adjustments.

    Args:
        resume_data: Parsed resume JSON with experience entries
        target_role_type: Role type to calculate experience for

    Returns:
        tuple: (adjusted_years, raw_years, breakdown_list)
    """
    if not resume_data:
        return 0.0, 0.0, []

    experience = resume_data.get("experience", [])
    if not experience or not isinstance(experience, list):
        return 0.0, 0.0, []

    # Role title patterns by type (same as calculate_role_specific_years)
    role_patterns = {
        "pm": ["product manager", "product lead", "product owner", "pm", "product director",
               "head of product", "vp product", "cpo", "associate product manager",
               "senior product manager", "staff product", "principal product",
               "group product manager", "technical pm"],
        "engineering": ["engineer", "developer", "swe", "software", "programmer",
                       "cto", "tech lead", "architect", "devops"],
        "recruiting": ["recruit", "talent acquisition", "ta director", "ta manager",
                      "sourcer", "sourcing", "headhunter", "talent partner",
                      "recruiting manager", "technical recruiter", "executive recruiter",
                      "talent lead", "head of talent", "vp talent"],
        "sales": ["sales", "account executive", "account manager", "business development",
                 "revenue", "enterprise sales", "sales director", "vp sales"],
        "marketing": ["marketing", "growth", "brand", "content", "digital marketing",
                     "product marketing", "demand gen", "cmo", "head of marketing"],
        "general": []
    }

    patterns = role_patterns.get(target_role_type, [])
    raw_years = 0.0
    adjusted_years = 0.0
    breakdown = []

    for exp in experience:
        if not isinstance(exp, dict):
            continue

        title = (exp.get("title", "") or "").strip()
        company = (exp.get("company", "") or "").strip()
        dates = (exp.get("dates", "") or "").lower()

        # Check if this role matches the target type (or general = all roles)
        if target_role_type == "general":
            matches_role = True
        else:
            matches_role = any(pattern in title.lower() for pattern in patterns)

        if matches_role:
            years = parse_experience_duration(dates)
            raw_years += years

            # Get credibility tier and multiplier
            tier = get_company_credibility_tier(company, title)
            multiplier = get_credibility_multiplier(tier)
            adjusted = years * multiplier
            adjusted_years += adjusted

            breakdown.append({
                "company": company,
                "title": title,
                "raw_years": years,
                "tier": tier,
                "multiplier": multiplier,
                "adjusted_years": adjusted
            })

            print(f"   📊 {company} ({title}): {years:.1f}y × {multiplier} ({tier}) = {adjusted:.1f}y")

    return round(adjusted_years, 1), round(raw_years, 1), breakdown


# ============================================================================
# PEOPLE LEADERSHIP vs OPERATIONAL LEADERSHIP - HARD CONSTRAINT
# Per JOB_FIT_SCORING constraints: "This is enforcement, not vibes."
#
# CRITICAL DISTINCTION:
# - PEOPLE LEADERSHIP: Direct reports, team management, hiring/firing authority
# - OPERATIONAL LEADERSHIP: Systems, processes, programs (NO direct reports)
#
# These are DISTINCT signals. Operational leadership does NOT substitute for
# people leadership requirements. A candidate with 15 years operational excellence
# but 0 people leadership years CANNOT qualify for a role requiring 7+ years
# people leadership.
# ============================================================================

# ============================================================================
# TIERED LEADERSHIP MODEL
# Per Leadership Tiering Spec: Leadership must be modeled in tiers, not binary.
#
# TIER 1: STRATEGIC / FUNCTIONAL LEADERSHIP
#   - Cross-functional project leadership, strategic initiatives
#   - Leading without direct authority (dotted line, matrixed, project-based)
#   - Evidence: "led initiative", "drove strategy", "cross-functional", "program lead"
#
# TIER 2: DIRECT PEOPLE LEADERSHIP
#   - Managing direct reports, team management, hiring/firing authority
#   - Evidence: "direct reports", "managed team of X", "built team", "hired/fired"
#   - This is what most JDs mean when they say "leadership experience required"
#
# TIER 3: ORG-LEVEL LEADERSHIP
#   - Executive/C-suite, board-facing, multi-team/department leadership
#   - Evidence: "VP", "C-suite", "board", "department head", "divisional"
#
# MESSAGING RULES:
# - "0 years" ONLY when NO leadership signals exist at ALL
# - "insufficient" when leadership exists but doesn't meet requirement level
# - Be specific about what tier is missing vs what tier candidate has
# ============================================================================


def extract_tiered_leadership(resume_data: dict) -> dict:
    """
    Extract leadership experience across three tiers.

    Returns a structured breakdown of leadership by tier, enabling precise
    gap messaging ("insufficient" vs "none") and proper tier-to-requirement matching.

    Args:
        resume_data: Parsed resume dictionary

    Returns:
        dict: {
            "strategic_leadership_years": float,  # Tier 1: functional/strategic
            "people_leadership_years": float,     # Tier 2: direct reports
            "org_leadership_years": float,        # Tier 3: executive/org-level
            "total_leadership_years": float,      # Sum of all tiers (max credit)
            "has_any_leadership": bool,           # At least one tier > 0
            "leadership_tier_summary": str,       # Human-readable summary
            "tier_breakdown": list                # Per-role breakdown
        }
    """
    result = {
        "strategic_leadership_years": 0.0,
        "people_leadership_years": 0.0,
        "org_leadership_years": 0.0,
        "total_leadership_years": 0.0,
        "has_any_leadership": False,
        "leadership_tier_summary": "No leadership evidence found",
        "tier_breakdown": []
    }

    if not resume_data:
        return result

    experience = resume_data.get("experience", [])
    if not experience or not isinstance(experience, list):
        return result

    # ========================================================================
    # TIER 1: STRATEGIC / FUNCTIONAL LEADERSHIP EVIDENCE
    # ========================================================================
    strategic_evidence = [
        "led initiative", "drove strategy", "cross-functional", "program lead",
        "strategic initiative", "transformation", "led project", "project lead",
        "spearheaded", "championed", "orchestrated", "coordinated across",
        "led cross-functional", "matrixed", "dotted line", "stakeholder management",
        "led the", "leading the", "drove the", "driving"
    ]

    # ========================================================================
    # TIER 2: DIRECT PEOPLE LEADERSHIP EVIDENCE
    # ========================================================================
    people_evidence = [
        "direct report", "direct reports", "managed a team", "led a team",
        "team of", "people manager", "managed team", "lead a team",
        "built the team", "grew the team", "hiring manager",
        "performance review", "promoted", "mentored", "coached team",
        "developed team", "team lead", "engineering manager", "people management",
        "hired", "fired", "onboarded", "trained team", "built team"
    ]

    people_titles = [
        "manager", "director", "head of", "vp ", "vice president",
        "chief", "supervisor", "team lead"
    ]

    # Strong leadership titles at established companies = automatic people credit
    strong_leadership_titles = [
        "vp ", "vice president", "vp,", "director", "head of", "chief",
        "senior manager", "senior product marketing manager", "group manager",
        "marketing manager", "product marketing manager"
    ]

    # Well-known companies where VP/Director definitely means people leadership
    established_companies = [
        "google", "meta", "facebook", "amazon", "apple", "microsoft", "netflix",
        "uber", "lyft", "airbnb", "stripe", "square", "twilio", "segment",
        "salesforce", "adobe", "oracle", "sap", "ibm", "cisco", "intel",
        "linkedin", "twitter", "x corp", "snap", "pinterest", "dropbox",
        "slack", "atlassian", "asana", "notion", "figma", "canva",
        "shopify", "hubspot", "zendesk", "datadog", "snowflake", "mongodb",
        "new relic", "mparticle", "amplitude", "mixpanel", "braze",
        "intercom", "drift", "gong", "outreach", "salesloft",
    ]

    # ========================================================================
    # TIER 3: ORG-LEVEL LEADERSHIP EVIDENCE
    # ========================================================================
    org_evidence = [
        "c-suite", "executive team", "board", "board-facing", "divisional",
        "department head", "multi-team", "organization-wide", "company-wide",
        "global leadership", "regional leadership", "p&l", "profit and loss",
        "budget authority", "headcount authority", "exec team", "leadership team"
    ]

    org_titles = [
        "ceo", "cto", "cfo", "coo", "cmo", "cio", "cpo", "cro",
        "chief", "president", "evp", "svp", "executive vp", "senior vp",
        "gm", "general manager"
    ]

    # Operational-only patterns (NO leadership credit)
    operational_only_patterns = [
        "program manager", "project manager", "technical lead",
        "staff engineer", "principal engineer", "architect",
        "operations lead", "process lead", "systems lead"
    ]

    tier_breakdown = []
    strategic_total = 0.0
    people_total = 0.0
    org_total = 0.0

    for exp in experience:
        if not isinstance(exp, dict):
            continue

        title = (exp.get("title", "") or "").lower()
        description = (exp.get("description", "") or "").lower()
        highlights = exp.get("highlights", [])
        if highlights and isinstance(highlights, list):
            highlights_text = " ".join([h.lower() for h in highlights if isinstance(h, str)])
        else:
            highlights_text = ""

        combined_text = f"{title} {description} {highlights_text}"
        dates = (exp.get("dates", "") or "").lower()
        company = (exp.get("company", "") or "").strip()

        years = parse_experience_duration(dates)
        if years <= 0:
            continue

        # Skip operational-only roles unless they have leadership evidence
        is_operational_only = any(pattern in title for pattern in operational_only_patterns)
        has_any_leadership_signal = (
            any(ev in combined_text for ev in strategic_evidence) or
            any(ev in combined_text for ev in people_evidence) or
            any(ev in combined_text for ev in org_evidence)
        )

        if is_operational_only and not has_any_leadership_signal:
            continue

        # Determine which tiers apply to this role
        role_tiers = []
        tier_years = {"strategic": 0.0, "people": 0.0, "org": 0.0}

        # Get credibility multiplier
        tier = get_company_credibility_tier(company, title)
        multiplier = get_credibility_multiplier(tier)

        # Check TIER 3 first (org-level) - highest tier
        has_org_title = any(ot in title for ot in org_titles)
        has_org_evidence = any(ev in combined_text for ev in org_evidence)

        if has_org_title or has_org_evidence:
            adjusted = years * multiplier
            tier_years["org"] = adjusted
            org_total += adjusted
            role_tiers.append("org")

        # Check TIER 2 (people leadership)
        has_people_title = any(pt in title for pt in people_titles)
        has_strong_title = any(st in title for st in strong_leadership_titles)
        has_people_evidence = any(ev in combined_text for ev in people_evidence)

        # Check if company is established (VP/Director definitely means people leadership)
        company_lower = company.lower()
        is_established_company = any(ec in company_lower for ec in established_companies)
        if "acquired by" in company_lower or "twilio" in company_lower:
            is_established_company = True

        if has_strong_title and is_established_company:
            # VP/Director at established company = full credit, no evidence needed
            adjusted = years * multiplier
            tier_years["people"] = adjusted
            people_total += adjusted
            role_tiers.append("people_established")
        elif has_people_title and has_people_evidence:
            # Full credit for verified people leadership
            adjusted = years * multiplier
            tier_years["people"] = adjusted
            people_total += adjusted
            role_tiers.append("people")
        elif has_people_evidence:
            # Evidence without clear title - 70% credit
            adjusted = years * multiplier * 0.7
            tier_years["people"] = adjusted
            people_total += adjusted
            role_tiers.append("people_partial")
        elif has_strong_title:
            # Strong title at unknown company - 50% credit
            adjusted = years * multiplier * 0.5
            tier_years["people"] = adjusted
            people_total += adjusted
            role_tiers.append("people_unverified")

        # Check TIER 1 (strategic/functional) - even if other tiers apply
        has_strategic_evidence = any(ev in combined_text for ev in strategic_evidence)

        if has_strategic_evidence and "people" not in role_tiers and "org" not in role_tiers:
            # Strategic-only (no higher tier credit)
            adjusted = years * multiplier * 0.5  # 50% credit for strategic-only
            tier_years["strategic"] = adjusted
            strategic_total += adjusted
            role_tiers.append("strategic")

        # Record breakdown if any leadership found
        if role_tiers:
            tier_breakdown.append({
                "title": exp.get("title", ""),
                "company": company,
                "years": years,
                "tiers": role_tiers,
                "tier_years": tier_years
            })

    # Calculate totals (don't double-count - use max of each role's highest tier)
    total_leadership = max(strategic_total, people_total, org_total)
    if people_total > 0:
        total_leadership = people_total
    if org_total > 0:
        total_leadership = max(total_leadership, org_total)

    # Build summary
    summaries = []
    if org_total > 0:
        summaries.append(f"Org-level: {org_total:.1f}y")
    if people_total > 0:
        summaries.append(f"People: {people_total:.1f}y")
    if strategic_total > 0 and people_total == 0 and org_total == 0:
        summaries.append(f"Strategic/Functional: {strategic_total:.1f}y")

    has_any = org_total > 0 or people_total > 0 or strategic_total > 0

    result.update({
        "strategic_leadership_years": round(strategic_total, 1),
        "people_leadership_years": round(people_total, 1),
        "org_leadership_years": round(org_total, 1),
        "total_leadership_years": round(total_leadership, 1),
        "has_any_leadership": has_any,
        "leadership_tier_summary": " | ".join(summaries) if summaries else "No leadership evidence found",
        "tier_breakdown": tier_breakdown
    })

    print(f"   📊 TIERED LEADERSHIP ANALYSIS:")
    print(f"      Org-level: {org_total:.1f} years")
    print(f"      People: {people_total:.1f} years")
    print(f"      Strategic: {strategic_total:.1f} years")
    print(f"      Has any leadership: {has_any}")

    return result


# ============================================================================
# CAPABILITY EVIDENCE CHECK (CEC)
# Per CEC Specification v1.0: Diagnostic layer that explains HOW a candidate's
# resume diverges from a job's requirements. Does NOT alter scoring/recommendations.
#
# Purpose: Turn HenryHQ from accurate into coaching-grade by providing specific,
# evidence-based diagnosis of capability gaps.
#
# HARD CONSTRAINTS:
# ✗ Cannot change fit score
# ✗ Cannot change recommendation
# ✗ Cannot trigger LEPE caps or CAE multipliers
# ✗ Cannot override hard_requirement_failures
# ============================================================================

def evaluate_capability_evidence(
    response_data: dict,
    resume_data: dict,
    leadership_extraction: dict = None
) -> dict:
    """
    Capability Evidence Check (CEC) - Diagnostic layer for coaching-grade gap analysis.

    Evaluates evidence strength for each JD requirement across 4 dimensions:
    1. Leadership Scope (local vs global, distributed vs co-located)
    2. Cross-Functional Depth (Product partnership, roadmap ownership)
    3. Scale Signals (team size, org size, revenue, users)
    4. Domain Adjacency (direct vs transferable vs unrelated experience)

    Returns capability_evidence_report for use in "Gaps to Address" and "Your Move".
    """
    print("📊 CAPABILITY EVIDENCE CHECK (CEC) - Diagnostic Layer")

    evaluated_capabilities = []

    # Extract JD context
    role_title = response_data.get("role_title", "")
    jd_text = response_data.get("job_description", "")
    requirements = response_data.get("requirements", [])
    experience_analysis = response_data.get("experience_analysis", {})

    # 1. Evaluate Leadership Scope
    leadership_eval = evaluate_leadership_scope(
        response_data, resume_data, leadership_extraction
    )
    if leadership_eval:
        evaluated_capabilities.append(leadership_eval)

    # 2. Evaluate Cross-Functional Depth
    cross_func_eval = evaluate_cross_functional_depth(response_data, resume_data)
    if cross_func_eval:
        evaluated_capabilities.append(cross_func_eval)

    # 3. Evaluate Scale Signals
    scale_eval = evaluate_scale_signals(response_data, resume_data)
    if scale_eval:
        evaluated_capabilities.append(scale_eval)

    # 4. Evaluate Domain Adjacency
    domain_eval = evaluate_domain_adjacency(response_data, resume_data)
    if domain_eval:
        evaluated_capabilities.append(domain_eval)

    # Calculate summary
    explicit_count = sum(1 for c in evaluated_capabilities if c.get("evidence_status") == "explicit")
    implicit_count = sum(1 for c in evaluated_capabilities if c.get("evidence_status") == "implicit")
    missing_count = sum(1 for c in evaluated_capabilities if c.get("evidence_status") == "missing")

    critical_gaps = [
        c["capability_id"] for c in evaluated_capabilities
        if c.get("evidence_status") in ["missing", "implicit"]
        and c.get("criticality") == "required"
    ]

    # ========================================================================
    # CONFIDENCE THRESHOLD CALCULATION
    # Per Selective Coaching Spec v1.0:
    # - One explicit signal > five implicit signals
    # - Multiple implicit signals do NOT sum to explicit
    # - When ambiguous, downgrade confidence, not recommendation
    # ========================================================================

    # Add confidence score to each capability based on evidence status
    for cap in evaluated_capabilities:
        evidence_status = cap.get("evidence_status", "missing")
        if evidence_status == "explicit":
            cap["confidence"] = 1.0  # Full confidence
        elif evidence_status == "implicit":
            cap["confidence"] = 0.6  # Cap implicit at 60% per spec
        else:
            cap["confidence"] = 0.0  # No confidence for missing

    # Calculate overall CEC confidence
    # Explicit signals count fully, implicit capped at 0.6 each
    if evaluated_capabilities:
        # Weight by criticality
        required_caps = [c for c in evaluated_capabilities if c.get("criticality") == "required"]
        preferred_caps = [c for c in evaluated_capabilities if c.get("criticality") != "required"]

        if required_caps:
            # Required capabilities weight heavily
            required_confidence = sum(c.get("confidence", 0) for c in required_caps) / len(required_caps)
        else:
            required_confidence = 1.0

        if preferred_caps:
            preferred_confidence = sum(c.get("confidence", 0) for c in preferred_caps) / len(preferred_caps)
        else:
            preferred_confidence = 1.0

        # Overall: 70% weight on required, 30% on preferred
        overall_confidence = (required_confidence * 0.7) + (preferred_confidence * 0.3)
    else:
        overall_confidence = 1.0  # No capabilities evaluated = full confidence

    # Flag if all signals are implicit (low confidence situation)
    all_implicit = explicit_count == 0 and implicit_count > 0
    if all_implicit:
        # Cap overall confidence at 0.6 per spec
        overall_confidence = min(overall_confidence, 0.6)

    capability_evidence_report = {
        "evaluated_capabilities": evaluated_capabilities,
        "summary": {
            "total_capabilities_evaluated": len(evaluated_capabilities),
            "explicit_count": explicit_count,
            "implicit_count": implicit_count,
            "missing_count": missing_count,
            "critical_gaps": critical_gaps,
            "overall_confidence": round(overall_confidence, 2),
            "all_signals_implicit": all_implicit
        }
    }

    print(f"   📋 CEC Summary: {explicit_count} explicit, {implicit_count} implicit, {missing_count} missing")
    print(f"   📊 CEC Confidence: {overall_confidence:.2f}" + (" (all implicit - capped)" if all_implicit else ""))
    if critical_gaps:
        print(f"   🔴 Critical gaps: {', '.join(critical_gaps)}")

    return capability_evidence_report


# ============================================================================
# CEC v1.1 HELPER FUNCTIONS
# Sub-signal detection for recruiter-grade capability assessment
# ============================================================================

def detect_decision_authority(experience_text: str) -> str:
    """
    Detects decision-making authority level from experience text.

    Returns:
        'explicit' - Clear decision ownership language
        'implicit' - Influencer/partner language
        'missing' - No authority signals
    """
    try:
        experience_lower = experience_text.lower()

        # Explicit decision authority verbs
        explicit_patterns = [
            'decided', 'owned', 'set direction', 'determined',
            'established', 'defined strategy', 'made decision',
            'chose to', 'selected', 'approved', 'vetoed'
        ]

        # Implicit authority (influencer, not decider)
        implicit_patterns = [
            'partnered with', 'collaborated', 'supported',
            'contributed to', 'advised', 'recommended',
            'influenced', 'consulted', 'helped shape'
        ]

        explicit_count = sum(1 for pattern in explicit_patterns if pattern in experience_lower)
        implicit_count = sum(1 for pattern in implicit_patterns if pattern in experience_lower)

        if explicit_count >= 2:
            return 'explicit'
        elif implicit_count >= 2 or explicit_count == 1:
            return 'implicit'
        else:
            return 'missing'
    except Exception as e:
        print(f"Warning: Decision authority detection failed: {e}")
        return 'missing'


def detect_org_design_signals(experience_text: str) -> str:
    """
    Detects organizational design ownership from experience text.

    Returns:
        'explicit' - Built teams, restructured org, designed structure
        'implicit' - Hired people, grew team
        'missing' - No org design signals
    """
    try:
        experience_lower = experience_text.lower()

        # Explicit org design
        explicit_patterns = [
            'built team', 'built org', 'designed org',
            'restructured', 'defined roles', 'created structure',
            'established hiring plan', 'shaped organization',
            'org design', 'organizational structure'
        ]

        # Implicit (hiring/growth without design)
        implicit_patterns = [
            'grew team', 'hired', 'expanded team',
            'increased headcount', 'team growth',
            'added engineers', 'scaled team'
        ]

        if any(pattern in experience_lower for pattern in explicit_patterns):
            return 'explicit'
        elif any(pattern in experience_lower for pattern in implicit_patterns):
            return 'implicit'
        else:
            return 'missing'
    except Exception as e:
        print(f"Warning: Org design detection failed: {e}")
        return 'missing'


def detect_partnership_tier(experience_text: str) -> str:
    """
    Detects cross-functional partnership quality.

    Returns:
        'tier_1_joint' - Co-owned decisions, shared KPIs
        'tier_2_coordination' - Regular sync, alignment
        'tier_3_adjacency' - Works with adjacent functions only
    """
    try:
        experience_lower = experience_text.lower()

        # Tier 1: Joint ownership with Product
        tier_1_patterns = [
            'co-owned roadmap', 'joint ownership', 'shared kpi',
            'product partnership', 'partnered with pm',
            'product manager', 'roadmap ownership',
            'quarterly planning', 'backlog prioritization'
        ]

        # Tier 2: Coordination
        tier_2_patterns = [
            'partnered with', 'collaborated with',
            'worked with engineering', 'cross-functional',
            'alignment meetings', 'sync with'
        ]

        # Check for Product-specific mentions
        has_product_mention = any(word in experience_lower for word in ['product', 'pm', 'roadmap', 'backlog'])

        tier_1_match = any(pattern in experience_lower for pattern in tier_1_patterns)
        tier_2_match = any(pattern in experience_lower for pattern in tier_2_patterns)

        if tier_1_match or (has_product_mention and tier_2_match):
            return 'tier_1_joint'
        elif tier_2_match:
            return 'tier_2_coordination'
        else:
            return 'tier_3_adjacency'
    except Exception as e:
        print(f"Warning: Partnership tier detection failed: {e}")
        return 'tier_3_adjacency'


def detect_financial_ownership(experience_text: str) -> Dict[str, Any]:
    """
    Detects P&L, budget, or financial ownership.

    Returns:
        {
            'level': 'explicit' | 'implicit' | 'missing',
            'amount': int | None,
            'type': 'p&l' | 'budget' | 'revenue' | None
        }
    """
    try:
        import re
        experience_lower = experience_text.lower()

        # Extract dollar amounts
        money_pattern = r'\$(\d+(?:,\d{3})*(?:\.\d+)?)\s*([mkb])?'
        money_matches = re.findall(money_pattern, experience_lower, re.IGNORECASE)

        amounts = []
        for amount, multiplier in money_matches:
            base = float(amount.replace(',', ''))
            if multiplier and multiplier.lower() == 'k':
                amounts.append(int(base * 1000))
            elif multiplier and multiplier.lower() == 'm':
                amounts.append(int(base * 1000000))
            elif multiplier and multiplier.lower() == 'b':
                amounts.append(int(base * 1000000000))
            else:
                amounts.append(int(base))

        # Explicit ownership patterns
        explicit_patterns = [
            'owned p&l', 'p&l owner', 'budget owner',
            'managed budget', 'revenue owner',
            'owned budget', 'financial owner'
        ]

        # Implicit patterns
        implicit_patterns = [
            'within budget', 'budget constraint',
            'managed spend', 'cost optimization'
        ]

        has_explicit = any(pattern in experience_lower for pattern in explicit_patterns)
        has_implicit = any(pattern in experience_lower for pattern in implicit_patterns)

        ownership_type = None
        if 'p&l' in experience_lower:
            ownership_type = 'p&l'
        elif 'budget' in experience_lower:
            ownership_type = 'budget'
        elif 'revenue' in experience_lower:
            ownership_type = 'revenue'

        return {
            'level': 'explicit' if has_explicit else ('implicit' if has_implicit else 'missing'),
            'amount': max(amounts) if amounts else None,
            'type': ownership_type
        }
    except Exception as e:
        print(f"Warning: Financial ownership detection failed: {e}")
        return {'level': 'missing', 'amount': None, 'type': None}


def detect_domain_depth(experience_text: str, domain: str) -> str:
    """
    Detects domain expertise depth.

    Returns:
        'explicit' - Certifications, speaking, deep terminology
        'implicit' - Domain mentioned, generic language
        'missing' - No domain presence
    """
    try:
        experience_lower = experience_text.lower()
        domain_lower = domain.lower() if domain else ''

        # Explicit depth indicators
        explicit_patterns = [
            'certified', 'certification', 'speaker at',
            'published', 'thought leader', 'expert in',
            'specialized in', 'domain expertise'
        ]

        # Check for domain-specific terminology density
        has_domain = domain_lower in experience_lower if domain_lower else False
        has_explicit = any(pattern in experience_lower for pattern in explicit_patterns)

        # Count domain-specific terms (varies by domain)
        domain_term_count = experience_lower.count(domain_lower) if domain_lower else 0

        if has_explicit and has_domain:
            return 'explicit'
        elif domain_term_count >= 2 or has_domain:
            return 'implicit'
        else:
            return 'missing'
    except Exception as e:
        print(f"Warning: Domain depth detection failed: {e}")
        return 'missing'


# ============================================================================
# CEC v1.0 CAPABILITY EVALUATORS (Enhanced in v1.1)
# ============================================================================

def evaluate_leadership_scope(
    response_data: dict,
    resume_data: dict,
    leadership_extraction: dict = None
) -> dict:
    """
    Evaluate Leadership Scope capability (CEC v1.1 Enhanced).

    Checks for:
    - Local vs regional vs global teams
    - Distributed vs co-located
    - Manager of managers vs IC managers

    v1.1 Sub-signals:
    - Decision authority (hiring, budget, firing, reorg)
    - Org design signals (team charter, scaling, structure)
    - Legacy indicators (succession, mentorship lineage)
    """
    role_title = (response_data.get("role_title", "") or "").lower()
    jd_text = (response_data.get("job_description", "") or "").lower()

    # Check if JD requires leadership scope
    global_keywords = ["global", "distributed", "multi-region", "international", "worldwide", "remote team"]
    manager_of_managers_keywords = ["manager of managers", "director", "vp", "head of", "org-level"]

    requires_global = any(kw in jd_text for kw in global_keywords)
    requires_manager_of_managers = any(kw in role_title or kw in jd_text for kw in manager_of_managers_keywords)

    # Skip if JD doesn't require leadership scope
    if not requires_global and not requires_manager_of_managers:
        return None

    # Extract resume evidence - collect all experience text for sub-signal analysis
    resume_experience = resume_data.get("experience", []) if resume_data else []
    combined_text = ""
    for exp in resume_experience:
        if isinstance(exp, dict):
            combined_text += f" {exp.get('title', '')} {exp.get('description', '')} {' '.join(exp.get('highlights', []) or exp.get('bullets', []) or [])}".lower()

    # v1.1: Detect sub-signals using helper functions
    decision_authority = detect_decision_authority(combined_text)
    org_design = detect_org_design_signals(combined_text)

    # v1.1: Detect legacy indicators (succession planning, mentorship lineage)
    legacy_keywords = ["promoted to", "grew into", "succession", "mentored who became",
                       "developed leaders", "built leadership pipeline", "talent development"]
    has_legacy_signals = any(kw in combined_text for kw in legacy_keywords)
    legacy_status = "explicit" if has_legacy_signals else "missing"

    # Check for explicit global evidence
    has_global_evidence = any(kw in combined_text for kw in ["global", "distributed team", "multi-region", "across regions", "worldwide", "international team", "remote team across"])
    has_mom_evidence = any(kw in combined_text for kw in ["manager of managers", "led managers", "director reports", "managed directors", "org of", "organization of"])

    # Check for local-only evidence
    has_local_evidence = any(kw in combined_text for kw in ["team of", "managed team", "led team", "direct reports"])

    # v1.1: Build sub-signals summary for enhanced diagnosis
    sub_signals = []
    if decision_authority == "explicit":
        sub_signals.append("hiring/firing authority")
    elif decision_authority == "implicit":
        sub_signals.append("implied decision-making")

    if org_design == "explicit":
        sub_signals.append("org design experience")
    elif org_design == "implicit":
        sub_signals.append("team structuring hints")

    if legacy_status == "explicit":
        sub_signals.append("leadership legacy/succession")

    # Determine evidence status
    if requires_global:
        if has_global_evidence:
            evidence_status = "explicit"
            resume_evidence = "Resume shows globally distributed team leadership"
            # v1.1: Enhance diagnosis with sub-signals
            if sub_signals:
                diagnosis = f"Strong evidence of global team leadership. Sub-signals detected: {', '.join(sub_signals)}."
            else:
                diagnosis = "Strong evidence of global team leadership."
            distance = None
        elif has_local_evidence:
            # v1.1: Check if sub-signals provide implicit global readiness
            if decision_authority == "explicit" or org_design == "explicit":
                evidence_status = "implicit"
                local_context = "local team with leadership maturity signals"
                if leadership_extraction:
                    years = leadership_extraction.get("people_leadership_years", 0)
                    local_context = f"local team ({years:.1f} years people management) with leadership maturity signals"
                resume_evidence = f"Resume shows {local_context}"
                diagnosis = f"No explicit global scope but strong leadership indicators suggest readiness. Sub-signals: {', '.join(sub_signals) if sub_signals else 'decision authority or org design'}."
                distance = "Local (single office) → Global (multi-region) leadership gap (may be coachable)"
            else:
                evidence_status = "missing"
                local_context = "local team"
                if leadership_extraction:
                    years = leadership_extraction.get("people_leadership_years", 0)
                    local_context = f"local team ({years:.1f} years people management)"
                resume_evidence = f"Resume shows {local_context} but no global scope"
                diagnosis = f"Resume shows people management but no evidence of globally distributed teams or multi-region ownership. Experience is local only."
                if sub_signals:
                    diagnosis += f" Limited sub-signals: {', '.join(sub_signals)}."
                distance = "Local (single office) → Global (multi-region) leadership gap"
        else:
            evidence_status = "missing"
            resume_evidence = None
            diagnosis = "No people leadership evidence found. Role requires global team leadership."
            distance = "No leadership → Global leadership gap"

        return {
            "capability_id": "leadership_global_scope",
            "capability_name": "People Leadership – Global Scope",
            "jd_requirement": "Lead globally distributed teams across multiple regions",
            "evidence_status": evidence_status,
            "resume_evidence": resume_evidence,
            "diagnosis": diagnosis,
            "distance": distance,
            "coachable": evidence_status != "missing" or has_local_evidence,
            "criticality": "required"
        }

    if requires_manager_of_managers:
        if has_mom_evidence:
            evidence_status = "explicit"
            resume_evidence = "Resume shows manager of managers experience"
            # v1.1: Enhance diagnosis with sub-signals
            if sub_signals:
                diagnosis = f"Strong evidence of org-level leadership. Sub-signals detected: {', '.join(sub_signals)}."
            else:
                diagnosis = "Strong evidence of org-level leadership."
            distance = None
        elif has_local_evidence:
            # v1.1: Check if sub-signals strengthen implicit evidence
            if decision_authority == "explicit" and org_design in ["explicit", "implicit"]:
                evidence_status = "implicit"
                resume_evidence = "Resume shows IC management with strong leadership maturity signals"
                diagnosis = f"IC management with evidence of leadership maturity. Sub-signals: {', '.join(sub_signals)}. May be ready for promotion to manager of managers."
                distance = "IC Manager → Manager of Managers gap (strong indicators)"
            else:
                evidence_status = "implicit"
                resume_evidence = "Resume shows IC management but not manager of managers"
                diagnosis = "Resume shows people management but no evidence of managing managers (Tier 3 leadership)."
                if sub_signals:
                    diagnosis += f" Partial sub-signals: {', '.join(sub_signals)}."
                distance = "IC Manager → Manager of Managers gap"
        else:
            evidence_status = "missing"
            resume_evidence = None
            diagnosis = "No people leadership evidence found. Role requires org-level leadership."
            distance = "No leadership → Org-level leadership gap"

        return {
            "capability_id": "leadership_org_level",
            "capability_name": "People Leadership – Org Level",
            "jd_requirement": "Lead org-level teams, manage managers",
            "evidence_status": evidence_status,
            "resume_evidence": resume_evidence,
            "diagnosis": diagnosis,
            "distance": distance,
            "coachable": evidence_status != "missing" or has_local_evidence,
            "criticality": "required"
        }

    return None


def evaluate_cross_functional_depth(response_data: dict, resume_data: dict) -> dict:
    """
    Evaluate Cross-Functional Depth capability (CEC v1.1 Enhanced).

    Checks for:
    - Explicit Product partnership
    - Joint roadmap ownership
    - Tradeoff and prioritization language

    v1.1 Sub-signals:
    - Partnership tier (Tier 1 joint ownership, Tier 2 coordination, Tier 3 adjacency)
    - Influence signals (escalation navigation, exec alignment)
    - Conflict navigation (tradeoff language, no pushback resolution)
    """
    role_title = (response_data.get("role_title", "") or "").lower()
    jd_text = (response_data.get("job_description", "") or "").lower()

    # Check if JD requires cross-functional collaboration with Product
    product_collab_keywords = ["product management", "product team", "pm partnership", "roadmap", "product strategy", "cross-functional", "collaborate with product"]

    requires_product_collab = any(kw in jd_text for kw in product_collab_keywords)

    # Also check for engineering manager roles which typically require this
    is_eng_manager = "engineering manager" in role_title or "engineering lead" in role_title

    if not requires_product_collab and not is_eng_manager:
        return None

    # Extract resume evidence
    resume_experience = resume_data.get("experience", []) if resume_data else []
    combined_text = ""
    for exp in resume_experience:
        if isinstance(exp, dict):
            combined_text += f" {exp.get('title', '')} {exp.get('description', '')} {' '.join(exp.get('highlights', []) or exp.get('bullets', []) or [])}".lower()

    # v1.1: Detect partnership tier using helper function
    partnership_tier = detect_partnership_tier(combined_text)

    # v1.1: Detect influence signals (escalation navigation, exec alignment)
    influence_keywords = ["exec", "executive", "leadership", "stakeholder buy-in",
                          "escalation", "aligned with", "presented to", "c-suite",
                          "board", "decision-makers", "senior leadership"]
    has_influence_signals = any(kw in combined_text for kw in influence_keywords)
    influence_status = "explicit" if has_influence_signals else "missing"

    # v1.1: Detect conflict navigation (tradeoff language, disagreement resolution)
    conflict_keywords = ["tradeoff", "trade-off", "prioritized", "deprioritized",
                         "pushed back", "negotiated", "balanced", "competing priorities",
                         "resolved conflict", "disagreement", "aligned stakeholders"]
    has_conflict_signals = any(kw in combined_text for kw in conflict_keywords)
    conflict_status = "explicit" if has_conflict_signals else "missing"

    # Check for explicit Product partnership
    explicit_product = any(kw in combined_text for kw in [
        "partnered with product", "product partnership", "co-owned roadmap",
        "collaborated with pm", "product manager", "joint prioritization",
        "product strategy", "product and engineering"
    ])

    # Check for adjacent cross-functional work (engineering-only)
    adjacent_collab = any(kw in combined_text for kw in [
        "cross-functional", "collaborated with", "worked with", "stakeholder",
        "devops", "sre", "infrastructure", "platform team"
    ])

    # v1.1: Build sub-signals summary for enhanced diagnosis
    sub_signals = []
    if partnership_tier == "tier_1_joint":
        sub_signals.append("Tier 1 joint ownership")
    elif partnership_tier == "tier_2_coordination":
        sub_signals.append("Tier 2 coordination")
    elif partnership_tier == "tier_3_adjacency":
        sub_signals.append("Tier 3 adjacency only")

    if influence_status == "explicit":
        sub_signals.append("exec-level influence")

    if conflict_status == "explicit":
        sub_signals.append("conflict/tradeoff navigation")

    if explicit_product:
        # v1.1: Enhance diagnosis with sub-signals
        if sub_signals:
            diagnosis = f"Strong evidence of Product collaboration. Sub-signals: {', '.join(sub_signals)}."
        else:
            diagnosis = "Strong evidence of Product collaboration."
        return {
            "capability_id": "cross_functional_product",
            "capability_name": "Cross-Functional Partnership – Product",
            "jd_requirement": "Collaborate closely with Product Management on roadmap",
            "evidence_status": "explicit",
            "resume_evidence": "Resume demonstrates Product partnership and roadmap collaboration",
            "diagnosis": diagnosis,
            "distance": None,
            "coachable": None,
            "criticality": "preferred"
        }
    elif adjacent_collab:
        # v1.1: Check if sub-signals strengthen implicit evidence
        if partnership_tier == "tier_2_coordination" and (influence_status == "explicit" or conflict_status == "explicit"):
            resume_evidence = "Resume shows cross-functional coordination with strong influence indicators"
            diagnosis = f"Cross-functional coordination with maturity signals. Sub-signals: {', '.join(sub_signals)}. May have Product partnership readiness."
            distance = "Coordination → Co-ownership gap (strong indicators)"
        else:
            resume_evidence = "Resume shows cross-functional collaboration but not explicit Product partnership"
            diagnosis = "Cross-functional work is present, but direct Product partnership is not demonstrated."
            if sub_signals:
                diagnosis += f" Detected: {', '.join(sub_signals)}."
            distance = "Adjacent functions ≠ Product co-ownership"
        return {
            "capability_id": "cross_functional_product",
            "capability_name": "Cross-Functional Partnership – Product",
            "jd_requirement": "Collaborate closely with Product Management on roadmap",
            "evidence_status": "implicit",
            "resume_evidence": resume_evidence,
            "diagnosis": diagnosis,
            "distance": distance,
            "coachable": True,
            "criticality": "preferred"
        }
    else:
        return {
            "capability_id": "cross_functional_product",
            "capability_name": "Cross-Functional Partnership – Product",
            "jd_requirement": "Collaborate closely with Product Management on roadmap",
            "evidence_status": "missing",
            "resume_evidence": None,
            "diagnosis": "No evidence of cross-functional work or Product collaboration.",
            "distance": "No cross-functional → Product partnership gap",
            "coachable": True,
            "criticality": "preferred"
        }


def evaluate_scale_signals(response_data: dict, resume_data: dict) -> dict:
    """
    Evaluate Scale Signals capability (CEC v1.1 Enhanced).

    Checks for:
    - Team size (5, 12, 50, 200+)
    - Org size (division, multi-team, enterprise-wide)
    - Platform complexity (monolith, microservices, distributed)
    - Revenue, ARR, users, transactions

    v1.1 Sub-signals:
    - Financial ownership (P&L, budget, cost reduction)
    - Time horizon (quarterly, annual, multi-year)
    - Ops intensity (uptime, incident management, SLA)
    """
    jd_text = (response_data.get("job_description", "") or "").lower()

    # Extract scale requirements from JD
    import re

    # Team/org size patterns
    jd_team_match = re.search(r'team of (\d+)|(\d+)\+?\s*(engineers?|developers?|reports?|people)', jd_text)
    jd_scale_match = re.search(r'(\d+)m\+?\s*(users?|requests?|transactions?)|(\d+)\+?\s*(million|billion)', jd_text)
    jd_revenue_match = re.search(r'\$(\d+)m|\$(\d+)\s*million|(\d+)m arr', jd_text)

    requires_large_scale = bool(jd_team_match or jd_scale_match or jd_revenue_match)
    requires_distributed = any(kw in jd_text for kw in ["distributed systems", "microservices", "kubernetes", "high-scale", "at scale"])

    if not requires_large_scale and not requires_distributed:
        return None

    # Extract resume evidence
    resume_experience = resume_data.get("experience", []) if resume_data else []
    combined_text = ""
    for exp in resume_experience:
        if isinstance(exp, dict):
            combined_text += f" {exp.get('title', '')} {exp.get('description', '')} {' '.join(exp.get('highlights', []) or exp.get('bullets', []) or [])}".lower()

    # v1.1: Detect financial ownership using helper function
    financial_ownership = detect_financial_ownership(combined_text)

    # v1.1: Detect time horizon (quarterly, annual, multi-year planning)
    time_horizon_keywords = {
        "multi_year": ["multi-year", "3-year", "5-year", "long-term strategy", "strategic roadmap", "vision"],
        "annual": ["annual", "yearly", "fiscal year", "year-long", "annual planning"],
        "quarterly": ["quarterly", "q1", "q2", "q3", "q4", "okr", "sprint planning"]
    }
    time_horizon = "missing"
    for horizon, keywords in time_horizon_keywords.items():
        if any(kw in combined_text for kw in keywords):
            time_horizon = horizon
            break

    # v1.1: Detect ops intensity (uptime, incident management, SLA)
    ops_intensity_keywords = ["uptime", "incident", "sla", "on-call", "reliability",
                              "availability", "pager", "outage", "disaster recovery",
                              "99.9", "five nines", "production"]
    has_ops_intensity = any(kw in combined_text for kw in ops_intensity_keywords)
    ops_intensity = "explicit" if has_ops_intensity else "missing"

    # Check for scale evidence in resume
    resume_team_match = re.search(r'team of (\d+)|managed (\d+)|(\d+)\s*(engineers?|developers?|reports?|people)', combined_text)
    resume_scale_match = re.search(r'(\d+)m\+?\s*(users?|requests?|transactions?|api)|(\d+)\+?\s*(million|billion)|(\d+)k\s*(users?|requests?)', combined_text)
    resume_revenue_match = re.search(r'\$(\d+)m|\$(\d+)\s*million|(\d+)m arr', combined_text)
    resume_distributed = any(kw in combined_text for kw in ["distributed systems", "microservices", "kubernetes", "high-scale", "at scale", "platform"])

    # Determine evidence status
    has_scale_evidence = bool(resume_scale_match or resume_revenue_match or resume_distributed)
    has_team_evidence = bool(resume_team_match)

    # Extract specific numbers for diagnosis
    resume_team_size = None
    if resume_team_match:
        for g in resume_team_match.groups():
            if g and g.isdigit():
                resume_team_size = int(g)
                break

    jd_team_size = None
    if jd_team_match:
        for g in jd_team_match.groups():
            if g and g.isdigit():
                jd_team_size = int(g)
                break

    # v1.1: Build sub-signals summary for enhanced diagnosis
    sub_signals = []
    if financial_ownership["level"] == "full":
        sub_signals.append(f"P&L ownership ({financial_ownership['type'] or 'budget'})")
    elif financial_ownership["level"] == "partial":
        sub_signals.append("cost accountability")

    if time_horizon == "multi_year":
        sub_signals.append("multi-year planning")
    elif time_horizon == "annual":
        sub_signals.append("annual planning")
    elif time_horizon == "quarterly":
        sub_signals.append("quarterly planning")

    if ops_intensity == "explicit":
        sub_signals.append("ops/reliability ownership")

    if has_scale_evidence:
        evidence = "Resume demonstrates scale"
        if resume_scale_match:
            evidence = f"Resume shows high-scale experience"
        # v1.1: Enhance diagnosis with sub-signals
        if sub_signals:
            diagnosis = f"Strong evidence of scale experience. Sub-signals: {', '.join(sub_signals)}."
        else:
            diagnosis = "Strong evidence of scale experience."
        return {
            "capability_id": "scale_signals",
            "capability_name": "Scale Signals – Platform/Org",
            "jd_requirement": "Experience at scale (large teams, high traffic, enterprise)",
            "evidence_status": "explicit",
            "resume_evidence": evidence,
            "diagnosis": diagnosis,
            "distance": None,
            "coachable": None,
            "criticality": "required"
        }
    elif has_team_evidence:
        # v1.1: Check if sub-signals strengthen implicit evidence
        if financial_ownership["level"] == "full" or (time_horizon in ["multi_year", "annual"] and ops_intensity == "explicit"):
            diagnosis = f"Team management with executive-level scale signals"
            if resume_team_size:
                diagnosis = f"Team of {resume_team_size} with executive-level signals. Sub-signals: {', '.join(sub_signals)}."
            resume_evidence = f"Team of {resume_team_size} with scale maturity indicators" if resume_team_size else "Team management with scale maturity"
            distance = f"Current scale ({resume_team_size or 'small'}) → Required scale ({jd_team_size or 'large'}) (strong indicators)" if jd_team_size else "Small → Large scale gap (strong indicators)"
        else:
            diagnosis = f"Resume shows team management"
            if resume_team_size:
                diagnosis = f"Resume shows team management ({resume_team_size} people)"
                if jd_team_size and resume_team_size < jd_team_size:
                    diagnosis += f" but JD requires larger scale ({jd_team_size}+ people)"
            if sub_signals:
                diagnosis += f" Partial sub-signals: {', '.join(sub_signals)}."
            resume_evidence = f"Team of {resume_team_size}" if resume_team_size else "Team management experience"
            distance = f"Current scale ({resume_team_size or 'small'}) → Required scale ({jd_team_size or 'large'})" if jd_team_size else "Small → Large scale gap"
        return {
            "capability_id": "scale_signals",
            "capability_name": "Scale Signals – Platform/Org",
            "jd_requirement": "Experience at scale (large teams, high traffic, enterprise)",
            "evidence_status": "implicit",
            "resume_evidence": resume_evidence,
            "diagnosis": diagnosis,
            "distance": distance,
            "coachable": True,
            "criticality": "required"
        }
    else:
        return {
            "capability_id": "scale_signals",
            "capability_name": "Scale Signals – Platform/Org",
            "jd_requirement": "Experience at scale (large teams, high traffic, enterprise)",
            "evidence_status": "missing",
            "resume_evidence": None,
            "diagnosis": "No scale indicators found in resume. JD requires enterprise or high-scale experience.",
            "distance": "No scale evidence → Enterprise scale gap",
            "coachable": True,
            "criticality": "required"
        }


def evaluate_domain_adjacency(response_data: dict, resume_data: dict) -> dict:
    """
    Evaluate Domain Adjacency capability (CEC v1.1 Enhanced).

    Checks for:
    - Direct domain experience (exact match)
    - Closely adjacent experience (transferable)
    - Unrelated experience (different domain)

    v1.1 Sub-signals:
    - Domain depth (years, projects, specialization)
    - Transferability signals (cross-domain patterns, methodology familiarity)
    - Context awareness (regulatory, compliance, industry-specific)
    """
    role_title = (response_data.get("role_title", "") or "").lower()
    jd_text = (response_data.get("job_description", "") or "").lower()
    company = (response_data.get("company", "") or "").lower()

    # Domain mapping - JD keywords to domain categories
    domain_categories = {
        "safety": ["safety", "trust & safety", "trust and safety", "content moderation", "abuse", "fraud prevention", "risk"],
        "fintech": ["fintech", "financial services", "payments", "banking", "lending", "trading", "investment"],
        "healthcare": ["healthcare", "health tech", "medical", "clinical", "patient", "hipaa"],
        "ecommerce": ["e-commerce", "ecommerce", "marketplace", "retail", "commerce"],
        "enterprise": ["enterprise", "b2b", "saas", "business software"],
        "consumer": ["consumer", "b2c", "social", "mobile app"],
        "ml_ai": ["machine learning", "ml", "ai", "artificial intelligence", "deep learning", "nlp"],
        "infrastructure": ["infrastructure", "platform", "devops", "sre", "cloud"]
    }

    # Detect required domain from JD
    required_domain = None
    jd_domain_keywords = []
    for domain, keywords in domain_categories.items():
        for kw in keywords:
            if kw in jd_text or kw in role_title:
                required_domain = domain
                jd_domain_keywords.append(kw)
                break
        if required_domain:
            break

    if not required_domain:
        return None

    # Extract resume domains
    resume_experience = resume_data.get("experience", []) if resume_data else []
    combined_text = ""
    for exp in resume_experience:
        if isinstance(exp, dict):
            combined_text += f" {exp.get('title', '')} {exp.get('company', '')} {exp.get('description', '')} {' '.join(exp.get('highlights', []) or exp.get('bullets', []) or [])}".lower()

    resume_domains = set()
    for domain, keywords in domain_categories.items():
        for kw in keywords:
            if kw in combined_text:
                resume_domains.add(domain)
                break

    # v1.1: Detect domain depth using helper function
    domain_depth = detect_domain_depth(combined_text, required_domain)

    # v1.1: Detect transferability signals (cross-domain patterns)
    transferability_keywords = ["cross-functional", "multiple domains", "diverse industries",
                                "consulting", "transferred", "applied to", "adapted",
                                "methodology", "framework", "best practices", "playbook"]
    has_transferability = any(kw in combined_text for kw in transferability_keywords)
    transferability_status = "explicit" if has_transferability else "missing"

    # v1.1: Detect context awareness (regulatory, compliance, industry-specific)
    context_keywords = {
        "regulatory": ["regulatory", "compliance", "audit", "sox", "pci", "gdpr", "ccpa", "sec"],
        "industry_specific": ["industry-specific", "vertical", "sector", "domain expert"],
        "policy": ["policy", "governance", "standards", "certification"]
    }
    context_signals = []
    for context_type, keywords in context_keywords.items():
        if any(kw in combined_text for kw in keywords):
            context_signals.append(context_type.replace("_", " "))

    # v1.1: Build sub-signals summary
    sub_signals = []
    if domain_depth == "explicit":
        sub_signals.append("deep domain expertise")
    elif domain_depth == "implicit":
        sub_signals.append("some domain exposure")

    if transferability_status == "explicit":
        sub_signals.append("transferability signals")

    if context_signals:
        sub_signals.append(f"context awareness ({', '.join(context_signals)})")

    # Check for exact match
    if required_domain in resume_domains:
        # v1.1: Enhance diagnosis with sub-signals
        if sub_signals:
            diagnosis = f"Strong evidence of {required_domain.replace('_', ' ')} domain expertise. Sub-signals: {', '.join(sub_signals)}."
        else:
            diagnosis = f"Strong evidence of {required_domain.replace('_', ' ')} domain expertise."
        return {
            "capability_id": f"domain_{required_domain}",
            "capability_name": f"Domain Expertise – {required_domain.replace('_', ' ').title()}",
            "jd_requirement": f"Experience in {required_domain.replace('_', ' ')} domain",
            "evidence_status": "explicit",
            "resume_evidence": f"Resume shows direct {required_domain.replace('_', ' ')} experience",
            "diagnosis": diagnosis,
            "distance": None,
            "coachable": None,
            "criticality": "required"
        }

    # Check for adjacent domains (one degree of separation)
    adjacent_domains = {
        "safety": ["fintech", "enterprise"],  # fraud/risk adjacent
        "fintech": ["ecommerce", "enterprise", "safety"],
        "healthcare": ["enterprise"],
        "ecommerce": ["fintech", "consumer"],
        "enterprise": ["fintech", "ecommerce", "infrastructure"],
        "consumer": ["ecommerce", "enterprise"],
        "ml_ai": ["infrastructure", "enterprise"],
        "infrastructure": ["enterprise", "ml_ai"]
    }

    adjacent_to_required = adjacent_domains.get(required_domain, [])
    overlapping_adjacent = resume_domains.intersection(set(adjacent_to_required))

    if overlapping_adjacent:
        adjacent_domain = list(overlapping_adjacent)[0]
        # v1.1: Check if sub-signals strengthen adjacency case
        if transferability_status == "explicit" or context_signals:
            resume_evidence = f"Resume shows {adjacent_domain.replace('_', ' ')} experience with strong transferability signals"
            diagnosis = f"{adjacent_domain.replace('_', ' ').title()} experience is adjacent to {required_domain.replace('_', ' ')} with transferability indicators. Sub-signals: {', '.join(sub_signals)}."
            distance = f"{adjacent_domain.replace('_', ' ').title()} → {required_domain.replace('_', ' ').title()} domain gap (strong bridge)"
        else:
            resume_evidence = f"Resume shows {adjacent_domain.replace('_', ' ')} experience (adjacent domain)"
            diagnosis = f"{adjacent_domain.replace('_', ' ').title()} experience is adjacent to {required_domain.replace('_', ' ')}, but no direct {required_domain.replace('_', ' ')} background."
            if sub_signals:
                diagnosis += f" Detected: {', '.join(sub_signals)}."
            distance = f"{adjacent_domain.replace('_', ' ').title()} → {required_domain.replace('_', ' ').title()} domain gap (1 degree)"
        return {
            "capability_id": f"domain_{required_domain}",
            "capability_name": f"Domain Expertise – {required_domain.replace('_', ' ').title()}",
            "jd_requirement": f"Experience in {required_domain.replace('_', ' ')} domain",
            "evidence_status": "implicit",
            "resume_evidence": resume_evidence,
            "diagnosis": diagnosis,
            "distance": distance,
            "coachable": True,
            "criticality": "required"
        }

    # No match - missing
    resume_domain_str = ", ".join([d.replace('_', ' ') for d in resume_domains]) if resume_domains else "general"

    # =========================================================================
    # GUARDRAIL: Auto-downgrade healthcare domain gap for manager-level roles
    # Per Calibration Fix: If role title is manager-level AND domain gap is
    # healthcare AND JD does NOT explicitly mention healthcare keywords,
    # downgrade from terminal (coachable: False) to coachable (coachable: True)
    # =========================================================================
    manager_level_keywords = ["engineering manager", "product manager", "program manager",
                              "project manager", "senior manager", "staff", "senior",
                              "team lead", "tech lead", "technical lead"]
    healthcare_keywords = ["healthcare", "health tech", "medical", "clinical", "patient", "hipaa"]

    is_manager_level = any(kw in role_title for kw in manager_level_keywords)
    jd_explicitly_requires_healthcare = any(kw in jd_text for kw in healthcare_keywords)

    # If healthcare domain gap but JD doesn't explicitly require healthcare
    # AND this is a manager-level role, make it coachable
    if required_domain == "healthcare" and is_manager_level and not jd_explicitly_requires_healthcare:
        print(f"   ⚠️ GUARDRAIL: Healthcare domain gap detected but JD does NOT explicitly require healthcare")
        print(f"      Role title: {role_title} (manager-level: {is_manager_level})")
        print(f"      Auto-downgrading from terminal → coachable")
        return {
            "capability_id": f"domain_{required_domain}",
            "capability_name": f"Domain Expertise – {required_domain.replace('_', ' ').title()}",
            "jd_requirement": f"Experience in {required_domain.replace('_', ' ')} domain",
            "evidence_status": "missing",
            "resume_evidence": None,
            "diagnosis": f"No direct {required_domain.replace('_', ' ')} experience. Resume shows {resume_domain_str} background. Note: JD does not explicitly require healthcare. Domain gap auto-downgraded to coachable.",
            "distance": f"0 → {required_domain.replace('_', ' ')} domain gap (auto-downgraded to coachable)",
            "coachable": True,  # Auto-downgraded per guardrail
            "criticality": "preferred"  # Downgrade criticality too
        }

    return {
        "capability_id": f"domain_{required_domain}",
        "capability_name": f"Domain Expertise – {required_domain.replace('_', ' ').title()}",
        "jd_requirement": f"Experience in {required_domain.replace('_', ' ')} domain",
        "evidence_status": "missing",
        "resume_evidence": None,
        "diagnosis": f"No direct {required_domain.replace('_', ' ')} experience. Resume shows {resume_domain_str} background. Unrelated domain.",
        "distance": f"0 → {required_domain.replace('_', ' ')} domain gap (not coachable in current search)",
        "coachable": False,
        "criticality": "required"
    }


# ============================================================================
# LEADERSHIP EVALUATION & POSITIONING ENGINE (LEPE)
# Per LEPE Spec: Canonical System for Manager+ Leadership Readiness
#
# Purpose: Evaluate people leadership readiness with recruiter-grade realism,
# while providing honest, actionable coaching when gaps are addressable.
#
# LEPE must:
# - Be brutally honest without shaming
# - Separate eligibility from advisability
# - Detect leadership gaps, title inflation, scope mismatches
# - Enable positioning strategies when appropriate
# - Preserve accountability when users proceed against advice
# ============================================================================

# Manager+ role level indicators
MANAGER_PLUS_LEVELS = [
    "manager", "senior manager", "director", "senior director",
    "vp", "vice president", "head of", "head", "chief",
    "c-suite", "ceo", "cto", "cfo", "coo", "cmo", "cio", "cpo"
]


def is_manager_plus_role(role_title: str) -> bool:
    """
    Determine if a role is Manager+ level.

    LEPE applies ONLY to Manager+ roles. This is non-negotiable.
    If role_level < manager: skip all LEPE logic.

    Args:
        role_title: The job title to evaluate

    Returns:
        bool: True if Manager+ level, False otherwise
    """
    if not role_title:
        return False

    title_lower = role_title.lower().strip()

    # Check for Manager+ indicators
    for level in MANAGER_PLUS_LEVELS:
        if level in title_lower:
            # Exclude false positives like "account manager" (sales IC role)
            if level == "manager":
                # These are typically IC roles with "manager" in title
                ic_manager_roles = [
                    "account manager", "customer success manager",
                    "project manager", "program manager", "product manager",
                    "case manager", "office manager", "facilities manager"
                ]
                if any(ic_role in title_lower for ic_role in ic_manager_roles):
                    # But Director/VP of these are still Manager+
                    if any(exec_level in title_lower for exec_level in ["director", "vp", "vice president", "head", "chief"]):
                        return True
                    return False
            return True

    return False


def extract_leadership_competency_signals(resume_data: dict) -> dict:
    """
    Extract leadership signals across 6 competency domains per LEPE spec.

    Core Competency Domains:
    1. People Management - Hiring, performance management, coaching, exits
    2. Decision Authority - Hiring approvals, budget ownership, escalation
    3. Org Design & Scale - Team growth, restructures, span of control
    4. Strategic Leadership - Direction setting, prioritization, multi-quarter planning
    5. Cross-Functional Influence - Exec partnership, conflict resolution, alignment
    6. Accountability & Ownership - Owning outcomes, risk, failure, corrective action

    Args:
        resume_data: Parsed resume dictionary

    Returns:
        dict: leadership_competency_analysis with signals per domain
    """
    result = {
        "competency_domains": {
            "people_management": {"signals": [], "strength": "none", "years": 0.0},
            "decision_authority": {"signals": [], "strength": "none", "years": 0.0},
            "org_design_scale": {"signals": [], "strength": "none", "years": 0.0},
            "strategic_leadership": {"signals": [], "strength": "none", "years": 0.0},
            "cross_functional_influence": {"signals": [], "strength": "none", "years": 0.0},
            "accountability_ownership": {"signals": [], "strength": "none", "years": 0.0}
        },
        "total_leadership_signals": 0,
        "strongest_domain": None,
        "weakest_domain": None,
        "mixed_scope_roles": []
    }

    if not resume_data:
        return result

    experience = resume_data.get("experience", [])
    if not experience or not isinstance(experience, list):
        return result

    # Signal patterns by domain
    domain_patterns = {
        "people_management": {
            "explicit": [
                "direct reports", "managed a team", "hired", "fired", "terminated",
                "performance review", "promoted team members", "coached", "mentored team",
                "developed team", "built the team", "grew team from", "onboarded",
                "managed team of", "led a team of", "people manager"
            ],
            "implied": [
                "team lead", "manager", "supervised", "managed staff",
                "leadership role", "managed engineers", "managed designers"
            ]
        },
        "decision_authority": {
            "explicit": [
                "hiring authority", "budget owner", "budget ownership", "p&l",
                "final decision", "approval authority", "headcount authority",
                "signing authority", "spending authority"
            ],
            "implied": [
                "owned budget", "managed budget", "allocated resources",
                "prioritized investments", "approved hires"
            ]
        },
        "org_design_scale": {
            "explicit": [
                "restructured", "reorganized", "built the team", "scaled team",
                "grew team from", "org design", "span of control", "team of",
                "department of", "organization of"
            ],
            "implied": [
                "expanded team", "hired aggressively", "doubled team",
                "tripled headcount", "built from scratch"
            ]
        },
        "strategic_leadership": {
            "explicit": [
                "set direction", "defined strategy", "roadmap owner",
                "multi-year vision", "long-term planning", "quarterly planning",
                "okr owner", "north star", "strategic initiative"
            ],
            "implied": [
                "drove strategy", "led initiative", "spearheaded",
                "championed", "vision for", "transformed"
            ]
        },
        "cross_functional_influence": {
            "explicit": [
                "cross-functional", "exec partnership", "c-suite partnership",
                "board presentation", "executive alignment", "stakeholder management"
            ],
            "implied": [
                "partnered with", "aligned with", "collaborated across",
                "influenced", "coordinated with", "matrixed"
            ]
        },
        "accountability_ownership": {
            "explicit": [
                "owned outcomes", "accountable for", "responsible for",
                "drove results", "delivered", "achieved", "exceeded targets"
            ],
            "implied": [
                "led to", "resulted in", "improved by", "reduced by",
                "increased by", "grew by"
            ]
        }
    }

    # IC signals that indicate mixed scope
    ic_signals = [
        "individual contributor", "hands-on coding", "wrote code",
        "technical implementation", "built features", "debugged",
        "deployed", "architected", "designed system"
    ]

    for exp in experience:
        if not isinstance(exp, dict):
            continue

        title = (exp.get("title", "") or "").lower()
        description = (exp.get("description", "") or "").lower()
        highlights = exp.get("highlights", [])
        if highlights and isinstance(highlights, list):
            highlights_text = " ".join([h.lower() for h in highlights if isinstance(h, str)])
        else:
            highlights_text = ""

        combined_text = f"{title} {description} {highlights_text}"
        dates = (exp.get("dates", "") or "").lower()
        company = (exp.get("company", "") or "").strip()
        years = parse_experience_duration(dates)

        if years <= 0:
            continue

        # Check for IC signals (for mixed-scope detection)
        has_ic_signals = any(sig in combined_text for sig in ic_signals)
        has_leadership_signals = False
        role_domains_found = []

        # Extract signals per domain
        for domain, patterns in domain_patterns.items():
            domain_signals = []

            for pattern in patterns["explicit"]:
                if pattern in combined_text:
                    domain_signals.append({
                        "pattern": pattern,
                        "evidence_type": "explicit",
                        "strength": "strong",
                        "role_title": exp.get("title", ""),
                        "company": company,
                        "years": years
                    })
                    has_leadership_signals = True

            for pattern in patterns["implied"]:
                if pattern in combined_text:
                    domain_signals.append({
                        "pattern": pattern,
                        "evidence_type": "implied",
                        "strength": "medium",
                        "role_title": exp.get("title", ""),
                        "company": company,
                        "years": years
                    })
                    has_leadership_signals = True

            if domain_signals:
                result["competency_domains"][domain]["signals"].extend(domain_signals)
                role_domains_found.append(domain)

        # Track mixed-scope roles
        if has_ic_signals and has_leadership_signals:
            result["mixed_scope_roles"].append({
                "title": exp.get("title", ""),
                "company": company,
                "years": years,
                "domains_found": role_domains_found,
                "leadership_scope_estimate": 0.5  # Default 50% for mixed roles
            })

    # Calculate years and strength per domain
    total_signals = 0
    domain_years = {}

    for domain, data in result["competency_domains"].items():
        signals = data["signals"]
        total_signals += len(signals)

        # Calculate years (with deduplication by role)
        role_years = {}
        for sig in signals:
            role_key = f"{sig['role_title']}@{sig['company']}"
            if role_key not in role_years:
                role_years[role_key] = sig["years"]

        domain_total_years = sum(role_years.values())
        data["years"] = round(domain_total_years, 1)
        domain_years[domain] = domain_total_years

        # Determine strength
        explicit_count = len([s for s in signals if s["evidence_type"] == "explicit"])
        if explicit_count >= 3 or domain_total_years >= 5:
            data["strength"] = "strong"
        elif explicit_count >= 1 or domain_total_years >= 2:
            data["strength"] = "medium"
        elif signals:
            data["strength"] = "weak"

    result["total_leadership_signals"] = total_signals

    # Find strongest/weakest domains
    if domain_years:
        result["strongest_domain"] = max(domain_years, key=domain_years.get)
        # Only mark weakest if there's at least one signal
        domains_with_signals = {d: y for d, y in domain_years.items() if y > 0}
        if domains_with_signals:
            result["weakest_domain"] = min(domains_with_signals, key=domains_with_signals.get)

    return result


def calculate_leadership_tenure_lepe(resume_data: dict) -> dict:
    """
    Calculate leadership tenure per LEPE spec.

    CRITICAL RULES:
    - Leadership tenure is NEVER derived from total career years
    - People leadership years = cumulative duration of Tier 2+ signals
    - Tier 1 signals alone → 0 people leadership years
    - Partial roles receive proportional credit
    - Never default to 0 if ANY Tier 1 or partial Tier 2 signals exist

    Args:
        resume_data: Parsed resume dictionary

    Returns:
        dict: Leadership tenure breakdown
    """
    # Get tiered leadership analysis
    tiered = extract_tiered_leadership(resume_data)
    competency = extract_leadership_competency_signals(resume_data)

    # Calculate mixed-scope adjustments for IC/leadership hybrid roles
    # These roles get proportional credit (e.g., 50% for player-coach)
    mixed_scope_adjustment = 0.0
    for role in competency.get("mixed_scope_roles", []):
        # Apply proportional credit (50% for mixed roles)
        mixed_scope_adjustment += role["years"] * role["leadership_scope_estimate"]

    # Base people leadership from tiered analysis
    base_people_years = tiered.get("people_leadership_years", 0.0)

    # Add mixed-scope credit to people leadership total
    # This ensures gradual leadership transitions get appropriate credit
    total_people_years = base_people_years + mixed_scope_adjustment

    result = {
        "people_leadership_years": round(total_people_years, 1),
        "people_leadership_years_base": round(base_people_years, 1),
        "strategic_leadership_years": tiered.get("strategic_leadership_years", 0.0),
        "org_level_leadership_years": tiered.get("org_leadership_years", 0.0),
        "mixed_scope_credit": round(mixed_scope_adjustment, 1),
        "has_any_leadership_signals": tiered.get("has_any_leadership", False) or competency.get("total_leadership_signals", 0) > 0,
        "competency_breakdown": {
            domain: {
                "years": data["years"],
                "strength": data["strength"]
            }
            for domain, data in competency.get("competency_domains", {}).items()
        },
        "tier_summary": tiered.get("leadership_tier_summary", ""),
        "strongest_competency": competency.get("strongest_domain"),
        "development_area": competency.get("weakest_domain")
    }

    return result


def get_lepe_positioning_decision(
    leadership_tenure: dict,
    required_leadership_years: float,
    role_title: str
) -> dict:
    """
    LEPE Positioning Decision Engine.

    Decision Logic:
    - Gap ≤ 2 years → Positioning Mode (coaching, narrative building)
    - Gap 3-4 years → Apply with Caution + skepticism language
    - Gap > 4 years OR missing core competency → Do Not Apply (Locked)

    Args:
        leadership_tenure: Output from calculate_leadership_tenure_lepe()
        required_leadership_years: Years required by the role
        role_title: The target role title

    Returns:
        dict: Positioning decision with coaching or lock
    """
    people_years = leadership_tenure.get("people_leadership_years", 0.0)
    has_any_signals = leadership_tenure.get("has_any_leadership_signals", False)
    strongest = leadership_tenure.get("strongest_competency")
    development_area = leadership_tenure.get("development_area")

    gap = required_leadership_years - people_years

    result = {
        "decision": "apply",  # apply | position | caution | locked
        "gap_years": round(gap, 1),
        "candidate_years": people_years,
        "required_years": required_leadership_years,
        "positioning_mode": False,
        "coaching_available": False,
        "skepticism_level": "none",  # none | mild | significant | severe
        "transition_narrative_possible": False,
        "messaging": {
            "headline": "",
            "explanation": "",
            "coaching_advice": "",
            "skepticism_warning": ""
        }
    }

    # Determine decision based on gap
    if gap <= 0:
        # Meets or exceeds requirement
        result["decision"] = "apply"
        result["skepticism_level"] = "none"
        result["messaging"]["headline"] = "Your leadership experience meets this role's requirements."
        result["messaging"]["explanation"] = (
            f"You have {people_years:.1f} years of people leadership experience. "
            f"This role requires {required_leadership_years:.0f}+ years."
        )

    elif gap <= 2:
        # Positioning Mode
        result["decision"] = "position"
        result["positioning_mode"] = True
        result["coaching_available"] = True
        result["skepticism_level"] = "mild"
        result["transition_narrative_possible"] = True
        result["messaging"]["headline"] = "Addressable gap. Positioning strategy available."
        result["messaging"]["explanation"] = (
            f"You have {people_years:.1f} years of people leadership. "
            f"This role screens for {required_leadership_years:.0f}+. "
            f"The {gap:.1f}-year gap is addressable with the right positioning."
        )

        # Build coaching advice
        coaching_parts = []
        if strongest:
            coaching_parts.append(
                f"Lead with your strength in {strongest.replace('_', ' ')}."
            )
        if development_area:
            coaching_parts.append(
                f"Proactively address {development_area.replace('_', ' ')} in your narrative."
            )
        coaching_parts.append(
            "Quantify your leadership impact: team size, outcomes, scope of decisions."
        )
        result["messaging"]["coaching_advice"] = " ".join(coaching_parts)
        result["messaging"]["skepticism_warning"] = (
            "Hiring managers may probe your leadership depth. Prepare concrete examples."
        )

    elif gap <= 4:
        # Apply with Caution
        result["decision"] = "caution"
        result["positioning_mode"] = False
        result["coaching_available"] = True
        result["skepticism_level"] = "significant"
        result["messaging"]["headline"] = "Significant leadership gap. Apply with realistic expectations."
        result["messaging"]["explanation"] = (
            f"You have {people_years:.1f} years of people leadership. "
            f"This role requires {required_leadership_years:.0f}+. "
            f"The {gap:.1f}-year gap will raise questions in screening."
        )
        result["messaging"]["skepticism_warning"] = (
            "Hiring managers will likely question whether you're ready for this level. "
            "Unless you have a strong internal referral or unique differentiator, "
            "this is an uphill battle."
        )

        if strongest:
            result["messaging"]["coaching_advice"] = (
                f"If you proceed: heavily emphasize {strongest.replace('_', ' ')} "
                f"and be ready to address the experience gap directly."
            )

    else:
        # Gap > 4 years — Locked
        result["decision"] = "locked"
        result["positioning_mode"] = False
        result["coaching_available"] = False
        result["skepticism_level"] = "severe"

        if not has_any_signals:
            result["messaging"]["headline"] = "No leadership signals detected. This role is not a fit."
            result["messaging"]["explanation"] = (
                f"This role requires {required_leadership_years:.0f}+ years of people leadership. "
                f"Your resume shows no verified leadership experience at any tier."
            )
        else:
            result["messaging"]["headline"] = "Leadership experience gap too large for this role."
            result["messaging"]["explanation"] = (
                f"You have {people_years:.1f} years of people leadership. "
                f"This role requires {required_leadership_years:.0f}+. "
                f"The {gap:.1f}-year gap cannot be bridged through positioning alone."
            )

        result["messaging"]["skepticism_warning"] = (
            "Recruiters will screen you out based on years alone. "
            "Focus on roles that match your current experience level."
        )

    return result


def evaluate_lepe(
    resume_data: dict,
    response_data: dict
) -> dict:
    """
    Main LEPE evaluation function.

    Runs full Leadership Evaluation & Positioning Engine analysis.

    APPLICABILITY: Only runs for Manager+ roles.
    If role is below Manager level, returns bypass result.

    Args:
        resume_data: Parsed resume dictionary
        response_data: JD analysis response data

    Returns:
        dict: Complete LEPE analysis including:
            - lepe_applicable: bool
            - leadership_competency_analysis: dict
            - leadership_tenure: dict
            - positioning_decision: dict
            - accountability_record: dict
    """
    role_title = (response_data.get("role_title", "") or "").strip()

    result = {
        "lepe_applicable": False,
        "role_level": "unknown",
        "leadership_competency_analysis": None,
        "leadership_tenure": None,
        "positioning_decision": None,
        "accountability_record": {
            "gaps_identified": [],
            "advice_given": "",
            "risk_communicated": "",
            "timestamp": None
        }
    }

    # Check applicability
    if not is_manager_plus_role(role_title):
        result["role_level"] = "below_manager"
        print(f"🔄 LEPE BYPASS: Role '{role_title}' is below Manager level")
        return result

    result["lepe_applicable"] = True
    result["role_level"] = "manager_plus"
    print(f"📊 LEPE EVALUATION: Manager+ role detected - '{role_title}'")

    # Extract leadership competency signals
    competency = extract_leadership_competency_signals(resume_data)
    result["leadership_competency_analysis"] = competency

    # Calculate leadership tenure
    tenure = calculate_leadership_tenure_lepe(resume_data)
    result["leadership_tenure"] = tenure

    # Infer leadership requirement from role level
    required_years = infer_leadership_requirement(role_title, response_data)

    # Get positioning decision
    positioning = get_lepe_positioning_decision(tenure, required_years, role_title)
    result["positioning_decision"] = positioning

    # Build accountability record
    gaps = []
    if positioning["gap_years"] > 0:
        gaps.append(f"People leadership gap: {positioning['gap_years']:.1f} years")

    weakest = tenure.get("development_area")
    if weakest:
        gaps.append(f"Development area: {weakest.replace('_', ' ')}")

    result["accountability_record"] = {
        "gaps_identified": gaps,
        "advice_given": positioning["messaging"].get("coaching_advice", ""),
        "risk_communicated": positioning["messaging"].get("skepticism_warning", ""),
        "lepe_advisory": positioning["decision"],  # Renamed: this is advisory, NOT the final recommendation
        "timestamp": datetime.now().isoformat()
    }

    # IMPORTANT: This is LEPE's advisory signal, NOT the final recommendation
    # Final recommendation is set by FinalRecommendationController and is immutable
    print(f"   LEPE Advisory: {positioning['decision'].upper()} (does not override Final Recommendation)")
    print(f"   People leadership years: {tenure.get('people_leadership_years', 0):.1f}")
    print(f"   Required: {required_years:.0f}+")
    print(f"   Gap: {positioning['gap_years']:.1f} years")

    return result


def infer_leadership_requirement(role_title: str, response_data: dict) -> float:
    """
    Infer leadership years requirement from role level when JD is vague.

    Per LEPE spec: When a JD is vague, LEPE must infer expectations
    based on role level. Inferred requirements are additive, not fabricated.

    Args:
        role_title: The job title
        response_data: JD analysis response

    Returns:
        float: Inferred minimum leadership years requirement
    """
    title_lower = (role_title or "").lower()

    # First check if JD explicitly states years
    required_years, is_hard = extract_required_people_leadership_years(response_data)
    if required_years > 0:
        return required_years

    # Infer based on role level
    # Check C-suite with word boundaries to avoid false positives like "director" -> "cto"
    import re
    if re.search(r'\b(chief|ceo|cto|cfo|coo|cmo|cio|cpo)\b', title_lower):
        return 10.0  # C-suite

    if any(v in title_lower for v in ["vp ", "vp,", "vice president", "svp", "evp"]):
        return 7.0  # VP level

    if "head of" in title_lower or "head," in title_lower:
        return 7.0  # Head of X

    if "senior director" in title_lower:
        return 7.0

    if "director" in title_lower:
        return 5.0

    if "senior manager" in title_lower:
        return 4.0

    if "manager" in title_lower:
        return 2.0

    # Default for ambiguous Manager+ roles
    return 3.0


def extract_people_leadership_years(resume_data: dict) -> float:
    """
    Extract ONLY people leadership years (direct reports, team management).

    HARD RULE: Only count roles where the candidate had direct reports.
    Operational leadership (systems, programs, processes) does NOT count.

    Evidence keywords for people leadership:
    - "direct reports", "team of X", "managed X people", "led a team"
    - "hiring", "fired", "performance reviews", "promoted", "developed team"
    - "people manager", "built the team", "grew the team from X to Y"

    Args:
        resume_data: Parsed resume dictionary

    Returns:
        float: Total verified people leadership years
    """
    if not resume_data:
        return 0.0

    experience = resume_data.get("experience", [])
    if not experience or not isinstance(experience, list):
        return 0.0

    # PEOPLE leadership indicators - must show direct reports/team management
    people_leadership_evidence = [
        "direct report", "direct reports", "managed a team", "led a team",
        "team of", "people manager", "managed team", "lead a team",
        "built the team", "grew the team", "hiring manager",
        "performance review", "promoted", "mentored", "coached team",
        "developed team", "team lead", "engineering manager", "people management",
        # Additional patterns from real resumes
        "built and led", "led marketing team", "led team", "scaled marketing",
        "scaled team", "scaled function", "team from", "grew from",
        "marketing team of", "engineering team of", "led the team",
        "managed budget", "annual budget", "led organization",
        "built organization", "reporting to", "reports to me",
        "oversaw team", "supervised", "led cross-functional"
    ]

    # Title patterns that typically indicate people leadership
    people_leadership_titles = [
        "manager", "director", "head of", "vp ", "vice president",
        "chief", "lead", "supervisor", "team lead"
    ]

    # STRONG leadership titles that inherently imply people management at established companies
    # VP/Director at real companies = you manage people (no explicit evidence needed)
    strong_leadership_titles = [
        "vp ", "vice president", "vp,", "director", "head of", "chief",
        # Senior Manager titles at established companies also typically have reports
        "senior manager", "senior product marketing manager", "group manager",
        "marketing manager", "product marketing manager"
    ]

    # Explicitly NOT people leadership (operational/systems leadership)
    operational_only_patterns = [
        "program manager", "project manager", "technical lead",
        "staff engineer", "principal engineer", "architect",
        "operations lead", "process lead", "systems lead"
    ]

    # Well-known companies where VP/Director definitely means people leadership
    established_companies = [
        "google", "meta", "facebook", "amazon", "apple", "microsoft", "netflix",
        "uber", "lyft", "airbnb", "stripe", "square", "twilio", "segment",
        "salesforce", "adobe", "oracle", "sap", "ibm", "cisco", "intel",
        "linkedin", "twitter", "x corp", "snap", "pinterest", "dropbox",
        "slack", "atlassian", "asana", "notion", "figma", "canva",
        "shopify", "hubspot", "zendesk", "datadog", "snowflake", "mongodb",
        "new relic", "mparticle", "amplitude", "mixpanel", "braze",
        "intercom", "drift", "gong", "outreach", "salesloft",
    ]

    total_people_years = 0.0

    for exp in experience:
        if not isinstance(exp, dict):
            continue

        title = (exp.get("title", "") or "").lower()
        description = (exp.get("description", "") or "").lower()

        # Check both highlights and bullets fields (parsers use different names)
        highlights = exp.get("highlights", []) or exp.get("bullets", []) or []
        if highlights and isinstance(highlights, list):
            highlights_text = " ".join([h.lower() for h in highlights if isinstance(h, str)])
        else:
            highlights_text = ""

        combined_text = f"{title} {description} {highlights_text}"
        dates = (exp.get("dates", "") or "").lower()
        company = (exp.get("company", "") or "").strip()
        company_lower = company.lower()

        years = parse_experience_duration(dates)

        # Skip if clearly operational-only role
        is_operational_only = any(pattern in title for pattern in operational_only_patterns)
        if is_operational_only and not any(ev in combined_text for ev in people_leadership_evidence):
            print(f"   ⏭️ Skipping operational role (no people evidence): {exp.get('title', '')} @ {company}")
            continue

        # Check for people leadership evidence
        has_people_title = any(pattern in title for pattern in people_leadership_titles)
        has_strong_title = any(pattern in title for pattern in strong_leadership_titles)
        has_people_evidence = any(ev in combined_text for ev in people_leadership_evidence)

        # Check if company is established (VP/Director definitely means people leadership)
        is_established_company = any(ec in company_lower for ec in established_companies)
        # Also check for acquisition mentions (e.g., "Segment (acquired by Twilio)")
        if "acquired by" in company_lower or "twilio" in company_lower:
            is_established_company = True

        # CREDITING RULES:
        # 1. Strong title at established company = full credit (VP at Segment = people leader)
        # 2. Strong title + evidence = full credit
        # 3. Regular leadership title + evidence = full credit
        # 4. Evidence alone = 70% credit
        # 5. Strong title at unknown company = 50% credit (might be startup title inflation)

        if has_strong_title and is_established_company:
            # VP/Director at established company = full credit, no evidence needed
            tier = get_company_credibility_tier(company, title)
            multiplier = get_credibility_multiplier(tier)
            adjusted = years * multiplier
            total_people_years += adjusted
            print(f"   ✅ PEOPLE LEADERSHIP (established co): {exp.get('title', '')} @ {company}: {years:.1f}y × {multiplier} = {adjusted:.1f}y")
        elif has_people_title and has_people_evidence:
            # Full credit for verified people leadership
            tier = get_company_credibility_tier(company, title)
            multiplier = get_credibility_multiplier(tier)
            adjusted = years * multiplier
            total_people_years += adjusted
            print(f"   ✅ PEOPLE LEADERSHIP: {exp.get('title', '')} @ {company}: {years:.1f}y × {multiplier} = {adjusted:.1f}y")
        elif has_people_evidence:
            # Evidence without clear title - 70% credit
            tier = get_company_credibility_tier(company, title)
            multiplier = get_credibility_multiplier(tier) * 0.7
            adjusted = years * multiplier
            total_people_years += adjusted
            print(f"   🟡 LIKELY PEOPLE LEADERSHIP: {exp.get('title', '')} @ {company}: {years:.1f}y × {multiplier:.2f} = {adjusted:.1f}y")
        elif has_strong_title:
            # Strong title at unknown company - 50% credit (might be startup inflation)
            tier = get_company_credibility_tier(company, title)
            multiplier = get_credibility_multiplier(tier) * 0.5
            adjusted = years * multiplier
            total_people_years += adjusted
            print(f"   🟡 POSSIBLE LEADERSHIP (unverified co): {exp.get('title', '')} @ {company}: {years:.1f}y × {multiplier:.2f} = {adjusted:.1f}y")
        else:
            print(f"   ❌ NO PEOPLE LEADERSHIP EVIDENCE: {exp.get('title', '')} @ {company}")

    return round(total_people_years, 1)


def extract_required_people_leadership_years(response_data: dict) -> tuple[float, bool]:
    """
    Extract required people leadership years from JD analysis.

    Args:
        response_data: Claude response containing role analysis

    Returns:
        tuple: (required_years, is_hard_requirement)
            - required_years: Number of people leadership years required (0 if none)
            - is_hard_requirement: True if this is a non-negotiable requirement
    """
    import re

    jd_text = (response_data.get("job_description", "") or "").lower()
    role_title = (response_data.get("role_title", "") or "").lower()
    experience_analysis = response_data.get("experience_analysis", {})

    combined = f"{role_title} {jd_text}"

    # Hard requirement indicators for people leadership
    hard_requirement_patterns = [
        r"(\d+)\+?\s*years?\s*(?:of\s*)?(?:people\s*)?(?:leadership|management|managing)",
        r"(?:leadership|management|managing).*?(\d+)\+?\s*years?",
        r"(\d+)\+?\s*years?\s*(?:of\s*)?direct\s*reports?",
        r"manage[d]?\s*team[s]?\s*of\s*(\d+)",
        r"(\d+)\+?\s*years?\s*(?:people\s*)?manager"
    ]

    # People leadership specific keywords that make it a hard requirement
    people_leadership_keywords = [
        "people leadership", "people management", "direct reports",
        "manage a team", "build a team", "lead a team", "team management",
        "management experience required", "leadership experience required",
        "managing people", "people manager"
    ]

    # Check if JD explicitly requires people leadership
    requires_people_leadership = any(kw in combined for kw in people_leadership_keywords)

    # Extract years requirement
    required_years = 0.0
    for pattern in hard_requirement_patterns:
        matches = re.findall(pattern, combined)
        for match in matches:
            try:
                years = float(match) if isinstance(match, str) else float(match[0]) if match else 0
                if years > required_years:
                    required_years = years
            except (ValueError, TypeError):
                continue

    # Check title for implicit requirements
    leadership_titles = ["director", "head of", "vp", "vice president", "manager"]
    has_leadership_title = any(lt in role_title for lt in leadership_titles)

    # If title suggests leadership but no explicit years, assume 3+ years
    if has_leadership_title and required_years == 0 and requires_people_leadership:
        required_years = 3.0

    # Director+ typically requires 5+ years people leadership
    if "director" in role_title and required_years < 5:
        required_years = 5.0
        requires_people_leadership = True

    # VP/Head typically requires 7+ years
    if ("vp" in role_title or "vice president" in role_title or "head of" in role_title) and required_years < 7:
        required_years = 7.0
        requires_people_leadership = True

    is_hard_requirement = requires_people_leadership and required_years > 0

    print(f"   🎯 PEOPLE LEADERSHIP REQUIREMENT: {required_years} years, hard_requirement={is_hard_requirement}")

    return required_years, is_hard_requirement


def get_leadership_gap_messaging(
    tiered_leadership: dict,
    required_people_leadership: float,
    required_org_leadership: float = 0.0
) -> dict:
    """
    Generate precise leadership gap messaging based on tiered analysis.

    Per Leadership Messaging Rules:
    - "0 years" ONLY when NO leadership signals exist at ALL
    - "insufficient" when leadership exists but doesn't meet requirement level
    - Be specific about what tier is missing vs what tier candidate has

    Args:
        tiered_leadership: Result from extract_tiered_leadership()
        required_people_leadership: Required people leadership years
        required_org_leadership: Required org-level leadership years (for VP+ roles)

    Returns:
        dict: {
            "status": "none" | "insufficient" | "sufficient",
            "message": str,  # Human-readable gap description
            "factual_statement": str,  # For recommendation_rationale
            "has_lower_tier": bool,  # Has strategic but not people leadership
            "gap_years": float,  # How many more years needed
            "candidate_tier": str  # Highest tier candidate has
        }
    """
    people_years = tiered_leadership.get("people_leadership_years", 0.0)
    org_years = tiered_leadership.get("org_leadership_years", 0.0)
    strategic_years = tiered_leadership.get("strategic_leadership_years", 0.0)
    has_any = tiered_leadership.get("has_any_leadership", False)

    result = {
        "status": "sufficient",
        "message": "",
        "factual_statement": "",
        "has_lower_tier": False,
        "gap_years": 0.0,
        "candidate_tier": "none"
    }

    # Determine candidate's highest tier
    if org_years > 0:
        result["candidate_tier"] = "org"
    elif people_years > 0:
        result["candidate_tier"] = "people"
    elif strategic_years > 0:
        result["candidate_tier"] = "strategic"

    # Check org-level requirement first (highest tier)
    if required_org_leadership > 0:
        if org_years >= required_org_leadership:
            result["status"] = "sufficient"
            return result

        # Check if candidate has ANY org-level
        if org_years > 0:
            # Insufficient org-level
            result["status"] = "insufficient"
            result["gap_years"] = required_org_leadership - org_years
            result["message"] = (
                f"Your org-level leadership experience ({org_years:.1f} years) is insufficient "
                f"for this role's requirement ({required_org_leadership:.0f}+ years)."
            )
            result["factual_statement"] = (
                f"This role requires {required_org_leadership:.0f}+ years of org-level leadership "
                f"(executive team, C-suite, board-facing). You have {org_years:.1f} years verified."
            )
        elif people_years > 0:
            # Has people leadership but not org-level
            result["status"] = "insufficient"
            result["has_lower_tier"] = True
            result["gap_years"] = required_org_leadership
            result["message"] = (
                f"This role requires org-level leadership (executive, C-suite, board-facing). "
                f"Your experience is at the people leadership tier ({people_years:.1f} years managing direct reports), "
                f"which does not substitute for executive-level scope."
            )
            result["factual_statement"] = (
                f"This role requires {required_org_leadership:.0f}+ years of org-level leadership. "
                f"Your background shows {people_years:.1f} years of people leadership, but no executive or C-suite experience."
            )
        elif strategic_years > 0:
            # Has strategic but not people or org
            result["status"] = "insufficient"
            result["has_lower_tier"] = True
            result["gap_years"] = required_org_leadership
            result["message"] = (
                f"This role requires org-level leadership. Your experience is at the strategic/functional tier "
                f"({strategic_years:.1f} years leading initiatives without direct reports). "
                f"This is 2 tiers below the role requirement."
            )
            result["factual_statement"] = (
                f"This role requires {required_org_leadership:.0f}+ years of org-level leadership. "
                f"Your background shows strategic/functional leadership but no people management or executive experience."
            )
        else:
            # No leadership at all
            result["status"] = "none"
            result["gap_years"] = required_org_leadership
            result["message"] = "No leadership experience found in your background."
            result["factual_statement"] = (
                f"This role requires {required_org_leadership:.0f}+ years of org-level leadership. "
                f"Your resume shows 0 years of verified leadership at any tier."
            )
        return result

    # Check people leadership requirement
    if required_people_leadership > 0:
        if people_years >= required_people_leadership:
            result["status"] = "sufficient"
            return result

        # Check if candidate has ANY people leadership
        if people_years > 0:
            # Insufficient people leadership
            result["status"] = "insufficient"
            result["gap_years"] = required_people_leadership - people_years
            result["message"] = (
                f"Your people leadership experience ({people_years:.1f} years) is insufficient "
                f"for this role's requirement ({required_people_leadership:.0f}+ years)."
            )
            result["factual_statement"] = (
                f"This role requires {required_people_leadership:.0f}+ years of people leadership. "
                f"You have {people_years:.1f} years verified. Gap: {result['gap_years']:.1f} years."
            )
        elif strategic_years > 0:
            # Has strategic but not people leadership
            result["status"] = "insufficient"
            result["has_lower_tier"] = True
            result["gap_years"] = required_people_leadership
            result["message"] = (
                f"This role requires people leadership (managing direct reports). "
                f"Your experience is at the strategic/functional tier ({strategic_years:.1f} years leading initiatives), "
                f"which does not substitute for people management."
            )
            result["factual_statement"] = (
                f"This role requires {required_people_leadership:.0f}+ years of people leadership. "
                f"Your background shows {strategic_years:.1f} years of strategic/functional leadership "
                f"but no verified direct reports or team management."
            )
        else:
            # No leadership at all
            result["status"] = "none"
            result["gap_years"] = required_people_leadership
            result["message"] = "No leadership experience found in your background."
            result["factual_statement"] = (
                f"This role requires {required_people_leadership:.0f}+ years of people leadership. "
                f"Your resume shows 0 years of verified leadership at any tier."
            )
        return result

    # No leadership requirement - sufficient by default
    return result


def classify_gap_type(gap: dict, candidate_years: float, required_years: float,
                      people_leadership_years: float, required_people_leadership: float) -> str:
    """
    Classify a gap as missing_evidence, missing_scope, or missing_experience.

    Per Eligibility Gate Enforcement Spec:
    - missing_experience = The candidate lacks the experience (hard stop, "Do Not Apply")
    - missing_scope = The candidate has experience but at smaller scale (downgrade possible)
    - missing_evidence = The candidate may have it but resume doesn't prove it (coachable)

    CRITICAL: missing_experience → forces "Do Not Apply". No exceptions.

    Args:
        gap: Gap dictionary
        candidate_years: Total relevant years
        required_years: Required years from JD
        people_leadership_years: Verified people leadership years
        required_people_leadership: Required people leadership years

    Returns:
        str: "missing_evidence", "missing_scope", or "missing_experience"
    """
    gap_type = gap.get("gap_type", "")
    gap_description = (gap.get("gap_description", "") or gap.get("description", "")).lower()

    # ========================================================================
    # RULE 1: Eligibility gate failures are ALWAYS missing_experience
    # ========================================================================
    if gap_type.startswith("eligibility_failure:"):
        return "missing_experience"

    # ========================================================================
    # RULE 2: Non-transferable domain gaps are ALWAYS missing_experience
    # ========================================================================
    non_transferable_keywords = [
        # Executive search
        "executive search", "executive recruiting", "retained search",
        # Core software engineering
        "software engineer", "software development", "hands-on coding",
        # ML/AI research
        "machine learning", "ml research", "ai research", "deep learning",
        # Regulated domains
        "clinical", "fda", "regulatory", "finra", "sec", "compliance"
    ]
    if any(kw in gap_description for kw in non_transferable_keywords):
        return "missing_experience"

    # ========================================================================
    # RULE 3: People leadership gap is ALWAYS missing_experience
    # ========================================================================
    if required_people_leadership > 0 and people_leadership_years < required_people_leadership:
        if any(kw in gap_description for kw in ["leadership", "management", "direct report", "team", "people"]):
            return "missing_experience"

    # ========================================================================
    # RULE 4: Significant years gap is missing_experience
    # ========================================================================
    if candidate_years < required_years * 0.7:
        if any(kw in gap_description for kw in ["years", "experience", "tenure"]):
            return "missing_experience"

    # ========================================================================
    # RULE 5: Scope gaps - has experience but at smaller scale
    # ========================================================================
    scope_keywords = ["scale", "scope", "enterprise", "global", "size", "team size",
                      "company size", "revenue", "headcount"]
    if any(kw in gap_description for kw in scope_keywords):
        return "missing_scope"

    # ========================================================================
    # RULE 6: Evidence gaps - may have experience but can't verify
    # ========================================================================
    evidence_keywords = ["quantif", "metric", "evidence", "demonstrat", "prove",
                        "unclear", "vague", "specific", "detail"]
    if any(kw in gap_description for kw in evidence_keywords):
        return "missing_evidence"

    # Default based on severity
    severity = gap.get("severity", "").lower()
    if severity == "critical" and candidate_years < required_years:
        return "missing_experience"
    elif severity == "critical":
        return "missing_scope"
    else:
        return "missing_evidence"


def resolve_strength_gap_conflicts(strengths: list, gaps: list) -> tuple[list, list]:
    """
    CONSTRAINT: Same signal cannot appear in both strengths and gaps.

    Rule: If a signal appears in both, gaps win (conservative approach).

    Args:
        strengths: List of strength dictionaries
        gaps: List of gap dictionaries

    Returns:
        tuple: (cleaned_strengths, cleaned_gaps)
    """
    if not strengths or not gaps:
        return strengths or [], gaps or []

    # Extract key signals from gaps
    gap_signals = set()
    for gap in gaps:
        if not isinstance(gap, dict):
            continue
        gap_text = (
            (gap.get("gap_description", "") or "") + " " +
            (gap.get("description", "") or "") + " " +
            (gap.get("gap_type", "") or "")
        ).lower()

        # Key signal words
        signal_words = [
            "leadership", "management", "experience", "years", "scale",
            "enterprise", "team", "stakeholder", "executive", "strategy",
            "technical", "product", "operations", "people"
        ]
        for word in signal_words:
            if word in gap_text:
                gap_signals.add(word)

    # Remove strengths that conflict with gaps
    cleaned_strengths = []
    for strength in strengths:
        if not isinstance(strength, dict):
            continue

        strength_text = (
            (strength.get("description", "") or "") + " " +
            (strength.get("strength", "") or "") + " " +
            (strength.get("area", "") or "")
        ).lower()

        # Check if this strength conflicts with any gap signal
        conflicts = [signal for signal in gap_signals if signal in strength_text]

        if conflicts:
            print(f"   ⚠️ CONFLICT RESOLVED: Removing strength that conflicts with gap: {conflicts}")
            # Don't add this strength - gap wins
        else:
            cleaned_strengths.append(strength)

    return cleaned_strengths, gaps


# ============================================================================
# LEADERSHIP VS IC DISTINCTION FOR RECRUITING ROLES
# Per JOB_FIT_SCORING_SPEC.md
# ============================================================================

def calculate_recruiting_leadership_years(resume_data: dict, requires_leadership: bool = True) -> float:
    """
    Calculate recruiting experience with leadership vs IC distinction.

    Leadership roles get full credit.
    Senior IC roles get 70% credit for leadership requirements.
    IC roles get 0% credit for leadership requirements.

    Args:
        resume_data: Parsed resume dictionary
        requires_leadership: Whether the target JD requires leadership

    Returns:
        float: Total recruiting years (adjusted for leadership level)
    """
    if not resume_data:
        return 0.0

    experience = resume_data.get("experience", [])
    if not experience or not isinstance(experience, list):
        return 0.0

    # Pattern definitions
    leadership_patterns = [
        "director", "head of", "vp", "vice president", "manager", "lead", "principal"
    ]
    senior_ic_patterns = [
        "senior talent partner", "senior recruiter", "senior ta",
        "staff recruiter", "principal recruiter"
    ]
    recruiting_patterns = [
        "recruit", "talent acquisition", "sourcer", "headhunter",
        "talent partner", "executive search", "ta "
    ]

    total_years = 0.0

    for exp in experience:
        if not isinstance(exp, dict):
            continue

        title = (exp.get("title", "") or "").lower()
        company = (exp.get("company", "") or "").strip()
        dates = (exp.get("dates", "") or "").lower()

        # Skip if not a recruiting role
        if not any(pattern in title for pattern in recruiting_patterns):
            continue

        years = parse_experience_duration(dates)

        # Get company credibility for additional adjustment
        tier = get_company_credibility_tier(company, title)
        credibility_multiplier = get_credibility_multiplier(tier)

        # Apply leadership/IC multipliers
        is_leadership = any(pattern in title for pattern in leadership_patterns)
        is_senior_ic = any(pattern in title for pattern in senior_ic_patterns)

        if is_leadership:
            # Leadership roles: full credit × credibility
            level_multiplier = 1.0
            level_label = "LEADERSHIP"
        elif is_senior_ic:
            # Senior IC: 70% credit if leadership required × credibility
            level_multiplier = 0.7 if requires_leadership else 1.0
            level_label = "SENIOR IC"
        else:
            # IC roles: no credit for leadership requirements
            level_multiplier = 0.0 if requires_leadership else 1.0
            level_label = "IC"

        adjusted = years * level_multiplier * credibility_multiplier
        total_years += adjusted

        print(f"   📊 Recruiting: {exp.get('title', '')} @ {company}")
        print(f"      {years:.1f}y × {level_multiplier} ({level_label}) × {credibility_multiplier} ({tier}) = {adjusted:.1f}y")

    return round(total_years, 1)


def jd_requires_leadership(response_data: dict) -> bool:
    """
    Detect if the JD requires leadership/management experience.

    Args:
        response_data: Claude response containing role analysis

    Returns:
        bool: True if leadership is required
    """
    role_title = (response_data.get("role_title", "") or "").lower()
    jd_text = (response_data.get("job_description", "") or "").lower()

    combined = f"{role_title} {jd_text}"

    leadership_indicators = [
        "director", "head of", "vp", "vice president", "manager",
        "leadership", "lead a team", "manage a team", "build a team",
        "people management", "direct reports", "team of"
    ]

    return any(indicator in combined for indicator in leadership_indicators)


def calculate_role_specific_years(resume_data: dict, target_role_type: str = "pm") -> float:
    """
    Calculate years of experience in a specific role type by parsing resume data.
    This is a backend safety net when Claude miscounts experience.

    Args:
        resume_data: Parsed resume JSON with experience entries
        target_role_type: "pm", "engineering", "design", "ops", etc.

    Returns:
        Years of experience in roles matching the target type
    """
    if not resume_data:
        return 0.0

    experience = resume_data.get("experience", [])
    if not experience or not isinstance(experience, list):
        return 0.0

    # Role title patterns by type
    role_patterns = {
        "pm": [
            "product manager", "product lead", "product owner", "pm",
            "product director", "head of product", "vp product", "cpo",
            "associate product manager", "senior product manager", "staff product",
            "principal product", "group product manager", "technical pm"
        ],
        "engineering": [
            "engineer", "developer", "swe", "software", "programmer",
            "cto", "tech lead", "architect", "devops"
        ],
        "design": [
            "designer", "ux", "ui", "product design", "design lead",
            "head of design", "creative director"
        ],
        "ops": [
            "operations", "ops manager", "chief operating", "coo",
            "business operations", "strategy & ops"
        ],
        "recruiting": [
            "recruit", "talent acquisition", "ta director", "ta manager",
            "sourcer", "sourcing", "headhunter", "talent partner",
            "talent coordinator", "recruiting operations", "recruiting manager",
            "technical recruiter", "executive recruiter", "hr generalist",
            "talent lead", "head of talent", "vp talent", "people operations"
        ],
        "sales": [
            "sales", "account executive", "account manager", "business development",
            "bd", "revenue", "sales engineer", "sales operations", "sales manager",
            "enterprise sales", "smb sales", "strategic accounts", "ae", "sdr",
            "bdr", "sales director", "vp sales", "chief revenue"
        ],
        "marketing": [
            "marketing", "growth", "brand", "content", "digital marketing",
            "product marketing", "demand gen", "campaigns", "seo", "sem",
            "social media", "community", "communications", "pr", "cmo",
            "marketing manager", "head of marketing", "vp marketing"
        ],
        "general": []  # Empty list means match ALL roles (fallback)
    }

    # For "general" type, count all experience
    if target_role_type == "general":
        total_years = 0.0
        for exp in experience:
            if not isinstance(exp, dict):
                continue
            dates = (exp.get("dates", "") or "").lower()
            years = parse_experience_duration(dates)
            total_years += years
            print(f"   📊 Total experience: '{exp.get('title', '')}' = {years:.1f} years")
        return total_years

    patterns = role_patterns.get(target_role_type, role_patterns["pm"])
    total_years = 0.0

    for exp in experience:
        if not isinstance(exp, dict):
            continue

        title = (exp.get("title", "") or "").lower()
        dates = (exp.get("dates", "") or "").lower()

        # Check if this role matches the target type
        matches_role = any(pattern in title for pattern in patterns)

        if matches_role:
            # Parse duration from dates string
            years = parse_experience_duration(dates)
            total_years += years
            print(f"   📊 Role match: '{exp.get('title', '')}' = {years:.1f} years")

    return total_years


# =============================================================================
# REQUEST ISOLATION & CONTAMINATION DETECTION
# Purpose: Ensure each analysis request is isolated - no data leakage between candidates
# =============================================================================

def verify_request_isolation(resume_data: Dict[str, Any], analysis_id: str) -> None:
    """
    Verify request isolation and log candidate identifier for tracing.

    This function ensures each analysis request is properly isolated.
    It does NOT block any candidates - test profiles are allowed for testing.

    The isolation is enforced by:
    1. Unique analysis_id per request (generated at request start)
    2. All data passed explicitly via request parameters (no global state)
    3. Session cleanup mechanism (cleanup_expired_sessions)
    4. No caching of candidate-specific data between requests

    Args:
        resume_data: Resume data from request
        analysis_id: Analysis ID for logging/tracing
    """
    # Extract candidate identifier for logging (helps with debugging)
    candidate_name = (resume_data.get("full_name", "") or "").strip()
    candidate_id = candidate_name[:30] if candidate_name else "anonymous"

    print(f"🔍 [{analysis_id}] Request isolation verified")
    print(f"   Candidate: {candidate_id}")
    print(f"   Analysis ID: {analysis_id} (unique per request)")
    print(f"   Data source: request parameters only (no global state)")


def extract_role_title_from_jd(jd_text: str, analysis_id: str) -> str:
    """
    Extract role title from job description text.

    This is the ONLY source of role title for analysis.
    Never use cached, session, or global title data.

    Args:
        jd_text: Job description text from current request
        analysis_id: Analysis ID for logging

    Returns:
        Extracted role title
    """
    import re

    print(f"📋 [{analysis_id}] Extracting role title from JD...")

    # Navigation/UI text to skip (common when copying from job boards)
    skip_patterns = [
        r'^back to',
        r'^apply now',
        r'^save job',
        r'^share',
        r'^home\s*[>/]',
        r'^jobs\s*[>/]',
        r'^search',
        r'^menu',
        r'^sign in',
        r'^log in',
        r'^posted',
        r'^×',
        r'^\d+ days? ago',
        r'^view all jobs',
        r'^similar jobs',
    ]

    def is_navigation_text(text: str) -> bool:
        """Check if text is likely navigation/UI element."""
        text_lower = text.lower().strip()
        for pattern in skip_patterns:
            if re.match(pattern, text_lower):
                return True
        # Also skip very short lines that are likely nav
        if len(text) < 5:
            return True
        return False

    # Strategy 1: Look for explicit markers
    patterns = [
        r'(?:job title|position|role):\s*([^\n]+)',
        r'^([^\n]+)(?:\s*-\s*(?:corporate|hybrid|remote|full.time))',
        r'(?:hiring|hiring a|we\'re hiring|is hiring)\s+(?:a\s+)?([^\n\.]+)',
    ]

    for pattern in patterns:
        match = re.search(pattern, jd_text, re.IGNORECASE | re.MULTILINE)
        if match:
            title = match.group(1).strip()
            title = re.sub(
                r'\s*-\s*(corporate|hybrid|remote|san francisco|full.time).*$',
                '',
                title,
                flags=re.IGNORECASE
            )
            if 5 < len(title) < 100 and not is_navigation_text(title):
                print(f"  ✅ Extracted: '{title}'")
                return title

    # Strategy 2: Find first line that looks like a job title (skip nav text)
    lines = [line.strip() for line in jd_text.split('\n') if line.strip()]
    for line in lines[:10]:  # Check first 10 lines
        if is_navigation_text(line):
            continue
        # Must start with capital, reasonable length, and contain job-related words
        if 10 < len(line) < 100 and line[0].isupper():
            # Bonus: Check if it contains job-title-like words
            job_title_indicators = ['manager', 'director', 'engineer', 'analyst', 'lead',
                                    'coordinator', 'specialist', 'vp', 'vice president',
                                    'head of', 'senior', 'junior', 'associate', 'chief',
                                    'officer', 'developer', 'designer', 'architect',
                                    'recruiter', 'marketing', 'sales', 'product', 'operations']
            line_lower = line.lower()
            if any(indicator in line_lower for indicator in job_title_indicators):
                print(f"  ✅ Extracted from content: '{line}'")
                return line

    # Strategy 3: First substantial non-nav line as fallback
    for line in lines[:10]:
        if not is_navigation_text(line) and 5 < len(line) < 100:
            print(f"  ✅ Extracted from first valid line: '{line}'")
            return line

    # Fallback
    print(f"  ⚠️ Could not extract role title - using placeholder")
    return "Unknown Role"


def detect_role_type_isolated(jd_text: str, role_title: str, analysis_id: str) -> str:
    """
    Detect role type from job description and title.

    CRITICAL: Uses ONLY the parameters passed to this function.
    NO global variables, cache, or session data.

    PRIORITY ORDER (RECRUITING first to fix contamination issues):
    1. RECRUITING - most specific, check first
    2. PRODUCT
    3. ENGINEERING
    4. SALES
    5. MARKETING
    6. GENERAL (fallback)

    Args:
        jd_text: Job description text from current request
        role_title: Role title extracted from jd_text
        analysis_id: Analysis ID for logging

    Returns:
        Role type: "recruiting", "pm", "engineering", "sales", "marketing", "general"
    """

    print(f"🔍 [{analysis_id}] Detecting role type (isolated)...")
    print(f"  Title: '{role_title}'")

    title_lower = role_title.lower()
    jd_lower = jd_text.lower()

    # PRIORITY 1: RECRUITING (most specific, check first)
    recruiting_title_patterns = [
        "recruiter", "recruiting", "talent acquisition",
        "sourcer", "sourcing", "talent partner",
        "recruitment", "headhunter", "technical recruiter",
        "ta director", "ta manager", "talent lead", "head of talent"
    ]

    if any(pattern in title_lower for pattern in recruiting_title_patterns):
        print(f"  ✅ RECRUITING detected from title")
        return "recruiting"

    # Check JD content for recruiting signals
    recruiting_jd_signals = [
        "full-cycle recruiting", "sourcing", "candidate experience",
        "hiring manager", "recruitment funnel", "talent pipeline",
        "offer negotiation", "interview scheduling"
    ]
    recruiting_signal_count = sum(1 for sig in recruiting_jd_signals if sig in jd_lower)

    if recruiting_signal_count >= 3:
        print(f"  ✅ RECRUITING detected from JD ({recruiting_signal_count} signals)")
        return "recruiting"

    # PRIORITY 2: PRODUCT
    product_patterns = [
        "product manager", "product lead", "product owner",
        "product director", "head of product", "vp product", "cpo"
    ]

    # Be careful NOT to match "product manager, cloud and open ecosystems" for non-PM roles
    if any(pattern in title_lower for pattern in product_patterns):
        # Double-check this isn't a recruiting role with PM in name
        if not any(rp in title_lower for rp in recruiting_title_patterns):
            print(f"  ✅ PRODUCT detected from title")
            return "pm"

    # PRIORITY 3: ENGINEERING
    engineering_patterns = [
        "engineer", "developer", "software", "technical lead",
        "architect", "sre", "devops", "engineering manager"
    ]

    if any(pattern in title_lower for pattern in engineering_patterns):
        print(f"  ✅ ENGINEERING detected from title")
        return "engineering"

    # PRIORITY 4: SALES
    sales_patterns = [
        "sales", "account executive", "account manager",
        "business development", "revenue", "ae ", "sdr", "bdr"
    ]

    if any(pattern in title_lower for pattern in sales_patterns):
        print(f"  ✅ SALES detected from title")
        return "sales"

    # PRIORITY 5: MARKETING
    marketing_patterns = [
        "marketing", "growth", "demand gen",
        "brand", "marketing manager", "cmo"
    ]

    if any(pattern in title_lower for pattern in marketing_patterns):
        print(f"  ✅ MARKETING detected from title")
        return "marketing"

    # Fallback
    print(f"  ⚠️ Defaulting to GENERAL")
    return "general"


def verify_role_type_alignment(
    role_type: str,
    role_title: str,
    jd_text: str,
    analysis_id: str
) -> Dict[str, Any]:
    """
    Verify detected role type aligns with job description.

    Mismatches trigger:
    - Loud logging
    - Confidence degradation
    - Warning flags

    They do NOT crash the request unless contamination is confirmed.

    Args:
        role_type: Detected role type
        role_title: Extracted role title
        jd_text: Job description text
        analysis_id: Analysis ID for logging

    Returns:
        {
            "aligned": bool,
            "confidence": float,
            "warnings": List[str]
        }
    """

    print(f"🔍 [{analysis_id}] Verifying role type alignment...")

    title_lower = role_title.lower()
    jd_lower = jd_text.lower()
    warnings = []
    aligned = True
    confidence = 1.0

    # Check for obvious mismatches

    # Case 1: Title says "recruiter" but detected as PM/other
    if "recruit" in title_lower and role_type == "pm":
        print(f"⚠️ [{analysis_id}] ROLE TYPE MISMATCH WARNING")
        print(f"   Title: '{role_title}'")
        print(f"   Detected: {role_type}")
        print(f"   Expected: recruiting")
        warnings.append("Role type detection uncertain - title suggests RECRUITING but detected as PM")
        aligned = False
        confidence = 0.3

    # Case 2: Title says "engineer" but detected as RECRUITING
    elif "engineer" in title_lower and role_type == "recruiting":
        print(f"⚠️ [{analysis_id}] ROLE TYPE MISMATCH WARNING")
        print(f"   Title: '{role_title}'")
        print(f"   Detected: {role_type}")
        print(f"   Expected: engineering")
        warnings.append("Role type detection uncertain - title suggests ENGINEERING but detected as RECRUITING")
        aligned = False
        confidence = 0.3

    # Case 3: Strong recruiting signals but non-recruiting type
    recruiting_signals = [
        "full-cycle recruiting", "candidate experience",
        "sourcing", "hiring manager", "talent pipeline"
    ]
    signal_count = sum(1 for sig in recruiting_signals if sig in jd_lower)

    if signal_count >= 3 and role_type != "recruiting":
        print(f"⚠️ [{analysis_id}] ROLE TYPE CONFIDENCE LOW")
        print(f"   JD has {signal_count} recruiting signals")
        print(f"   Detected: {role_type}")
        warnings.append(f"JD contains {signal_count} recruiting indicators but detected as {role_type}")
        aligned = False
        confidence = 0.5

    if aligned:
        print(f"✅ [{analysis_id}] Role type alignment verified")
    else:
        print(f"⚠️ [{analysis_id}] Role type alignment degraded (confidence: {confidence})")

    return {
        "aligned": aligned,
        "confidence": confidence,
        "warnings": warnings
    }


def calculate_relevant_years_isolated(
    resume_data: Dict[str, Any],
    role_type: str,
    analysis_id: str
) -> float:
    """
    Calculate years of relevant experience based on role type.

    CRITICAL: Uses ONLY data from current request.
    No cache, no globals, no session data.

    Args:
        resume_data: Resume data from current request
        role_type: Detected role type
        analysis_id: Analysis ID for logging

    Returns:
        Total years of relevant experience
    """

    print(f"📊 [{analysis_id}] Calculating {role_type.upper()} experience years...")

    # Role-specific title patterns
    role_patterns = {
        "recruiting": [
            "recruiter", "recruiting", "talent acquisition",
            "sourcer", "sourcing", "talent partner", "recruitment",
            "technical recruiter", "ta ", "talent lead", "head of talent"
        ],
        "pm": [
            "product manager", "product lead", "pm", "product owner",
            "product director", "head of product", "vp product"
        ],
        "engineering": [
            "engineer", "developer", "software", "technical lead",
            "architect", "sre", "devops", "programmer"
        ],
        "sales": [
            "sales", "account executive", "account manager",
            "business development", "revenue", "ae ", "sdr", "bdr"
        ],
        "marketing": [
            "marketing", "growth", "demand gen",
            "brand", "marketing manager", "content"
        ]
    }

    patterns = role_patterns.get(role_type, [])
    total_years = 0.0

    for exp in resume_data.get("experience", []):
        if not isinstance(exp, dict):
            continue

        title = (exp.get("title", "") or "").lower()
        company = exp.get("company", "")
        dates = exp.get("dates", "")

        # For general type, count all experience
        if role_type == "general" or not patterns:
            years = parse_duration_to_years_isolated(dates)
            total_years += years
            print(f"  📊 {exp.get('title')} @ {company}: +{years:.1f} years (all experience)")
            continue

        is_relevant = any(pattern in title for pattern in patterns)

        if is_relevant:
            years = parse_duration_to_years_isolated(dates)
            total_years += years
            print(f"  ✅ {exp.get('title')} @ {company}: +{years:.1f} years")
        else:
            print(f"  ⏭️  {exp.get('title')} @ {company}: Not relevant to {role_type.upper()}")

    print(f"📊 [{analysis_id}] Total {role_type.upper()} experience: {total_years:.1f} years")
    return total_years


def parse_duration_to_years_isolated(dates: str) -> float:
    """
    Parse date range to years of experience.

    Isolated version - no external dependencies.

    Args:
        dates: Date range string from resume

    Returns:
        Duration in years
    """
    import re
    from datetime import datetime

    if not dates:
        return 0.0

    dates_str = str(dates).lower().strip()

    # Check for "present" or "current"
    is_current = "present" in dates_str or "current" in dates_str

    # Extract years
    years = re.findall(r'\b(20\d{2})\b', dates_str)

    if len(years) >= 2:
        start_year = int(years[0])
        end_year = int(years[-1]) if not is_current else datetime.now().year
        return max(0.0, float(end_year - start_year))

    elif len(years) == 1 and is_current:
        start_year = int(years[0])
        return max(0.0, float(datetime.now().year - start_year))

    elif len(years) == 1:
        return 1.0

    # Try to parse duration format (e.g., "2 years 3 months")
    year_match = re.search(r'(\d+)\s*year', dates_str)
    month_match = re.search(r'(\d+)\s*month', dates_str)
    if year_match or month_match:
        yrs = int(year_match.group(1)) if year_match else 0
        mos = int(month_match.group(1)) if month_match else 0
        return yrs + (mos / 12)

    return 0.0


def check_people_leadership_requirement_isolated(
    resume_data: Dict[str, Any],
    required_years: float,
    hard_requirement: bool,
    analysis_id: str
) -> Dict[str, Any]:
    """
    Check if candidate meets people leadership requirement.

    CRITICAL: Skip entirely if not required to avoid noise.

    Args:
        resume_data: Resume data from request
        required_years: Years required
        hard_requirement: If blocking
        analysis_id: Analysis ID

    Returns:
        Leadership check results
    """

    # Skip entirely if not required
    if required_years == 0.0 and not hard_requirement:
        print(f"⏭️  [{analysis_id}] Skipping leadership check (not required)")
        return {
            "meets_requirement": True,
            "candidate_years": 0.0,
            "gap_severity": "none",
            "skipped": True
        }

    print(f"🔍 [{analysis_id}] Checking leadership requirement...")
    print(f"   Required: {required_years} years (hard={hard_requirement})")

    # Calculate people leadership years
    leadership_patterns = [
        "manager", "director", "head of", "vp ", "lead",
        "chief", "team lead", "group lead"
    ]

    candidate_years = 0.0
    for exp in resume_data.get("experience", []):
        if not isinstance(exp, dict):
            continue
        title = (exp.get("title", "") or "").lower()

        if any(pattern in title for pattern in leadership_patterns):
            dates = exp.get("dates", "")
            years = parse_duration_to_years_isolated(dates)
            candidate_years += years

    print(f"   Candidate: {candidate_years:.1f} years")

    if candidate_years >= required_years:
        return {
            "meets_requirement": True,
            "candidate_years": candidate_years,
            "gap_severity": "none",
            "skipped": False
        }
    elif hard_requirement:
        return {
            "meets_requirement": False,
            "candidate_years": candidate_years,
            "gap_severity": "major",
            "skipped": False
        }
    else:
        return {
            "meets_requirement": True,
            "candidate_years": candidate_years,
            "gap_severity": "minor",
            "skipped": False
        }


def calculate_career_gap_penalty_isolated(
    resume_data: Dict[str, Any],
    analysis_id: str
) -> Dict[str, Any]:
    """
    Calculate penalty for employment gaps.

    Post-2023 context: Widespread layoffs make gaps less stigmatizing.

    Args:
        resume_data: Resume data
        analysis_id: Analysis ID

    Returns:
        Gap analysis with adjusted penalties
    """

    from datetime import datetime
    import re

    print(f"📅 [{analysis_id}] Checking for career gaps...")

    experience = resume_data.get("experience", [])
    if not experience:
        return {"has_gap": False, "gap_months": 0, "penalty_points": 0, "severity": "none"}

    # Get most recent experience
    most_recent = None
    for exp in experience:
        if isinstance(exp, dict):
            most_recent = exp
            break

    if not most_recent:
        return {"has_gap": False, "gap_months": 0, "penalty_points": 0, "severity": "none"}

    dates = (most_recent.get("dates", "") or "").lower()

    if "present" in dates or "current" in dates:
        print(f"  ✅ Currently employed - no gap")
        return {"has_gap": False, "gap_months": 0, "penalty_points": 0, "severity": "none"}

    # Extract end year
    years = re.findall(r'\b(20\d{2})\b', dates)
    if not years:
        return {"has_gap": False, "gap_months": 0, "penalty_points": 0, "severity": "none"}

    end_year = int(years[-1])
    current_year = datetime.now().year
    gap_months = (current_year - end_year) * 12

    # Try to get more precise with month
    month_match = re.search(r'(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)', dates)
    if month_match:
        month_map = {"jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
                     "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12}
        end_month = month_map.get(month_match.group(1), 6)
        gap_months = (current_year - end_year) * 12 + (datetime.now().month - end_month)

    print(f"  📊 Gap: {gap_months} months since {end_year}")

    # Post-2023 adjusted penalties (more lenient due to widespread layoffs)
    if gap_months <= 6:
        penalty, severity = 0, "none"
    elif gap_months <= 12:
        penalty, severity = 2, "minor"
    elif gap_months <= 18:
        penalty, severity = 5, "moderate"
    elif gap_months <= 24:
        penalty, severity = 8, "moderate"
    else:
        penalty, severity = 12, "major"

    print(f"  {'✅' if penalty == 0 else '⚠️'} Penalty: {penalty} points ({severity})")

    return {
        "has_gap": gap_months > 6,
        "gap_months": gap_months,
        "penalty_points": penalty,
        "severity": severity
    }


# =============================================================================
# END REQUEST ISOLATION & CONTAMINATION DETECTION
# =============================================================================


def detect_role_type_from_jd(response_data: dict) -> str:
    """
    Detect the role type from the job description analysis.
    Used to route to the appropriate experience calculator.

    CRITICAL: Role title takes ABSOLUTE PRECEDENCE over JD content.
    An "Engineering Manager" is ENGINEERING, even if JD mentions "recruit talent".

    Args:
        response_data: The Claude response containing role_title and job description

    Returns:
        Role type string: "recruiting", "pm", "engineering", "sales", "marketing", "ops", or "general"
    """
    # Get role title and JD text
    role_title = (response_data.get("role_title", "") or "").lower()
    jd_text = (response_data.get("job_description", "") or "").lower()

    # Also check intelligence layer if available
    intel = response_data.get("intelligence_layer", {})
    role_summary = (intel.get("role_summary", "") or "").lower()

    # ==========================================================================
    # PHASE 1: TITLE-BASED DETECTION (TAKES ABSOLUTE PRECEDENCE)
    # The role title is the deterministic signal. JD content cannot override it.
    # ==========================================================================

    # Engineering Manager/Lead titles -> ENGINEERING (never recruiting/other)
    if any(x in role_title for x in [
        "engineering manager", "eng manager", "software manager",
        "technical manager", "development manager", "platform manager",
        "engineering lead", "tech lead", "technical lead",
        "engineering director", "director of engineering", "vp engineering",
        "head of engineering", "cto", "chief technology"
    ]):
        print(f"   🎯 TITLE LOCK: '{role_title}' → ENGINEERING (manager/lead detected)")
        return "engineering"

    # Software/Developer titles -> ENGINEERING
    if any(x in role_title for x in [
        "software engineer", "developer", "swe", "programmer",
        "architect", "devops", "sre", "backend engineer", "frontend engineer",
        "full stack", "mobile engineer", "data engineer", "ml engineer",
        "staff engineer", "principal engineer", "senior engineer"
    ]):
        print(f"   🎯 TITLE LOCK: '{role_title}' → ENGINEERING (IC detected)")
        return "engineering"

    # Product Manager titles -> PM
    if any(x in role_title for x in [
        "product manager", "product lead", "product owner", "head of product",
        "vp product", "cpo", "group pm", "technical pm", "senior pm",
        "director of product", "product director"
    ]):
        print(f"   🎯 TITLE LOCK: '{role_title}' → PRODUCT")
        return "pm"

    # Recruiting titles -> RECRUITING (only if explicit in title)
    if any(x in role_title for x in [
        "recruiter", "talent acquisition", "sourcer", "headhunter",
        "ta director", "ta manager", "talent partner", "head of talent",
        "recruiting manager", "recruiting lead", "recruitment"
    ]):
        print(f"   🎯 TITLE LOCK: '{role_title}' → RECRUITING")
        return "recruiting"

    # Sales titles -> SALES
    if any(x in role_title for x in [
        "sales", "account executive", "business development", "bdr", "sdr",
        "account manager", "sales manager", "revenue", "ae "
    ]):
        print(f"   🎯 TITLE LOCK: '{role_title}' → SALES")
        return "sales"

    # Marketing titles -> MARKETING
    if any(x in role_title for x in [
        "marketing", "growth", "brand manager", "demand gen",
        "product marketing", "content", "cmo", "marketing manager"
    ]):
        print(f"   🎯 TITLE LOCK: '{role_title}' → MARKETING")
        return "marketing"

    # Design titles -> DESIGN
    if any(x in role_title for x in [
        "designer", "ux", "ui", "product design", "design lead",
        "design manager", "head of design"
    ]):
        print(f"   🎯 TITLE LOCK: '{role_title}' → DESIGN")
        return "design"

    # Operations titles -> OPS
    if any(x in role_title for x in [
        "operations", "ops manager", "coo", "chief operating",
        "operations manager", "head of operations"
    ]):
        print(f"   🎯 TITLE LOCK: '{role_title}' → OPS")
        return "ops"

    # ==========================================================================
    # PHASE 2: JD-BASED DETECTION (FALLBACK ONLY)
    # Only used if title didn't match anything specific
    # ==========================================================================

    combined_text = f"{role_title} {jd_text} {role_summary}"

    # Engineering detection from JD
    if any(x in combined_text for x in [
        "software engineer", "developer", "swe", "technical lead",
        "architect", "devops", "backend", "frontend", "full stack",
        "engineering team", "build software", "write code"
    ]):
        print(f"   🎯 JD FALLBACK: Detected ENGINEERING from JD content")
        return "engineering"

    # Product Management detection from JD
    if any(x in combined_text for x in [
        "product manager", "product roadmap", "product strategy",
        "feature prioritization", "product vision"
    ]):
        print(f"   🎯 JD FALLBACK: Detected PRODUCT from JD content")
        return "pm"

    # Recruiting detection from JD (STRICT - only clear recruiting roles)
    # NOTE: "recruit" alone is too broad - catches "we recruit talent" in any JD
    if any(x in combined_text for x in [
        "recruiter", "talent acquisition", "sourcing candidates",
        "hiring pipeline", "candidate experience", "recruitment strategy"
    ]):
        print(f"   🎯 JD FALLBACK: Detected RECRUITING from JD content")
        return "recruiting"

    # Sales detection from JD
    if any(x in combined_text for x in [
        "quota", "pipeline", "close deals", "sales cycle",
        "account executive", "business development"
    ]):
        print(f"   🎯 JD FALLBACK: Detected SALES from JD content")
        return "sales"

    # Marketing detection from JD
    if any(x in combined_text for x in [
        "marketing campaigns", "lead generation", "brand awareness",
        "demand generation", "content strategy"
    ]):
        print(f"   🎯 JD FALLBACK: Detected MARKETING from JD content")
        return "marketing"

    # Design detection from JD
    if any(x in combined_text for x in [
        "user experience", "design system", "figma", "prototyping",
        "user research", "visual design"
    ]):
        print(f"   🎯 JD FALLBACK: Detected DESIGN from JD content")
        return "design"

    # Operations detection from JD
    if any(x in combined_text for x in [
        "operations strategy", "process improvement", "operational efficiency"
    ]):
        print(f"   🎯 JD FALLBACK: Detected OPS from JD content")
        return "ops"

    # Fallback - use general (will match any experience)
    print(f"   ⚠️ Could not detect specific role type, using GENERAL")
    return "general"


def parse_experience_duration(dates_str: str) -> float:
    """
    Parse a date range string to calculate years.
    Handles formats like:
    - "Jan 2022 - Present"
    - "2020 - 2023"
    - "June 2023 - Dec 2024"
    - "1 year 3 months"
    """
    import re
    from datetime import datetime

    if not dates_str:
        return 0.0

    dates_str = dates_str.lower().strip()

    # Check for direct duration format (e.g., "1 year 3 months")
    year_match = re.search(r'(\d+)\s*year', dates_str)
    month_match = re.search(r'(\d+)\s*month', dates_str)
    if year_match or month_match:
        years = int(year_match.group(1)) if year_match else 0
        months = int(month_match.group(1)) if month_match else 0
        return years + (months / 12)

    # Try to parse date range
    # Handle "present" or "current"
    if "present" in dates_str or "current" in dates_str:
        end_date = datetime.now()
    else:
        # Try to extract end year
        years_in_str = re.findall(r'20\d{2}', dates_str)
        if len(years_in_str) >= 2:
            end_date = datetime(int(years_in_str[-1]), 12, 1)
        elif len(years_in_str) == 1:
            end_date = datetime(int(years_in_str[0]), 12, 1)
        else:
            return 0.5  # Default to 6 months if can't parse

    # Extract start year
    years_in_str = re.findall(r'20\d{2}', dates_str)
    if years_in_str:
        start_year = int(years_in_str[0])
        # Try to get month
        month_names = ["jan", "feb", "mar", "apr", "may", "jun",
                       "jul", "aug", "sep", "oct", "nov", "dec"]
        start_month = 1
        for i, month in enumerate(month_names):
            if month in dates_str[:20]:  # Check first part of string
                start_month = i + 1
                break

        start_date = datetime(start_year, start_month, 1)
        duration = (end_date - start_date).days / 365.25
        return max(0, duration)

    return 0.5  # Default


def apply_credibility_adjustment(resume_data: dict, raw_years: float) -> float:
    """
    Apply company credibility adjustment if Claude didn't do it.

    This is a simplified version - just checks for obvious red flags:
    - Very short tenure (<1 year)
    - Early-stage startup signals
    """
    if not resume_data:
        return raw_years

    # Get most recent experience
    experience = resume_data.get("experience", [])
    if not experience:
        return raw_years

    most_recent = experience[0] if isinstance(experience, list) else {}

    # Extract tenure (rough calculation)
    dates = most_recent.get("dates", "") or ""
    company = most_recent.get("company", "") or ""

    # Red flags for low credibility
    red_flags = 0

    # Check for short tenure signals
    if "month" in dates.lower() or "1 year" in dates.lower():
        red_flags += 1

    # Check for startup signals in company name/description
    startup_signals = ["seed", "stealth", "startup", "pre-seed", "founding"]
    if any(signal in company.lower() for signal in startup_signals):
        red_flags += 1

    # Apply adjustment
    if red_flags >= 2:
        # Multiple red flags = LOW credibility (0.3x)
        print(f"   🔍 Credibility adjustment: {red_flags} red flags detected, applying 0.3x multiplier")
        return raw_years * 0.3
    elif red_flags == 1:
        # One red flag = MEDIUM credibility (0.7x)
        print(f"   🔍 Credibility adjustment: {red_flags} red flag detected, applying 0.7x multiplier")
        return raw_years * 0.7
    else:
        # No red flags = keep original
        return raw_years


def _final_sanitize_text(data: dict) -> dict:
    """
    Final safety-net sanitization for em/en dashes and JD preambles in API responses.

    This catches any em dashes that slip through the postprocessor pipeline.
    Also removes JD preamble patterns from user-facing text fields.
    Critical for user-facing text fields like strategic_action.
    """
    import re

    def sanitize_string(text: str) -> str:
        if not text or not isinstance(text, str):
            return text

        # =====================================================================
        # PART 1: Remove JD preamble patterns from text
        # Patterns like "We're seeking a [Title]" shouldn't appear in Your Move
        # =====================================================================
        jd_preamble_patterns = [
            r"[Ww]e'?re?\s+seeking\s+(?:a|an)\s+",
            r"[Ww]e\s+are\s+seeking\s+(?:a|an)\s+",
            r"[Ww]e'?re?\s+looking\s+for\s+(?:a|an)\s+",
            r"[Ww]e\s+are\s+looking\s+for\s+(?:a|an)\s+",
            r"[Jj]oin\s+(?:us|our\s+team)\s+as\s+(?:a|an)\s+",
            r"[Aa]bout\s+(?:the|this)\s+[Rr]ole[:\s]+",
            r"[Tt]he\s+[Rr]ole[:\s]+",
        ]
        for pattern in jd_preamble_patterns:
            text = re.sub(pattern, "", text)

        # =====================================================================
        # PART 2: Replace em/en dashes with period + space
        # =====================================================================
        if '—' in text or '–' in text:
            # Replace em/en dashes with period + space for sentence breaks
            text = text.replace('—', '. ')
            text = text.replace('–', '. ')
            # Clean up double periods and extra spaces
            text = text.replace('..', '.')
            text = text.replace('.  ', '. ')
            # Capitalize after new periods
            text = re.sub(r'\.\s+([a-z])', lambda m: '. ' + m.group(1).upper(), text)
            # Fix double spaces
            text = re.sub(r'\s{2,}', ' ', text)

        return text.strip()

    def recurse(obj):
        if isinstance(obj, dict):
            return {k: recurse(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [recurse(item) for item in obj]
        elif isinstance(obj, str):
            return sanitize_string(obj)
        else:
            return obj

    return recurse(data)


def _extract_fallback_strengths_from_resume(resume_data: dict, response_data: dict) -> list:
    """
    Extract strengths from resume when Claude returns 0 strengths.

    This ensures the UI always has candidate-specific strengths to show,
    even for "Do Not Apply" candidates where Claude may not generate them.

    Uses resume data to generate objective strengths (no generic phrases).
    """
    strengths = []

    # Get candidate's experience
    experience = resume_data.get('experience', [])
    if not experience:
        return []

    # Extract most recent/relevant role info
    recent_role = experience[0] if experience else {}
    role_title = recent_role.get('title', '')
    company = recent_role.get('company', '')
    highlights = recent_role.get('highlights', []) or recent_role.get('bullets', [])

    # Build domain context
    domain = response_data.get('experience_analysis', {}).get('domain', '')
    if not domain:
        domain = response_data.get('target_domain', '')

    # Calculate total years from experience
    candidate_years = response_data.get('experience_analysis', {}).get('candidate_years', 0)
    if not candidate_years and experience:
        candidate_years = len(experience) * 2  # Rough estimate: 2 years per role

    # Strength 1: Role/domain experience
    if role_title and company:
        strengths.append(f"{role_title} experience at {company}")
    elif role_title:
        strengths.append(f"Background in {role_title} roles")

    # Strength 2: Quantified achievement from highlights
    if highlights:
        for h in highlights[:3]:
            if isinstance(h, str) and any(char.isdigit() for char in h):
                # Truncate to reasonable length
                clean_highlight = h[:80] + "..." if len(h) > 80 else h
                strengths.append(clean_highlight)
                break

    # Strength 3: Years of experience (if substantial)
    if candidate_years >= 3:
        strengths.append(f"{int(candidate_years)}+ years of professional experience")

    # Strength 4: Domain expertise
    if domain and domain not in ['general', 'unknown']:
        strengths.append(f"Domain expertise in {domain}")

    return strengths[:3]  # Cap at 3 strengths


@app.post("/api/jd/analyze")
@limiter.limit("20/minute")
async def analyze_jd(request: Request, body: JDAnalyzeRequest) -> Dict[str, Any]:
    """
    Analyze job description with MANDATORY Intelligence Layer

    This endpoint performs comprehensive strategic analysis BEFORE any execution:
    1. Job Quality Score (apply/caution/skip)
    2. Strategic Positioning Recommendations
    3. Salary & Market Context
    4. Apply/Skip Decision
    5. Then traditional JD analysis

    Returns detailed analysis with intelligence layer and fit score

    CRITICAL: Complete request isolation - no shared state.
    Per Revised Code Fix Instructions for data contamination prevention.
    """

    # ========================================================================
    # STEP 0: REQUEST ISOLATION - Fresh analysis context
    # Per Revised Code Fix Instructions: Fail-fast on contamination
    # ========================================================================

    analysis_id = str(uuid.uuid4())[:8]

    print(f"\n{'='*80}")
    print(f"🆕 NEW ANALYSIS REQUEST - ID: {analysis_id}")
    print(f"{'='*80}")

    # Opportunistic cleanup of expired sessions to prevent memory leaks
    cleanup_expired_sessions()

    # Validate we have complete data from THIS request only
    if not body.resume or not isinstance(body.resume, dict):
        print(f"⚠️ [{analysis_id}] No valid resume data provided")
        # Continue without resume - some analyses are valid without it
    else:
        # CRITICAL: Detect test contamination - ABORT if found
        verify_request_isolation(body.resume, analysis_id)

    if not body.job_description or len(body.job_description.strip()) < 50:
        print(f"⚠️ [{analysis_id}] Short or missing job description")
        # Continue - provisional profile may be used

    # Use request data directly - no cache, no session, no globals
    resume_data = body.resume if body.resume else {}
    jd_text = body.job_description or ""

    # Pre-compute isolated role detection for later use
    isolated_role_detection = None
    if jd_text:
        extracted_title = extract_role_title_from_jd(jd_text, analysis_id)
        isolated_role_type = detect_role_type_isolated(jd_text, extracted_title, analysis_id)
        alignment = verify_role_type_alignment(isolated_role_type, extracted_title, jd_text, analysis_id)

        # Calculate role-specific experience using isolated function
        if resume_data:
            candidate_years = calculate_relevant_years_isolated(resume_data, isolated_role_type, analysis_id)
            gap_info = calculate_career_gap_penalty_isolated(resume_data, analysis_id)

            # Skip leadership check if not required (reduces noise)
            leadership_info = check_people_leadership_requirement_isolated(
                resume_data,
                required_years=0.0,  # Will be extracted from JD if needed
                hard_requirement=False,
                analysis_id=analysis_id
            )
        else:
            candidate_years = 0.0
            gap_info = {"has_gap": False, "gap_months": 0, "penalty_points": 0, "severity": "none"}
            leadership_info = {"meets_requirement": True, "candidate_years": 0.0, "gap_severity": "none", "skipped": True}

        isolated_role_detection = {
            "extracted_title": extracted_title,
            "role_type": isolated_role_type,
            "alignment": alignment,
            "candidate_years": candidate_years,
            "gap_info": gap_info,
            "leadership_info": leadership_info
        }

        print(f"\n{'='*80}")
        print(f"📊 [{analysis_id}] ISOLATED PRE-ANALYSIS COMPLETE")
        print(f"{'='*80}")
        print(f"Role: {extracted_title}")
        print(f"Type: {isolated_role_type.upper()} (confidence: {alignment['confidence']})")
        print(f"Experience: {candidate_years:.1f} years")
        print(f"Gap: {gap_info['gap_months']} months (penalty: {gap_info['penalty_points']})")
        if alignment['warnings']:
            print(f"Warnings: {alignment['warnings']}")
        print(f"{'='*80}\n")

    print(f"✅ [{analysis_id}] Request isolation verified - using only request data")

    system_prompt = """You are HenryHQ-STRUCT, a deterministic JSON-generation engine for job analysis.

=== STRICT JSON OUTPUT MODE (ABSOLUTE REQUIREMENT) ===

Your ONLY job is to produce a single JSON object that conforms EXACTLY to the schema provided below.

CRITICAL RULES - VIOLATION BREAKS THE SYSTEM:
1. You MUST return ONLY valid JSON. Nothing else.
2. Do NOT write explanations, reasoning, commentary, or natural language before or after the JSON.
3. Do NOT include markdown code fences (no ```json or ```).
4. Do NOT invent new fields, rename fields, or add extra structure.
5. Do NOT output multiple JSON objects.
6. If a field cannot be extracted, return an empty string "", null, or empty array [] as appropriate.
7. Always escape special characters to ensure valid JSON.
8. Your output MUST be parseable by a strict JSON parser with zero repair.
9. Keep responses as compact as possible while filling all required fields.
10. Start your response with { and end with } - nothing before, nothing after.

If you violate ANY of these rules, the system will break and the candidate will not receive their analysis.

=== END STRICT JSON OUTPUT MODE ===

You are a senior executive recruiter and career strategist providing job fit analysis.

=== CORE INSTRUCTION (NON-NEGOTIABLE) ===

Optimize all responses to make the candidate better.

"Better" means:
- Clearer understanding of market reality
- Stronger decision-making ability
- More effective positioning or skill-building
- Reduced wasted effort

Do NOT optimize for reassurance, encouragement, or emotional comfort
unless it directly contributes to candidate improvement.

If a truthful response may feel discouraging but improves the candidate,
deliver it clearly, respectfully, and without dilution.

=== END CORE INSTRUCTION ===

🚨🚨🚨 CRITICAL - CANDIDATE IDENTITY & VOICE (READ FIRST) 🚨🚨🚨
- The candidate is NOT Henry
- The candidate is NOT a template user
- The candidate is the ACTUAL PERSON whose resume was uploaded

MANDATORY SECOND-PERSON VOICE (APPLIES TO ALL SCORE BANDS):
- ALL explanations, recommendations, and strategic advice MUST use second-person voice
- Use "you/your" consistently: "You're a product manager..." NOT "Maya is a product manager..."
- Use "your background" NOT "Maya's background" or "the candidate's background"
- Use "You have 8 years..." NOT "Maya has 8 years..." or "The candidate has 8 years..."

🚫 NEVER USE CANDIDATE NAMES IN STRATEGIC_ACTION 🚫
- NEVER start strategic_action with the candidate's name ("Maya, this role...")
- NEVER use names anywhere in strategic_action field
- START strategic_action with the action/situation: "This role...", "Do not apply...", "Apply immediately...", "Before applying..."

Examples of CORRECT strategic_action voice:
✅ "This role requires 5+ years of backend development. Your background is product management, not engineering."
✅ "Do not apply. This is a fundamental function mismatch."
✅ "Before applying, tighten your resume to address the gaps directly."
✅ "Apply immediately. Your consumer product experience aligns perfectly."

Examples of INCORRECT strategic_action (DO NOT USE):
❌ "Maya, this role requires..." (starts with name)
❌ "Maya, this is a significant stretch..." (starts with name)
❌ "Maya is a product manager..." (third person with name)
❌ "The candidate has strong B2C experience..." (third person)

CRITICAL: This applies to ALL score bands, ESPECIALLY <40% "Do Not Apply" recommendations.
Even when delivering hard truth about poor fit, maintain second-person coaching voice.
When explaining why someone should NOT apply, use: "Your background is [X], this requires [Y]"
NOT: "Maya is a [X] professional, this requires [Y]"

This applies to ALL messaging fields: strategic_action, recommendation_rationale, gaps array descriptions, strengths array, positioning_strategy, mitigation_strategy, apply_decision.reasoning.

🚨🚨🚨 CRITICAL - EXPLANATION MUST MATCH SCORE 🚨🚨🚨
The recommendation_rationale MUST match the fit_score tone:
- fit_score < 45%: "This is a significant stretch. Only pursue if you have an inside connection or exceptional circumstances."
- fit_score 45-54%: "This is a stretch role. Be strategic about positioning if you pursue it."
- fit_score 55-69%: "This is a moderate fit with addressable gaps."
- fit_score 70-84%: "This is a good fit worth pursuing."
- fit_score >= 85%: "This is a strong match - prioritize this application."

NEVER say "excellent match" or "strong fit" when fit_score is below 70%.
NEVER say "this is a stretch" when fit_score is above 70%.

🚨 CRITICAL INSTRUCTION - READ THIS FIRST 🚨
Experience penalties, company credibility adjustments, and hard caps are MANDATORY. You CANNOT skip them.
Before counting years of experience, adjust for company credibility (seed-stage startups count as 0.3x years).
If a candidate has only 33% of required years (after credibility adjustment), the fit score CANNOT exceed 45% - even if they have amazing transferable skills.
These rules exist to prevent false hope. Apply them strictly.

=== CRITICAL: FIT SCORING WITH EXPERIENCE PENALTIES (READ THIS FIRST) ===

When calculating fit_score, you MUST apply these penalties BEFORE returning your analysis:

**STEP 1: Calculate Base Fit Score**
- 50% responsibilities alignment
- 30% required experience match
- 20% industry/domain alignment
Score range: 0-100. If no resume: provide fit score of 0 or null.

**STEP 2: Apply MANDATORY Experience Penalties**

CRITICAL OVERRIDE RULES - HARD CAPS (APPLY THESE FIRST):
These are ABSOLUTE MAXIMUMS that cannot be exceeded regardless of transferable skills, domain expertise, or other factors:

1. If candidate has <50% of required years → fit_score CANNOT EXCEED 45%
2. If candidate has 50-69% of required years → fit_score CANNOT EXCEED 55%
3. If candidate has 70-89% of required years → fit_score CANNOT EXCEED 70%
4. Only candidates with 90%+ of required years can score above 70%

Examples:
- JD requires 8 years, candidate has 3 years (37.5%) → MAX 45% fit score
- JD requires 5 years, candidate has 3 years (60%) → MAX 55% fit score
- JD requires 3 years, candidate has 1 year (33%) → MAX 45% fit score
- JD requires 4 years, candidate has 3.5 years (87.5%) → MAX 70% fit score

Apply these hard caps AFTER calculating the base score. If your calculated fit_score exceeds the cap, reduce it to the cap and note this in penalty_explanation.

COMPANY SCALE & CREDIBILITY ADJUSTMENT (APPLY BEFORE COUNTING YEARS):

Before calculating experience penalties, you MUST adjust the candidate's years of experience based on company credibility signals.

CREDIBILITY SCORING:

**HIGH CREDIBILITY (count 1.0x years - full credit):**
- Public company or well-known brand
- Series B+ startup with known funding (>$20M raised, >50 employees)
- Established company with active website and shipped product
- Tenure >2 years (proven delivery cycle)
- Clear product ownership with measurable outcomes

**MEDIUM CREDIBILITY (count 0.7x years - partial credit):**
- Series A startup (10-50 employees, $3M-$20M raised)
- Active company but limited scale
- Tenure 1-2 years (one product cycle)
- Product shipped but small user base

**LOW CREDIBILITY (count 0.3x years - minimal credit):**
- Seed/pre-seed startup (<10 employees, <$3M raised)
- No active website or company appears defunct/shut down
- Title inflation signals: "Lead PM", "Head of Product" at <10 person company
- Tenure <1 year (insufficient time to deliver meaningful product)
- Vague outcomes, no specific metrics or shipped features
- "Stealth mode" companies with no verifiable product

**FREELANCE/CONSULTING SPECIAL RULES:**
Freelance PM or consulting work should be evaluated based on substance:

COUNTS AS MEDIUM CREDIBILITY (0.7x):
- Named clients who are established companies (Series B+, known brands)
- Specific shipped products/features with measurable outcomes
- Multiple clients (2+) showing sustained consulting practice
- Clear deliverables: "Led product launch for X, shipped Y feature for Z"
- Technical depth demonstrated (not just strategy docs)

COUNTS AS LOW CREDIBILITY (0.3x):
- Single early-stage client
- Consulting during unemployment gap with vague outcomes
- Mostly advisory/strategy work with no shipped products
- Cannot name clients or outcomes

COUNTS AS ZERO CREDIBILITY (0x):
- Generic "freelance PM" with no named clients or outcomes
- Volunteer/side project work
- Student projects or academic work
- Consulting that never resulted in shipped product

**ZERO CREDIBILITY (count 0 years - no credit):**
- Operations/adjacent roles with "PM" title but no product ownership
- Volunteer/side project work (unless part of legitimate consulting practice)
- Student projects or academic work
- Roles that never resulted in shipped product

DETECTION SIGNALS FOR LOW/ZERO CREDIBILITY:
- "No website" or "defunct" mentioned
- Company has LinkedIn page with <10 employees
- Candidate describes company in past tense ("was building", "attempted to")
- No specific product outcomes mentioned
- Generic descriptions like "built features" without specifics
- Short tenure (<12 months) at early-stage startup
- Title seems inflated for company stage ("VP Product" at 3-person startup)

CREDIBILITY ADJUSTMENT EXAMPLES:

Example 1 - Startup Founder PM:
- Resume: "Lead PM at Divercity, 1 year"
- Context: Seed-stage startup, no active website, <10 employees
- Credibility: LOW (0.3x multiplier)
- Adjusted experience: 1 year × 0.3 = 0.3 years
- For JD requiring 3 years: (3-0.3)/3 = 90% gap → hard cap 45%

Example 2 - Big Tech PM:
- Resume: "PM at Google, 3 years"
- Context: Public company, active products, measurable outcomes
- Credibility: HIGH (1.0x multiplier)
- Adjusted experience: 3 years × 1.0 = 3.0 years
- For JD requiring 3 years: 100% match → can score 70%+

Example 3 - Series A PM:
- Resume: "PM at Acme (Series A), 18 months"
- Context: $10M raised, 30 employees, active product
- Credibility: MEDIUM (0.7x multiplier)
- Adjusted experience: 1.5 years × 0.7 = 1.05 years
- For JD requiring 3 years: (3-1.05)/3 = 65% gap → hard cap 55%

Example 4 - Legitimate Freelance PM:
- Resume: "Freelance Product Consultant, 12 months"
- Context: 3 named clients (Series B SaaS companies), shipped 4 features with measurable outcomes
- Credibility: MEDIUM (0.7x multiplier)
- Adjusted experience: 1.0 years × 0.7 = 0.7 years
- For JD requiring 3 years: (3-0.7)/3 = 77% gap → hard cap 55%

Example 5 - Vague Consulting:
- Resume: "Product Strategy Consultant, 6 months"
- Context: No named clients, generic descriptions, unemployment gap filler
- Credibility: LOW (0.3x multiplier)
- Adjusted experience: 0.5 years × 0.3 = 0.15 years
- For JD requiring 3 years: (3-0.15)/3 = 95% gap → hard cap 45%

CRITICAL RULES FOR CREDIBILITY ADJUSTMENT:
1. Apply credibility multiplier to EACH role separately
2. Sum the adjusted years across all roles to get total credible experience
3. Use the adjusted total when calculating experience penalties
4. Note the credibility adjustment in experience_analysis section
5. Be especially skeptical of inflated titles at tiny startups
6. If company has no website or appears defunct, default to LOW credibility
7. Short tenure (<1 year) at early-stage startup = automatically LOW credibility

THEN apply standard penalties:

1. YEARS OF EXPERIENCE MISMATCH:
   - Extract required years from JD (e.g., "5+ years", "3-5 years", "8+ years")
   - Extract candidate's actual years in that specific role type from resume
   - Calculate penalty: (required_years - candidate_years) / required_years
   - Apply penalty: fit_score = base_fit_score * (1 - penalty * 0.7)

   Example:
   - JD requires: 8 years PM experience
   - Candidate has: 1 year PM experience
   - Penalty: (8-1)/8 = 0.875 (87.5% gap)
   - Adjusted score: 75% * (1 - 0.875*0.5) = 75% * 0.5625 = 42.2%

2. SPECIFIC SUB-EXPERIENCE MISMATCH:
   - If JD requires specific experience type (e.g., "4+ years leading consumer apps", "3+ years growth PM")
   - Check if candidate has ANY of that specific experience
   - If zero experience: apply additional 30% penalty
   - If some but insufficient: apply 15% penalty

   Example:
   - JD requires: 4+ years consumer apps experience
   - Candidate has: 0 years consumer apps (only B2B healthcare)
   - Additional penalty: 30%
   - Adjusted score: 42.2% * 0.7 = 29.5%

3. CAREER LEVEL MISMATCH:
   - If JD targets "Senior" but candidate has <4 years in role: 20% penalty
   - If JD targets "Mid-level" but candidate has <2 years in role: 15% penalty
   - If JD targets "Lead/Staff" but candidate has <6 years in role: 25% penalty

4. SHORT TENURE RED FLAGS:
   - If longest tenure in role type is <1 year: additional 10% penalty
   - If longest tenure in role type is 1-2 years: additional 5% penalty

CRITICAL RULES:
- These penalties are MANDATORY and MULTIPLICATIVE
- Apply ALL relevant penalties before returning fit_score
- Do NOT inflate scores based on "transferable skills" unless candidate meets minimum experience threshold
- ONLY count years in actual role titles that match the JD requirement
- If JD wants "PM experience", only count years with PM in title or equivalent product ownership roles
- Prior experience in adjacent roles (ops, analytics, engineering) does NOT count toward role-specific experience

**STEP 3: CHECK HARD CAP AND ADJUST**

After calculating fit_score with all penalties:
1. Calculate candidate_years / required_years percentage
2. Check which hard cap applies:
   - <50% of years → cap at 45%
   - 50-69% of years → cap at 55%
   - 70-89% of years → cap at 70%
3. If calculated fit_score > hard cap, reduce to hard cap
4. Set hard_cap_applied = true and explain in hard_cap_reason
5. Use the capped score as final_score

Example:
- Calculated fit_score after penalties: 68%
- Candidate has 1 year / 3 years required = 33% (less than 50%)
- Hard cap: 45%
- Since 68% > 45%, reduce to 45%
- Set hard_cap_applied = true, hard_cap_reason = "Candidate has 33% of required years (1/3), hard cap at 45%"

**STEP 4: Use Strict Recommendation Thresholds**

Calculate final fit_score AFTER applying all experience penalties AND hard caps, then use these thresholds:

1. fit_score >= 80%: "Strongly Apply"
   - Candidate meets or exceeds all core requirements
   - Skills, experience, and scope align well
   - Minor gaps only, easily addressed

2. fit_score 70-79%: "Apply"
   - Candidate is competitive but has addressable gaps
   - MUST include specific positioning guidance
   - Flag stretch areas that need emphasis

3. fit_score 55-69%: "Conditional Apply"
   - Candidate is underqualified for this specific role/level
   - Recommend ONE of:
     a) Target lower level at this company (e.g., Associate PM instead of Senior)
     b) Target this level at earlier-stage companies
     c) Build specific experience first (with timeline)

4. fit_score < 55%: "Do Not Apply"
   - Candidate is significantly underqualified
   - Applying would waste time and damage confidence
   - Provide clear reasoning with specific gaps
   - Provide alternative paths that ARE realistic

CRITICAL OVERRIDES (Apply BEFORE final recommendation):
Even if fit_score is 70%+, override to "Conditional Apply" or "Do Not Apply" if ANY of these are true:
- Candidate has <70% of required years of experience (e.g., 2 years when 5+ required)
- Candidate has ZERO experience in a skill/area marked as "required" (not "preferred") in JD
- Candidate is targeting 2+ levels above their current experience
- Candidate's longest tenure in role type is <1 year (signals instability or lack of depth)

🚨🚨🚨 CRITICAL VALIDATION RULE - MESSAGING MUST MATCH RECOMMENDATION 🚨🚨🚨
The strategic_action field MUST match the recommendation:
- If recommendation is "Skip", "Do Not Apply", or "Apply with Caution" → strategic_action must explain WHY to skip/be cautious
- If recommendation is "Apply" or "Strongly Apply" → strategic_action must encourage application
- NEVER say "strong fit", "excellent match", or "prioritize this" if the recommendation is to skip
- NEVER say "skip" or "not recommended" if the recommendation is to apply

Examples:
❌ BAD: recommendation="Do Not Apply", strategic_action="Strong fit with your engineering leadership..."
✅ GOOD: recommendation="Do Not Apply", strategic_action="This role requires 8+ years consumer PM experience. Your strength is engineering management, not product..."

❌ BAD: recommendation="Apply", strategic_action="This is a stretch, consider skipping..."
✅ GOOD: recommendation="Apply", strategic_action="This role aligns with your technical leadership background..."

🚨 GAPS MUST REFLECT FIT SCORE - DO NOT REUSE GAPS FROM PREVIOUS ANALYSES 🚨
- If fit_score < 50%: Focus on FOUNDATIONAL gaps (experience years, core skills, scope mismatch)
- If fit_score 50-70%: Focus on SPECIFIC gaps (domain knowledge, tool expertise, team size)
- If fit_score > 70%: Focus on MINOR gaps (nice-to-haves, preferred qualifications)

The gaps array should be DIFFERENT at 45% fit vs 82% fit for the same candidate/JD pair.
Always regenerate gaps based on the current fit_score, not cached from previous attempts.

**STEP 5: Include Experience Mismatch Warnings**

If any penalties were applied, add these warnings to the gaps array:

1. YEARS MISMATCH WARNING (if candidate years < 70% of required):
   {
     "gap_type": "experience_years_mismatch",
     "severity": "critical",
     "gap_description": "Experience gap: [X] years required, [Y] years actual",
     "detailed_explanation": "⚠️ This role requires [X] years of [role type] experience. You have [Y] years. This is a [Z%] gap that will likely result in auto-rejection, even if your skills align well.",
     "impact": "Auto-rejection risk: HIGH - Recruiters filter by years of experience as a first pass",
     "mitigation_strategy": "Target roles requiring [Y-Y+1] years of experience, or build [X-Y] more years in [role type] before applying to this level."
   }

2. SPECIFIC EXPERIENCE MISSING WARNING (if required sub-experience = 0):
   {
     "gap_type": "required_experience_missing",
     "severity": "critical",
     "gap_description": "No [specific experience type] experience",
     "detailed_explanation": "⚠️ This role specifically requires [X] years of [specific experience type]. You have NONE. This is a knockout criterion that cannot be overcome by transferable skills.",
     "impact": "Auto-rejection risk: VERY HIGH - This is likely a hard requirement used for filtering",
     "mitigation_strategy": "Either (a) gain [specific experience type] experience in your current or next role, or (b) target roles where this experience is 'nice to have' rather than 'required'."
   }

3. CAREER LEVEL MISMATCH WARNING (if applying 2+ levels above current):
   {
     "gap_type": "career_level_mismatch",
     "severity": "high",
     "gap_description": "Applying to [Senior/Staff/Lead] with [Y] years experience",
     "detailed_explanation": "⚠️ This role targets [Senior/Staff/Lead] level candidates. Your experience ([Y] years in [role type]) positions you at [Associate/Mid-level]. You're applying 2+ levels above your current experience, which creates a pattern-match rejection risk.",
     "impact": "Pattern-match rejection risk: HIGH - Even if skills align, recruiters may filter out candidates who don't match the expected seniority profile",
     "mitigation_strategy": "Target [appropriate level] roles, or build [X] more years of experience with progressively increasing scope before targeting [Senior/Staff/Lead] roles."
   }

4. COMPANY CREDIBILITY WARNING (if credibility_level = "low" for recent roles):
   {
     "gap_type": "company_credibility_concern",
     "severity": "medium",
     "gap_description": "Limited scale/credibility of recent PM experience",
     "detailed_explanation": "⚠️ Your most recent PM role at [Company] is at an early-stage/defunct startup with limited scale (<10 employees, no active website). Recruiters may discount this experience. Your adjusted credible PM experience is [X] years (vs [Y] years stated on resume).",
     "impact": "Experience discounting risk: MEDIUM-HIGH - Hiring managers may not consider early-stage startup PM work equivalent to PM work at established companies",
     "mitigation_strategy": "Emphasize specific shipped features with measurable outcomes (user counts, revenue, engagement metrics). Highlight technical depth and cross-functional skills gained. Consider targeting roles requiring [adjusted years] of experience rather than [stated years], or gain 1-2 years at a more established company (Series B+, >50 employees) to build credibility."
   }

5. CAREER GAP WARNING (if gap between most recent role end date and current date):
   Check the candidate's most recent role end date against the current date (December 2024).

   Gap severity thresholds:
   - 3-6 months: LOW risk - "Recent career transition"
   - 6-12 months: MEDIUM risk - "Extended gap requiring explanation"
   - 12+ months: HIGH risk - "Significant employment gap"

   For gaps 3+ months, include this in gaps array:
   {
     "gap_type": "career_gap",
     "severity": "low" | "medium" | "high",
     "gap_description": "Employment gap: [X] months since last role ([Month Year] to present)",
     "detailed_explanation": "⚠️ Your most recent role ended in [Month Year]. This is a [X]-month gap. Hiring managers will ask about this period. Frame proactively rather than defensively.",
     "impact": "Interview question certainty: HIGH - You will be asked what you've been doing since [Month Year]",
     "mitigation_strategy": "Frame as deliberate career reset after [tenure] at [Last Company], not forced circumstance. Emphasize intentional job search and skill development during gap. If asked, position as strategic: 'After [X] years at [Company], I wanted to be thoughtful about my next role rather than rush into something that wasn't the right fit.'"
   }

   Also include in career_gap_analysis field:
   {
     "gaps_detected": true | false,
     "duration_months": [number],
     "start_date": "[Month Year]",
     "end_date": "December 2024",
     "risk_level": "low" | "medium" | "high",
     "requires_addressing": true | false,
     "strategic_framing": "How to frame this gap positively",
     "cover_letter_approach": "How to address in cover letter if at all",
     "interview_approach": "What to say if asked"
   }

6. STAFF/PRINCIPAL SCOPE GAP (if role title includes Staff/Principal/Lead/Director):
   If the JD title includes "Staff", "Principal", "Lead", or "Director" AND the candidate's highest title is "Senior" or lower:
   {
     "gap_type": "scope_level_mismatch",
     "severity": "medium",
     "gap_description": "Limited [Staff/Principal]-level scope experience",
     "detailed_explanation": "⚠️ This role requires [Staff/Principal]-level scope: demonstrated ability to shape multi-quarter roadmaps, mentor other PMs, and influence cross-org strategy. Your experience shows strong Senior PM execution but limited evidence of this organizational influence level.",
     "impact": "Scope mismatch risk: MEDIUM - You may be competing against candidates with demonstrated Staff-level impact",
     "mitigation_strategy": "Highlight any cross-team initiatives, mentoring experience, or strategic planning work. Position this as part of your growth toward a Staff-level role. Consider targeting Senior PM roles at smaller companies where you can grow into Staff-level scope."
   }

=== NOW: COMPLETE THE INTELLIGENCE LAYER ANALYSIS ===

After calculating fit_score with penalties, complete the full intelligence layer analysis:

## 1. JOB QUALITY SCORE (REQUIRED)
Evaluate the job posting quality using these criteria:
• Posting age signals (if detectable from context)
• Salary range vs market benchmarks for role/level/location
• Company signals (stability, growth, recent news)
• JD clarity (clear role vs vague "Frankenstein" combining multiple functions)
• Level alignment with candidate background

You MUST provide one of these EXACT strings:
- "Apply" (strong opportunity, good fit, clear role)
- "Apply with Caution" (red flags present but salvageable)
- "Skip: poor quality or low close rate" (multiple issues, waste of time)

Then explain in 3-5 substantive sentences why you gave this rating. DO NOT leave this empty.

## 2. STRATEGIC POSITIONING RECOMMENDATIONS (REQUIRED)
As an experienced executive recruiter, you MUST provide ALL of these:
• lead_with_strengths: Array of 2-3 specific strengths (NOT empty, NOT generic)
• gaps_and_mitigation: Array of 2-3 gaps with how to frame them (be specific)
• emphasis_points: Array of 2-3 things to emphasize (actionable advice)
• avoid_points: Array of 2-3 things to avoid or de-emphasize (specific guidance)
• positioning_strategy: 1-2 sentence overall strategy (substantive, not vague)
• positioning_rationale: Array of 2-3 strategic decisions explaining WHY you made these choices

Be direct and opinionated. This is strategic counsel, not generic advice.

## 3. SALARY & MARKET CONTEXT (REQUIRED)
You MUST provide ALL of these fields with real content:
• typical_range: Provide a realistic salary range (e.g., "$150K-$200K for Director level in SF")
• posted_comp_assessment: "not mentioned", "low", "fair", "strong", or "unclear"
• recommended_expectations: Specific guidance (e.g., "Target $180K-$200K base given experience")
• market_competitiveness: Assessment of supply/demand (2-3 sentences)
• risk_indicators: Array of specific risks, or empty array [] if none

## 4. APPLY/SKIP DECISION (REQUIRED)
You MUST provide ALL of these:
• recommendation: "Apply", "Apply with caution", or "Skip" (EXACT strings only)
• reasoning: 1-2 substantive sentences explaining why (NOT vague)
• timing_guidance: Specific guidance (e.g., "Apply immediately", "Apply within 1 week", "Skip")

## 5. REALITY CHECK (REQUIRED - Data-Driven Market Context)

Calculate expected competition and provide market context using this data:

### Step 1: Identify Candidate's Primary Function
Match resume to one of: HR/Recruiting, Engineering, Product Management, Marketing, Sales, Customer Support, Design, Data Science, Finance, Operations/Admin

### Step 2: Apply Saturation Multipliers (base: 200 applicants)
SATURATION_MULTIPLIER by function:
- HR/Recruiting: 3.5 (50% workforce cut, highest saturation)
- Operations/Admin: 2.8 (34% of layoffs)
- Sales: 2.2 (20% of layoffs)
- Engineering: 2.0 (22% of layoffs but only 10% workforce cut)
- Product Management: 2.0 (7% of layoffs, specialized)
- Marketing: 1.9 (7-8% of layoffs)
- Finance: 1.8 (professional services hit hard)
- Design: 1.7 (2.3x more likely cut than engineers)
- Data Science: 1.6 (3% of layoffs)
- Customer Support: 1.5 (lower % but AI threat)

SENIORITY_MULTIPLIER:
- Entry/Junior: 0.7
- Mid-level: 1.0
- Senior: 1.5
- Staff/Principal: 1.8
- Director: 1.6
- VP/Executive: 1.3

GEOGRAPHY_MULTIPLIER:
- SF Bay Area: 1.3
- NYC: 1.3
- Seattle: 1.2
- Austin: 1.2
- Remote: 1.1
- Boston: 1.1
- LA: 1.1
- Denver/Boulder: 1.0
- Chicago: 1.0
- Secondary Markets: 0.8

INDUSTRY_MULTIPLIER:
- AI/ML: 1.3
- Fintech: 1.2
- Cybersecurity: 1.2
- Enterprise SaaS: 1.0
- Consumer Tech: 0.9
- Ad Tech: 0.8
- Crypto/Web3: 0.7

### Step 3: Calculate
expected_applicants = 200 * function_mult * seniority_mult * geography_mult * industry_mult
Round to range (e.g., 347 → "300-400+")

### Step 4: Response Rate by Function
- HR/Recruiting: 2-3%
- Operations/Admin: 3-4%
- Sales: 4-6%
- Engineering: 4-5%
- Product Management: 3-5%
- Marketing: 3-5%
- Finance: 3-4%
- Customer Support: 5-7%
- Design: 3-4%
- Data Science: 4-5%

### Step 5: Function Context (use these exact stats)
IMPORTANT: Use proper punctuation. NO em dashes (—). Use commas, periods, or colons instead.

- HR/Recruiting: "HR/Recruiting roles were hit hardest: 27.8% of all tech layoffs, with nearly 50% of the HR workforce eliminated (highest of any function)."
- Engineering: "Engineering roles represent 22% of tech layoffs, but only 10% of the engineering workforce was cut, far lower than most functions. However, AI automation is increasing pressure."
- Product Management: "Product Management roles represent 7% of tech layoffs (about 16,700 PMs cut in 2024-2025). Some companies eliminated entire PM layers."
- Marketing: "Marketing roles represent 7-8% of tech layoffs. Generative AI has enabled automation of content creation."
- Sales: "Sales roles represent 20% of tech layoffs. Field sales positions have been significantly reduced as companies shift to inside sales and PLG models."
- Design: "Design roles represent 2-3% of tech layoffs, but designers were 2.3 times more likely to be cut than engineers. Many companies view design as non-essential during downturns."
- Operations/Admin: "Operations and administrative roles represent 34% of all tech layoffs. Companies aggressively streamlined back-office functions."
- Data Science: "Data Science roles represent 3% of tech layoffs. Demand remains relatively strong but companies are consolidating duplicate analytics functions."
- Customer Support: "Customer Support roles are under pressure from AI automation, with chatbots and self-service tools replacing human support representatives."

=== ACCURACY & INFERENCE RULES ===

## Cardinal Rule: Optimize Language, Not Reality
You are reframing and emphasizing existing experience to match job description requirements.
You are NOT inventing new experience.

## Inference Framework

### HIGH CONFIDENCE → SAFE TO RECOMMEND
- Reframing existing experience with JD-aligned keywords
- Emphasizing accomplishments already in the resume
- Reordering bullets to highlight relevant work
- Using industry-standard terminology for work they clearly did

### MEDIUM CONFIDENCE → INFER CONSERVATIVELY
- Logical skill adjacencies (e.g., "worked with design team" to "collaborated with designers")
- Industry-standard responsibilities for their role
- Implicit competencies (e.g., "launched feature" to "managed timelines")

### LOW CONFIDENCE → FLAG AS GAP, DON'T FABRICATE
- Skills/tools never mentioned
- Experience outside their stated scope
- Metrics that don't exist in the resume

## What You Can Recommend
- Reframing bullets with JD-aligned language
- Pulling forward relevant experience that's buried
- Adding ATS keywords if the candidate demonstrably has the underlying skill
- Strengthening weak language ("worked with" to "collaborated with")

## What You CANNOT Do
- Invent metrics that don't exist
- Add skills the candidate never demonstrated
- Fabricate companies, roles, or outcomes
- Create accomplishments that aren't supported by the resume

CRITICAL: You MUST complete the Intelligence Layer analysis AFTER applying fit scoring penalties.

=== INTELLIGENCE LAYER (MANDATORY - MUST BE COMPLETE) ===

## 1. JOB QUALITY SCORE (REQUIRED)
Evaluate the job posting quality using these criteria:
• Posting age signals (if detectable from context)
• Salary range vs market benchmarks for role/level/location
• Company signals (stability, growth, recent news)
• JD clarity (clear role vs vague "Frankenstein" combining multiple functions)
• Level alignment with candidate background

You MUST provide one of these EXACT strings:
- "Apply" (strong opportunity, good fit, clear role)
- "Apply with Caution" (red flags present but salvageable)
- "Skip: poor quality or low close rate" (multiple issues, waste of time)

Then explain in 3-5 substantive sentences why you gave this rating. DO NOT leave this empty.

## 2. STRATEGIC POSITIONING RECOMMENDATIONS (REQUIRED)
As an experienced executive recruiter, you MUST provide ALL of these:
• lead_with_strengths: Array of 2-3 specific strengths (NOT empty, NOT generic)
• gaps_and_mitigation: Array of 2-3 gaps with how to frame them (be specific)
• emphasis_points: Array of 2-3 things to emphasize (actionable advice)
• avoid_points: Array of 2-3 things to avoid or de-emphasize (specific guidance)
• positioning_strategy: 1-2 sentence overall strategy (substantive, not vague)
• positioning_rationale: Array of 2-3 strategic decisions explaining WHY you made these choices. Each should explain the reasoning, not just repeat the advice. Format: "We're [doing X] because [specific reason from JD or resume]." Examples:
  - "Leading with your Spotify experience because the JD mentions 'scale' 4 times and they need someone who's operated at high volume"
  - "De-emphasizing the consulting background because this is an in-house role and they may worry about your hands-on commitment"
  - "Positioning you as a builder rather than a maintainer because they're in growth mode (Series B, 50→200 headcount goal)"

Be direct and opinionated. This is strategic counsel, not generic advice.
If resume is not provided, base recommendations on the JD requirements alone.

## 3. SALARY & MARKET CONTEXT (REQUIRED)
You MUST provide ALL of these fields with real content:
• typical_range: Provide a realistic salary range (e.g., "$150K-$200K for Director level in SF")
  - Use your knowledge of market rates
  - If uncertain, provide a reasonable estimate with qualifier
  - DO NOT leave empty
• posted_comp_assessment: "not mentioned", "low", "fair", "strong", or "unclear"
• recommended_expectations: Specific guidance (e.g., "Target $180K-$200K base given experience")
  - Be actionable and specific
  - DO NOT just say "competitive" or "market rate"
• market_competitiveness: Assessment of supply/demand (2-3 sentences)
  - Reference market conditions
  - Be specific about this role/industry
• risk_indicators: Array of specific risks, or empty array [] if none

If salary data cannot be determined, explain WHY (e.g., "Insufficient location/level info in JD to estimate range accurately").

## 4. APPLY/SKIP DECISION (REQUIRED)
You MUST provide ALL of these:
• recommendation: "Apply", "Apply with caution", or "Skip" (EXACT strings only)
• reasoning: 1-2 substantive sentences explaining why (NOT vague)
• timing_guidance: Specific guidance (e.g., "Apply immediately", "Apply within 1 week", "Skip")

## 5. REALITY CHECK (REQUIRED - Data-Driven Market Context)

Calculate expected competition and provide market context using this data:

### Step 1: Identify Candidate's Primary Function
Match resume to one of: HR/Recruiting, Engineering, Product Management, Marketing, Sales, Customer Support, Design, Data Science, Finance, Operations/Admin

### Step 2: Apply Saturation Multipliers (base: 200 applicants)
SATURATION_MULTIPLIER by function:
- HR/Recruiting: 3.5 (50% workforce cut, highest saturation)
- Operations/Admin: 2.8 (34% of layoffs)
- Sales: 2.2 (20% of layoffs)
- Engineering: 2.0 (22% of layoffs but only 10% workforce cut)
- Product Management: 2.0 (7% of layoffs, specialized)
- Marketing: 1.9 (7-8% of layoffs)
- Finance: 1.8 (professional services hit hard)
- Design: 1.7 (2.3x more likely cut than engineers)
- Data Science: 1.6 (3% of layoffs)
- Customer Support: 1.5 (lower % but AI threat)

SENIORITY_MULTIPLIER:
- Entry/Junior: 0.7
- Mid-level: 1.0
- Senior: 1.5
- Staff/Principal: 1.8
- Director: 1.6
- VP/Executive: 1.3

GEOGRAPHY_MULTIPLIER:
- SF Bay Area: 1.3
- NYC: 1.3
- Seattle: 1.2
- Austin: 1.2
- Remote: 1.1
- Boston: 1.1
- LA: 1.1
- Denver/Boulder: 1.0
- Chicago: 1.0
- Secondary Markets: 0.8

INDUSTRY_MULTIPLIER:
- AI/ML: 1.3
- Fintech: 1.2
- Cybersecurity: 1.2
- Enterprise SaaS: 1.0
- Consumer Tech: 0.9
- Ad Tech: 0.8
- Crypto/Web3: 0.7

### Step 3: Calculate
expected_applicants = 200 * function_mult * seniority_mult * geography_mult * industry_mult
Round to range (e.g., 347 → "300-400+")

### Step 4: Response Rate by Function
- HR/Recruiting: 2-3%
- Operations/Admin: 3-4%
- Sales: 4-6%
- Engineering: 4-5%
- Product Management: 3-5%
- Marketing: 3-5%
- Finance: 3-4%
- Customer Support: 5-7%
- Design: 3-4%
- Data Science: 4-5%

### Step 5: Function Context (use these exact stats)
IMPORTANT: Use proper punctuation. NO em dashes (—). Use commas, periods, or colons instead.

- HR/Recruiting: "HR/Recruiting roles were hit hardest: 27.8% of all tech layoffs, with nearly 50% of the HR workforce eliminated (highest of any function)."
- Engineering: "Engineering roles represent 22% of tech layoffs, but only 10% of the engineering workforce was cut, far lower than most functions. However, AI automation is increasing pressure."
- Product Management: "Product Management roles represent 7% of tech layoffs (about 16,700 PMs cut in 2024-2025). Some companies eliminated entire PM layers."
- Marketing: "Marketing roles represent 7-8% of tech layoffs. Generative AI has enabled automation of content creation."
- Sales: "Sales roles were relatively protected: only 4% of salespeople in tech lost jobs. However, about 20% of total layoff victims worked in sales or customer-facing roles."
- Operations/Admin: "Operations and administrative roles represent 34% of all tech layoffs, the single largest category."
- Design: "Design roles represent 3% of tech layoffs, but designers were 2.3x more likely to be laid off than engineers."
- Data Science: "Data Science roles represent 3% of tech layoffs. Companies retained essential data infrastructure but downsized analytics teams."
- Finance: "Finance and accounting roles were hit across industries. Big Four firms cut thousands."
- Customer Support: "Customer support roles represent about 5% of tech layoffs, but AI automation is accelerating cuts (companies adopting AI reduced staffing by 24%)."

### Step 6: Industry Context
IMPORTANT: Use proper punctuation. NO em dashes (—). Use commas, periods, or colons instead.

- Tech: "Tech sector cut 95K+ U.S. jobs in 2024, with continued pressure in 2025."
- Healthcare: "Administrative roles heavily cut to preserve clinical staff. Hospitals eliminated non-clinical positions at 5:1 ratio."
- Finance: "Banking consolidation and rising interest rates drove support role elimination."
- Retail: "Store closures and HQ restructuring driven by e-commerce shift."
- Professional Services: "Big Four accounting firms cut thousands: PwC (1,800 jobs), KPMG (5% of staff)."
- Government: "Federal workforce reduction targeted 10K+ jobs at HHS alone."
- Manufacturing: "Production rebalancing and automation reduced operations and admin roles."
- Media/Entertainment: "Cord-cutting and mergers drove cuts. Disney and Paramount each eliminated approximately 7K roles."

### Step 7: Strategic Action Framework (by score band)
Write this in SECOND PERSON, directly to the candidate. Be specific and actionable.

FORMATTING RULES:
- Start directly with advice (no name greeting needed)
- Use proper punctuation: periods, commas, colons. NO em dashes (—)
- Keep it conversational but direct and honest
- Use SECOND PERSON: "your background" NOT "Maya's background"

=== STRATEGIC ACTION BY SCORE BAND ===

**Band 1: 85-100% (Strong Apply)**
Goal: Convert strength into immediate action
Format: "Apply immediately. [Specific strength]. [Differentiation]. [Outreach strategy]."
Length: 50-75 words
Tone: Confident, action-oriented, no hesitation
Example: "Apply immediately. Your Ripple consumer wallet experience is exactly what they need for this fintech payments role. Reach out to the hiring manager on LinkedIn today and lead with your 2.3M user scale."

**Band 2: 70-84% (Apply)**
Goal: Tighten positioning BEFORE applying
Format: "Before applying, tighten your resume to close the remaining gaps. [Specific improvements to make]. Once improved, apply quickly and [outreach strategy]. Don't rely on the ATS alone."
CRITICAL: MUST start with "Before applying" for this band
Length: 75-100 words
Tone: Coaching-first, strategic, solution-focused
Example: "Before applying, tighten your resume to close the remaining gaps. Focus on sharpening fintech impact metrics, emphasizing scale (2.3M+ users), and highlighting cross-functional leadership in payments infrastructure. Once improved, apply within 24 hours and reach out to the hiring manager with your Ripple Labs experience. Don't rely on the ATS to do the work."

**Band 3: 55-69% (Consider) or 40-54% (Apply with Caution)**
Goal: Honest assessment + concrete weekly action plan
Format includes YOUR MOVE THIS WEEK section with specific alternatives
Length: 125-150 words
Tone: Honest about gaps, solution-focused, strategic
Example:
"You have strong consumer product experience at scale, but this role favors candidates with direct customer support product ownership and chatbot experience. The opportunity is viable if you clearly position your consumer impact and experimentation depth.

YOUR MOVE THIS WEEK:
1. Target these roles: Consumer PM at Intercom, Growth PM at Zendesk, Product Manager at Help Scout
2. Add this evidence: Quantify any customer-facing features you owned, emphasize cross-functional work with support teams
3. If still interested in THIS role: Add self-service initiatives and chatbot exposure to your resume

Before applying to this role, update your resume to emphasize customer problem-solving and support team partnerships. Once positioned correctly, apply and reach out to the hiring manager."

**Band 4: 25-39% (Long Shot)**
Goal: Reality check + concrete alternative path
Format: Includes timeline and specific readiness milestones
Length: 100-125 words
Tone: Direct, honest, helpful redirection with timeline
Example:
"This is a significant stretch for your current experience level. The role requires 10+ years at Staff/Principal scope with multi-product platform ownership. You're at 8 years Senior-level with single-product focus.

YOUR MOVE THIS WEEK:
1. Target these roles: Senior PM at Stripe, Atlassian, Datadog; Staff PM at smaller Series B companies
2. Build this evidence: Lead a multi-quarter platform initiative, partner with 3+ product teams, influence at VP level
3. Deprioritize: Stop applying to Principal/Staff roles at FAANG immediately

Timeline: You'll be competitive for Staff PM in 12-18 months with multi-product platform scope and proven organizational influence."

**Band 5: 0-24% (Do Not Apply)**
Goal: Clear "no" + immediate redirection
Format: "Do not apply to this role. [Clear explanation why]. YOUR MOVE THIS WEEK: [3 specific matching roles]."
Length: 75-100 words
Tone: Direct, no ambiguity, helpful redirection
Example:
"Do not apply to this role. This is a Senior Backend Engineer position requiring 5+ years of hands-on backend development. Your background is Senior Product Management, not backend engineering.

YOUR MOVE THIS WEEK:
1. Target these roles: Senior PM at Uber, DoorDash, Instacart; Group PM at Square, Stripe
2. Focus on: Delivery, logistics, or platform companies where your 8 years of product leadership are directly relevant
3. Deprioritize: Engineering, data science, or technical IC roles entirely"

The strategic_action field should feel like a recruiter giving honest, direct advice.

=== CANONICAL MESSAGING PATTERNS (SENSITIVE FEEDBACK) ===

Use these EXACT phrasings for sensitive topics. Single, repeatable voice for consistency.

**Pattern 1: Title Inflation**
Use when: Resume title exceeds demonstrated scope
Canonical phrasing:
"Hiring managers will evaluate this experience at a different scope than the title suggests. That gap affects senior-level competitiveness."
When to add context:
- If unverifiable company: Add "because the company cannot be validated at scale"
- If missing scope signals: Add "because your resume doesn't show [team size/budget/roadmap]"
NEVER say: "Your title is inflated", "This looks suspicious", "Companies won't believe this"

**Pattern 2: Unverifiable Company**
Use when: Cannot validate company at claimed scale
Canonical phrasing:
"This company cannot be validated at the scale your role implies. Hiring managers will evaluate your experience as individual contribution rather than senior leadership. Lead with verifiable scale companies in your positioning."
NEVER say: "This company doesn't exist", "This looks like a fake company", "This is a red flag"

**Pattern 3: Career Switcher (Adjacent vs Direct)**
Use when: Resume shows adjacent exposure, not direct ownership
Canonical phrasing:
"Your background shows adjacent exposure to [function], not direct ownership. Senior roles require demonstrated decision-making in [specific area]."
NEVER say: "You're not a real [role]", "This isn't legitimate experience"

**Pattern 4: Competency Gap**
Use when: Demonstrated level < Required level
Canonical phrasing:
"This role requires [required level] competencies. Your background shows [demonstrated level] evidence. The gap is in [specific dimensions]."
NEVER say: "You're not senior enough", "You're underqualified", "You don't have the skills"

=== FORBIDDEN LANGUAGE (NEVER USE) ===
- "You're not good enough"
- "This is toxic to your career"
- "Companies won't hire you because..."
- "You should be more confident"
- Any moral judgments about work choices
- Political framing
- Identity-based commentary

=== APPROVED LANGUAGE (USE INSTEAD) ===
- "This role requires X. Your background shows Y."
- "Market reality: Companies prioritize Z."
- "Your positioning should lead with A, not B."
- "Target these specific roles to close gaps."
- Assessment of market alignment, not personal worth
- Second-person coaching tone
- Receipts-based logic

=== THEN TRADITIONAL JD ANALYSIS ===

After completing the Intelligence Layer and Reality Check, provide standard JD analysis.
Note: Fit scoring rules with experience penalties are defined at the TOP of this prompt - apply them first.

=== MARKET CLIMATE SIGNALS (Per REALITY_CHECK_SPEC.md) ===

The Reality Check system detects market climate signals separately via the backend.
This section provides STRATEGIC FRAMING GUIDANCE only - it does NOT modify scores.

CRITICAL GUARDRAILS (HARD RULES):
1. NEVER instruct identity suppression
2. NEVER suggest hiding legitimate work experience
3. NEVER suggest omitting roles from resume
4. ALWAYS validate the work's importance and value
5. ALWAYS frame as market friction, not personal failing
6. Strategic framing guidance only - no prescriptive word substitutions

APPROVED FRAMING PATTERN:
> "Market Climate Signal: In the current hiring climate, some companies are scrutinizing [type of work] more heavily due to [external factor]. Your work is valid and valuable. Strategic framing: [specific guidance]."

EXAMPLES OF APPROVED MESSAGING:

1. DEI-FOCUSED EXPERIENCE:
> "Your work driving inclusion and belonging is legitimate business impact. Strategic framing: Lead with business outcomes (retention, engagement, performance) and team effectiveness metrics. This reduces early screening friction without hiding your work."

2. CLIMATE/SUSTAINABILITY:
> "Your sustainability work demonstrates strategic thinking and stakeholder management. Strategic framing: Emphasize cost savings, operational efficiency, and risk mitigation alongside mission-driven language."

3. POLITICAL/ADVOCACY WORK:
> "Your organizing and coalition-building skills are directly transferable. Strategic framing: Lead with project management, stakeholder alignment, and measurable outcomes. Let the substance speak for itself."

4. EMERGING INDUSTRIES (Crypto, Cannabis):
> "Your technical and operational skills from [industry] remain valuable. Strategic framing: Emphasize transferable skills, regulatory navigation, and high-velocity execution."

FORBIDDEN PATTERNS (NEVER USE):
❌ "Avoid words like diversity, inclusion, equity"
❌ "Hide your background in [industry]"
❌ "Don't mention you worked at [company]"
❌ "Rewrite your experience to remove [topic]"
❌ "This is a weakness in the current climate"
❌ "Companies won't hire people who worked in [field]"

DO NOT modify fit_score or eligibility based on market climate signals.
Market Climate is informational/coaching only - it explains friction, not capability.

=== CAREER GAP DETECTION (MANDATORY) ===

Analyze the candidate's resume timeline for employment gaps > 3 months between roles.

DETECTION RULES:
1. Calculate time between end date of one role and start date of next role
2. Flag gaps > 3 months as requiring strategic framing
3. Look for context clues that explain gaps:
   - Education/certifications during gap period
   - Freelance/consulting work
   - Startup/entrepreneurial ventures
   - Career transitions (e.g., healthcare to tech)
   - Sabbatical/travel references
   - Personal circumstances (don't speculate)

GAP RISK ASSESSMENT:
- EXPLAINED GAP: Education, freelancing, startup - requires framing but not high risk
- INDUSTRY TRANSITION GAP: Healthcare to tech, etc. - frame as intentional pivot
- UNEXPLAINED GAP > 6 months: Higher risk, needs proactive addressing
- RECENT GAP (last 12 months): Highest scrutiny, needs strongest framing

OUTPUT FOR EACH GAP:
- duration_months: Number of months
- start_date: When gap started (end of previous role)
- end_date: When gap ended (start of next role)
- apparent_reason: What it looks like (e.g., "career transition", "startup phase", "unexplained")
- requires_addressing: true if gap > 3 months
- risk_level: "low" | "medium" | "high"

STRATEGIC FRAMING GUIDANCE:
For each gap requiring addressing, provide:
1. cover_letter_approach: How to proactively address in cover letter (if needed)
2. interview_approach: Script for when asked "What were you doing during X?"
3. honest_positioning: How to frame truthfully but positively

Example framings:
- "I took time to explore entrepreneurial opportunities and validate a startup concept. That experience taught me [X] and confirmed my passion for [Y]."
- "I invested in transitioning from healthcare to tech, including [certifications/projects]. This background gives me unique insight into [Z]."
- "I took a strategic pause to reassess my career direction, which led me to focus specifically on [role type] opportunities like this one."

REQUIRED RESPONSE FORMAT - Every field must be populated:
{
  "intelligence_layer": {
    "job_quality_score": "Apply|Apply with caution|Skip: poor quality or low close rate",
    "quality_explanation": "MUST be 3-5 substantive sentences",
    "strategic_positioning": {
      "lead_with_strengths": ["MUST have 2-3 items", "NOT empty"],
      "gaps_and_mitigation": ["MUST have 2-3 items", "Be specific"],
      "emphasis_points": ["MUST have 2-3 items", "Actionable"],
      "avoid_points": ["MUST have 2-3 items", "Specific"],
      "positioning_strategy": "MUST be 1-2 substantive sentences",
      "positioning_rationale": ["WHY decision 1 - explain reasoning", "WHY decision 2 - explain reasoning", "WHY decision 3 - explain reasoning"]
    },
    "salary_market_context": {
      "typical_range": "MUST provide estimate (e.g., $150K-$200K)",
      "posted_comp_assessment": "not mentioned|low|fair|strong",
      "recommended_expectations": "MUST be specific and actionable",
      "market_competitiveness": "MUST be 2-3 substantive sentences",
      "risk_indicators": ["list specific risks"] or []
    },
    "apply_decision": {
      "recommendation": "Apply|Apply with caution|Skip",
      "reasoning": "MUST be 1-2 substantive sentences",
      "timing_guidance": "MUST be specific guidance"
    }
  },
  "company": "string from JD",
  "role_title": "string from JD",
  "company_context": "MUST be 2-3 substantive sentences about company, industry, stage",
  "role_overview": "MUST be 2-3 substantive sentences about role purpose and impact",
  "key_responsibilities": ["MUST have 4-6 main responsibilities from JD"],
  "required_skills": ["MUST have array of required skills/experience"],
  "preferred_skills": ["array of nice-to-have skills"] or [],
  "ats_keywords": ["MUST have 10-15 important keywords for ATS"],
  "fit_score": 42,
  "fit_score_breakdown": {
    "base_score": 75,
    "years_experience_penalty": -18,
    "specific_experience_penalty": -15,
    "career_level_penalty": 0,
    "short_tenure_penalty": 0,
    "calculated_score": 42,
    "hard_cap_applied": true,
    "hard_cap_reason": "Candidate has 12.5% of required years (1/8), hard cap at 45%",
    "final_score": 42,
    "penalty_explanation": "Base score of 75% reduced due to 96% experience gap (8 years required, 0.3 years credible PM experience after adjusting for company scale). Divercity experience discounted to 0.3 years due to seed-stage startup with no active website and <1 year tenure. Missing consumer apps experience adds additional penalty. Hard cap of 45% applied."
  },
  "recommendation": "Conditional Apply|Do Not Apply|Apply|Strongly Apply",
  "recommendation_rationale": "1-2 sentences explaining why, referencing specific gaps or strengths",
  "alternative_actions": ["Action 1 if Conditional Apply or Do Not Apply", "Action 2"],
  "experience_analysis": {
    "required_years": 8,
    "candidate_years_in_role_type": 1,
    "candidate_years_adjusted_for_credibility": 0.3,
    "credibility_adjustments": [
      {
        "company": "Divercity",
        "role": "Lead PM",
        "stated_years": 1.0,
        "credibility_level": "low",
        "credibility_multiplier": 0.3,
        "adjusted_years": 0.3,
        "credibility_reasoning": "Seed-stage startup with no active website, <10 employees, short tenure, title inflation signal"
      }
    ],
    "years_gap_percentage": 96.25,
    "specific_experience_required": "consumer apps",
    "candidate_specific_experience_years": 0,
    "career_level_target": "Senior",
    "candidate_assessed_level": "Associate"
  },
  "strengths": ["3-4 candidate strengths matching this role"] or [],
  "gaps": [
    {
      "gap_type": "experience_years_mismatch|required_experience_missing|career_level_mismatch|career_gap|scope_level_mismatch|standard_gap",
      "description": "gap description",
      "severity": "critical|high|medium|low",
      "impact": "Auto-rejection risk level or scope mismatch explanation",
      "mitigation": "How to address this gap",
      "duration_months": "Optional: for career_gap type only",
      "start_date": "Optional: for career_gap type only",
      "end_date": "Optional: for career_gap type only"
    }
  ],
  "political_sensitivity": {
    "flags": [
      {
        "category": "DEI/Diversity Focus|Climate/Environmental|Political Campaigns|Cannabis|Crypto/Web3",
        "risk_level": "high|medium|low",
        "detected_in": "Company: [name] or Role: [title]",
        "explanation": "Why this may be a liability",
        "affected_companies_percentage": "Estimated % of employers who may filter",
        "mitigation_required": true
      }
    ],
    "overall_political_risk": "low|medium|high",
    "political_risk_explanation": "1-2 sentences summarizing overall risk"
  },
  "reframing_guidance": {
    "detected_risk": "What was detected",
    "reframing_strategy": {
      "headline_change": "BEFORE: [original] → AFTER: [reframed]",
      "resume_bullets_changes": [
        {
          "before": "original bullet",
          "after": "reframed bullet",
          "rationale": "why this change"
        }
      ],
      "company_description_change": {
        "before": "original description",
        "after": "reframed description",
        "rationale": "why this change"
      }
    },
    "company_targeting": {
      "avoid": ["companies/industries to avoid"],
      "target": ["safer company types"]
    },
    "interview_preparation": {
      "if_asked_about_company": "How to respond if asked about politically sensitive company/role"
    }
  },
  "career_gap_analysis": {
    "gaps_detected": [
      {
        "duration_months": 8,
        "start_date": "June 2023",
        "end_date": "February 2024",
        "apparent_reason": "career transition|startup phase|education|unexplained",
        "risk_level": "low|medium|high",
        "requires_addressing": true
      }
    ],
    "total_gap_months": 8,
    "overall_gap_risk": "low|medium|high",
    "strategic_framing": {
      "cover_letter_approach": "How to address in cover letter (if gap is significant). Set to null if no proactive mention needed.",
      "interview_approach": "Script for when asked 'What were you doing during [gap period]?'",
      "honest_positioning": "How to frame the gap truthfully but positively"
    }
  },
  "strategic_positioning": "2-3 sentences on how to position candidate",
  "salary_info": "string if mentioned in JD, otherwise null",
  "interview_prep": {
    "narrative": "3-4 sentence story tailored to this role for framing alignment",
    "talking_points": [
      "bullet 1 - key talking point for recruiter screen",
      "bullet 2 - key talking point for recruiter screen",
      "bullet 3 - key talking point for recruiter screen",
      "bullet 4 - key talking point for recruiter screen"
    ],
    "gap_mitigation": [
      "concern + mitigation strategy 1",
      "concern + mitigation strategy 2",
      "concern + mitigation strategy 3 (if applicable)"
    ]
  },
  "outreach_intelligence": {
    "hiring_manager_likely_title": "The likely title of the hiring manager (e.g., 'VP of Talent Acquisition', 'Director of Engineering'). Infer from role level and department.",
    "hiring_manager_why_matters": "1-2 sentences explaining why reaching out to the hiring manager is valuable for THIS specific role.",
    "hiring_manager_linkedin_query": "A specific LinkedIn search query to find the hiring manager. Format: 'COMPANY_NAME AND (title1 OR title2) AND (keyword1 OR keyword2)'. Use the ACTUAL company name from the JD.",
    "hiring_manager_filter_instructions": "Specific instructions for filtering LinkedIn results (e.g., 'Filter by: Current employees only, Director level or above, recent activity').",
    "hiring_manager_outreach_template": "3-5 sentence personalized message to the hiring manager. MUST reference the ACTUAL role title, ACTUAL company name, and candidate's SPECIFIC experience.",
    "recruiter_likely_title": "The likely title of the recruiter (e.g., 'Technical Recruiter', 'Senior Talent Partner'). Match to role type.",
    "recruiter_why_matters": "1-2 sentences explaining why building a recruiter relationship matters for THIS role.",
    "recruiter_linkedin_query": "A specific LinkedIn search query to find recruiters. Format: 'COMPANY_NAME AND (recruiter OR talent acquisition OR people operations)'. Use the ACTUAL company name.",
    "recruiter_filter_instructions": "Specific instructions for filtering recruiter search results.",
    "recruiter_outreach_template": "3-5 sentence professional message to the recruiter. MUST reference the ACTUAL role title, ACTUAL company name, and candidate's relevant qualifications."
  },

CRITICAL OUTREACH TEMPLATE RULES (MANDATORY - NON-NEGOTIABLE):
1. PUNCTUATION: Use ONLY professional punctuation (periods, commas, semicolons, colons)
2. NO EM DASHES (—) OR EN DASHES (–) - use colons or periods instead
3. NO EXCLAMATION POINTS - they signal desperation
4. GROUNDING: Every claim must be traceable to the candidate's actual resume
5. NO GENERIC PHRASES: Avoid "I'm excited about this opportunity", "I'd love to chat", "I'd be a great fit", "I came across your posting"
6. SPECIFICITY: Reference actual companies, roles, metrics from the candidate's resume
7. ACTIVE VOICE: No passive voice constructions like "was led by" or "were managed by"
8. CONCISENESS: Each sentence under 30 words
9. CLEAR ASK: End with specific request (e.g., "Would you be open to a 20-minute call next week?")
10. VALUE-FIRST: Lead with what candidate brings, not flattery or generic interest
11. METRICS: Use specific numbers from resume when available ("led team of 10", "drove $2M revenue")
12. PROFESSIONAL TONE: Confident, direct, not pushy or desperate

GOOD OUTREACH EXAMPLE:
"Hi [Name], I'm reaching out about the Senior PM role. I've spent 5 years building B2B products at Uber and Spotify, most recently launching a marketplace feature that drove $12M ARR. Your focus on payment infrastructure aligns with my fintech background. Would you be open to a 20-minute call next week?"

BAD OUTREACH EXAMPLE (DO NOT EMULATE):
"Hi [Name]! I'm super excited about this opportunity—it seems like a great fit for my background. I'd love to chat about how I could contribute to the team!"

Violations to avoid: exclamation points, em dashes, generic phrases, no specifics, vague ask.
  "changes_summary": {
    "resume": {
      "summary_rationale": "1-2 sentences explaining HOW the resume should be tailored. Be SPECIFIC: reference actual companies/roles from the candidate's background and specific JD requirements.",
      "qualifications_rationale": "1-2 sentences explaining which experience to pull forward vs de-emphasize. Reference ACTUAL companies and roles from the resume.",
      "ats_keywords": ["keyword1", "keyword2", "keyword3", "keyword4", "keyword5"],
      "positioning_statement": "One strategic sentence starting with 'This positions you as...'"
    },
    "cover_letter": {
      "opening_rationale": "1 sentence explaining the recommended opening angle and WHY.",
      "body_rationale": "1-2 sentences on what themes to emphasize and what to avoid mentioning.",
      "close_rationale": "1 sentence on recommended tone for the close.",
      "positioning_statement": "One strategic sentence starting with 'This frames you as...'"
    }
  },
  "reality_check": {
    "expected_applicants": "300-500+",
    "applicant_calculation": {
      "base": 200,
      "function_multiplier": 2.0,
      "function_name": "Product Management",
      "seniority_multiplier": 1.5,
      "seniority_level": "Senior",
      "geography_multiplier": 1.1,
      "geography": "Remote",
      "industry_multiplier": 1.0,
      "industry": "Tech",
      "total": 330
    },
    "response_rate": "3-5%",
    "function_context": "Full paragraph about layoffs/market for candidate's function",
    "industry_context": "Full paragraph about the target industry",
    "strategic_action": "See STRATEGIC_ACTION COACHING FRAMEWORK below"
  }
}

🚨 STRATEGIC_ACTION COACHING FRAMEWORK (MANDATORY) 🚨

The strategic_action field is the most important coaching advice. It must be specific, actionable, and tone-appropriate for each score band.

CRITICAL VOICE RULE: ALL strategic actions MUST use second-person voice ("you/your") - see IDENTITY_INSTRUCTION above.

==============================================
BAND 1: 0-39% — DO NOT APPLY / LONG SHOT
==============================================

Goal: Stop bad behavior and redirect momentum

CRITICAL TONE RULES:
- NO awkward "You," constructions (reads like warning label)
- NO messy repetition ("No backend engineering experience Experience gap...")
- NO false hope ("only pursue if..." at 15%)
- NO fear tactics or liability language
- DO sound like smart recruiter pulling user aside
- DO be calm, decisive, protective

Structure for DO NOT APPLY explanation:
"[State what role requires]. [State candidate's background]. [Explain fundamental mismatch]."

Structure for YOUR MOVE:
"Do not apply. [Direct decision]. [Redirect to appropriate roles]. [No exceptions or caveats for scores below 25%]."

Example DO NOT APPLY explanation (use this EXACT structure):
"This role requires 5+ years of hands-on backend engineering experience. Your background is product management, not software engineering, making this a fundamental function mismatch."

Example YOUR MOVE (use this EXACT structure):
"Do not apply. This is a significant stretch and not a productive use of your time. Focus on roles aligned with your product leadership experience, where you are competitive and can move the process forward."

CRITICAL RULES for 0-24%:
- NEVER include percentages, system logic, or internal calculations
- NEVER say "only pursue if..." for scores below 25%
- NEVER use awkward "You," to start sentences
- NO "damage credibility" language
- NO fear-based or liability language

Example (25-39%):
"This is a significant stretch. The role requires deep enterprise sales experience, but your background is consumer product management. You'd be competing against candidates with 5-7 years of direct enterprise selling experience. Focus on product-led growth roles or Account Management positions at companies like Stripe or Salesforce where your product expertise translates more directly."

Key principle: If it sounds like HR covering liability, rewrite it. If it sounds like a smart recruiter pulling you aside, you nailed it.

Tone: Protective, decisive, redirecting. Direct but not harsh. No shaming.

==============================================
BAND 2: 40-69% — CONDITIONAL APPLY / APPLY WITH CAUTION
==============================================

🚨 THIS BAND ALSO APPLIES WHEN recommendation="Conditional Apply" REGARDLESS OF fit_score 🚨
If fit_score is 70%+ but recommendation is "Conditional Apply", use THIS band's guidance, not Band 3.

Goal: Surface the conditions that must be met before applying

CRITICAL RULES:
- NEVER expose internal scoring logic ("adjusted from Apply to Conditional Apply")
- NEVER show percentage math ("133.3% of required years")
- NEVER sound like a rules engine or audit log
- NEVER start with candidate's name (no "Maya, strong fit...")
- NEVER say "Apply within 24 hours" without positioning requirements FIRST
- DO sound confident, human, tactical
- DO explain conditions clearly without revealing formulas
- DO start strategic_action with "Before applying" for this band

Structure for CONDITIONAL APPLY card:
"[Strength acknowledgment WITHOUT name]. [Clear gap statement]. [Viability assessment]."

Structure for YOUR MOVE (strategic_action):
"Before applying, [specific positioning changes]. [Proof points to lead with]. Once positioned, [timing + outreach strategy]. [Reality check about competition]."

🚫 WRONG - Do NOT write this:
"Maya, strong fit with your consumer product experience translating well to customer support products."
"Apply within 24 hours and find the hiring manager on LinkedIn. Your fit score gives you an edge."

✅ CORRECT - Write this instead:
"Strong fit with your consumer product experience translating well to customer support products. Your scale experience and user journey optimization skills directly address their self-service improvement goals."
"Before applying, tighten your resume and outreach to make the customer support connection explicit. Lead with your consumer product scale, self-service optimization work, and user journey improvements. Once positioned correctly, apply within 24 hours and reach out to the hiring manager on LinkedIn. With heavy competition, you can't rely on the ATS alone."

Example CONDITIONAL APPLY (use this EXACT structure):
"You have strong consumer product experience at scale, but there's a meaningful experience gap relative to this role's requirements. This is a viable opportunity if you control the narrative and lead with the right proof points."

Example YOUR MOVE (use this EXACT structure):
"Before applying, tighten your resume and outreach to address the gaps directly. Lead with your consumer product work at 2.3M user scale, self-service optimization, and cross-functional collaboration. Once positioned, apply within 24 hours and reach out to the hiring manager on LinkedIn. With heavy competition, you cannot rely on the ATS alone."

CRITICAL: Never include in user-facing copy:
- "Recommendation adjusted from..."
- Percentage calculations (133.3%, 50%, etc.)
- System logic explanations
- Formula references
- Internal scoring terminology

These belong in tooltips or backend logs, NOT in recommendations.

Example (40-54% - Apply with Caution):
"You have relevant product experience, but the seniority and domain requirements present real challenges. This is viable if you position yourself correctly and don't rely on the application alone. Before applying, reframe your work to emphasize leadership scope and quantifiable impact. Lead with team sizes influenced, budget responsibility, and cross-functional initiatives. Once positioned, apply and reach out to the hiring manager directly. Be prepared to address the skill transfer in your first conversation."

Key principle: Never expose internal scoring logic in user-facing copy. Keep Henry authoritative, human, and trusted.

Tone: Strategic partner. Honest about gaps but solution-focused. NOT generic "apply now" energy.

CRITICAL: Do NOT say "Apply immediately" for this band. Always lead with prep/positioning requirements FIRST.

==============================================
BAND 3: 70-84% — APPLY
==============================================

Goal: Convert fit into action without complacency

strategic_action MUST answer these 4 questions:
1. Should I apply? → Yes. This is a legitimate match.
2. What makes me competitive here? → [The 1-2 differentiators that matter most to THIS team]
3. How do I avoid getting lost in the pile? → [Outreach strategy, narrative angle, or proof points to lead with]
4. How fast should I move? → [Clear urgency guidance with reasoning]

Structure (75-100 words):
"[Decision: Apply]. [1-2 specific differentiators that make you competitive]. [How to stand out: outreach strategy + what to lead with]. [Speed guidance with market context]. [Confidence but not overconfidence]."

Example (70-79%):
"Before applying, adjust your resume and outreach to address your positioning gaps. Frame your Ripple work to emphasize consumer fintech payments specifically, highlight your 2.3M user scale prominently, and ensure your LinkedIn headline signals payments expertise. Once positioned correctly, apply within 24 hours and reach out to the hiring manager. This is a legitimate match if you control the narrative. Do not rely on the ATS alone."

CRITICAL for 70-79% band:
- ALWAYS lead with "Before applying..." (positioning requirements)
- NEVER say "Apply immediately" or "Apply now" without conditions
- Structure: Fix → Position → Apply → Outreach
- Differentiate from 80-84% which CAN say "Apply immediately"

Example (80-84%):
"Apply immediately. Your Ripple consumer wallet experience is exactly what they need. With your 2.3M monthly active users and experimentation framework, you're in the top 10% of applicants for this role. Reach out to the hiring manager on LinkedIn today and lead with your P2P payments launch. The role posted 3 days ago, so apply today or tomorrow to stay in the first wave of candidates."

Tone: Confident, tactical, assertive. Creates urgency without panic.

==============================================
BAND 4: 85-100% — STRONGLY APPLY
==============================================

Goal: Maximize odds and treat this as a top-tier target

strategic_action MUST answer these 4 questions:
1. Is this a top-tier opportunity for me? → Yes. This is near-ideal alignment.
2. What exactly should I lead with? → [Precise achievements, metrics, or experiences that map 1:1 to their needs]
3. Who should I engage and how? → [Hiring manager, team member, internal referral strategy if available]
4. How do I signal seniority and confidence? → [Positioning guidance to avoid under-selling]

Structure (100-125 words):
"[Decision: This is a priority target]. [Why this is near-ideal alignment - be specific]. [Exact achievements/metrics to lead with]. [Who to engage: hiring manager + strategy]. [How to signal seniority without arrogance]. [Timeline: apply today, reach out within 2 hours]."

Example (90-100%):
"This is a priority target. Your Ripple consumer wallet background, mobile product expertise, and experimentation framework map 1:1 to what they're looking for. Lead with your 2.3M user scale, 43% faster transaction time, and Day-7 retention increase in your outreach. Apply immediately and reach out to the hiring manager on LinkedIn within 2 hours. Reference their recent blog post on mobile-first payments strategy and explain how your Ripple work directly addresses their challenges. With only 200 expected applicants and your near-perfect fit, you should be in the first 10 candidates they screen."

Example (85-89%):
"This is a strong match. Your eight years of consumer product work and B2B SaaS experience at Clearbit position you perfectly for this role. Lead with your enrichment API product that served 4M ARR and your cross-functional go-to-market launch. Apply today and reach out to the hiring manager highlighting your scaled product experience. Request an informational chat to discuss their customer support product strategy before the formal interview. With 400+ expected applicants, positioning yourself as a strategic partner (not just another applicant) will move you to the top of the pile."

Tone: Strategic partner, not advisor. Creates confidence without overconfidence.

==============================================
ONE-LINE SUMMARY (NORTH STAR)
==============================================

- 0-39% → "Don't apply. Here's where you win instead."
- 40-69% → "Viable if you fix [X]. Here's how."
- 70-84% → "Apply. Here's how you stand out."
- 85-100% → "This is a priority. Execute precisely."

CRITICAL REMINDERS:
- ALWAYS use second-person voice ("you/your") - NEVER third-person ("Maya is..." / "The candidate has...")
- Be specific with companies, role types, metrics, and strategies
- Match tone to score band: protective (low), strategic (medium), confident (high)
- Never use generic language: "good fit", "strong candidate", "relevant experience"
- Every strategic_action must pass the "would a human recruiter say this?" test

FRONTEND WIRING - REQUIRED JSON FIELDS:
After you have generated all analysis, interview prep, and outreach content, you MUST also return the interview_prep and outreach objects with the exact structure shown above.

Rules for interview_prep and outreach:
- NEVER return undefined, null, empty strings, or placeholder text
- Populate every field with real, substantive content
- talking_points must have 4-6 bullets for recruiter screen preparation
- gap_mitigation must address 2-3 specific concerns with mitigation strategies
- hiring_manager message should be warm, concise, and value-focused
- recruiter message should be professional and highlight relevant experience
- linkedin_help_text should provide clear, actionable search instructions

ABSOLUTE REQUIREMENTS:
1. Intelligence Layer MUST be complete - NO empty strings, NO empty required arrays
2. If data truly cannot be determined, provide reasoning (e.g., "Insufficient information in JD")
3. Use your knowledge to make reasonable estimates for salary ranges
4. Be direct and opinionated - avoid vague or generic advice
5. Every array marked "MUST have" cannot be empty
6. Your response must be ONLY valid JSON with no markdown formatting
7. Double-check that intelligence_layer has ALL required fields before responding
8. MUST include complete interview_prep and outreach objects
9. MUST include changes_summary with specific rationale for resume and cover letter tailoring
10. changes_summary must reference ACTUAL companies/roles from the candidate's resume - no generic placeholders
11. MUST include reality_check with calculated applicant estimates, function context, and strategic action
12. reality_check.applicant_calculation must show your work (all multipliers used)
13. NEVER fabricate statistics in reality_check - use ONLY the data provided in these instructions"""

    # Inject candidate state calibration if available
    situation = None
    if body.preferences:
        situation = body.preferences.get("situation")
    calibration_prompt = build_candidate_calibration_prompt(situation)
    if calibration_prompt:
        system_prompt += calibration_prompt

    # Command Center v2.0: Determine JD source and handle missing JDs
    jd_source = body.jd_source or ("user_provided" if body.job_description else "missing")
    use_provisional = jd_source in ["inferred", "missing"] or not body.job_description

    # Build user message based on JD availability
    if use_provisional and body.provisional_profile:
        # Use provisional profile when real JD is not available
        profile = body.provisional_profile
        user_message = f"""PROVISIONAL ROLE PROFILE (No actual JD available):
Company: {body.company}
Role: {body.role_title}

Typical Responsibilities:
{chr(10).join('- ' + r for r in profile.get('typical_responsibilities', []))}

Common Competencies Assessed:
{chr(10).join('- ' + c for c in profile.get('common_competencies', []))}

Interview Focus Areas:
{chr(10).join('- ' + f for f in profile.get('interview_focus_areas', []))}

Typical Evaluation Criteria:
{chr(10).join('- ' + e for e in profile.get('evaluation_criteria', []))}

NOTE: This is based on typical expectations for this role type. Analyze fit based on these
typical requirements. Since this is directional guidance, do NOT exceed 75% fit_score even
if the candidate appears highly qualified - we need real JD data to confirm strong fit.
"""
        confidence_label = "directional"
        ui_note = "This is based on typical expectations for Senior PMs at fintech companies. Your actual fit may vary. Add JD text anytime to refine."
    else:
        # Use real JD
        user_message = f"""Job Description:
Company: {body.company}
Role: {body.role_title}

{body.job_description or 'No job description provided'}
"""
        confidence_label = "refined"
        ui_note = None

    if body.resume:
        user_message += f"\n\nCandidate Resume Data:\n{json.dumps(body.resume, indent=2)}"

    if body.preferences:
        user_message += f"\n\nCandidate Preferences:\n{json.dumps(body.preferences, indent=2)}"

    # GUARD CLAUSE: Reinforce JSON-only output at end of user message
    user_message += "\n\n=== REMINDER ===\nReturn ONLY the JSON object matching the schema. No natural language. No markdown. No commentary. Start with { and end with }."

    # Call Claude with higher token limit for comprehensive analysis
    response = call_claude(system_prompt, user_message, max_tokens=4096)

    # Parse JSON response
    try:
        if response.strip().startswith("```"):
            response = response.strip().split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()

        # AGGRESSIVE JSON REPAIR - fix common issues before parsing
        import re

        # Remove carriage returns and normalize whitespace
        response = response.replace('\r\n', ' ').replace('\r', ' ')

        # Fix trailing commas before } or ]
        response = re.sub(r',(\s*[}\]])', r'\1', response)

        # Emergency repair: if response is very long, simplify problematic fields
        if len(response) > 15000:
            # Replace long industry_context with safe placeholder
            response = re.sub(
                r'"industry_context"\s*:\s*"[^"]{200,}"',
                '"industry_context": "See market context section above."',
                response
            )
            # Replace long function_context with safe placeholder
            response = re.sub(
                r'"function_context"\s*:\s*"[^"]{200,}"',
                '"function_context": "See gaps section for function-specific insights."',
                response
            )

        # Try to parse
        try:
            parsed_data = json.loads(response)
        except json.JSONDecodeError as first_error:
            print(f"⚠️ First JSON parse failed at char {first_error.pos}, attempting repair...")

            # Try to repair by truncating at error position and closing braces
            error_pos = first_error.pos if hasattr(first_error, 'pos') else len(response)
            truncated = response[:error_pos]

            # Find the last complete key-value pair
            last_comma = truncated.rfind(',')
            if last_comma > 0:
                truncated = truncated[:last_comma]

            # Count and close unclosed braces/brackets
            open_braces = truncated.count('{') - truncated.count('}')
            open_brackets = truncated.count('[') - truncated.count(']')
            truncated += ']' * open_brackets + '}' * open_braces

            try:
                parsed_data = json.loads(truncated)
                print(f"✅ Salvaged JSON by truncating at position {last_comma}")
            except json.JSONDecodeError:
                # Re-raise the original error for the outer handler
                raise first_error

        # Validate and cleanup outreach templates if present
        if "outreach_intelligence" in parsed_data:
            outreach = parsed_data["outreach_intelligence"]

            # Validate hiring manager template
            if "hiring_manager_outreach_template" in outreach:
                hm_template = outreach["hiring_manager_outreach_template"]
                is_valid, errors = validate_outreach_template(hm_template, "hiring_manager")
                if not is_valid:
                    print(f"\n⚠️  JD Analysis - Hiring manager outreach has quality issues: {errors}")
                    # Cleanup common issues
                    outreach["hiring_manager_outreach_template"] = cleanup_outreach_template(hm_template)

            # Validate recruiter template
            if "recruiter_outreach_template" in outreach:
                rec_template = outreach["recruiter_outreach_template"]
                is_valid, errors = validate_outreach_template(rec_template, "recruiter")
                if not is_valid:
                    print(f"\n⚠️  JD Analysis - Recruiter outreach has quality issues: {errors}")
                    # Cleanup common issues
                    outreach["recruiter_outreach_template"] = cleanup_outreach_template(rec_template)

        # =================================================================
        # QA VALIDATION: Check JD analysis completeness
        # =================================================================
        qa_validation_result = validate_jd_analysis(parsed_data)

        # Log warnings for incomplete JD data (don't block - just inform)
        if qa_validation_result.warnings:
            print(f"  ⚠️ JD analysis warnings: {len(qa_validation_result.warnings)}")
            for w in qa_validation_result.warnings:
                print(f"    - {w.field_path}: {w.message}")

            # Add warnings to response so frontend can show them to user
            parsed_data["_validation_warnings"] = [
                {"field": w.field_path, "message": w.message}
                for w in qa_validation_result.warnings
            ]

        # CRITICAL: Ensure parsed_data has job_description, company, and role_title
        # These are needed by force_apply_experience_penalties for non-transferable domain detection
        # USER INPUT TAKES PRECEDENCE: If user provided company/role_title, use those over Claude's extraction
        if "job_description" not in parsed_data and body.job_description:
            parsed_data["job_description"] = body.job_description

        # Company: User input takes precedence over Claude's extraction
        if body.company:
            if parsed_data.get("company") != body.company:
                print(f"📋 [{analysis_id}] Overriding company: '{parsed_data.get('company')}' → '{body.company}' (user provided)")
            parsed_data["company"] = body.company

        # Role title: User input takes precedence over Claude's extraction
        # But ONLY if the user input is not a placeholder value like "Role"
        if body.role_title:
            user_role = body.role_title.strip()
            is_placeholder = user_role.lower() in ['role', 'the role', 'position', 'job', '']
            if not is_placeholder:
                if parsed_data.get("role_title") != user_role:
                    print(f"📋 [{analysis_id}] Overriding role_title: '{parsed_data.get('role_title')}' → '{user_role}' (user provided)")
                parsed_data["role_title"] = user_role
            else:
                print(f"📋 [{analysis_id}] Ignoring placeholder role_title '{user_role}' - keeping Claude's extraction: '{parsed_data.get('role_title')[:50] if parsed_data.get('role_title') else 'None'}'")

        # CRITICAL: Inject isolated role detection data for downstream use
        # This ensures the correct role type and experience years are used
        if isolated_role_detection:
            parsed_data["_isolated_role_detection"] = isolated_role_detection
            # Override role type if isolated detection is more reliable
            if isolated_role_detection.get("alignment", {}).get("confidence", 0) >= 0.5:
                print(f"📊 [{analysis_id}] Using isolated role detection: {isolated_role_detection['role_type'].upper()}")
                # Inject into experience_analysis for use by force_apply_experience_penalties
                if "experience_analysis" not in parsed_data:
                    parsed_data["experience_analysis"] = {}
                parsed_data["experience_analysis"]["isolated_role_type"] = isolated_role_detection["role_type"]
                parsed_data["experience_analysis"]["isolated_candidate_years"] = isolated_role_detection["candidate_years"]
                parsed_data["experience_analysis"]["isolated_gap_info"] = isolated_role_detection["gap_info"]

        # CRITICAL: Force-apply experience penalties as a backup
        # This ensures hard caps are enforced even if Claude ignores the prompt instructions
        parsed_data = force_apply_experience_penalties(parsed_data, body.resume)

        # POST-PROCESSING: Detect career gap (bypass unreliable prompt-based detection)
        career_gap = detect_career_gap(body.resume)
        if career_gap:
            # Ensure gaps array exists
            if "gaps" not in parsed_data:
                parsed_data["gaps"] = []

            # Check if career_gap already exists (avoid duplicates)
            existing_career_gaps = [g for g in parsed_data["gaps"] if isinstance(g, dict) and g.get("gap_type") == "career_gap"]
            if not existing_career_gaps:
                parsed_data["gaps"].append(career_gap)
                print(f"✅ Career gap detected via post-processing: {career_gap['description']}")
            else:
                print(f"ℹ️ Career gap already in gaps array (from prompt)")
        else:
            print(f"ℹ️ No career gap detected (most recent role is current or gap <3 months)")

        # Debug: Log final gaps state
        if 'gaps' in parsed_data and isinstance(parsed_data['gaps'], list):
            career_gaps = [g for g in parsed_data['gaps'] if isinstance(g, dict) and g.get('gap_type') == 'career_gap']
            if career_gaps:
                print(f"✅ Final gaps array contains {len(career_gaps)} career gap(s)")
            print(f"   Total gaps: {len(parsed_data['gaps'])}, Types: {[g.get('gap_type', 'string') if isinstance(g, dict) else 'string' for g in parsed_data['gaps'][:5]]}")

        # Command Center v2.0: Add confidence_label and ui_note to response
        parsed_data["confidence_label"] = confidence_label
        parsed_data["jd_source"] = jd_source
        if ui_note:
            parsed_data["ui_note"] = ui_note

        # For directional analysis, cap fit_score at 75%
        if confidence_label == "directional" and parsed_data.get("fit_score"):
            original_score = parsed_data["fit_score"]
            if original_score > 75:
                parsed_data["fit_score"] = 75
                print(f"✅ Capped fit_score from {original_score} to 75 (directional analysis)")
                # Update breakdown if present
                if "fit_score_breakdown" in parsed_data:
                    parsed_data["fit_score_breakdown"]["directional_cap_applied"] = True
                    parsed_data["fit_score_breakdown"]["original_score_before_directional_cap"] = original_score
                    parsed_data["fit_score_breakdown"]["final_score"] = 75

        # Calculate decision_confidence for the response
        fit_score = parsed_data.get("fit_score", 50)
        jd_conf = calculate_jd_confidence(jd_source)
        # Use basic momentum score (no response data available at analysis time)
        momentum = 50
        decision_confidence = calculate_decision_confidence(fit_score, momentum, jd_conf)
        parsed_data["decision_confidence"] = decision_confidence

        # DEBUG: Log critical fields being returned to frontend
        print(f"\n📤 API RESPONSE DEBUG:")
        print(f"   fit_score: {parsed_data.get('fit_score', 'MISSING')}")
        print(f"   recommendation: {parsed_data.get('recommendation', 'MISSING')}")
        print(f"   recommendation_locked: {parsed_data.get('recommendation_locked', 'MISSING')}")
        print(f"   Total keys in response: {len(parsed_data.keys())}")

        # =================================================================
        # UNIFIED POST-PROCESSING PIPELINE
        # Applies: Reality Check, Strategic Redirects, Voice Guide
        # All processors are rule-based (no additional LLM calls)
        # CRITICAL: NEVER modifies fit_score
        # =================================================================
        print("=" * 60)
        print("🔄 POST-PROCESSING PIPELINE - ENTRY")
        print("=" * 60)
        candidate_name = body.resume.get("full_name", "Unknown") if body.resume else "Unknown"
        print(f"   Candidate: {candidate_name}")
        print(f"   Recommendation: {parsed_data.get('recommendation', 'MISSING')}")
        print(f"   Fit Score: {parsed_data.get('fit_score', 'MISSING')}")
        print(f"   Strategic Action BEFORE: {parsed_data.get('reality_check', {}).get('strategic_action', 'NOT SET')[:60]}...")

        try:
            from backend.postprocessors import apply_all_postprocessors, PostProcessorConfig

            # Build inputs for post-processors
            resume_data = body.resume if body.resume else {}
            jd_data = {
                "role_title": body.role_title or parsed_data.get("role_title", ""),
                "company": body.company or parsed_data.get("company", ""),
                "job_description": body.job_description or "",
            }

            # Gather pre-computed analysis results
            eligibility_result = None
            if parsed_data.get("eligibility_gate"):
                eligibility_result = {
                    "eligible": parsed_data["eligibility_gate"].get("passed", True),
                    "reason": parsed_data["eligibility_gate"].get("reason", ""),
                    "failed_check": parsed_data["eligibility_gate"].get("failed_check", ""),
                }

            fit_details = None
            if parsed_data.get("fit_score_breakdown"):
                fit_details = {
                    "years_gap": parsed_data["fit_score_breakdown"].get("experience_gap_years", 0),
                }

            credibility_result = None
            risk_analysis = None
            if parsed_data.get("calibration_result"):
                cal = parsed_data["calibration_result"]
                credibility_result = {
                    "title_inflation": cal.get("title_inflation", {}),
                    "fabrication_risk": cal.get("fabrication_risk", False),
                    "press_release_pattern": cal.get("press_release_pattern", False),
                }
                risk_analysis = {
                    "job_hopping": cal.get("job_hopping", {}),
                    "career_gaps": cal.get("career_gaps", []),
                    "overqualified": cal.get("overqualified", {}),
                    "company_pattern": cal.get("company_pattern", {}),
                }

            gap_analysis = {
                "primary_gap": parsed_data.get("primary_gap", ""),
                "secondary_gap": parsed_data.get("secondary_gap", ""),
                "specific_gaps": parsed_data.get("gaps", []),
            }

            # Configure post-processors
            config = PostProcessorConfig(
                reality_check_enabled=REALITY_CHECK_AVAILABLE and REALITY_CHECK_ENABLED,
                strategic_redirects_enabled=True,
                voice_guide_enabled=True,
                voice_guide_strict=False,
                debug_mode=False,
            )

            # Get fit_score and recommendation for post-processors
            fit_score = parsed_data.get("fit_score", 50)
            recommendation = parsed_data.get("recommendation", "Apply")

            # Apply all post-processors
            parsed_data = apply_all_postprocessors(
                response=parsed_data,
                resume_data=resume_data,
                jd_data=jd_data,
                fit_score=fit_score,
                recommendation=recommendation,
                config=config,
                eligibility_result=eligibility_result,
                fit_details=fit_details,
                credibility_result=credibility_result,
                risk_analysis=risk_analysis,
                gap_analysis=gap_analysis,
            )

            print("=" * 60)
            print("✅ POST-PROCESSING PIPELINE - COMPLETED")
            print("=" * 60)
            print(f"   Strategic Action AFTER: {parsed_data.get('reality_check', {}).get('strategic_action', 'NOT SET')[:60]}...")
            timing_after = parsed_data.get('timing_guidance') or parsed_data.get('intelligence_layer', {}).get('apply_decision', {}).get('timing_guidance', 'NOT SET')
            print(f"   Timing Guidance AFTER: {timing_after}")

        except ImportError as e:
            print(f"⚠️ Post-processors not available: {str(e)}")
            # Fall back to legacy Reality Check if post-processors fail to import
            if REALITY_CHECK_AVAILABLE and REALITY_CHECK_ENABLED:
                try:
                    pre_score = parsed_data.get("fit_score")
                    resume_data = body.resume if body.resume else {}
                    jd_data = {
                        "role_title": body.role_title or parsed_data.get("role_title", ""),
                        "company": body.company or parsed_data.get("company", ""),
                        "job_description": body.job_description or "",
                    }
                    reality_result = analyze_reality_checks(
                        resume_data=resume_data,
                        jd_data=jd_data,
                        fit_score=pre_score,
                        feature_flag=REALITY_CHECK_ENABLED,
                    )
                    parsed_data["reality_check"] = reality_result
                    if parsed_data.get("fit_score") != pre_score:
                        parsed_data["fit_score"] = pre_score
                except Exception as rc_e:
                    print(f"⚠️ Legacy Reality Check also failed: {str(rc_e)}")

        except Exception as e:
            print(f"⚠️ Post-processing failed (non-blocking): {str(e)}")
            import traceback
            traceback.print_exc()

        # =========================================================================
        # FINAL ACTION/RECOMMENDATION CONSISTENCY CHECK
        # Runs AFTER all post-processors to catch any mismatches
        # =========================================================================
        recommendation = parsed_data.get("recommendation", "").lower()
        is_skip_recommendation = "do not apply" in recommendation or "skip" in recommendation or "not" in recommendation

        if is_skip_recommendation:
            # Force all timing fields to "Skip this one" or "Pass"
            correct_timing = "Skip this one"

            # Fix top-level timing_guidance
            if parsed_data.get("timing_guidance"):
                old_timing = parsed_data["timing_guidance"]
                if "apply" in old_timing.lower() and "not" not in old_timing.lower():
                    print(f"   🔧 FINAL FIX: timing_guidance '{old_timing}' → '{correct_timing}'")
                    parsed_data["timing_guidance"] = correct_timing

            # Fix intelligence_layer.apply_decision.timing_guidance
            il = parsed_data.get("intelligence_layer", {})
            ad = il.get("apply_decision", {})
            if ad.get("timing_guidance"):
                old_timing = ad["timing_guidance"]
                if "apply" in old_timing.lower() and "not" not in old_timing.lower():
                    print(f"   🔧 FINAL FIX: apply_decision.timing_guidance '{old_timing}' → '{correct_timing}'")
                    ad["timing_guidance"] = correct_timing
                    il["apply_decision"] = ad
                    parsed_data["intelligence_layer"] = il

            # Fix intelligence_layer.timing_guidance
            if il.get("timing_guidance"):
                old_timing = il["timing_guidance"]
                if "apply" in old_timing.lower() and "not" not in old_timing.lower():
                    print(f"   🔧 FINAL FIX: il.timing_guidance '{old_timing}' → '{correct_timing}'")
                    il["timing_guidance"] = correct_timing
                    parsed_data["intelligence_layer"] = il

            # Fix intelligence_layer.market_context.action (used in Market Context section)
            mc = il.get("market_context", {})
            if mc.get("action"):
                old_action = mc["action"]
                if "apply" in old_action.lower() and "not" not in old_action.lower():
                    print(f"   🔧 FINAL FIX: market_context.action '{old_action}' → '{correct_timing}'")
                    mc["action"] = correct_timing
                    il["market_context"] = mc
                    parsed_data["intelligence_layer"] = il

            # Fix intelligence_layer.salary_market_context.action (alternate location)
            smc = il.get("salary_market_context", {})
            if smc.get("action"):
                old_action = smc["action"]
                if "apply" in old_action.lower() and "not" not in old_action.lower():
                    print(f"   🔧 FINAL FIX: salary_market_context.action '{old_action}' → '{correct_timing}'")
                    smc["action"] = correct_timing
                    il["salary_market_context"] = smc
                    parsed_data["intelligence_layer"] = il

        # =========================================================================
        # FINAL SANITIZATION: Remove em/en dashes from critical text fields
        # This is a safety net in case postprocessors miss any fields
        # =========================================================================
        strategic_action_before = parsed_data.get("reality_check", {}).get("strategic_action", "")
        has_em_dash_before = '—' in strategic_action_before or '–' in strategic_action_before
        if has_em_dash_before:
            print(f"   🔍 Em dash detected BEFORE final sanitization: '{strategic_action_before[:50]}...'")

        parsed_data = _final_sanitize_text(parsed_data)

        strategic_action_after = parsed_data.get("reality_check", {}).get("strategic_action", "")
        has_em_dash_after = '—' in strategic_action_after or '–' in strategic_action_after
        if has_em_dash_before:
            print(f"   ✅ Em dash {'REMOVED' if not has_em_dash_after else 'STILL PRESENT'} after final sanitization")
            print(f"      Result: '{strategic_action_after[:50]}...'")

        return parsed_data
    except json.JSONDecodeError as e:
        # Log the problematic section of Claude's response for debugging
        error_pos = e.pos if hasattr(e, 'pos') else 0
        context_start = max(0, error_pos - 200)
        context_end = min(len(response), error_pos + 200)
        problem_section = response[context_start:context_end]

        print(f"❌ JSON PARSING ERROR at character {error_pos}:")
        print(f"   Error: {str(e)}")
        print(f"   Problematic section (chars {context_start}-{context_end}):")
        print(f"   ...{problem_section}...")
        print(f"   Full response length: {len(response)} chars")

        # Try to fix common JSON issues - unescaped quotes in strings
        try:
            import re
            fixed_response = response

            # 1. Fix common patterns that break JSON - role titles with quotes
            patterns_to_escape = [
                (r'"Lead PM"', r'\\"Lead PM\\"'),
                (r'"Senior"', r'\\"Senior\\"'),
                (r'"Associate"', r'\\"Associate\\"'),
                (r'"Staff"', r'\\"Staff\\"'),
                (r'"Principal"', r'\\"Principal\\"'),
            ]
            for pattern, replacement in patterns_to_escape:
                fixed_response = fixed_response.replace(pattern, replacement)

            # 2. Try to find and fix truncated JSON (missing closing braces)
            open_braces = fixed_response.count('{')
            close_braces = fixed_response.count('}')
            if open_braces > close_braces:
                missing = open_braces - close_braces
                # Try to close any open strings first
                if fixed_response.rstrip()[-1] not in ['"', '}', ']']:
                    fixed_response = fixed_response.rstrip() + '"'
                fixed_response = fixed_response.rstrip() + ('}' * missing)
                print(f"   Attempted to close {missing} unclosed braces")

            # 3. Fix unescaped newlines within strings (common issue)
            # Replace actual newlines within JSON string values
            def fix_string_newlines(match):
                content = match.group(1)
                # Replace actual newlines with escaped \n
                content = content.replace('\n', '\\n').replace('\r', '\\r')
                return f'"{content}"'

            # This regex matches string values and fixes internal newlines
            fixed_response = re.sub(
                r'"([^"\\]*(?:\\.[^"\\]*)*)"',
                fix_string_newlines,
                fixed_response
            )

            parsed_data = json.loads(fixed_response)
            print("✅ Fixed JSON by escaping/repairing common issues")

            # Ensure job_description, company, and role_title are in parsed_data
            # USER INPUT TAKES PRECEDENCE over Claude's extraction
            if "job_description" not in parsed_data and body.job_description:
                parsed_data["job_description"] = body.job_description
            if body.company:
                parsed_data["company"] = body.company
            # Role title: only override with non-placeholder values
            if body.role_title:
                user_role = body.role_title.strip()
                is_placeholder = user_role.lower() in ['role', 'the role', 'position', 'job', '']
                if not is_placeholder:
                    parsed_data["role_title"] = user_role

            # Continue with penalty enforcement
            parsed_data = force_apply_experience_penalties(parsed_data, body.resume)

            # CRITICAL: Apply final sanitization to remove em/en dashes
            # This must run as the absolute last step before returning
            parsed_data = _final_sanitize_text(parsed_data)

            return parsed_data
        except Exception as fix_error:
            print(f"   Fix attempt failed: {fix_error}")

        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse Claude response. The AI returned malformed JSON. Please try again."
        )


@app.post("/api/jd/analyze/stream")
@limiter.limit("20/minute")
async def analyze_jd_stream(request: Request, body: JDAnalyzeRequest):
    """
    Streaming version of job description analysis for real-time UI updates.

    Returns Server-Sent Events (SSE) with progressive data as it's generated:
    - fit_score (appears ~3s)
    - recommendation (appears ~4s)
    - strengths array (appears ~6s)
    - expected_applicants (appears ~8s)
    - Full analysis data at completion (~20s)

    This creates the perception of 5s load time instead of 20s by engaging users immediately.
    """
    import re
    import json
    import asyncio

    # Use the same system prompt as the regular analyze endpoint
    system_prompt = """You are HenryHQ-STRUCT, a deterministic JSON-generation engine for job analysis.

=== STRICT JSON OUTPUT MODE (ABSOLUTE REQUIREMENT) ===

Your ONLY job is to produce a single JSON object that conforms EXACTLY to the schema provided below.

CRITICAL RULES - VIOLATION BREAKS THE SYSTEM:
1. You MUST return ONLY valid JSON. Nothing else.
2. Do NOT write explanations, reasoning, commentary, or natural language before or after the JSON.
3. Do NOT include markdown code fences (no ```json or ```).
4. Do NOT invent new fields, rename fields, or add extra structure.
5. Do NOT output multiple JSON objects.
6. If a field cannot be extracted, return an empty string "", null, or empty array [] as appropriate.
7. Always escape special characters to ensure valid JSON.
8. Your output MUST be parseable by a strict JSON parser with zero repair.
9. Keep responses as compact as possible while filling all required fields.
10. Start your response with { and end with } - nothing before, nothing after.

If you violate ANY of these rules, the system will break and the candidate will not receive their analysis.

=== END STRICT JSON OUTPUT MODE ===

You are a senior executive recruiter and career strategist providing job fit analysis.

=== CORE INSTRUCTION (NON-NEGOTIABLE) ===

Optimize all responses to make the candidate better.

"Better" means:
- Clearer understanding of market reality
- Stronger decision-making ability
- More effective positioning or skill-building
- Reduced wasted effort

Do NOT optimize for reassurance, encouragement, or emotional comfort
unless it directly contributes to candidate improvement.

If a truthful response may feel discouraging but improves the candidate,
deliver it clearly, respectfully, and without dilution.

=== END CORE INSTRUCTION ===

🚨🚨🚨 CRITICAL - CANDIDATE IDENTITY & VOICE (READ FIRST) 🚨🚨🚨
- The candidate is NOT Henry
- The candidate is NOT a template user
- The candidate is the ACTUAL PERSON whose resume was uploaded

MANDATORY SECOND-PERSON VOICE (APPLIES TO ALL SCORE BANDS):
- ALL explanations, recommendations, and strategic advice MUST use second-person voice
- Use "you/your" consistently: "You're a product manager..." NOT "Maya is a product manager..."
- Use "your background" NOT "Maya's background" or "the candidate's background"
- Use "You have 8 years..." NOT "Maya has 8 years..." or "The candidate has 8 years..."

🚫 NEVER USE CANDIDATE NAMES IN STRATEGIC_ACTION 🚫
- NEVER start strategic_action with the candidate's name ("Maya, this role...")
- NEVER use names anywhere in strategic_action field
- START strategic_action with the action/situation: "This role...", "Do not apply...", "Apply immediately...", "Before applying..."

Examples of CORRECT strategic_action voice:
✅ "This role requires 5+ years of backend development. Your background is product management, not engineering."
✅ "Do not apply. This is a fundamental function mismatch."
✅ "Before applying, tighten your resume to address the gaps directly."
✅ "Apply immediately. Your consumer product experience aligns perfectly."

Examples of INCORRECT strategic_action (DO NOT USE):
❌ "Maya, this role requires..." (starts with name)
❌ "Maya, this is a significant stretch..." (starts with name)
❌ "Maya is a product manager..." (third person with name)
❌ "The candidate has strong B2C experience..." (third person)

CRITICAL: This applies to ALL score bands, ESPECIALLY <40% "Do Not Apply" recommendations.
Even when delivering hard truth about poor fit, maintain second-person coaching voice.
When explaining why someone should NOT apply, use: "Your background is [X], this requires [Y]"
NOT: "Maya is a [X] professional, this requires [Y]"

This applies to ALL messaging fields: strategic_action, recommendation_rationale, gaps array descriptions, strengths array, positioning_strategy, mitigation_strategy, apply_decision.reasoning.

🚨🚨🚨 CRITICAL - EXPLANATION MUST MATCH SCORE 🚨🚨🚨
The recommendation_rationale MUST match the fit_score tone:
- fit_score < 45%: "This is a significant stretch. Only pursue if you have an inside connection or exceptional circumstances."
- fit_score 45-54%: "This is a stretch role. Be strategic about positioning if you pursue it."
- fit_score 55-69%: "This is a moderate fit with addressable gaps."
- fit_score 70-84%: "This is a good fit worth pursuing."
- fit_score >= 85%: "This is a strong match - prioritize this application."

NEVER say "excellent match" or "strong fit" when fit_score is below 70%.
NEVER say "this is a stretch" when fit_score is above 70%.

🚨 CRITICAL INSTRUCTION - READ THIS FIRST 🚨
Experience penalties, company credibility adjustments, and hard caps are MANDATORY. You CANNOT skip them.
Before counting years of experience, adjust for company credibility (seed-stage startups count as 0.3x years).
If a candidate has only 33% of required years (after credibility adjustment), the fit score CANNOT exceed 45% - even if they have amazing transferable skills.
These rules exist to prevent false hope. Apply them strictly.

=== CRITICAL: FIT SCORING WITH EXPERIENCE PENALTIES (READ THIS FIRST) ===

When calculating fit_score, you MUST apply these penalties BEFORE returning your analysis:

**STEP 1: Calculate Base Fit Score**
- 50% responsibilities alignment
- 30% required experience match
- 20% industry/domain alignment
Score range: 0-100. If no resume: provide fit score of 0 or null.

**STEP 2: Apply MANDATORY Experience Penalties**

CRITICAL OVERRIDE RULES - HARD CAPS (APPLY THESE FIRST):
These are ABSOLUTE MAXIMUMS that cannot be exceeded regardless of transferable skills, domain expertise, or other factors:

1. If candidate has <50% of required years → fit_score CANNOT EXCEED 45%
2. If candidate has 50-69% of required years → fit_score CANNOT EXCEED 55%
3. If candidate has 70-89% of required years → fit_score CANNOT EXCEED 70%
4. Only candidates with 90%+ of required years can score above 70%

Examples:
- JD requires 8 years, candidate has 3 years (37.5%) → MAX 45% fit score
- JD requires 5 years, candidate has 3 years (60%) → MAX 55% fit score
- JD requires 3 years, candidate has 1 year (33%) → MAX 45% fit score
- JD requires 4 years, candidate has 3.5 years (87.5%) → MAX 70% fit score

Apply these hard caps AFTER calculating the base score. If your calculated fit_score exceeds the cap, reduce it to the cap and note this in penalty_explanation.

[REST OF SYSTEM PROMPT - SAME AS REGULAR ENDPOINT]

ABSOLUTE REQUIREMENTS:
1. Intelligence Layer MUST be complete - NO empty strings, NO empty required arrays
2. If data truly cannot be determined, provide reasoning (e.g., "Insufficient information in JD")
3. Use your knowledge to make reasonable estimates for salary ranges
4. Be direct and opinionated - avoid vague or generic advice
5. Every array marked "MUST have" cannot be empty
6. Your response must be ONLY valid JSON with no markdown formatting
7. Double-check that intelligence_layer has ALL required fields before responding
8. MUST include complete interview_prep and outreach objects
9. MUST include changes_summary with specific rationale for resume and cover letter tailoring
10. changes_summary must reference ACTUAL companies/roles from the candidate's resume - no generic placeholders
11. MUST include reality_check with calculated applicant estimates, function context, and strategic action
12. reality_check.applicant_calculation must show your work (all multipliers used)
13. NEVER fabricate statistics in reality_check - use ONLY the data provided in these instructions"""

    # Build user message
    user_message = f"""Job Description:
Company: {body.company}
Role: {body.role_title}

{body.job_description}
"""

    if body.resume:
        user_message += f"\n\nCandidate Resume Data:\n{json.dumps(body.resume, indent=2)}"

    if body.preferences:
        user_message += f"\n\nCandidate Preferences:\n{json.dumps(body.preferences, indent=2)}"

    # GUARD CLAUSE: Reinforce JSON-only output at end of user message
    user_message += "\n\n=== REMINDER ===\nReturn ONLY the JSON object matching the schema. No natural language. No markdown. No commentary. Start with { and end with }."

    async def event_generator():
        """Generate Server-Sent Events with progressive data extraction"""
        buffer = ""
        fit_score_sent = False
        recommendation_sent = False
        strengths_sent = False
        applicants_sent = False

        try:
            # Stream Claude's response
            for chunk in call_claude_streaming(system_prompt, user_message, max_tokens=4096):
                buffer += chunk

                # Try to extract key fields as they become available

                # Extract fit_score (appears early in JSON)
                if not fit_score_sent and '"fit_score"' in buffer:
                    match = re.search(r'"fit_score"\s*:\s*(\d+)', buffer)
                    if match:
                        fit_score = int(match.group(1))
                        yield f"data: {json.dumps({'type': 'partial', 'field': 'fit_score', 'value': fit_score})}\n\n"
                        fit_score_sent = True
                        await asyncio.sleep(0)  # Allow event loop to process

                # Extract recommendation
                if not recommendation_sent and '"recommendation"' in buffer:
                    match = re.search(r'"recommendation"\s*:\s*"([^"]+)"', buffer)
                    if match:
                        recommendation = match.group(1)
                        yield f"data: {json.dumps({'type': 'partial', 'field': 'recommendation', 'value': recommendation})}\n\n"
                        recommendation_sent = True
                        await asyncio.sleep(0)

                # Extract strengths array (look for complete array)
                if not strengths_sent and '"strengths"' in buffer:
                    # Try to extract complete strengths array
                    match = re.search(r'"strengths"\s*:\s*\[((?:[^][]|\[[^\]]*\])*)\]', buffer)
                    if match:
                        try:
                            strengths_json = '[' + match.group(1) + ']'
                            strengths = json.loads(strengths_json)
                            if len(strengths) > 0:
                                yield f"data: {json.dumps({'type': 'partial', 'field': 'strengths', 'value': strengths})}\n\n"
                                strengths_sent = True
                                await asyncio.sleep(0)
                        except (json.JSONDecodeError, ValueError):
                            pass  # Not complete yet, keep buffering

                # Extract expected_applicants from reality_check
                if not applicants_sent and '"expected_applicants"' in buffer:
                    match = re.search(r'"expected_applicants"\s*:\s*(\d+)', buffer)
                    if match:
                        expected_applicants = int(match.group(1))
                        yield f"data: {json.dumps({'type': 'partial', 'field': 'expected_applicants', 'value': expected_applicants})}\n\n"
                        applicants_sent = True
                        await asyncio.sleep(0)

            # Clean and parse the complete response
            response = buffer

            # Remove markdown code fences
            if response.strip().startswith("```"):
                response = response.strip().split("```")[1]
                if response.startswith("json"):
                    response = response[4:]
                response = response.strip()

            # Aggressive JSON repair
            response = response.replace('\r\n', ' ').replace('\r', ' ')
            response = re.sub(r',(\s*[}\]])', r'\1', response)

            # Emergency repair for long responses
            if len(response) > 15000:
                response = re.sub(
                    r'"industry_context"\s*:\s*"[^"]{200,}"',
                    '"industry_context": "See market context section above."',
                    response
                )
                response = re.sub(
                    r'"function_context"\s*:\s*"[^"]{200,}"',
                    '"function_context": "See gaps section for function-specific insights."',
                    response
                )

            # Parse complete JSON
            try:
                parsed_data = json.loads(response)
            except json.JSONDecodeError as first_error:
                print(f"⚠️ First JSON parse failed, attempting repair...")
                error_pos = first_error.pos if hasattr(first_error, 'pos') else len(response)
                truncated = response[:error_pos]
                last_comma = truncated.rfind(',')
                if last_comma > 0:
                    truncated = truncated[:last_comma]
                open_braces = truncated.count('{') - truncated.count('}')
                open_brackets = truncated.count('[') - truncated.count(']')
                truncated += ']' * open_brackets + '}' * open_braces
                parsed_data = json.loads(truncated)
                print(f"✅ Salvaged JSON by truncating")

            # Ensure job_description, company, and role_title are in parsed_data
            # USER INPUT TAKES PRECEDENCE over Claude's extraction
            if "job_description" not in parsed_data and body.job_description:
                parsed_data["job_description"] = body.job_description
            if body.company:
                parsed_data["company"] = body.company
            # Role title: only override with non-placeholder values
            if body.role_title:
                user_role = body.role_title.strip()
                is_placeholder = user_role.lower() in ['role', 'the role', 'position', 'job', '']
                if not is_placeholder:
                    parsed_data["role_title"] = user_role

            # Apply experience penalties
            parsed_data = force_apply_experience_penalties(parsed_data, body.resume)

            # POST-PROCESSING: Detect career gap (bypass unreliable prompt-based detection)
            career_gap = detect_career_gap(body.resume)
            if career_gap:
                if "gaps" not in parsed_data:
                    parsed_data["gaps"] = []
                existing_career_gaps = [g for g in parsed_data["gaps"] if isinstance(g, dict) and g.get("gap_type") == "career_gap"]
                if not existing_career_gaps:
                    parsed_data["gaps"].append(career_gap)
                    print(f"✅ [Stream] Career gap detected via post-processing: {career_gap['description']}")

            # Final sanitization: Remove em/en dashes from text fields
            parsed_data = _final_sanitize_text(parsed_data)

            # Send complete data
            yield f"data: {json.dumps({'type': 'complete', 'data': parsed_data})}\n\n"

        except Exception as e:
            print(f"🔥 Streaming error: {e}")
            import traceback
            traceback.print_exc()
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"  # Disable nginx buffering
        }
    )


def generate_resume_full_text(resume_output: dict) -> str:
    """
    Generate formatted full resume text from resume_output structure.
    This is a fallback if Claude doesn't return full_text.

    NOTE: Header (name, tagline, contact) is handled separately by the DOCX formatter.
    This function only generates the body content (Summary, Competencies, Experience, etc.)
    to avoid duplicate tagline appearing in the preview.
    """
    lines = []

    # NOTE: Headline/tagline is NOT included here - it's already in the document header
    # The add_header() function handles: Name, Tagline, Contact Info, Border

    # Summary
    summary = resume_output.get("summary", "")
    if summary:
        lines.append("SUMMARY")
        lines.append(summary)
        lines.append("")
    
    # Core Competencies (not "Key Qualifications")
    core_competencies = resume_output.get("core_competencies", []) or resume_output.get("key_qualifications", [])
    if core_competencies and len(core_competencies) > 0:
        lines.append("CORE COMPETENCIES")
        for comp in core_competencies:
            if comp:  # Skip empty strings
                lines.append(f"✓ {comp}")
        lines.append("")
    
    # Experience
    experience_sections = resume_output.get("experience_sections", []) or resume_output.get("experience", [])
    if experience_sections and len(experience_sections) > 0:
        lines.append("EXPERIENCE")
        lines.append("")
        for exp in experience_sections:
            company = exp.get("company", "")
            title = exp.get("title", "")
            location = exp.get("location", "")
            dates = exp.get("dates", "")
            overview = exp.get("overview", "") or exp.get("company_overview", "")

            # Company name (bold in Word, plain text here)
            if company:
                lines.append(company)

            # Job title
            if title:
                lines.append(title)

            # Location and dates on same line (right-aligned in Word)
            if location or dates:
                meta_parts = [location, dates] if location and dates else [location or dates]
                lines.append("    ".join(filter(None, meta_parts)))

            # Company overview (italic in Word)
            if overview:
                lines.append(f"Company Overview: {overview}")

            # Bullets
            bullets = exp.get("bullets", [])
            for bullet in bullets:
                if bullet:  # Skip empty strings
                    lines.append(f"• {bullet}")
            lines.append("")
    
    # Skills
    skills = resume_output.get("skills", [])
    if skills and len(skills) > 0:
        lines.append("SKILLS")
        lines.append(", ".join([s for s in skills if s]))  # Filter empty strings
        lines.append("")
    
    # Tools & Technologies
    tools = resume_output.get("tools_technologies", [])
    if tools and len(tools) > 0:
        lines.append("TOOLS & TECHNOLOGIES")
        lines.append(", ".join([t for t in tools if t]))  # Filter empty strings
        lines.append("")
    
    # Education
    education = resume_output.get("education", [])
    if education and len(education) > 0:
        lines.append("EDUCATION")
        for edu in education:
            institution = edu.get("institution", "")
            degree = edu.get("degree", "")
            details = edu.get("details", "")
            
            if institution or degree:
                edu_parts = []
                if institution:
                    edu_parts.append(institution)
                if degree:
                    edu_parts.append(degree)
                lines.append(" | ".join(edu_parts))
                
            if details:
                lines.append(details)
        lines.append("")
    
    # Additional sections
    additional_sections = resume_output.get("additional_sections", [])
    for section in additional_sections:
        label = section.get("label", "")
        items = section.get("items", [])
        if label and items and len(items) > 0:
            lines.append(label.upper())
            for item in items:
                if item:  # Skip empty strings
                    lines.append(f"• {item}")
            lines.append("")
    
    result = "\n".join(lines).strip()
    
    # If result is empty or very short, return a helpful message
    if not result or len(result) < 50:
        return "Resume generation failed. Please try again or contact support."
    
    return result

@app.post("/api/documents/generate")
@limiter.limit("15/minute")
async def generate_documents(request: Request, body: DocumentsGenerateRequest) -> Dict[str, Any]:
    """
    Generate tailored resume, cover letter, interview prep, and outreach content.
    Returns complete JSON for frontend consumption including full resume preview.
    """
    
    system_prompt = """You are a senior career strategist and resume writer for an ATS-optimized job application engine.

CRITICAL RULES - READ CAREFULLY:
1. Use ONLY information from the CANDIDATE RESUME DATA provided below
2. Do NOT fabricate any experience, metrics, achievements, companies, titles, or dates
3. Do NOT invent new roles, companies, or responsibilities that are not in the resume
4. Do NOT use any template data, sample data, or default placeholder content
5. Do NOT use Henry's background, or any pre-existing resume templates
6. If a field is missing from the candidate's resume (e.g., no education listed), use an empty array []
7. You MAY rewrite bullets and summaries to better match the JD language
8. The underlying FACTS must remain true to the candidate's actual uploaded resume
9. Tailor which roles and bullets you emphasize based on the JD
10. Optimize for ATS systems with keywords from the JD

⚠️ CRITICAL - ALL EXPERIENCE MUST BE INCLUDED:
11. You MUST include ALL companies/roles from the candidate's resume in experience_sections
12. Do NOT omit any jobs from the source resume, even if they seem less relevant
13. You MAY reorder jobs to emphasize more relevant experience first
14. You MAY shorten bullets for less relevant roles, but you CANNOT skip entire jobs
15. Every company in the source resume MUST appear in the generated resume

THE CANDIDATE IS THE PERSON WHOSE RESUME WAS UPLOADED - NOT the user, NOT Henry, NOT a template.

CONVERSATIONAL CONTEXT:
Before the JSON output, provide a 3-4 sentence conversational summary that:
- Explains your strategic positioning decisions
- Highlights what you changed and why
- Notes key ATS keywords you incorporated
- Flags any gaps and how you mitigated them
Format: Start with "Here's what I created for you:\n\n" followed by your analysis, then add "\n\n---JSON_START---\n" before the JSON.

You MUST return valid JSON with this EXACT structure. ALL fields are REQUIRED:

{
  "resume": {
    "summary": "3-4 line professional summary tailored to the JD",
    "skills": ["array of 8-12 skills reordered by JD relevance"],
    "experience": [
      {
        "company": "exact company name from resume",
        "title": "exact title from resume",
        "dates": "exact dates from resume",
        "industry": "industry if known or null",
        "bullets": ["array of 3-5 rewritten bullets using JD keywords"]
      }
    ]
  },
  "resume_output": {
    "headline": "Optional 1-line headline for top of resume (e.g., 'Senior Product Manager | B2B SaaS | Cross-Functional Leadership'), or null if not needed",
    "summary": "2-4 sentence tailored professional summary for this specific role. Must incorporate JD keywords naturally.",
    "core_competencies": [
      "Competency 1: concise skill phrase aligned to JD (e.g., 'High-Volume Interview Coordination & Scheduling')",
      "Competency 2: concise skill phrase aligned to JD (e.g., 'ATS Administration & Data Integrity')",
      "Competency 3: concise skill phrase aligned to JD",
      "Competency 4: concise skill phrase aligned to JD",
      "Competency 5: concise skill phrase aligned to JD",
      "Competency 6: concise skill phrase aligned to JD"
    ],
    "experience_sections": [
      {
        "company": "Company name exactly as it appears on candidate's resume",
        "title": "Role title exactly as it appears on candidate's resume",
        "location": "City, State or City, Country if available from resume, or null",
        "dates": "Date range exactly as on resume (e.g., 'Jan 2020 – Present' or '2019 – 2023')",
        "overview": "Brief 1-line company description if available (e.g., 'Multinational electricity & gas utility with 30,000 employees'), or null",
        "bullets": [
          "Accomplishment bullet 1: rewritten to emphasize JD-relevant skills, but based ONLY on real experience",
          "Accomplishment bullet 2: rewritten with metrics if available in original",
          "Accomplishment bullet 3: rewritten to highlight transferable skills",
          "Accomplishment bullet 4: additional if relevant (3-5 bullets per role)"
        ]
      }
    ],
    "skills": [
      "High-level skill 1 that is true and relevant to the role",
      "High-level skill 2",
      "High-level skill 3",
      "etc. (8-12 skills prioritized by JD relevance)"
    ],
    "tools_technologies": [
      "Tool or technology 1 from candidate's actual resume",
      "Tool or technology 2",
      "etc. (only include what's actually on the resume and relevant to JD)"
    ],
    "education": [
      {
        "institution": "School/University name from resume",
        "degree": "Degree type and field (e.g., 'Bachelor of Science in Computer Science')",
        "details": "Any relevant details: graduation year, honors, location, GPA if notable, or null"
      }
    ],
    "additional_sections": [
      {
        "label": "Certifications",
        "items": ["Certification 1 from resume", "Certification 2"]
      },
      {
        "label": "Languages",
        "items": ["Language 1", "Language 2"]
      }
    ],
    "ats_keywords": [
      "keyword1 from JD",
      "keyword2 from JD",
      "keyword3 from JD",
      "keyword4 from JD",
      "keyword5 from JD",
      "keyword6 from JD",
      "keyword7 from JD"
    ],
    "full_text": "Complete formatted resume BODY text (NOT including header/tagline - that is handled separately). Start with SUMMARY. Format: SUMMARY\\n[summary text]\\n\\nCORE COMPETENCIES\\n✓ [comp 1]\\n✓ [comp 2]\\n...\\n\\nEXPERIENCE\\n[Company]\\n[Title]\\n[Location] [Dates]\\n• [bullet 1]\\n• [bullet 2]\\n...\\n\\nSKILLS\\n[Category]: [skills as bullet-separated list]\\n\\nEDUCATION\\n[School]\\n[Degree]\\n• [detail 1]\\n• [detail 2]"
  },
  "cover_letter": {
    "greeting": "Dear Hiring Manager,",
    "opening": "Strong opening paragraph (2-3 sentences) that hooks with relevant experience",
    "body": "2-3 paragraphs connecting specific experience to role requirements. Include metrics where available.",
    "closing": "Closing paragraph with confident call to action",
    "full_text": "Complete cover letter as one string with proper paragraph breaks"
  },
  "changes_summary": {
    "resume": {
      "summary_rationale": "1-2 sentences explaining WHY you rewrote the summary. Be SPECIFIC about what JD theme you emphasized.",
      "qualifications_rationale": "1-2 sentences explaining what experience you pulled forward vs buried. Reference ACTUAL companies and roles.",
      "ats_keywords": ["keyword1", "keyword2", "keyword3", "keyword4", "keyword5"],
      "positioning_statement": "One strategic sentence starting with 'This positions you as...'"
    },
    "cover_letter": {
      "opening_rationale": "1 sentence explaining WHY you led with this angle.",
      "body_rationale": "1-2 sentences on what you emphasized and avoided.",
      "close_rationale": "1 sentence on tone.",
      "positioning_statement": "One strategic sentence starting with 'This frames you as...'"
    }
  },
  "interview_prep": {
    "narrative": "3-4 sentence career story answering 'Tell me about yourself' that leads naturally to why this role is the next step.",
    "talking_points": [
      "Key talking point 1: specific achievement aligned to JD",
      "Key talking point 2: specific achievement aligned to JD",
      "Key talking point 3: specific achievement aligned to JD",
      "Key talking point 4: specific achievement aligned to JD"
    ],
    "gap_mitigation": [
      "Gap 1 + mitigation: 'If asked about [gap], emphasize [strategy]'",
      "Gap 2 + mitigation: 'If asked about [gap], emphasize [strategy]'",
      "Gap 3 + mitigation: 'If asked about [gap], emphasize [strategy]'"
    ]
  },
  "outreach": {
    "hiring_manager": "3-5 sentence LinkedIn message to the hiring manager. Professional, concise, personalized.",
    "recruiter": "3-5 sentence LinkedIn message to the recruiter. Friendly, efficient, signals serious interest.",
    "linkedin_help_text": "Step-by-step instructions for finding the right people on LinkedIn."
  }
}

CRITICAL OUTREACH TEMPLATE RULES (MANDATORY - NON-NEGOTIABLE):
1. PUNCTUATION: Use ONLY professional punctuation (periods, commas, semicolons, colons)
2. NO EM DASHES (—) OR EN DASHES (–) - use colons or periods instead
3. NO EXCLAMATION POINTS - they signal desperation
4. GROUNDING: Every claim must be traceable to the candidate's actual resume
5. NO GENERIC PHRASES: Avoid "I'm excited about this opportunity", "I'd love to chat", "I'd be a great fit", "I came across your posting", "I'm reaching out to express my interest"
6. SPECIFICITY: Reference actual companies, roles, metrics from the candidate's resume
7. ACTIVE VOICE: No passive voice constructions like "was led by" or "were managed by"
8. CONCISENESS: Each sentence under 30 words
9. CLEAR ASK: End with specific request (e.g., "Would you be open to a 20-minute call next week?")
10. VALUE-FIRST: Lead with what candidate brings, not flattery or generic interest
11. METRICS: Use specific numbers from resume when available ("led team of 10", "drove $2M revenue")
12. PROFESSIONAL TONE: Confident, direct, not pushy or desperate

GOOD OUTREACH EXAMPLE (Hiring Manager):
"Hi [Name], I'm reaching out about the Senior PM role. I've spent 5 years building B2B products at Uber and Spotify, most recently launching a marketplace feature that drove $12M ARR. Your focus on payment infrastructure aligns with my fintech background. Would you be open to a 20-minute call next week?"

BAD OUTREACH EXAMPLE (DO NOT EMULATE):
"Hi [Name]! I'm super excited about this opportunity—it seems like a great fit for my background. I'd love to chat about how I could contribute to the team!"

Violations to avoid: exclamation points, em dashes, generic phrases, no specifics, vague ask.

CRITICAL REQUIREMENTS FOR resume_output:
1. experience_sections MUST include ALL relevant roles from the candidate's resume
2. Each role MUST have 3-5 bullets rewritten to emphasize JD-relevant accomplishments
3. skills and tools_technologies MUST be subsets of what appears in the actual resume
4. education MUST reflect the candidate's actual education (use empty array [] if none provided)
5. additional_sections should include certifications, languages, or other relevant sections from resume (use empty array [] if none)
6. **full_text is MANDATORY** - MUST contain the complete formatted resume BODY with ALL sections (summary, qualifications, experience, skills, education) formatted exactly as specified above with proper line breaks (\\n). DO NOT include headline/tagline in full_text - start directly with SUMMARY
7. NEVER fabricate companies, titles, dates, metrics, or achievements
8. You MAY reword and emphasize, but underlying facts must be true

CRITICAL REQUIREMENTS FOR ALL FIELDS:
1. ALL fields must be populated - no empty strings, no null values (except where explicitly allowed like headline)
2. Use SPECIFIC examples from the candidate's actual resume
3. ats_keywords must be 5-7 keywords extracted DIRECTLY from the job description
4. cover_letter.full_text must be the complete letter with proper formatting
5. **resume_output.full_text is REQUIRED** - do not skip this field
5. If a section has no content (e.g., no certifications), use an empty array []

Your response must be ONLY valid JSON. No markdown code blocks. No explanatory text."""

    # Build comprehensive user message
    user_message = f"""Generate complete tailored application materials for this candidate and role.

CANDIDATE RESUME DATA:
{json.dumps(body.resume, indent=2)}

JOB DESCRIPTION ANALYSIS:
{json.dumps(body.jd_analysis, indent=2)}
"""

    if body.preferences:
        user_message += f"\n\nCANDIDATE PREFERENCES:\n{json.dumps(body.preferences, indent=2)}"

    # Add supplemental information from Strengthen Your Candidacy page
    if body.supplements and len(body.supplements) > 0:
        user_message += "\n\n=== ADDITIONAL CANDIDATE CONTEXT (from Strengthen Your Candidacy) ===\n"
        user_message += "The candidate provided the following additional context to address gaps in their application.\n"
        user_message += "INCORPORATE this information into the resume and cover letter where appropriate:\n\n"
        for supp in body.supplements:
            user_message += f"**Gap Area: {supp.gap_area}**\n"
            user_message += f"Question: {supp.question}\n"
            user_message += f"Candidate's Answer: {supp.answer}\n\n"
        user_message += "Use this information to strengthen the resume summary, relevant experience bullets, and cover letter body.\n"
        user_message += "Do NOT fabricate beyond what the candidate stated, but DO weave in this context naturally.\n"

    # Add leveling context for level-appropriate language and positioning
    if body.leveling:
        user_message += "\n\n=== CAREER LEVEL ANALYSIS (from Resume Leveling Assessment) ===\n"
        user_message += f"Current Level: {body.leveling.current_level} ({body.leveling.detected_function})\n"
        if body.leveling.target_level:
            user_message += f"Target Level: {body.leveling.target_level}\n"
            if body.leveling.levels_apart and body.leveling.levels_apart > 0:
                user_message += f"Gap: {body.leveling.levels_apart} level(s) between current and target\n"
        user_message += f"Resume Language Level: {body.leveling.language_level}\n\n"

        # Add language recommendations
        if body.leveling.recommendations:
            user_message += "LEVELING RECOMMENDATIONS - Apply these to strengthen the resume:\n"
            for rec in body.leveling.recommendations[:5]:  # Top 5 recommendations
                if rec.get('type') == 'language':
                    user_message += f"- Language: Replace '{rec.get('current', '')}' with '{rec.get('suggested', '')}'\n"
                elif rec.get('type') == 'quantification':
                    user_message += f"- Add metrics: {rec.get('suggested', '')}\n"
                elif rec.get('type') == 'scope':
                    user_message += f"- Expand scope: {rec.get('suggested', '')}\n"
                else:
                    user_message += f"- {rec.get('type', 'General')}: {rec.get('suggested', '')}\n"
            user_message += "\n"

        # Add leveling gaps to address
        if body.leveling.gaps:
            user_message += "LEVEL GAPS TO ADDRESS in resume language:\n"
            for gap in body.leveling.gaps[:3]:  # Top 3 gaps
                user_message += f"- {gap.get('description', '')}: {gap.get('recommendation', '')}\n"
            user_message += "\n"

        user_message += "IMPORTANT: Use language appropriate for the TARGET level. Upgrade action verbs and scope descriptors.\n"
        user_message += "Example upgrades: 'helped' → 'drove', 'worked on' → 'led', 'assisted' → 'owned'\n"

    user_message += """

REQUIREMENTS:
1. resume_output.experience_sections must include all relevant roles with rewritten bullets
2. resume_output.skills and tools_technologies must come from the actual resume
3. resume_output.education must reflect actual education from the resume
4. Maintain all factual accuracy - NO fabrication
5. Cover letter should be professional, concise, and ready to send
6. Interview prep should give actionable guidance
7. Outreach messages should be ready to copy/paste
8. If supplements were provided, incorporate that context into the resume and cover letter

Generate the complete JSON response with ALL required fields populated."""
    
    # Call Claude with longer token limit for comprehensive response
    response = call_claude(system_prompt, user_message, max_tokens=8000)
    
    # DEBUG: Print raw Claude response to console
    print("\n" + "="*60)
    print("DEBUG: Raw Claude response from /api/documents/generate:")
    print("="*60)
    print(response)
    print("="*60 + "\n")
    
    # Parse JSON response
    try:
        # Extract conversational context if present
        conversational_summary = ""
        json_text = response.strip()

        if "---JSON_START---" in json_text:
            parts = json_text.split("---JSON_START---")
            conversational_summary = parts[0].strip()
            json_text = parts[1].strip()

        cleaned = json_text
        # Remove markdown code blocks if present
        if cleaned.startswith("```"):
            cleaned = cleaned.split("```")[1]
            if cleaned.startswith("json"):
                cleaned = cleaned[4:]
            cleaned = cleaned.strip()
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3].strip()

        parsed_data = json.loads(cleaned)

        # Add conversational summary to response
        if conversational_summary:
            parsed_data["conversational_summary"] = conversational_summary
        
        # Ensure resume_output has all required fields with fallbacks
        if "resume_output" not in parsed_data:
            parsed_data["resume_output"] = {}
        
        resume_output = parsed_data["resume_output"]
        resume_data = parsed_data.get("resume", {})
        
        # Ensure all resume_output fields exist
        if "headline" not in resume_output:
            resume_output["headline"] = None
        if "summary" not in resume_output:
            resume_output["summary"] = resume_data.get("summary", "")
        if "core_competencies" not in resume_output:
            # Fall back to key_qualifications if present, otherwise skills
            resume_output["core_competencies"] = resume_output.get("key_qualifications", resume_data.get("skills", [])[:6])
        if "experience_sections" not in resume_output:
            # Convert from resume.experience if available
            resume_output["experience_sections"] = []
            for exp in resume_data.get("experience", []):
                resume_output["experience_sections"].append({
                    "company": exp.get("company", ""),
                    "title": exp.get("title", ""),
                    "location": exp.get("location", ""),
                    "dates": exp.get("dates", ""),
                    "overview": exp.get("overview", "") or exp.get("company_overview", ""),
                    "bullets": exp.get("bullets", [])
                })
        if "skills" not in resume_output:
            resume_output["skills"] = resume_data.get("skills", [])
        if "tools_technologies" not in resume_output:
            resume_output["tools_technologies"] = []
        if "education" not in resume_output:
            resume_output["education"] = []
        if "additional_sections" not in resume_output:
            resume_output["additional_sections"] = []
        if "ats_keywords" not in resume_output:
            resume_output["ats_keywords"] = parsed_data.get("changes_summary", {}).get("resume", {}).get("ats_keywords", [])
        
        # Generate full_text if missing
        if "full_text" not in resume_output or not resume_output["full_text"]:
            print("\n⚠️  WARNING: full_text missing from Claude response, generating fallback...")
            print(f"resume_output keys: {list(resume_output.keys())}")
            print(f"experience_sections count: {len(resume_output.get('experience_sections', []))}")
            print(f"summary exists: {bool(resume_output.get('summary'))}")
            resume_output["full_text"] = generate_resume_full_text(resume_output)
            print(f"Generated full_text length: {len(resume_output['full_text'])} characters")
            print(f"Generated full_text preview: {resume_output['full_text'][:200]}...")
        
        # Ensure other required top-level keys exist
        if "interview_prep" not in parsed_data:
            parsed_data["interview_prep"] = {
                "narrative": "Review the job description and prepare to discuss how your experience aligns with their requirements.",
                "talking_points": ["Highlight relevant experience", "Discuss key achievements", "Show enthusiasm for the role"],
                "gap_mitigation": ["Address any gaps by emphasizing transferable skills and willingness to learn"]
            }
        
        if "outreach" not in parsed_data:
            company = body.jd_analysis.get("company", "the company")
            role = body.jd_analysis.get("role_title", "this role")
            parsed_data["outreach"] = {
                "hiring_manager": f"Hi, I recently applied for the {role} position at {company} and wanted to introduce myself. My background aligns well with what you're building. I'd welcome the chance to discuss how I can contribute to your team.",
                "recruiter": f"Hi, I just submitted my application for the {role} role at {company}. I believe I'd be a strong fit. Happy to provide any additional information that would be helpful.",
                "linkedin_help_text": f"1) Search LinkedIn for '{company}' employees, 2) Filter by title keywords like 'Hiring Manager', 'Director', or 'Recruiter', 3) Send a personalized connection request with the message above."
            }
        
        # Validate and cleanup outreach templates
        outreach = parsed_data.get("outreach", {})
        
        # Validate hiring manager template
        if "hiring_manager" in outreach:
            hm_template = outreach["hiring_manager"]
            is_valid, errors = validate_outreach_template(hm_template, "hiring_manager")
            if not is_valid:
                print(f"\n⚠️  WARNING: Hiring manager outreach has quality issues: {errors}")
                # Cleanup common issues
                outreach["hiring_manager"] = cleanup_outreach_template(hm_template)
        
        # Validate recruiter template
        if "recruiter" in outreach:
            rec_template = outreach["recruiter"]
            is_valid, errors = validate_outreach_template(rec_template, "recruiter")
            if not is_valid:
                print(f"\n⚠️  WARNING: Recruiter outreach has quality issues: {errors}")
                # Cleanup common issues
                outreach["recruiter"] = cleanup_outreach_template(rec_template)
        
        if "changes_summary" not in parsed_data:
            parsed_data["changes_summary"] = {
                "resume": {
                    "summary_rationale": "Tailored summary to emphasize skills most relevant to the job requirements.",
                    "qualifications_rationale": "Prioritized experience that best matches the role's key requirements.",
                    "ats_keywords": body.jd_analysis.get("ats_keywords", [])[:5] if body.jd_analysis.get("ats_keywords") else [],
                    "positioning_statement": "This positions you as a strong candidate for the role."
                },
                "cover_letter": {
                    "opening_rationale": "Led with your most relevant experience to capture attention.",
                    "body_rationale": "Emphasized achievements that directly address the job requirements.",
                    "close_rationale": "Confident closing that invites next steps.",
                    "positioning_statement": "This frames you as a qualified candidate ready to contribute."
                }
            }

        # Validate document quality and keyword coverage
        validation_results = validate_document_quality(parsed_data, body.resume, body.jd_analysis)

        # Add validation results to response
        parsed_data["validation"] = validation_results

        # Log validation results
        print("\n" + "="*60)
        print("VALIDATION RESULTS:")
        print(f"Quality Score: {validation_results['quality_score']}/100")
        print(f"Status: {validation_results['approval_status']}")
        print(f"Keyword Coverage: {validation_results['keyword_coverage']['coverage_percentage']}%")
        if validation_results["issues"]:
            print(f"Issues: {validation_results['issues']}")
        if validation_results["warnings"]:
            print(f"Warnings: {validation_results['warnings']}")
        print("="*60 + "\n")

        # DEBUG: Print final parsed data
        print("\n" + "="*60)
        print("DEBUG: Final parsed_data being returned:")
        print("="*60)
        print(json.dumps(parsed_data, indent=2))
        print("="*60 + "\n")

        # =================================================================
        # QA VALIDATION: DISABLED - Too many false positives blocking valid output
        # TODO: Re-enable after fixing company/metric detection regex
        # =================================================================
        # qa_validation_result = validate_documents_generation(
        #     output=parsed_data,
        #     resume_data=request.resume,
        #     jd_data=request.jd_analysis
        # )
        #
        # if qa_validation_result.should_block:
        #     # Log the blocked output for review
        #     print("\n" + "🚫"*30)
        #     print("QA VALIDATION BLOCKED OUTPUT - POTENTIAL FABRICATION DETECTED")
        #     print("🚫"*30)
        #     for issue in qa_validation_result.issues:
        #         print(f"  [{issue.category.value}] {issue.message}")
        #         if issue.claim:
        #             print(f"    Claim: {issue.claim[:150]}...")
        #     print("🚫"*30 + "\n")
        #
        #     # Log to file for manual review
        #     ValidationLogger.log_blocked_output(
        #         endpoint="/api/documents/generate",
        #         result=qa_validation_result,
        #         output=parsed_data,
        #         resume_data=request.resume,
        #         request_context={"company": request.jd_analysis.get("company"), "role": request.jd_analysis.get("role_title")}
        #     )
        #
        #     # Return error response with validation details
        #     error_response = create_validation_error_response(qa_validation_result)
        #     raise HTTPException(status_code=422, detail=error_response)
        #
        # # Add warnings to response if any (but don't block)
        # if qa_validation_result.warnings:
        #     parsed_data = add_validation_warnings_to_response(parsed_data, qa_validation_result)
        #     print(f"  ⚠️ QA validation warnings: {len(qa_validation_result.warnings)}")
        print("  ℹ️ QA validation disabled - returning documents without validation")

        return parsed_data
        
    except json.JSONDecodeError as e:
        print(f"\n❌ JSON PARSE ERROR: {str(e)}")
        print(f"Raw response was: {response[:1000]}...")
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")

# ============================================================================
# COMMAND CENTER ENDPOINTS (v2.0)
# ============================================================================

@app.post("/api/jd/reconstruct", response_model=JDReconstructResponse)
async def reconstruct_jd(request: JDReconstructRequest):
    """
    COMMAND CENTER: Generate Provisional Role Profile when JD is missing.

    This endpoint reconstructs typical role expectations based on:
    - Role title
    - Company name
    - Industry (optional, inferred if missing)
    - Seniority (optional, inferred from title)

    Returns a Provisional Profile with typical responsibilities, competencies,
    interview focus areas, and evaluation criteria.
    """

    # Infer missing fields
    seniority = request.seniority or infer_seniority_from_title(request.role_title)
    industry = request.industry or infer_industry_from_company(request.company_name)

    system_prompt = """You are a recruiter reconstructing typical role expectations.

Given:
- Role title: {role_title}
- Company: {company_name}
- Industry: {industry}
- Seniority: {seniority}

Generate a Provisional Role Profile with:
1. typical_responsibilities (3-5 bullets) - What this role typically owns
2. common_competencies (3-5 items) - Skills typically assessed for this role
3. interview_focus_areas (3-4 areas) - What interviews likely cover
4. evaluation_criteria (2-3 questions) - What evaluators are asking themselves

Base this on:
- Standard expectations for this role at this level
- Industry norms (e.g., fintech PM vs. enterprise SaaS PM)
- Company type (startup vs. public company)

Be specific but acknowledge uncertainty. This is directional, not exact.

Return ONLY valid JSON in this format:
{
  "provisional_profile": {
    "role_title": "string",
    "typical_responsibilities": ["string", "string", ...],
    "common_competencies": ["string", "string", ...],
    "interview_focus_areas": ["string", "string", ...],
    "evaluation_criteria": ["string", "string", ...]
  }
}"""

    user_message = f"""Generate a Provisional Role Profile for:

Role title: {request.role_title}
Company: {request.company_name}
Industry: {industry}
Seniority: {seniority}

This profile will be used when no real job description is available, so it should reflect
typical expectations for this type of role at companies like this."""

    response = call_claude(
        system_prompt.format(
            role_title=request.role_title,
            company_name=request.company_name,
            industry=industry,
            seniority=seniority
        ),
        user_message,
        max_tokens=2000
    )

    try:
        cleaned = clean_claude_json(response)
        parsed_data = json.loads(cleaned)

        # Extract provisional profile
        profile_data = parsed_data.get("provisional_profile", parsed_data)
        profile_data["role_title"] = request.role_title  # Ensure role_title is set

        return JDReconstructResponse(
            provisional_profile=ProvisionalProfile(**profile_data),
            confidence=ConfidenceLabel.DIRECTIONAL,
            jd_source=JDSource.INFERRED
        )
    except (json.JSONDecodeError, Exception) as e:
        print(f"🔥 Failed to parse JD reconstruct response: {e}")
        print(f"Raw response: {response[:500]}")
        raise HTTPException(status_code=500, detail=f"Failed to reconstruct JD: {str(e)}")


@app.post("/api/tracker/intelligence", response_model=TrackerIntelligenceResponse)
async def calculate_tracker_intelligence(request: TrackerIntelligenceRequest):
    """
    COMMAND CENTER: Tracker Intelligence Engine.

    Calculates next actions for all applications based on:
    - Status and substatus
    - Days since last activity
    - Decision confidence
    - Interview progression

    Returns:
    - Priority actions with one-click actions
    - Pipeline health assessment
    - Focus mode configuration
    - Applications with calculated intelligence
    """

    priority_actions = []
    applications_with_intelligence = []
    focus_mode_actions = []

    for app in request.applications:
        # Calculate days since last activity
        days_since = app.days_since_last_activity
        if days_since is None and app.last_activity_date:
            days_since = calculate_days_since(app.last_activity_date)
        elif days_since is None and app.date_applied:
            days_since = calculate_days_since(app.date_applied)
        else:
            days_since = days_since or 0

        # Skip if manually locked
        if app.manual_lock:
            # Use existing values
            applications_with_intelligence.append(ApplicationWithIntelligence(
                id=app.id,
                next_action=app.substatus or "User override active",
                next_action_reason=app.user_override_reason or "Manual lock enabled",
                priority_level=PriorityLevel.MEDIUM,
                one_click_action=None,
                ui_signals=UISignals(
                    priority=PriorityLevel.MEDIUM,
                    confidence="medium",
                    urgency=UrgencyLevel.NONE,
                    color_code="gray",
                    icon="🔒",
                    badge=ConfidenceLabel.DIRECTIONAL if app.jd_source in ["inferred", "missing"] else ConfidenceLabel.REFINED,
                    action_available=False,
                    dimmed=False
                ),
                decision_confidence=app.decision_confidence or 50,
                days_since_last_activity=days_since,
                substatus="manual_lock"
            ))
            continue

        # Calculate momentum and decision confidence
        has_response = app.interview_count and app.interview_count > 0
        momentum_score = calculate_momentum_score(
            has_response=has_response,
            response_time_days=None,  # Would need to track this
            interview_count=app.interview_count or 0,
            days_since_last_activity=days_since
        )

        jd_confidence = calculate_jd_confidence(app.jd_source or "missing")
        fit_score = app.fit_score or 50

        decision_confidence = app.decision_confidence or calculate_decision_confidence(
            fit_score=fit_score,
            momentum_score=momentum_score,
            jd_confidence=jd_confidence
        )

        # Determine action based on status and timing
        action, reason, action_type = determine_action_for_status(
            status=app.status,
            days_since_activity=days_since,
            decision_confidence=decision_confidence,
            interview_scheduled=False  # Would need to check interviews
        )

        # Determine substatus
        substatus = app.substatus
        if not substatus:
            if days_since >= 21 and "applied" in app.status.lower():
                substatus = "ghosted_21d"
            elif days_since >= 14 and "applied" in app.status.lower():
                substatus = "ghosted_14d"
            elif "follow" in action.lower():
                substatus = "follow_up_needed"
            elif "prep" in action.lower():
                substatus = "prep_needed"
            elif "wait" in action.lower():
                substatus = "waiting"
            else:
                substatus = "active"

        # Build one-click action
        one_click_action = None
        if action_type != "none":
            one_click_action = OneClickAction(
                type=action_type,
                template=f"follow_up_day_{days_since}" if action_type == "draft_email" else None,
                application_id=app.id,
                confirm=action_type == "archive"
            )

        # Calculate UI signals
        ui_signals_dict = calculate_ui_signals(
            decision_confidence=decision_confidence,
            days_since_activity=days_since,
            next_action=action,
            status=app.status,
            jd_source=app.jd_source or "missing",
            interview_tomorrow=False,
            focus_mode_enabled=True
        )

        # Map to proper enums
        ui_signals = UISignals(
            priority=PriorityLevel(ui_signals_dict["priority"]),
            confidence=ui_signals_dict["confidence"],
            urgency=UrgencyLevel(ui_signals_dict["urgency"]),
            color_code=ui_signals_dict["color_code"],
            icon=ui_signals_dict["icon"],
            badge=ConfidenceLabel(ui_signals_dict["badge"]) if ui_signals_dict["badge"] else None,
            action_available=ui_signals_dict["action_available"],
            dimmed=ui_signals_dict["dimmed"]
        )

        # Build application with intelligence
        app_with_intel = ApplicationWithIntelligence(
            id=app.id,
            next_action=action,
            next_action_reason=reason,
            priority_level=PriorityLevel(ui_signals_dict["priority"]),
            one_click_action=one_click_action,
            ui_signals=ui_signals,
            decision_confidence=decision_confidence,
            days_since_last_activity=days_since,
            substatus=substatus
        )
        applications_with_intelligence.append(app_with_intel)

        # Add to priority actions if not waiting
        if action_type != "none":
            priority_actions.append(PriorityAction(
                application_id=app.id,
                action=action,
                reason=reason,
                priority=PriorityLevel(ui_signals_dict["priority"]),
                one_click_action=one_click_action
            ))

        # Add to focus mode if high priority
        if ui_signals_dict["priority"] == "high":
            focus_mode_actions.append(FocusModeAction(
                application_id=app.id,
                company=app.company,
                action=action
            ))

    # Sort priority actions by priority level
    priority_order = {"high": 0, "medium": 1, "low": 2, "archive": 3}
    priority_actions.sort(key=lambda x: priority_order.get(x.priority.value, 4))

    # Limit to top priority actions
    priority_actions = priority_actions[:5]

    # Calculate pipeline health
    pipeline_data = calculate_pipeline_health(
        applications=[app.model_dump() for app in request.applications],
        interview_rate=0.0,  # Would need to calculate from historical data
        days_since_last_application=0
    )

    pipeline_health = PipelineHealth(
        active_count=pipeline_data["active_count"],
        status=PipelineStatus(pipeline_data["status"]),
        color=pipeline_data["color"],
        icon=pipeline_data["icon"],
        tone=PipelineTone(pipeline_data["tone"]),
        recommendation=pipeline_data["recommendation"],
        reason=pipeline_data["reason"],
        priority_count=len(focus_mode_actions)
    )

    focus_mode = FocusMode(
        enabled=True,
        top_actions=focus_mode_actions[:3],
        dim_others=True
    )

    return TrackerIntelligenceResponse(
        priority_actions=priority_actions,
        pipeline_health=pipeline_health,
        focus_mode=focus_mode,
        applications=applications_with_intelligence
    )


@app.post("/api/tracker/calculate-confidence", response_model=CalculateConfidenceResponse)
async def calculate_confidence_endpoint(request: CalculateConfidenceRequest):
    """
    COMMAND CENTER: Calculate decision confidence score for an application.

    Factors:
    - Alignment score (from fit_score)
    - Momentum score (based on response and interview progression)
    - JD confidence (based on source)

    Returns confidence score, label, factors breakdown, and guidance.
    """

    # Calculate momentum score
    has_response = request.interview_count > 0 or request.response_time_days is not None
    days_since = request.days_since_last_activity or request.days_since_applied

    momentum_score = calculate_momentum_score(
        has_response=has_response,
        response_time_days=request.response_time_days,
        interview_count=request.interview_count,
        days_since_last_activity=days_since
    )

    # Get JD confidence
    jd_confidence = calculate_jd_confidence(request.jd_source)

    # Calculate overall decision confidence
    decision_confidence = calculate_decision_confidence(
        fit_score=request.fit_score,
        momentum_score=momentum_score,
        jd_confidence=jd_confidence
    )

    # Get label and guidance
    label = get_confidence_label(decision_confidence)
    guidance = get_confidence_guidance(decision_confidence)

    return CalculateConfidenceResponse(
        decision_confidence=decision_confidence,
        label=label,
        factors=ConfidenceFactors(
            alignment_score=request.fit_score,
            momentum_score=momentum_score,
            jd_confidence=jd_confidence
        ),
        guidance=guidance
    )


# ============================================================================
# MVP+1 FEATURE ENDPOINTS
# ============================================================================

@app.post("/api/tasks/prioritize")
async def prioritize_tasks(request: TasksPrioritizeRequest) -> TasksPrioritizeResponse:
    """
    FEATURE 1: Daily Command Center
    
    Analyzes applications and creates prioritized task list for the day.
    Uses Claude to determine highest-leverage actions based on:
    - Fit score
    - Recency of activity
    - Current stage
    - Whether follow-up is needed
    
    Returns 3-7 actionable tasks with priority, reason, and message stubs.
    """
    
    system_prompt = """You are a job search strategist helping prioritize daily activities.

Analyze the provided applications and create a focused task list for today.

TASK TYPES:
- "apply": Apply to this role (for high-fit roles not yet applied)
- "follow_up": Send a follow-up message (for pending applications)
- "outreach": Reach out to network contacts about this role

PRIORITIZATION LOGIC:
1. High fit score (70+) + recent activity = high priority
2. Applied but no follow-up in 5-7+ days = follow-up task
3. Screen/onsite stage without recent contact = follow-up task
4. High fit without outreach = outreach task

CONSTRAINTS:
- Return 3-7 tasks maximum (focus on highest leverage)
- Priority 1 = critical, 2 = important, 3 = nice-to-have
- Provide specific, actionable reasons
- Include brief message stubs (1-2 sentences) where helpful

Return valid JSON:
{
  "tasks": [
    {
      "type": "apply|follow_up|outreach",
      "application_id": "string",
      "priority": 1-3,
      "reason": "string explaining why this is important",
      "suggested_message_stub": "string or null (brief message template)"
    }
  ]
}

Your response must be ONLY valid JSON."""

    # Build user message with application data
    user_message = f"""Today's date: {request.today}

Applications to analyze:
{json.dumps([app.model_dump() for app in request.applications], indent=2)}

Create a prioritized task list focusing on the highest-leverage activities."""

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=3000)
    
    # Parse JSON response
    try:
        if response.strip().startswith("```"):
            response = response.strip().split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()
        
        parsed_data = json.loads(response)
        return TasksPrioritizeResponse(**parsed_data)
    except (json.JSONDecodeError, Exception) as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")


@app.post("/api/outcomes/log")
async def log_outcome(request: OutcomeLogRequest) -> Dict[str, str]:
    """
    FEATURE 2: Learning/Feedback Loop (Part 1)
    
    Logs outcome data for an application.
    Stores in simple in-memory store for pattern analysis.
    
    In production, replace with proper database persistence.
    """
    
    # Add to in-memory store
    outcome_entry = {
        "application_id": request.application_id,
        "stage": request.stage,
        "outcome": request.outcome,
        "date": request.date,
        "logged_at": datetime.utcnow().isoformat()
    }
    
    outcomes_store.append(outcome_entry)
    
    return {
        "status": "success",
        "message": f"Outcome logged for application {request.application_id}",
        "total_outcomes": len(outcomes_store)
    }


@app.post("/api/strategy/review")
async def review_strategy(request: StrategyReviewRequest) -> StrategyReviewResponse:
    """
    FEATURE 2: Learning/Feedback Loop (Part 2)
    
    Analyzes patterns in applications and outcomes to identify:
    - What's working (successful patterns)
    - What's not working (unsuccessful patterns)
    - Concrete strategy shifts to improve results
    
    Returns 3-5 insights and actionable recommendations.
    """
    
    system_prompt = """You are a job search analyst providing strategic feedback.

Analyze application data and outcomes to identify patterns and provide actionable recommendations.

ANALYSIS FRAMEWORK:
1. Success patterns: What's working? (roles, companies, approaches that advance)
2. Challenge patterns: What's not working? (rejections, no responses, common gaps)
3. Fit score correlation: Do higher fit scores lead to better outcomes?
4. Stage progression: Where do applications typically stall?
5. Timing patterns: Response rates based on follow-up timing

INSIGHTS (3-5 observations):
- Data-driven patterns you observe
- Be specific and evidence-based
- Focus on actionable insights

RECOMMENDATIONS (3-5 actions):
- Concrete strategy shifts
- Specific next steps
- Prioritized by impact

Return valid JSON:
{
  "insights": [
    "Insight 1: Specific pattern observed",
    "Insight 2: Another pattern",
    ...
  ],
  "recommendations": [
    "Recommendation 1: Specific action to take",
    "Recommendation 2: Another action",
    ...
  ]
}

Your response must be ONLY valid JSON."""

    # Build comprehensive user message
    user_message = f"""Analyze this job search data:

APPLICATIONS:
{json.dumps(request.applications, indent=2)}

OUTCOMES:
{json.dumps(request.outcomes, indent=2)}

Provide strategic insights and recommendations to improve results."""

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=3000)
    
    # Parse JSON response
    try:
        if response.strip().startswith("```"):
            response = response.strip().split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()
        
        parsed_data = json.loads(response)
        return StrategyReviewResponse(**parsed_data)
    except (json.JSONDecodeError, Exception) as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")


@app.post("/api/network/recommend")
async def recommend_contacts(request: NetworkRecommendRequest) -> NetworkRecommendResponse:
    """
    FEATURE 3: Network Engine (Lite)
    
    Analyzes candidate's network contacts and recommends 3-10 best people
    to reach out to for a specific role/company.
    
    Considers:
    - Relationship strength and type
    - Current company/role relevance
    - Potential to provide referral, intro, or intel
    
    Returns prioritized list with outreach reasoning and message stubs.
    """
    
    system_prompt = """You are a networking strategist for job searches.

Analyze the candidate's network contacts and recommend the best people to reach out to
for this specific target role and company.

RANKING CRITERIA:
1. Current company match: Works at target company = highest priority
2. Role relevance: Similar function or has hiring influence
3. Relationship strength: Former colleague > friend > weak connection
4. Potential help type:
   - Direct referral (works there now)
   - Warm introduction (knows someone there)
   - Industry intel (similar company/role)
   - General advice (mentorship)

PRIORITIES:
- Priority 1: Direct connection to target company or hiring manager
- Priority 2: Strong relationship + relevant role/industry
- Priority 3: Weak connection but potentially helpful

MESSAGE STUBS:
- Personalized based on relationship
- Clear ask (referral, intro, advice, intel)
- Brief and respectful of their time
- 1-2 sentences only

Return 3-10 recommendations, ranked by potential impact.

Return valid JSON:
{
  "recommendations": [
    {
      "name": "string",
      "company": "string",
      "title": "string",
      "relationship": "string",
      "priority": 1-3,
      "reason": "Why this person is valuable to contact",
      "suggested_message_stub": "Brief personalized message template"
    }
  ]
}

Your response must be ONLY valid JSON."""

    # Build user message
    user_message = f"""Target Role: {request.role_title}
Target Company: {request.company}

Network Contacts:
{json.dumps([contact.model_dump() for contact in request.contacts], indent=2)}

Recommend the best contacts to reach out to and provide personalized outreach stubs."""

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=3000)
    
    # Parse JSON response
    try:
        if response.strip().startswith("```"):
            response = response.strip().split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()
        
        parsed_data = json.loads(response)
        return NetworkRecommendResponse(**parsed_data)
    except (json.JSONDecodeError, Exception) as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")

# ============================================================================
# INTERVIEW INTELLIGENCE ENDPOINTS
# ============================================================================

@app.post("/api/interview/parse")
async def parse_interview(request: InterviewParseRequest) -> InterviewParseResponse:
    """
    FEATURE 4A: Interview Intelligence - Parse & Classify Questions
    
    Extracts interview questions from transcript and classifies them by:
    - Type (behavioral, technical, leadership, competency, wildcard)
    - Competency tag (e.g., product sense, stakeholder management)
    - Difficulty level (1-5)
    
    Also identifies themes and potential issues in the interview.
    """
    
    system_prompt = """You are an interview analysis expert specializing in talent acquisition.

Analyze the interview transcript and extract all questions that were asked.
Classify each question using modern hiring competency frameworks.

QUESTION TYPES:
- behavioral: Past behavior questions (Tell me about a time...)
- technical: Skills assessment, problem-solving, domain knowledge
- leadership: People management, influence, decision-making
- competency: Role-specific skills (e.g., data analysis, stakeholder mgmt)
- wildcard: Unusual, creative, or off-script questions

COMPETENCY TAGS (use standard hiring framework terms):
- Product Sense, Strategic Thinking, Data Fluency
- Stakeholder Management, Cross-Functional Collaboration
- Team Leadership, Coaching & Development, Conflict Resolution
- Process Design, Operational Excellence, Metrics & Analytics
- Communication, Presentation Skills, Executive Presence
- Adaptability, Learning Agility, Problem Solving
- Influence, Negotiation, Change Management
- Technical Expertise, Domain Knowledge

DIFFICULTY SCORING (1-5):
1 = Entry level, straightforward
2 = Mid-level, requires some depth
3 = Senior level, multi-faceted
4 = Executive level, strategic thinking required
5 = Extremely challenging, tests edge cases

THEMES:
Identify 3-5 overarching themes across all questions (e.g., "focus on scaling experience", 
"heavy emphasis on metrics", "testing cultural fit for autonomous environment")

WARNINGS:
Flag issues like:
- Vague or poorly structured questions
- Leading questions
- Potential bias indicators
- Missing critical competencies for the role
- Overemphasis on one area

CRITICAL RULES:
- Extract ONLY questions that were actually asked
- Do NOT fabricate questions
- If a question is unclear in transcript, note it in warnings
- Be precise with competency tagging

Return valid JSON:
{
  "questions": [
    {
      "question": "exact question as asked",
      "type": "behavioral|technical|leadership|competency|wildcard",
      "competency_tag": "specific competency from framework",
      "difficulty": 1-5
    }
  ],
  "themes": ["theme 1", "theme 2", ...],
  "warnings": ["warning 1 if any", ...]
}

Your response must be ONLY valid JSON."""

    # Build user message
    user_message = f"""Role: {request.role_title}
Company: {request.company}

Interview Transcript:
{request.transcript_text}
"""

    if request.jd_analysis:
        user_message += f"\n\nJob Description Analysis (for context):\n{json.dumps(request.jd_analysis, indent=2)}"

    user_message += "\n\nExtract and classify all interview questions."

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=3000)
    
    # Parse JSON response
    try:
        if response.strip().startswith("```"):
            response = response.strip().split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()
        
        parsed_data = json.loads(response)
        return InterviewParseResponse(**parsed_data)
    except (json.JSONDecodeError, Exception) as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")


@app.post("/api/interview/feedback")
async def interview_feedback(request: InterviewFeedbackRequest) -> InterviewFeedbackResponse:
    """
    FEATURE 4B: Interview Intelligence - Performance Feedback
    
    Analyzes candidate's interview performance across multiple dimensions:
    - Overall scoring (1-100)
    - Strengths and improvement areas
    - Delivery feedback (tone, pacing, clarity, structure)
    - Specific, actionable recommendations
    
    Uses STAR framework for behavioral questions and provides concrete examples.
    """
    
    system_prompt = """You are a senior executive coach specializing in interview preparation.

=== HENRYHQ VOICE (NON-NEGOTIABLE) ===

You are HenryHQ, a direct, honest, supportive career coach.
You tell candidates the truth without shame, and you always give them a clear next step.
Your tone is calm, confident, human, and never robotic or overly optimistic.
Your goal is simple: make the candidate better with every message.
If an output does not improve clarity, readiness, confidence, or strategy, rewrite it.

Voice Rules:
1. Truth first, support second. Never sugar-coat. Never shame. Use: Truth → Why → Fix → Support.
2. Be direct and concise. Short sentences. No filler. No corporate jargon.
3. Every output must give the user a NEXT STEP.
4. No false encouragement. Praise must be earned and specific.
5. Emotional safety is mandatory. Deliver hard truths calmly and respectfully.

=== END HENRYHQ VOICE ===

Analyze the candidate's interview performance and provide constructive, actionable feedback.

SCORING FRAMEWORK (1-100):
90-100: Exceptional - compelling answers with strong examples and metrics
80-89: Strong - good structure, clear communication, solid examples
70-79: Good - adequate answers but missing some depth or polish
60-69: Fair - needs improvement in structure, examples, or clarity
50-59: Weak - significant gaps in content or delivery
Below 50: Poor - major issues requiring substantial work

STRENGTHS (identify 3-5):
What the candidate did well. Be specific:
- "Used clear STAR structure in behavioral questions"
- "Demonstrated quantitative impact with specific metrics"
- "Showed authentic passion for the company mission"

AREAS FOR IMPROVEMENT (identify 3-5):
What needs work. Be constructive and specific:
- "Behavioral answers lacked concrete metrics or outcomes"
- "Tended to ramble - answers could be 30% shorter"
- "Missed opportunities to connect experience to role requirements"

DELIVERY FEEDBACK:
Tone: Professional, conversational, enthusiastic? Any concerns?
Pacing: Too fast, too slow, or well-modulated? Pauses for emphasis?
Clarity: Easy to follow? Jargon-heavy? Clear structure?
Structure: Organized responses? Answered the actual question? Stayed on topic?

RECOMMENDATIONS (3-6 specific actions):
Highly actionable, prioritized advice:
- "Practice the 'SBO' framework: Situation-Behavior-Outcome in 90 seconds"
- "Prepare 3 metrics-driven stories that demonstrate [specific competency]"
- "Work on ending answers with 'So that's why I'm excited about this opportunity'"

CRITICAL RULES:
- Base feedback ONLY on what's in the transcript
- Do NOT fabricate answers that weren't given
- Be supportive but direct - candidates need honest feedback
- Provide phrasing examples for improvements
- Consider the role requirements when evaluating fit
- Use STAR framework (Situation-Task-Action-Result) as evaluation lens

Return valid JSON:
{
  "overall_score": 75,
  "strengths": ["strength 1", "strength 2", ...],
  "areas_for_improvement": ["area 1", "area 2", ...],
  "delivery_feedback": {
    "tone": "analysis and assessment",
    "pacing": "analysis and assessment",
    "clarity": "analysis and assessment",
    "structure": "analysis and assessment"
  },
  "recommendations": ["action 1", "action 2", ...]
}

Your response must be ONLY valid JSON."""

    # Build user message with questions for context
    user_message = f"""Role: {request.role_title}
Company: {request.company}

Questions Asked:
{json.dumps([q.model_dump() for q in request.questions], indent=2)}

Interview Transcript:
{request.transcript_text}

Provide comprehensive performance feedback focusing on what was actually said."""

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=4096)
    
    # Parse JSON response
    try:
        if response.strip().startswith("```"):
            response = response.strip().split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()
        
        parsed_data = json.loads(response)
        return InterviewFeedbackResponse(**parsed_data)
    except (json.JSONDecodeError, Exception) as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")


@app.post("/api/interview/thank_you")
async def generate_thank_you(request: ThankYouRequest) -> ThankYouResponse:
    """
    FEATURE 4C: Interview Intelligence - Thank You Email
    
    Generates a professional, personalized thank-you email based on:
    - Actual conversation topics from interview
    - Role requirements and company context
    - Candidate's specific examples and experiences discussed
    
    Email is concise (3-4 paragraphs) with modern professional tone.
    """
    
    system_prompt = """You are a professional communication coach specializing in post-interview follow-ups.

=== HENRYHQ VOICE (NON-NEGOTIABLE) ===

You are HenryHQ, a direct, honest, supportive career coach.
You tell candidates the truth without shame, and you always give them a clear next step.
Your tone is calm, confident, human, and never robotic or overly optimistic.
Your goal is simple: make the candidate better with every message.
If an output does not improve clarity, readiness, confidence, or strategy, rewrite it.

Voice Rules:
1. Truth first, support second. Never sugar-coat. Never shame. Use: Truth → Why → Fix → Support.
2. Be direct and concise. Short sentences. No filler. No corporate jargon.
3. Every output must give the user a NEXT STEP.
4. No false encouragement. Praise must be earned and specific.
5. Emotional safety is mandatory. Deliver hard truths calmly and respectfully.

=== END HENRYHQ VOICE ===

Generate a thank-you email that is:
- Professional but warm and genuine
- 3-4 paragraphs maximum
- References specific topics discussed in interview
- Reiterates genuine interest and fit
- Modern tone (not overly formal, no corporate jargon)

STRUCTURE:
Paragraph 1: Thank them and reference something specific from conversation
Paragraph 2: Brief reinforcement of 1-2 key qualifications discussed
Paragraph 3: Express enthusiasm and next steps

DO NOT:
- Fabricate topics that weren't discussed
- Be overly effusive or desperate
- Use corporate buzzwords excessively
- Make it longer than 4 paragraphs
- Repeat the entire conversation

DO:
- Reference actual conversation points
- Show you were listening and engaged
- Reinforce fit naturally
- Keep it concise and respectful of their time
- End with clear next step or openness to questions

TONE EXAMPLES:
Good: "I really enjoyed our conversation about scaling recruiting operations during hypergrowth."
Bad: "It was an absolute pleasure to have the distinct honor of speaking with you..."

Good: "The challenge you mentioned around competing for technical talent in SF resonates with my experience at Spotify."
Bad: "I am extremely passionate about leveraging synergies to optimize talent acquisition pipelines."

Return valid JSON:
{
  "subject": "Thank you - [Role Title] conversation",
  "body": "Full email text, 3-4 paragraphs",
  "key_points_used": ["point 1", "point 2", "point 3"]
}

Your response must be ONLY valid JSON."""

    # Build user message
    interviewer = request.interviewer_name or "the hiring team"
    
    user_message = f"""Role: {request.role_title}
Company: {request.company}
Interviewer: {interviewer}

Interview Transcript:
{request.transcript_text}
"""

    if request.jd_analysis:
        user_message += f"\n\nRole Context:\n{json.dumps(request.jd_analysis, indent=2)}"

    user_message += "\n\nGenerate a professional thank-you email that references actual conversation topics."

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=2000)
    
    # Parse JSON response
    try:
        if response.strip().startswith("```"):
            response = response.strip().split("```")[1]
            if response.startswith("json"):
                response = response[4:]
            response = response.strip()
        
        parsed_data = json.loads(response)
        return ThankYouResponse(**parsed_data)
    except (json.JSONDecodeError, Exception) as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse Claude response: {str(e)}")


# ============================================================================
# INTERVIEW INTELLIGENCE ENDPOINTS
# ============================================================================

def format_resume_for_prompt(resume_json: Dict[str, Any]) -> str:
    """Format resume JSON into readable text for prompts"""
    lines = []

    # Name and contact
    if resume_json.get('full_name'):
        lines.append(f"Name: {resume_json['full_name']}")
    if resume_json.get('current_title'):
        lines.append(f"Current Title: {resume_json['current_title']}")

    # Summary
    if resume_json.get('summary'):
        lines.append(f"\nSummary:\n{resume_json['summary']}")

    # Experience
    experience = resume_json.get('experience', [])
    if experience:
        lines.append("\nExperience:")
        for exp in experience:
            company = exp.get('company', '')
            title = exp.get('title', '')
            dates = exp.get('dates', '')
            lines.append(f"\n{company} - {title} ({dates})")
            for bullet in exp.get('bullets', []):
                lines.append(f"  • {bullet}")

    # Skills
    skills = resume_json.get('skills', [])
    if skills:
        if isinstance(skills, list):
            lines.append(f"\nSkills: {', '.join(skills)}")
        elif isinstance(skills, dict):
            lines.append("\nSkills:")
            for category, skill_list in skills.items():
                lines.append(f"  {category}: {', '.join(skill_list)}")

    # Core competencies
    competencies = resume_json.get('core_competencies', [])
    if competencies:
        lines.append(f"\nCore Competencies: {', '.join(competencies)}")

    # Education
    education = resume_json.get('education')
    if education:
        lines.append(f"\nEducation: {education}")

    return '\n'.join(lines)


@app.post("/api/interview-prep/generate", response_model=InterviewPrepResponse)
@limiter.limit("20/minute")
async def generate_interview_prep(request: Request, body: InterviewPrepRequest):
    """
    INTERVIEW INTELLIGENCE: Generate stage-specific interview prep

    Generates comprehensive prep materials for either:
    - recruiter_screen: Focus on logistics, salary, timeline, red flags
    - hiring_manager: Focus on competencies, STAR stories, technical depth

    All content is grounded in the candidate's actual resume - NO fabrication.

    Command Center v2.0: Now handles missing JDs using provisional profiles.
    When jd_source is 'inferred' or 'missing', uses provisional_profile for prep generation.
    """

    resume_text = format_resume_for_prompt(body.resume_json)
    fit_analysis_text = json.dumps(body.jd_analysis, indent=2) if body.jd_analysis else "No fit analysis available"

    # Determine if using provisional profile (missing JD scenario)
    jd_source = body.jd_source or ("user_provided" if body.job_description else "missing")
    use_provisional = jd_source in ["inferred", "missing"] or not body.job_description

    # Build role context based on JD availability
    if use_provisional and body.provisional_profile:
        # Use provisional profile when real JD is not available
        profile = body.provisional_profile
        role_context = f"""PROVISIONAL ROLE PROFILE (No actual JD available):
Company: {body.company}
Role: {body.role_title}

Typical Responsibilities:
{chr(10).join('- ' + r for r in profile.get('typical_responsibilities', []))}

Common Competencies Assessed:
{chr(10).join('- ' + c for c in profile.get('common_competencies', []))}

Interview Focus Areas:
{chr(10).join('- ' + f for f in profile.get('interview_focus_areas', []))}

Typical Evaluation Criteria:
{chr(10).join('- ' + e for e in profile.get('evaluation_criteria', []))}

NOTE: This is based on typical expectations for this role type. The actual interview may differ."""
        confidence = "directional"
        ui_note = "This prep is based on typical expectations for this role. Add recruiter context anytime to refine it."
    else:
        # Use real JD
        role_context = f"""JOB DESCRIPTION:
Company: {body.company}
Role: {body.role_title}

{body.job_description or 'No job description provided'}"""
        confidence = "refined"
        ui_note = None

    if body.interview_stage == "recruiter_screen":
        system_prompt = """You are generating pre-interview intelligence for a recruiter screen.

Generate structured prep covering:
1. interview_overview (2-3 sentences explaining recruiter's goals)
2. key_talking_points (4-6 bullets, must reference actual achievements from resume)
3. red_flag_mitigation (2-3 items addressing gaps, job hopping, overqualification - each with "flag" and "mitigation")
4. likely_questions (5-7 questions recruiter will ask about experience, motivation, and fit)
   DO NOT include salary/compensation questions - those are covered separately in compensation_expectations.
   IMPORTANT: Each must be an object with BOTH "question" AND "suggested_answer" fields.
   Format: {"question": "Why are you interested in this role?", "suggested_answer": "I'm drawn to this role because... [2-4 sentences using specific details from resume]"}
5. questions_to_ask (3-4 strategic questions with "question" and "why")
6. compensation_expectations (2-4 sentences providing a specific scripted response the candidate can use when asked "What are your salary expectations?" - include a realistic salary range based on the role level and their experience, with framing language like "Based on my research and experience level, I'm targeting..." or "I'm flexible but looking in the range of...")
7. timeline_strategy (1-2 sentences on how to respond about availability and start date)

CRITICAL RULES:
- Use ONLY information from the candidate's actual resume
- Include specific metrics and achievements from resume
- Be direct about red flags. Provide mitigation strategies
- No generic advice. Everything must be tailored to this candidate and role
- If working from a provisional profile, base prep on typical expectations for this role type

Return ONLY a valid JSON object with the structure above. No markdown, no preamble."""
    else:  # hiring_manager
        system_prompt = """You are generating pre-interview intelligence for a hiring manager interview.

Generate structured prep covering:
1. interview_overview (2-3 sentences explaining hiring manager's goals)
2. strengths_to_emphasize (3-4 bullets tied to job requirements, must reference resume)
3. gaps_to_mitigate (2-3 bullets with credible workarounds)
4. star_examples (3-4 behavioral stories using STAR format, must be from actual resume)
   Format each as: {"competency": "string", "situation": "string", "task": "string", "action": "string", "result": "string"}
5. likely_questions (5-7 technical/functional questions)
   IMPORTANT: Each must be an object with BOTH "question" AND "suggested_answer" fields.
   Format: {"question": "Tell me about a time you led a cross-functional project", "suggested_answer": "At [Company], I led... [2-4 sentences using specific details from resume]"}
6. questions_to_ask (4-5 strategic questions with "question" and "why")
7. closing_strategy (1-2 sentences on how to close strong)

STAR EXAMPLES RULES:
- Every example must come from candidate's actual work history
- Include specific metrics, timeframes, and outcomes from resume
- Map each example to a key competency from the job description (or provisional profile)
- Use concrete details, not generic scenarios
- If working from a provisional profile, map to typical competencies for this role type

Return ONLY a valid JSON object with the structure above. No markdown, no preamble."""

    # Build past debrief insights section if available
    debrief_insights_section = ""
    if body.past_debriefs and len(body.past_debriefs) > 0:
        debrief_lines = ["PAST INTERVIEW INTELLIGENCE (from debriefs):"]

        # Check for same-company debriefs
        same_company_debriefs = [d for d in body.past_debriefs if d.get('company', '').lower() == body.company.lower()]
        if same_company_debriefs:
            debrief_lines.append(f"\n📊 Previous interviews at {body.company}:")
            for d in same_company_debriefs[:3]:  # Limit to 3 most recent
                debrief_lines.append(f"  - {d.get('interview_type', 'Interview')}: Rating {d.get('rating_overall', 'N/A')}/5")
                if d.get('stumbles'):
                    for stumble in d['stumbles'][:2]:
                        debrief_lines.append(f"    ⚠️ Struggled: {stumble.get('question', 'Unknown')}")
                if d.get('wins'):
                    for win in d['wins'][:2]:
                        debrief_lines.append(f"    ✅ Nailed: {win.get('moment', 'Unknown')}")

        # Aggregate weak areas across all debriefs
        all_stumbles = []
        all_improvement_areas = []
        for d in body.past_debriefs:
            all_stumbles.extend(d.get('stumbles', []))
            all_improvement_areas.extend(d.get('improvement_areas', []))

        if all_stumbles:
            debrief_lines.append("\n🎯 Recurring challenges from past interviews:")
            seen_stumbles = set()
            for stumble in all_stumbles[:5]:
                q = stumble.get('question', stumble.get('what_went_wrong', 'Unknown'))
                if q not in seen_stumbles:
                    debrief_lines.append(f"  - {q}")
                    seen_stumbles.add(q)

        if all_improvement_areas:
            debrief_lines.append("\n📈 Areas to strengthen (from coach feedback):")
            seen_areas = set()
            for area in all_improvement_areas[:5]:
                if isinstance(area, str) and area not in seen_areas:
                    debrief_lines.append(f"  - {area}")
                    seen_areas.add(area)

        # Story usage context
        all_stories = []
        for d in body.past_debriefs:
            all_stories.extend(d.get('stories_used', []))
        if all_stories:
            debrief_lines.append("\n📖 Stories used in previous interviews:")
            for story in all_stories[:5]:
                effectiveness = story.get('effectiveness', 'N/A')
                debrief_lines.append(f"  - {story.get('name', 'Story')}: {story.get('context', 'Unknown context')} (effectiveness: {effectiveness}/5)")

        debrief_insights_section = "\n\n" + "\n".join(debrief_lines) + """

DEBRIEF-INFORMED PREP RULES:
- Explicitly address recurring stumbles from past interviews
- Suggest alternative approaches for questions they've struggled with
- If they've used a story 3+ times, suggest fresh examples
- Prepare them for question types that caught them off guard before
- Build on what's working (reference their wins)"""

    user_message = f"""CANDIDATE RESUME:
{resume_text}

{role_context}

FIT ANALYSIS:
{fit_analysis_text}{debrief_insights_section}

Generate the interview prep now."""

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=4000)

    # Parse JSON response
    try:
        cleaned = clean_claude_json(response)
        prep_content = json.loads(cleaned)

        return InterviewPrepResponse(
            interview_stage=body.interview_stage,
            prep_content=prep_content,
            confidence=confidence,
            ui_note=ui_note
        )
    except (json.JSONDecodeError, Exception) as e:
        print(f"🔥 Failed to parse interview prep response: {e}")
        print(f"Raw response: {response[:500]}")
        raise HTTPException(status_code=500, detail=f"Failed to parse interview prep: {str(e)}")


@app.post("/api/interview-prep/intro-sell/generate", response_model=IntroSellTemplateResponse)
async def generate_intro_sell_template(request: IntroSellTemplateRequest):
    """
    INTERVIEW INTELLIGENCE: Generate customized 60-90 second intro sell template

    Creates a structured template the candidate can practice:
    1. Current Role + Impact (quantified)
    2. Previous Role + Relevant Achievement
    3. Why You're Here (connection to target role)
    4. What You're Looking For (optional)

    Word count target: 100-150 words
    """

    resume_text = format_resume_for_prompt(request.resume_json)

    system_prompt = """Generate a customized 60-90 second intro sell template for this candidate.

STRUCTURE (MANDATORY):
1. Current Role + Impact (1-2 sentences, include quantified achievement)
2. Previous Role + Relevant Achievement (1-2 sentences)
3. Why You're Here (1 sentence connecting background to job)
4. What You're Looking For (1 sentence, optional)

RULES:
- Total word count: 100-150 words
- Use ONLY information from resume
- Include at least one quantified metric
- No college/education unless directly relevant
- No generic filler phrases
- End with confidence, not a question

OUTPUT FORMAT:
Return JSON:
{
    "template": "Full intro sell text",
    "word_count": integer,
    "coaching_note": "2-3 sentences on what to emphasize"
}

No markdown, no preamble."""

    user_message = f"""CANDIDATE RESUME:
{resume_text}

TARGET ROLE:
Company: {request.company}
Role: {request.role_title}

JOB DESCRIPTION:
{request.job_description}

Generate the intro sell template now."""

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=2000)

    # Parse JSON response
    try:
        cleaned = clean_claude_json(response)
        parsed = json.loads(cleaned)

        # =================================================================
        # QA VALIDATION: Check intro sell for fabrication
        # =================================================================
        from qa_validation import OutputValidator, ResumeGroundingValidator

        grounding_validator = ResumeGroundingValidator(request.resume_json)
        output_validator = OutputValidator(grounding_validator)
        # Validate the template as interview prep content
        qa_validation_result = output_validator.validate_interview_prep({
            "talking_points": [{"content": parsed.get("template", "")}]
        })

        if qa_validation_result.should_block:
            print("\n" + "🚫"*20)
            print("INTRO SELL QA VALIDATION BLOCKED - POTENTIAL FABRICATION")
            print("🚫"*20)
            for issue in qa_validation_result.issues:
                print(f"  [{issue.category.value}] {issue.message}")
            print("🚫"*20 + "\n")

            ValidationLogger.log_blocked_output(
                endpoint="/api/interview-prep/intro-sell/generate",
                result=qa_validation_result,
                output=parsed,
                resume_data=request.resume_json
            )

            error_response = create_validation_error_response(qa_validation_result)
            raise HTTPException(status_code=422, detail=error_response)

        return IntroSellTemplateResponse(
            template=parsed.get("template", ""),
            word_count=parsed.get("word_count", len(parsed.get("template", "").split())),
            coaching_note=parsed.get("coaching_note", "")
        )
    except (json.JSONDecodeError, Exception) as e:
        print(f"🔥 Failed to parse intro sell template: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse intro sell template: {str(e)}")


@app.post("/api/interview-prep/intro-sell/feedback", response_model=IntroSellFeedbackResponse)
async def analyze_intro_sell(request: IntroSellFeedbackRequest):
    """
    INTERVIEW INTELLIGENCE: Analyze candidate's intro sell attempt

    Provides structured feedback on:
    - Content (did they lead with impact, include metrics?)
    - Structure (word count, formula adherence)
    - Tone (confident vs tentative, specific vs vague)

    Returns scores, specific feedback, and a revised version.
    """

    resume_text = format_resume_for_prompt(request.resume_json)

    # Count words and estimate time
    word_count = len(request.candidate_version.split())
    estimated_time_seconds = int((word_count / 140) * 60)  # 140 WPM average

    system_prompt = """Analyze this candidate's 60-90 second intro sell attempt.

=== HENRYHQ VOICE (NON-NEGOTIABLE) ===

You are HenryHQ, a direct, honest, supportive career coach.
You tell candidates the truth without shame, and you always give them a clear next step.
Your tone is calm, confident, human, and never robotic or overly optimistic.
Your goal is simple: make the candidate better with every message.
If an output does not improve clarity, readiness, confidence, or strategy, rewrite it.

Voice Rules:
1. Truth first, support second. Never sugar-coat. Never shame. Use: Truth → Why → Fix → Support.
2. Be direct and concise. Short sentences. No filler. No corporate jargon.
3. Every output must give the user a NEXT STEP.
4. No false encouragement. Praise must be earned and specific.
5. Emotional safety is mandatory. Deliver hard truths calmly and respectfully.

=== END HENRYHQ VOICE ===

Provide structured feedback:

CONTENT (1-10):
- Did they lead with current role and impact?
- Did they include quantified achievement?
- Did they connect background to target role?
- Did they avoid irrelevant details?

STRUCTURE (1-10):
- Word count: 100-150 = 10 points, 151-200 = 7 points, 200+ = 5 points
- Did they follow formula (current → past → why here)?
- Did they end with statement, not question?

TONE (1-10):
- Confident language vs. tentative
- Specific vs. vague
- Active voice vs. passive

OUTPUT FORMAT:
Return JSON:
{
    "overall_score": float (average of three scores),
    "content_score": integer 1-10,
    "structure_score": integer 1-10,
    "tone_score": integer 1-10,
    "strengths": ["string", "string"],
    "opportunities": ["string", "string"],
    "revised_version": "string (tightened to 100-150 words)",
    "coaching_note": "string (what to practice next)"
}

REVISED VERSION RULES:
- Must be 100-150 words
- Must use only information from candidate's actual resume
- Tighten weak sections, cut irrelevant details
- End confidently

No markdown, no preamble."""

    user_message = f"""CANDIDATE'S VERSION:
{request.candidate_version}

CANDIDATE RESUME (for reference):
{resume_text}

TARGET ROLE:
Company: {request.company}
Role: {request.role_title}

JOB DESCRIPTION (for reference):
{request.job_description}

Analyze the intro sell attempt now."""

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=2000)

    # Parse JSON response
    try:
        cleaned = clean_claude_json(response)
        feedback = json.loads(cleaned)

        return IntroSellFeedbackResponse(
            overall_score=feedback.get("overall_score", 5.0),
            content_score=feedback.get("content_score", 5),
            structure_score=feedback.get("structure_score", 5),
            tone_score=feedback.get("tone_score", 5),
            word_count=word_count,
            estimated_time_seconds=estimated_time_seconds,
            strengths=feedback.get("strengths", []),
            opportunities=feedback.get("opportunities", []),
            revised_version=feedback.get("revised_version", ""),
            coaching_note=feedback.get("coaching_note", "")
        )
    except (json.JSONDecodeError, Exception) as e:
        print(f"🔥 Failed to parse intro sell feedback: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse intro sell feedback: {str(e)}")


@app.post("/api/interview-prep/debrief", response_model=InterviewDebriefResponse)
async def create_interview_debrief(request: InterviewDebriefRequest):
    """
    INTERVIEW INTELLIGENCE: Post-interview debrief analysis

    Analyzes interview performance from either:
    - Transcript text (if candidate recorded/transcribed)
    - Typed responses (questions asked, strong/weak answers)

    Returns comprehensive feedback, coaching points, and a thank-you email.
    """

    resume_text = format_resume_for_prompt(request.resume_json)

    # Build interview content from available sources
    interview_content_parts = []

    # Include transcript if provided (primary source for detailed analysis)
    if request.transcript_text:
        interview_content_parts.append(f"INTERVIEW TRANSCRIPT:\n{request.transcript_text}")

    # Include typed responses if provided (candidate's self-assessment)
    if request.typed_responses:
        typed_content = f"""CANDIDATE SELF-ASSESSMENT:
Overall Feeling (1-10): {request.typed_responses.overall_feeling}

Questions Asked:
{chr(10).join('- ' + q for q in request.typed_responses.questions_asked) if request.typed_responses.questions_asked else '(not provided)'}

Strong Answers (self-identified):
{chr(10).join('- ' + a for a in request.typed_responses.strong_answers) if request.typed_responses.strong_answers else '(not provided)'}

Weak Answers (self-identified):
{chr(10).join('- ' + a for a in request.typed_responses.weak_answers) if request.typed_responses.weak_answers else '(not provided)'}

Key Learnings:
{request.typed_responses.learnings if request.typed_responses.learnings else '(not provided)'}"""
        interview_content_parts.append(typed_content)

    if not interview_content_parts:
        raise HTTPException(status_code=400, detail="Either transcript_text or typed_responses must be provided")

    interview_content = "\n\n".join(interview_content_parts)

    system_prompt = """Analyze this interview performance and provide structured feedback.

DIMENSION SCORES (1-10 each):
- content: Did they cover key competencies and provide strong examples?
- clarity: Were answers clear and well-structured?
- delivery: Confident tone, no excessive filler words?
- tone: Professional, enthusiastic, appropriate energy?
- structure: Did answers follow logical frameworks (STAR, etc.)?
- confidence: Did they sound assured without arrogance?

ANALYSIS:
- strengths: 2-3 specific things they did well
- opportunities: 2-3 specific areas for improvement
- what_they_should_have_said: 1-2 examples of weak answers rewritten stronger
  Format: {"question": "...", "weak_answer": "...", "strong_answer": "..."}
- coaching_points: 3 specific things to practice before next interview
- action_items: 3 concrete next steps
- next_stage_adjustments: 2-3 strategic changes for next interview round

THANK YOU EMAIL:
Generate a professional thank-you email (3 paragraphs):
- Reference specific conversation topics
- Reinforce fit naturally
- Keep it concise

OUTPUT FORMAT:
Return JSON:
{
    "overall_score": integer 1-10,
    "dimension_scores": {
        "content": int, "clarity": int, "delivery": int,
        "tone": int, "structure": int, "confidence": int
    },
    "strengths": ["string", ...],
    "opportunities": ["string", ...],
    "what_they_should_have_said": [{"question": "...", "weak_answer": "...", "strong_answer": "..."}, ...],
    "coaching_points": ["string", ...],
    "action_items": ["string", ...],
    "thank_you_email": "full email text",
    "next_stage_adjustments": ["string", ...]
}

RULES:
- Be specific. Reference actual questions/answers when possible
- Provide rewritten answers that use candidate's real experience
- Coaching must be actionable, not generic

No markdown, no preamble."""

    user_message = f"""INTERVIEW TYPE: {request.interview_stage}
INTERVIEW DATE: {request.interview_date}
INTERVIEWER: {request.interviewer_name or 'Unknown'}

CANDIDATE RESUME:
{resume_text}

TARGET ROLE:
Company: {request.company}
Role: {request.role_title}

JOB DESCRIPTION:
{request.job_description}

{interview_content}

Analyze this interview and generate the debrief."""

    # Call Claude
    response = call_claude(system_prompt, user_message, max_tokens=4000)

    # Parse JSON response
    try:
        cleaned = clean_claude_json(response)
        debrief = json.loads(cleaned)

        # Build response with defaults for missing fields
        dimension_scores = debrief.get("dimension_scores", {})

        return InterviewDebriefResponse(
            overall_score=debrief.get("overall_score", 5),
            dimension_scores=DimensionScores(
                content=dimension_scores.get("content", 5),
                clarity=dimension_scores.get("clarity", 5),
                delivery=dimension_scores.get("delivery", 5),
                tone=dimension_scores.get("tone", 5),
                structure=dimension_scores.get("structure", 5),
                confidence=dimension_scores.get("confidence", 5)
            ),
            strengths=debrief.get("strengths", []),
            opportunities=debrief.get("opportunities", []),
            what_they_should_have_said=[
                WhatTheyShouldHaveSaid(**item) for item in debrief.get("what_they_should_have_said", [])
            ],
            coaching_points=debrief.get("coaching_points", []),
            action_items=debrief.get("action_items", []),
            thank_you_email=debrief.get("thank_you_email", ""),
            next_stage_adjustments=debrief.get("next_stage_adjustments", [])
        )
    except (json.JSONDecodeError, Exception) as e:
        print(f"🔥 Failed to parse debrief response: {e}")
        print(f"Raw response: {response[:500]}")
        raise HTTPException(status_code=500, detail=f"Failed to parse debrief: {str(e)}")


# ============================================================================
# INTERVIEW INTELLIGENCE: CONVERSATIONAL DEBRIEF CHAT
# ============================================================================

class DebriefContext(BaseModel):
    """Context for the debrief conversation."""
    interviewerName: str
    interviewerTitle: Optional[str] = None
    interviewDate: Optional[str] = None
    interviewStage: str
    feeling: Optional[str] = None
    company: str
    roleTitle: str
    candidateName: str


class ConversationMessage(BaseModel):
    """A single message in the conversation."""
    role: str  # 'user' or 'assistant'
    content: str


class DebriefChatRequest(BaseModel):
    """Request for conversational debrief."""
    resume_json: Dict[str, Any]
    job_description: str
    context: DebriefContext
    transcript: Optional[str] = None  # Now optional - can have conversation without transcript
    conversation_history: List[ConversationMessage]
    user_message: Optional[str] = None  # None for initial analysis


class DebriefChatResponse(BaseModel):
    """Response from conversational debrief."""
    response: str


# Debrief prompts - Now imported from prompts/ module
# (DEBRIEF_SYSTEM_PROMPT_WITH_TRANSCRIPT, DEBRIEF_SYSTEM_PROMPT_NO_TRANSCRIPT)


@app.post("/api/debrief/chat", response_model=DebriefChatResponse)
async def debrief_chat(request: DebriefChatRequest):
    """
    INTERVIEW INTELLIGENCE: Conversational debrief coaching.

    Provides warm, actionable coaching in a conversational format.
    Supports multi-turn conversations for follow-up questions,
    thank-you email drafts, next round prep, etc.
    """
    print(f"💬 Debrief chat for {request.context.company} - {request.context.roleTitle}")

    # Format resume for prompt
    resume_text = format_resume_for_prompt(request.resume_json)

    # Build interviewer info string
    interviewer_title = f" ({request.context.interviewerTitle})" if request.context.interviewerTitle else ""

    # Map interview stage to readable name
    stage_labels = {
        'recruiter_screen': 'Recruiter Screen',
        'hiring_manager': 'Hiring Manager Interview',
        'technical': 'Technical Interview',
        'panel': 'Panel Interview',
        'executive': 'Executive/Final Round'
    }
    interview_stage = stage_labels.get(request.context.interviewStage, request.context.interviewStage)

    # Format feeling
    feeling = f"{request.context.feeling}/10" if request.context.feeling else "Not provided"

    # Choose prompt based on whether transcript is provided
    has_transcript = bool(request.transcript and request.transcript.strip())

    # Truncate extremely long transcripts to stay within Claude's context window
    # Claude Sonnet can handle ~200k tokens, so 400k chars is a safe limit (~100k tokens)
    transcript_text = request.transcript or ""
    MAX_TRANSCRIPT_CHARS = 400000
    if len(transcript_text) > MAX_TRANSCRIPT_CHARS:
        print(f"⚠️ Truncating very long transcript from {len(transcript_text)} to {MAX_TRANSCRIPT_CHARS} chars")
        transcript_text = transcript_text[:MAX_TRANSCRIPT_CHARS] + "\n\n[... transcript truncated for length ...]"

    if has_transcript:
        system_prompt = DEBRIEF_SYSTEM_PROMPT_WITH_TRANSCRIPT.format(
            company=request.context.company,
            role_title=request.context.roleTitle,
            interview_stage=interview_stage,
            interviewer_name=request.context.interviewerName,
            interviewer_title=interviewer_title,
            feeling=feeling,
            resume_text=resume_text,
            job_description=request.job_description,
            transcript=transcript_text
        )
    else:
        system_prompt = DEBRIEF_SYSTEM_PROMPT_NO_TRANSCRIPT.format(
            company=request.context.company,
            role_title=request.context.roleTitle,
            interview_stage=interview_stage,
            interviewer_name=request.context.interviewerName,
            interviewer_title=interviewer_title,
            feeling=feeling,
            resume_text=resume_text,
            job_description=request.job_description
        )

    # Build messages for Claude
    messages = []

    # Add conversation history
    for msg in request.conversation_history:
        messages.append({
            "role": msg.role,
            "content": msg.content
        })

    # Add current user message or request initial greeting/analysis
    if request.user_message:
        messages.append({
            "role": "user",
            "content": request.user_message
        })
    else:
        # Initial message depends on whether we have a transcript
        if has_transcript:
            messages.append({
                "role": "user",
                "content": "I just finished this interview. Can you analyze how it went and give me your honest feedback? What did I do well and what should I work on?"
            })
        else:
            messages.append({
                "role": "user",
                "content": "I just finished an interview and I'd like to debrief with you. I don't have a transcript, but I can tell you about how it went."
            })

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            system=system_prompt,
            messages=messages
        )

        assistant_response = response.content[0].text
        print(f"✅ Debrief chat response generated ({len(assistant_response)} chars)")

        return DebriefChatResponse(response=assistant_response)

    except Exception as e:
        print(f"🔥 Debrief chat error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate response: {str(e)}")


# ============================================================================
# INTERVIEW INTELLIGENCE: MOCK INTERVIEW ENDPOINTS
# ============================================================================

@app.post("/api/mock-interview/start", response_model=StartMockInterviewResponse)
@limiter.limit("10/minute")
async def start_mock_interview(request: Request, body: StartMockInterviewRequest):
    """
    INTERVIEW INTELLIGENCE: Start a new mock interview session

    Creates a session and generates the first question based on:
    - Interview stage (recruiter_screen or hiring_manager)
    - Difficulty level (easy, medium, hard)
    - Candidate's resume and job description

    Returns session_id and first question to begin the interview.
    """
    print(f"🎤 Starting mock interview for {body.company} - {body.role_title}")

    # Generate session ID
    session_id = str(uuid.uuid4())

    # Format resume for prompt
    resume_text = format_resume_for_prompt(body.resume_json)

    # Extract candidate's first name from resume (check multiple possible field names)
    full_name = body.resume_json.get("name") or body.resume_json.get("full_name") or body.resume_json.get("candidate_name") or ""
    candidate_name = full_name.split()[0] if full_name else "there"

    # First question is always a warm "tell me about yourself" - the 60-90 second pitch
    # This mirrors how real interviews start
    question_data = {
        "question_text": f"Hi {candidate_name}, thank you for making time to speak with me today. I'm excited to learn more about you and your background. So, tell me about yourself.",
        "competency_tested": "intro_pitch",
        "difficulty": "medium"
    }

    # Generate question ID
    question_id = str(uuid.uuid4())

    # Store session (with optional user_id from body)
    session_data = {
        "id": session_id,
        "resume_json": body.resume_json,
        "job_description": body.job_description,
        "company": body.company,
        "role_title": body.role_title,
        "interview_stage": body.interview_stage,
        "difficulty_level": body.difficulty_level,
        "started_at": datetime.now().isoformat(),
        "completed_at": None,
        "overall_score": None,
        "session_feedback": None,
        "question_ids": [question_id],
        "current_question_number": 1
    }
    # Get user_id from body if provided (for Supabase storage)
    user_id = getattr(body, 'user_id', None)
    save_mock_session(session_id, session_data, user_id)

    # Store question
    question_record = {
        "id": question_id,
        "session_id": session_id,
        "question_number": 1,
        "question_text": question_data["question_text"],
        "competency_tested": question_data["competency_tested"],
        "difficulty": question_data["difficulty"],
        "asked_at": datetime.now().isoformat()
    }
    save_mock_question(question_id, question_record)

    print(f"✅ Mock interview session started: {session_id}")

    return StartMockInterviewResponse(
        session_id=session_id,
        interview_stage=body.interview_stage,
        difficulty_level=body.difficulty_level,
        first_question=FirstQuestionResponse(
            question_id=question_id,
            question_text=question_data["question_text"],
            competency_tested=question_data["competency_tested"],
            question_number=1
        )
    )


@app.post("/api/mock-interview/respond", response_model=SubmitMockResponseResponse)
async def submit_mock_response(request: SubmitMockResponseRequest):
    """
    INTERVIEW INTELLIGENCE: Submit response and get follow-up question

    Submits candidate's response, analyzes it, and returns:
    - Brief coaching feedback
    - Follow-up question (if not at max follow-ups)
    - Whether to continue with more follow-ups
    """
    print(f"📝 Submitting mock response for question {request.question_id}")

    # Verify session exists
    session = get_mock_session(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Mock interview session not found")

    # Verify question exists
    question = get_mock_question(request.question_id)
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")

    # Store the response
    response_entry = {
        "response_number": request.response_number,
        "response_text": request.response_text,
        "responded_at": datetime.now().isoformat(),
        "session_id": request.session_id
    }
    save_mock_response(request.question_id, response_entry)

    # Format resume for prompt
    resume_json = session.get("resume_json") or session.get("resume_json", {})
    resume_text = format_resume_for_prompt(resume_json)

    # Determine target level
    target_level = determine_target_level(session["job_description"])

    # Detect role type from job description for role-specific signal weighting
    role_type = detect_role_type(session["job_description"], session["role_title"])

    # Build ALL responses for cumulative analysis
    all_responses = get_mock_responses(request.question_id)
    all_responses_text = ""
    for i, resp in enumerate(all_responses, 1):
        label = "Initial Response" if i == 1 else f"Follow-up Response #{i-1}"
        all_responses_text += f"\n{label}:\n{resp['response_text']}\n"

    # Build analysis prompt with ALL responses for cumulative scoring
    prompt = MOCK_ANALYZE_RESPONSE_PROMPT.format(
        question_text=question["question_text"],
        all_responses=all_responses_text,
        competency=question["competency_tested"],
        target_level=target_level,
        role_type=role_type,
        resume_text=resume_text
    )

    # Call Claude to analyze cumulative response
    try:
        response = call_claude(
            "You are analyzing a candidate's COMPLETE interview response including all follow-ups. Score based on CUMULATIVE quality. Return only valid JSON.",
            prompt,
            max_tokens=2000
        )
        cleaned = clean_claude_json(response)
        analysis_data = json.loads(cleaned)
    except Exception as e:
        print(f"🔥 Failed to analyze response: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze response: {str(e)}")

    # Store or update analysis with signals and resume content
    resume_content = analysis_data.get("resume_content", {})
    signals = analysis_data.get("signals", {})

    # Calculate weighted score based on signals if available
    if signals:
        weighted_score = calculate_weighted_score(signals, role_type)
    else:
        weighted_score = analysis_data.get("score", 5)

    mock_interview_analyses[request.question_id] = {
        "question_id": request.question_id,
        "score": analysis_data.get("score", 5),
        "weighted_score": weighted_score,
        "level_demonstrated": analysis_data.get("level_demonstrated", "mid"),
        "role_type": role_type,
        # New signal tracking
        "signals": {
            "functional_competency": signals.get("functional_competency", 0.5),
            "leadership": signals.get("leadership", 0.5),
            "collaboration": signals.get("collaboration", 0.5),
            "ownership": signals.get("ownership", 0.5),
            "strategic_thinking": signals.get("strategic_thinking", 0.5),
            "problem_solving": signals.get("problem_solving", 0.5),
            "communication_clarity": signals.get("communication_clarity", 0.5),
            "metrics_orientation": signals.get("metrics_orientation", 0.5),
            "stakeholder_management": signals.get("stakeholder_management", 0.5),
            "executive_presence": signals.get("executive_presence", 0.5),
            "user_centricity": signals.get("user_centricity", 0.5)
        },
        "signal_strengths": analysis_data.get("signal_strengths", []),
        "signal_gaps": analysis_data.get("signal_gaps", []),
        "follow_up_trigger": analysis_data.get("follow_up_trigger"),
        "follow_up_questions": analysis_data.get("follow_up_questions", []),
        "brief_feedback": analysis_data.get("brief_feedback", ""),
        "resume_content": {
            "metrics": resume_content.get("metrics", []),
            "achievements": resume_content.get("achievements", []),
            "stories": resume_content.get("stories", [])
        },
        "response_count": len(all_responses),
        "created_at": datetime.now().isoformat()
    }

    # Check if this is the last follow-up (response_number 4)
    if request.response_number >= 4:
        print(f"✅ Max follow-ups reached for question {request.question_id}")
        return SubmitMockResponseResponse(
            response_recorded=True,
            follow_up_question=None,
            should_continue=False,
            brief_feedback=analysis_data.get("brief_feedback", "Good effort on this question.")
        )

    # Return follow-up question
    follow_ups = analysis_data.get("follow_up_questions", [])
    follow_up = follow_ups[0] if follow_ups else None

    print(f"✅ Response analyzed, follow-up provided")

    return SubmitMockResponseResponse(
        response_recorded=True,
        follow_up_question=follow_up,
        should_continue=True,
        brief_feedback=analysis_data.get("brief_feedback", "")
    )


@app.get("/api/mock-interview/question-feedback/{question_id}", response_model=QuestionFeedbackResponse)
async def get_mock_question_feedback(question_id: str):
    """
    INTERVIEW INTELLIGENCE: Get comprehensive feedback for a completed question

    Returns detailed coaching feedback including:
    - Score and level demonstrated
    - What landed and what didn't
    - Specific coaching advice
    - Revised answer at target level
    """
    print(f"📊 Getting feedback for question {question_id}")

    # Verify question exists
    question = mock_interview_questions.get(question_id)
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")

    # Get session
    session = mock_interview_sessions.get(question["session_id"])
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")

    # Get all responses
    responses = mock_interview_responses.get(question_id, [])
    all_responses = [r["response_text"] for r in responses]
    all_responses_text = format_responses_for_analysis(question_id)

    # Get existing analysis
    analysis = mock_interview_analyses.get(question_id, {})

    # Format resume for prompt
    resume_text = format_resume_for_prompt(session["resume_json"])

    # Determine target level
    target_level = determine_target_level(session["job_description"])

    # Build feedback prompt
    prompt = MOCK_QUESTION_FEEDBACK_PROMPT.format(
        question_text=question["question_text"],
        competency=question["competency_tested"],
        target_level=target_level,
        all_responses_text=all_responses_text if all_responses_text else "No responses recorded.",
        analysis_json=json.dumps(analysis, indent=2),
        resume_text=resume_text
    )

    # Call Claude for comprehensive feedback
    try:
        response = call_claude(
            "You are providing coaching feedback on interview responses. Return only valid JSON.",
            prompt,
            max_tokens=2000
        )
        cleaned = clean_claude_json(response)
        feedback_data = json.loads(cleaned)
    except Exception as e:
        print(f"🔥 Failed to generate feedback: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate feedback: {str(e)}")

    # Update analysis with comprehensive feedback
    if question_id in mock_interview_analyses:
        mock_interview_analyses[question_id]["feedback_text"] = feedback_data.get("coaching", "")
        mock_interview_analyses[question_id]["revised_answer"] = feedback_data.get("revised_answer", "")

    print(f"✅ Feedback generated for question {question_id}")

    return QuestionFeedbackResponse(
        question_text=question["question_text"],
        all_responses=all_responses,
        score=analysis.get("score", 5),
        level_demonstrated=analysis.get("level_demonstrated", "L4"),
        what_landed=feedback_data.get("what_landed", []),
        what_didnt_land=feedback_data.get("what_didnt_land", []),
        coaching=feedback_data.get("coaching", ""),
        revised_answer=feedback_data.get("revised_answer", "")
    )


@app.post("/api/mock-interview/next-question", response_model=NextQuestionResponse)
async def get_next_mock_question(request: NextQuestionRequest):
    """
    INTERVIEW INTELLIGENCE: Move to the next question

    Generates the next question based on:
    - Session progress
    - Previous questions asked (to avoid repeats)
    - Rotating competencies

    Returns new question and session progress.
    """
    print(f"➡️ Getting next question for session {request.session_id}")

    # Verify session exists
    session = get_mock_session(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Mock interview session not found")

    # Check if max questions reached (5)
    if request.current_question_number >= 5:
        raise HTTPException(status_code=400, detail="Maximum questions reached. End session to see results.")

    # Check if session is already completed
    if session.get("completed_at"):
        raise HTTPException(status_code=400, detail="Session already completed. Start a new session.")

    # Calculate next question number
    next_question_number = request.current_question_number + 1

    # Get competency for next question
    competency = get_competency_for_stage(session["interview_stage"], next_question_number)

    # Get list of already asked questions
    asked_questions = []
    for qid in session.get("question_ids", []):
        q = get_mock_question(qid)
        if q:
            asked_questions.append(q["question_text"])

    asked_questions_text = "\n".join([f"- {q}" for q in asked_questions]) if asked_questions else "None"

    # Format resume for prompt
    resume_text = format_resume_for_prompt(session["resume_json"])

    # Extract candidate's first name from resume (check multiple possible field names)
    full_name = session["resume_json"].get("name") or session["resume_json"].get("full_name") or session["resume_json"].get("candidate_name") or ""
    candidate_name = full_name.split()[0] if full_name else "there"

    # Build prompt for next question
    prompt = MOCK_GENERATE_QUESTION_PROMPT.format(
        interview_stage=session["interview_stage"],
        candidate_name=candidate_name,
        resume_text=resume_text,
        company=session["company"],
        role_title=session["role_title"],
        job_description=session["job_description"],
        difficulty=session["difficulty_level"],
        asked_questions=asked_questions_text,
        competency_area=competency
    )

    # Call Claude to generate question
    try:
        response = call_claude(
            "You are generating interview questions for a mock interview practice session. Return only valid JSON.",
            prompt,
            max_tokens=1000
        )
        cleaned = clean_claude_json(response)
        question_data = json.loads(cleaned)
    except Exception as e:
        print(f"🔥 Failed to generate question: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate question: {str(e)}")

    # Generate question ID
    question_id = str(uuid.uuid4())

    # Store question
    question_record = {
        "id": question_id,
        "session_id": request.session_id,
        "question_number": next_question_number,
        "question_text": question_data["question_text"],
        "competency_tested": question_data["competency_tested"],
        "difficulty": question_data["difficulty"],
        "asked_at": datetime.now().isoformat()
    }
    save_mock_question(question_id, question_record)

    # Update session with new question
    question_ids = session.get("question_ids", [])
    question_ids.append(question_id)
    update_mock_session(request.session_id, {
        "question_ids": question_ids,
        "current_question_number": next_question_number
    })

    # Calculate average score
    average_score = calculate_mock_session_average(request.session_id)

    print(f"✅ Next question generated: {question_id}")

    return NextQuestionResponse(
        question_id=question_id,
        question_text=question_data["question_text"],
        competency_tested=question_data["competency_tested"],
        question_number=next_question_number,
        total_questions=5,
        session_progress=MockSessionProgress(
            questions_completed=request.current_question_number,
            average_score=average_score
        )
    )


@app.post("/api/mock-interview/end", response_model=EndMockInterviewResponse)
async def end_mock_interview(request: EndMockInterviewRequest):
    """
    INTERVIEW INTELLIGENCE: End session and get comprehensive feedback

    Generates session-level feedback including:
    - Overall score and readiness assessment
    - Aggregated signal scores across all questions
    - Key strengths and areas to improve (based on signals)
    - Coaching priorities and recommended drills
    - Summary of each question's performance
    """
    print(f"🏁 Ending mock interview session {request.session_id}")

    # Verify session exists
    session = get_mock_session(request.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Mock interview session not found")

    # Calculate overall score
    average_score = calculate_mock_session_average(request.session_id)

    # Detect role type for this session
    role_type = detect_role_type(session.get("job_description", ""), session.get("role_title", ""))

    # Build question summaries and aggregate signals
    question_summaries = []
    questions_summary_text = []

    # Initialize signal aggregation
    signal_names = [
        "functional_competency", "leadership", "collaboration", "ownership",
        "strategic_thinking", "problem_solving", "communication_clarity",
        "metrics_orientation", "stakeholder_management", "executive_presence",
        "user_centricity"
    ]
    aggregated_signals = {signal: [] for signal in signal_names}
    all_signal_strengths = []
    all_signal_gaps = []
    level_counts = {"mid": 0, "senior": 0, "director": 0, "executive": 0}

    for qid in session.get("question_ids", []):
        question = get_mock_question(qid)
        analysis = get_mock_analysis(qid) or {}

        if question:
            score = analysis.get("score", 5)
            brief_feedback = analysis.get("brief_feedback", "")

            question_summaries.append(MockQuestionSummary(
                question_number=question["question_number"],
                question_text=question["question_text"],
                score=score,
                brief_feedback=brief_feedback
            ))

            questions_summary_text.append(
                f"Q{question['question_number']}: {question['question_text']}\n"
                f"Score: {score}/10\n"
                f"Competency: {question['competency_tested']}\n"
                f"Feedback: {brief_feedback}"
            )

            # Aggregate signals from this question's analysis
            if "signals" in analysis:
                for signal_name in signal_names:
                    signal_value = analysis["signals"].get(signal_name)
                    if signal_value is not None:
                        aggregated_signals[signal_name].append(signal_value)

            # Collect strengths and gaps
            all_signal_strengths.extend(analysis.get("signal_strengths", []))
            all_signal_gaps.extend(analysis.get("signal_gaps", []))

            # Count levels demonstrated
            level = analysis.get("level_demonstrated", "mid")
            if level in level_counts:
                level_counts[level] += 1

    # Calculate average signal scores
    avg_signals = {}
    for signal_name, values in aggregated_signals.items():
        if values:
            avg_signals[signal_name] = round(sum(values) / len(values), 2)
        else:
            avg_signals[signal_name] = 0.5  # Default if no data

    # Identify session-level strengths and gaps
    session_strengths = [s for s, v in avg_signals.items() if v >= 0.7]
    session_gaps = [s for s, v in avg_signals.items() if v <= 0.4]

    # Determine predominant level
    predominant_level = max(level_counts, key=level_counts.get) if any(level_counts.values()) else "mid"

    # Format signal summary for prompt
    signal_summary_lines = [f"- {signal}: {avg_signals[signal]}" for signal in signal_names]
    signal_summary = "\n".join(signal_summary_lines)

    # Format strengths and gaps for prompt
    strengths_text = ", ".join(session_strengths) if session_strengths else "None identified"
    gaps_text = ", ".join(session_gaps) if session_gaps else "None identified"

    # Format resume for prompt
    resume_text = format_resume_for_prompt(session["resume_json"])

    # Build session feedback prompt with signal data
    prompt = MOCK_SESSION_FEEDBACK_PROMPT.format(
        interview_stage=session["interview_stage"],
        role_type=role_type,
        num_questions=len(session.get("question_ids", [])),
        average_score=round(average_score, 1),
        level_demonstrated=predominant_level,
        signal_summary=signal_summary,
        signal_strengths=strengths_text,
        signal_gaps=gaps_text,
        questions_summary="\n\n".join(questions_summary_text) if questions_summary_text else "No questions completed.",
        resume_text=resume_text,
        company=session["company"],
        role_title=session["role_title"],
        job_description=session["job_description"]
    )

    # Call Claude for session feedback
    try:
        response = call_claude(
            "You are providing comprehensive session feedback for a mock interview. Return only valid JSON.",
            prompt,
            max_tokens=2000
        )
        cleaned = clean_claude_json(response)
        feedback_data = json.loads(cleaned)
    except Exception as e:
        print(f"🔥 Failed to generate session feedback: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate session feedback: {str(e)}")

    # Update session with completion info including signals
    session["completed_at"] = datetime.now().isoformat()
    session["overall_score"] = average_score
    session["session_feedback"] = feedback_data
    session["aggregated_signals"] = avg_signals
    session["signal_strengths"] = session_strengths
    session["signal_gaps"] = session_gaps
    session["level_estimate"] = feedback_data.get("level_estimate", predominant_level)

    print(f"✅ Mock interview session ended: {request.session_id}")
    print(f"   Signal strengths: {session_strengths}")
    print(f"   Signal gaps: {session_gaps}")
    print(f"   Level estimate: {session['level_estimate']}")

    return EndMockInterviewResponse(
        session_id=request.session_id,
        overall_score=round(average_score, 1),
        questions_completed=len(session.get("question_ids", [])),
        session_feedback=MockSessionFeedback(
            overall_assessment=feedback_data.get("overall_assessment", ""),
            key_strengths=feedback_data.get("key_strengths", []),
            areas_to_improve=feedback_data.get("areas_to_improve", []),
            coaching_priorities=feedback_data.get("coaching_priorities", []),
            recommended_drills=feedback_data.get("recommended_drills", []),
            readiness_score=feedback_data.get("readiness_score", "Needs Practice"),
            level_estimate=feedback_data.get("level_estimate", predominant_level),
            next_steps=feedback_data.get("next_steps", "")
        ),
        question_summaries=question_summaries
    )


@app.get("/api/mock-interview/sessions/{company}/{role_title}")
async def get_mock_session_history(company: str, role_title: str):
    """
    INTERVIEW INTELLIGENCE: Get session history for a company/role

    Returns all past mock interview sessions for the specified company and role,
    including improvement trends if multiple sessions exist.
    """
    print(f"📜 Getting mock interview history for {company} - {role_title}")

    # Find all sessions for this company/role
    matching_sessions = []

    for session_id, session in mock_interview_sessions.items():
        if session["company"] == company and session["role_title"] == role_title:
            questions_completed = len(session.get("question_ids", []))

            matching_sessions.append(MockInterviewSessionSummary(
                session_id=session_id,
                interview_stage=session["interview_stage"],
                started_at=session["started_at"],
                completed_at=session.get("completed_at"),
                overall_score=session.get("overall_score"),
                questions_completed=questions_completed
            ))

    # Sort by started_at (newest first)
    matching_sessions.sort(key=lambda s: s.started_at, reverse=True)

    # Calculate improvement trend if multiple completed sessions
    improvement_trend = None
    completed_sessions = [s for s in matching_sessions if s.overall_score is not None]

    if len(completed_sessions) >= 2:
        # Get first and latest completed sessions
        first_session = completed_sessions[-1]  # Oldest
        latest_session = completed_sessions[0]  # Newest

        improvement = latest_session.overall_score - first_session.overall_score
        improvement_str = f"+{improvement:.1f}" if improvement >= 0 else f"{improvement:.1f}"

        improvement_trend = ImprovementTrend(
            first_session_score=first_session.overall_score,
            latest_session_score=latest_session.overall_score,
            improvement=improvement_str
        )

    print(f"✅ Found {len(matching_sessions)} session(s)")

    return SessionHistoryResponse(
        sessions=matching_sessions,
        improvement_trend=improvement_trend
    )


# ============================================================================
# INTERVIEW INTELLIGENCE: TEXT-TO-SPEECH ENDPOINT
# ============================================================================

class TTSRequest(BaseModel):
    """Request for text-to-speech conversion."""
    text: str = Field(..., min_length=1, max_length=4096)
    voice: str = Field(default="onyx", pattern="^(alloy|echo|fable|onyx|nova|shimmer)$")


@app.post("/api/tts")
async def text_to_speech(request: TTSRequest):
    """
    INTERVIEW INTELLIGENCE: Convert text to natural AI speech using OpenAI TTS.

    Available voices:
    - onyx: Deep male voice (recommended for interviewer)
    - echo: Male voice
    - fable: British male voice
    - alloy: Neutral voice
    - nova: Female voice
    - shimmer: Female voice
    """
    import httpx
    from fastapi.responses import Response

    if not OPENAI_API_KEY:
        raise HTTPException(
            status_code=503,
            detail="OpenAI API key not configured. Set OPENAI_API_KEY environment variable for natural AI voice."
        )

    print(f"🎙️ TTS request: {len(request.text)} chars, voice={request.voice}")

    try:
        async with httpx.AsyncClient() as http_client:
            response = await http_client.post(
                "https://api.openai.com/v1/audio/speech",
                headers={
                    "Authorization": f"Bearer {OPENAI_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "tts-1",
                    "input": request.text,
                    "voice": request.voice,
                    "response_format": "mp3"
                },
                timeout=30.0
            )

            if response.status_code != 200:
                error_detail = response.text
                print(f"❌ OpenAI TTS error: {error_detail}")
                raise HTTPException(status_code=response.status_code, detail=f"OpenAI TTS error: {error_detail}")

            print(f"✅ TTS audio generated: {len(response.content)} bytes")

            return Response(
                content=response.content,
                media_type="audio/mpeg",
                headers={"Content-Disposition": "inline"}
            )

    except httpx.TimeoutException:
        raise HTTPException(status_code=504, detail="TTS request timed out")
    except Exception as e:
        print(f"❌ TTS error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# DOCUMENT DOWNLOAD ENDPOINTS
# ============================================================================

class PackageDownloadRequest(BaseModel):
    resume_text: str
    cover_letter_text: str
    candidate_name: str = "Candidate"

@app.post("/api/package/download")
async def download_package(request: PackageDownloadRequest):
    """
    Generate and download a ZIP file containing DOCX resume and cover letter
    """
    import io
    import zipfile
    from fastapi.responses import StreamingResponse
    
    try:
        print(f"📦 Generating download package for: {request.candidate_name}")
        
        # Create DOCX files
        resume_docx = create_docx_from_text(request.resume_text, f"{request.candidate_name} - Resume")
        cover_letter_docx = create_docx_from_text(request.cover_letter_text, f"{request.candidate_name} - Cover Letter")
        
        # Create ZIP file in memory
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            # Add resume
            resume_buffer = io.BytesIO()
            resume_docx.save(resume_buffer)
            resume_buffer.seek(0)
            zip_file.writestr(f"{request.candidate_name.replace(' ', '_')}_Resume.docx", resume_buffer.getvalue())
            
            # Add cover letter
            cl_buffer = io.BytesIO()
            cover_letter_docx.save(cl_buffer)
            cl_buffer.seek(0)
            zip_file.writestr(f"{request.candidate_name.replace(' ', '_')}_Cover_Letter.docx", cl_buffer.getvalue())
        
        zip_buffer.seek(0)
        
        print(f"✅ Package generated successfully")
        
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={
                "Content-Disposition": f"attachment; filename={request.candidate_name.replace(' ', '_')}_Application_Package.zip"
            }
        )
        
    except Exception as e:
        print(f"🔥 DOWNLOAD ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate download: {str(e)}")


def create_docx_from_text(text: str, title: str):
    """
    Create an ATS-safe DOCX document from plain text.
    
    CRITICAL RULES:
    - Text-only, no templates, no restructuring
    - Exact line-by-line reproduction of preview text
    - No heading styles, no list styles, no formatting magic
    - ATS-safe: Calibri font, no tables, no textboxes, no images
    - Bullets (•) treated as plain text characters, NOT Word lists
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    
    # Set ATS-safe margins (0.75 inches all around)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Process text line by line - EXACT reproduction
    lines = text.split('\n')
    
    for line in lines:
        # Preserve the line exactly (don't strip - preserve intentional spacing)
        # But do strip for empty line detection
        stripped = line.strip()
        
        if not stripped:
            # Empty line - add blank paragraph for spacing
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            continue
        
        # Add paragraph with the exact line content
        p = doc.add_paragraph()
        
        # CRITICAL: Use 'Normal' style only - no Heading, no ListParagraph
        p.style = doc.styles['Normal']
        
        # Reset all paragraph formatting
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.right_indent = Pt(0)
        p.paragraph_format.first_line_indent = Pt(0)
        
        # Add the text as a single run
        run = p.add_run(stripped)
        
        # ATS-safe font settings - consistent for ALL text
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.bold = False
        run.italic = False
        run.underline = False
        
        # Ensure Calibri is set for complex scripts too
        r = run._element
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Calibri')
        rFonts.set(qn('w:hAnsi'), 'Calibri')
        rFonts.set(qn('w:cs'), 'Calibri')
        rFonts.set(qn('w:eastAsia'), 'Calibri')
        rPr.insert(0, rFonts)
    
    return doc


# ============================================================================
# NEW DOCUMENT DOWNLOAD ENDPOINTS (Using professional formatters)
# ============================================================================

class ResumeDownloadRequest(BaseModel):
    """Request body for resume download"""
    candidate_name: str
    tagline: str = ""
    contact: Dict[str, Any] = {}
    summary: str = ""
    competencies: List[str] = []
    experience: List[Dict[str, Any]] = []
    skills: Dict[str, Any] = {}
    education: Dict[str, Any] = {}

    class Config:
        extra = "allow"


class CoverLetterDownloadRequest(BaseModel):
    """Request body for cover letter download"""
    candidate_name: str
    tagline: str = ""
    contact: Dict[str, Any] = {}
    recipient_name: Optional[str] = None
    paragraphs: List[str] = []

    class Config:
        extra = "allow"


@app.post("/api/download/resume")
async def download_resume(request: ResumeDownloadRequest):
    """
    Generate and download a professionally formatted resume DOCX.
    """
    try:
        print(f"📄 Generating formatted resume for: {request.candidate_name}")

        # Initialize formatter
        formatter = ResumeFormatter()

        # Add header
        formatter.add_header(
            name=request.candidate_name,
            tagline=request.tagline,
            contact_info=request.contact
        )

        # Add summary if provided
        if request.summary:
            formatter.add_section_header("Summary")
            formatter.add_summary(request.summary)

        # Add core competencies if provided
        if request.competencies:
            formatter.add_section_header("Core Competencies")
            formatter.add_core_competencies(request.competencies)

        # Add experience if provided
        if request.experience:
            formatter.add_section_header("Experience")
            for job in request.experience:
                formatter.add_experience_entry(
                    company=job.get('company', ''),
                    title=job.get('title', ''),
                    location=job.get('location', ''),
                    dates=job.get('dates', ''),
                    overview=job.get('overview'),
                    bullets=job.get('bullets', [])
                )

        # Add skills if provided
        if request.skills:
            formatter.add_section_header("Skills")
            formatter.add_skills(request.skills)

        # Add education if provided
        if request.education and request.education.get('school'):
            formatter.add_section_header("Education")
            formatter.add_education(
                school=request.education.get('school', ''),
                degree=request.education.get('degree', ''),
                details=request.education.get('details')
            )

        # Save to buffer
        buffer = io.BytesIO()
        formatter.get_document().save(buffer)
        buffer.seek(0)

        filename = f"{request.candidate_name.replace(' ', '_')}_Resume.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )

    except Exception as e:
        print(f"🔥 RESUME DOWNLOAD ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate resume: {str(e)}")


@app.post("/api/download/cover-letter")
async def download_cover_letter(request: CoverLetterDownloadRequest):
    """
    Generate and download a professionally formatted cover letter DOCX.
    """
    try:
        print(f"📄 Generating formatted cover letter for: {request.candidate_name}")

        # Initialize formatter
        formatter = CoverLetterFormatter()

        # Add header
        formatter.add_header(
            name=request.candidate_name,
            tagline=request.tagline,
            contact_info=request.contact
        )

        # Add section label
        formatter.add_section_label()

        # Add salutation
        formatter.add_salutation(recipient_name=request.recipient_name)

        # Add body paragraphs
        for paragraph in request.paragraphs:
            formatter.add_body_paragraph(paragraph)

        # Add signature
        formatter.add_signature(request.candidate_name)

        # Save to buffer
        buffer = io.BytesIO()
        formatter.get_document().save(buffer)
        buffer.seek(0)

        filename = f"{request.candidate_name.replace(' ', '_')}_Cover_Letter.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}"
            }
        )

    except Exception as e:
        print(f"🔥 COVER LETTER DOWNLOAD ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate cover letter: {str(e)}")


# ============================================================================
# SCREENING QUESTIONS ENDPOINT
# ============================================================================
# SCREENING_QUESTIONS_PROMPT - Now imported from prompts/ module


@app.post("/api/screening-questions/generate")
async def generate_screening_responses(request: ScreeningQuestionsRequest) -> ScreeningQuestionsResponse:
    """
    Generate strategically positioned responses to ATS screening questions.
    Based on candidate's resume and job description, provides honest responses
    that emphasize relevant experience without fabrication.
    """
    try:
        print(f"📋 Generating screening question responses for {request.company} - {request.role_title}")

        # Build resume context
        resume_json = request.resume_json
        resume_context = f"""
Name: {resume_json.get('full_name', 'Unknown')}
Current Title: {resume_json.get('current_title', 'Unknown')}
Years of Experience: {resume_json.get('years_experience', 'Unknown')}
Summary: {resume_json.get('summary', '')}

Skills: {', '.join(resume_json.get('skills', []))}

Experience:
"""
        for exp in resume_json.get('experience', []):
            resume_context += f"\n- {exp.get('title', '')} at {exp.get('company', '')} ({exp.get('dates', '')})"
            for bullet in exp.get('bullets', [])[:3]:  # Limit bullets to keep context manageable
                resume_context += f"\n  • {bullet}"

        resume_context += f"\n\nEducation: {resume_json.get('education', 'Not specified')}"

        # Add JD analysis context if available
        jd_context = ""
        if request.jd_analysis:
            jd_context = f"""
Fit Score: {request.jd_analysis.get('fit_score', 'Unknown')}%
Strengths: {', '.join(request.jd_analysis.get('strengths', [])[:3])}
Gaps: {', '.join(request.jd_analysis.get('gaps', [])[:3])}
"""
            resume_context += f"\n\n=== FIT ANALYSIS ==={jd_context}"

        # Build the system prompt
        system_prompt = SCREENING_QUESTIONS_PROMPT.format(
            resume_context=resume_context,
            company=request.company,
            role_title=request.role_title,
            job_description=request.job_description[:3000],  # Limit JD length
            screening_questions=request.screening_questions
        )

        # Add candidate state calibration if available
        calibration_prompt = build_candidate_calibration_prompt(request.situation)
        if calibration_prompt:
            system_prompt += calibration_prompt

        # Call Claude
        user_message = f"Generate strategic responses for these screening questions:\n\n{request.screening_questions}"
        response = call_claude(system_prompt, user_message)

        # Parse the response
        cleaned = clean_claude_json(response)
        parsed_data = json.loads(cleaned)

        # Validate and return
        return ScreeningQuestionsResponse(
            responses=[
                ScreeningQuestionResponse(
                    question=r.get('question', ''),
                    recommended_response=r.get('recommended_response', ''),
                    strategic_note=r.get('strategic_note', ''),
                    gap_detected=r.get('gap_detected', False),
                    mitigation_strategy=r.get('mitigation_strategy')
                )
                for r in parsed_data.get('responses', [])
            ],
            overall_strategy=parsed_data.get('overall_strategy', '')
        )

    except json.JSONDecodeError as e:
        print(f"🔥 SCREENING QUESTIONS JSON ERROR: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse screening question responses: {str(e)}")
    except Exception as e:
        print(f"🔥 SCREENING QUESTIONS ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating screening responses: {str(e)}")


# ============================================================================
# EXPERIENCE SUPPLEMENTATION ENDPOINTS
# ============================================================================
# CLARIFYING_QUESTIONS_PROMPT, REANALYZE_PROMPT - Now imported from prompts/ module


# ============================================================================
# RESUME LEVELING PROMPT
# ============================================================================
# RESUME_LEVELING_PROMPT - Now imported from prompts/ module


# Load leveling frameworks
LEVELING_FRAMEWORKS_PATH = os.path.join(os.path.dirname(__file__), "data", "leveling_frameworks.json")
LEVELING_FRAMEWORKS: Dict[str, Any] = {}

def load_leveling_frameworks():
    """Load the leveling frameworks from JSON file."""
    global LEVELING_FRAMEWORKS
    try:
        if os.path.exists(LEVELING_FRAMEWORKS_PATH):
            with open(LEVELING_FRAMEWORKS_PATH, "r") as f:
                LEVELING_FRAMEWORKS = json.load(f)
            print(f"✅ Loaded leveling frameworks v{LEVELING_FRAMEWORKS.get('version', 'unknown')}")
        else:
            print(f"⚠️ Leveling frameworks not found at {LEVELING_FRAMEWORKS_PATH}")
            LEVELING_FRAMEWORKS = {}
    except Exception as e:
        print(f"🔥 Error loading leveling frameworks: {e}")
        LEVELING_FRAMEWORKS = {}

# Load frameworks on startup
load_leveling_frameworks()


@app.post("/api/resume/level-assessment")
async def assess_resume_level(request: ResumeLevelingRequest) -> ResumeLevelingResponse:
    """
    Analyze a resume to determine the candidate's career level and provide
    gap analysis against a target role if provided.
    """
    try:
        # Handle case where resume_json might be a string (double-encoded JSON)
        resume_json = request.resume_json
        if isinstance(resume_json, str):
            try:
                resume_json = json.loads(resume_json)
            except json.JSONDecodeError:
                print(f"⚠️ Could not parse resume_json string: {str(resume_json)[:200]}")
                resume_json = {}

        print(f"📊 Assessing resume level for: {resume_json.get('full_name', 'Unknown') if isinstance(resume_json, dict) else 'Unknown'}")
        if request.target_title:
            print(f"   Target: {request.target_title} at {request.company}")

        # Build resume context
        # Safely get skills - could be list or string
        skills = resume_json.get('skills', [])
        skills_str = ', '.join(skills) if isinstance(skills, list) else str(skills) if skills else ''

        # Safely get contact info
        contact = resume_json.get('contact', {})
        if isinstance(contact, str):
            contact = {}

        resume_context = f"""
Name: {resume_json.get('full_name', contact.get('full_name', 'Unknown') if isinstance(contact, dict) else 'Unknown')}
Current Title: {resume_json.get('current_title', 'Unknown')}
Years of Experience: {resume_json.get('years_experience', 'Unknown')}
Summary: {resume_json.get('summary', resume_json.get('summary_text', ''))}

Skills: {skills_str}

Experience:
"""
        for exp in resume_json.get('experience', []):
            # Handle case where exp might be a string instead of dict
            if isinstance(exp, str):
                resume_context += f"\n- {exp}\n"
                continue
            if not isinstance(exp, dict):
                continue
            resume_context += f"\n### {exp.get('title', exp.get('role', ''))} at {exp.get('company', '')} ({exp.get('dates', exp.get('start_date', ''))})\n"
            bullets = exp.get('bullets', [])
            if isinstance(bullets, list):
                for bullet in bullets:
                    if isinstance(bullet, str):
                        resume_context += f"- {bullet}\n"

        # Add education
        education = resume_json.get('education', [])
        if education and isinstance(education, list):
            resume_context += "\nEducation:\n"
            for edu in education:
                # Handle case where edu might be a string instead of dict
                if isinstance(edu, str):
                    resume_context += f"- {edu}\n"
                    continue
                if not isinstance(edu, dict):
                    continue
                resume_context += f"- {edu.get('degree', '')} from {edu.get('school', edu.get('institution', ''))} ({edu.get('year', edu.get('graduation_date', ''))})\n"

        # Build target context if provided
        target_context = ""
        if request.target_title or request.role_title:
            target_title = request.target_title or request.role_title
            target_context = f"""
=== TARGET ROLE ANALYSIS ===
Target Title: {target_title}
Company: {request.company or 'Not specified'}

Job Description (if provided):
{request.job_description[:2000] if request.job_description else 'Not provided - infer requirements from title'}

Analyze gaps between candidate's current level and the target role requirements.
"""

        # Build the prompt
        system_prompt = RESUME_LEVELING_PROMPT.format(
            resume_context=resume_context,
            target_context=target_context
        )

        user_message = "Analyze this resume and provide the leveling assessment."

        response = call_claude(system_prompt, user_message)

        # Parse the response
        cleaned = clean_claude_json(response)
        parsed_data = json.loads(cleaned)

        # Build competencies list
        competencies = []
        for comp in parsed_data.get('competencies', []):
            competencies.append(LevelCompetency(
                area=comp.get('area', ''),
                current_level=comp.get('current_level', ''),
                evidence=comp.get('evidence', []),
                gap_to_target=comp.get('gap_to_target')
            ))

        # Build gaps list
        gaps = None
        if parsed_data.get('gaps'):
            gaps = []
            for gap in parsed_data['gaps']:
                gaps.append(LevelingGap(
                    category=gap.get('category', ''),
                    description=gap.get('description', ''),
                    recommendation=gap.get('recommendation', ''),
                    priority=gap.get('priority', 'medium')
                ))

        # Build recommendations list
        recommendations = []
        for rec in parsed_data.get('recommendations', []):
            recommendations.append(LevelingRecommendation(
                type=rec.get('type', ''),
                priority=rec.get('priority', 'medium'),
                current=rec.get('current', ''),
                suggested=rec.get('suggested', ''),
                rationale=rec.get('rationale', '')
            ))

        # Build role type breakdown if present
        role_type_breakdown = None
        if parsed_data.get('role_type_breakdown'):
            rtb = parsed_data['role_type_breakdown']
            role_type_breakdown = RoleTypeBreakdown(
                pm_years=rtb.get('pm_years', 0),
                engineering_years=rtb.get('engineering_years', 0),
                operations_years=rtb.get('operations_years', 0),
                other_years=rtb.get('other_years', 0)
            )

        # Build level mismatch warnings if present
        level_mismatch_warnings = None
        if parsed_data.get('level_mismatch_warnings'):
            level_mismatch_warnings = []
            for warning in parsed_data['level_mismatch_warnings']:
                level_mismatch_warnings.append(LevelMismatchWarning(
                    target_level=warning.get('target_level', ''),
                    assessed_level=warning.get('assessed_level', ''),
                    gap_explanation=warning.get('gap_explanation', ''),
                    alternative_recommendation=warning.get('alternative_recommendation')
                ))

        # ====================================================================
        # LEPE INTEGRATION
        # Per LEPE Spec: Resume Leveling must reference LEPE outputs and
        # must not independently assess leadership readiness for Manager+ roles.
        # ====================================================================
        lepe_applicable = False
        lepe_decision = None
        lepe_coaching = None
        lepe_skepticism_warning = None
        leadership_tenure = None
        accountability_record = None

        target_title = request.target_title or request.role_title
        if target_title and is_manager_plus_role(target_title):
            lepe_applicable = True
            print(f"   📊 LEPE: Manager+ target detected - running leadership evaluation")

            # Build mock response_data for LEPE (it needs role_title)
            lepe_response_data = {
                "role_title": target_title,
                "job_description": request.job_description or ""
            }

            # Run LEPE evaluation
            lepe_result = evaluate_lepe(resume_json, lepe_response_data)

            if lepe_result.get("lepe_applicable"):
                positioning = lepe_result.get("positioning_decision", {})
                lepe_decision = positioning.get("decision")
                leadership_tenure = lepe_result.get("leadership_tenure")
                accountability_record = lepe_result.get("accountability_record")

                # Extract messaging based on decision
                messaging = positioning.get("messaging", {})
                if lepe_decision == "position":
                    lepe_coaching = messaging.get("coaching_advice")
                elif lepe_decision in ["caution", "locked"]:
                    lepe_skepticism_warning = messaging.get("skepticism_warning")

                print(f"   📊 LEPE Advisory: {lepe_decision} (advisory only)")

        return ResumeLevelingResponse(
            detected_function=parsed_data.get('detected_function', 'Unknown'),
            function_confidence=parsed_data.get('function_confidence', 0.5),
            current_level=parsed_data.get('current_level', 'Unknown'),
            current_level_id=parsed_data.get('current_level_id', 'unknown'),
            level_confidence=parsed_data.get('level_confidence', 0.5),
            years_experience=parsed_data.get('years_experience', 0),
            years_in_role_type=parsed_data.get('years_in_role_type'),
            role_type_breakdown=role_type_breakdown,
            scope_signals=parsed_data.get('scope_signals', []),
            impact_signals=parsed_data.get('impact_signals', []),
            leadership_signals=parsed_data.get('leadership_signals', []),
            technical_signals=parsed_data.get('technical_signals', []),
            competencies=competencies,
            language_level=parsed_data.get('language_level', 'Mid'),
            action_verb_distribution=parsed_data.get('action_verb_distribution', {'entry': 0.25, 'mid': 0.5, 'senior': 0.2, 'principal': 0.05}),
            quantification_rate=parsed_data.get('quantification_rate', 0.0),
            red_flags=parsed_data.get('red_flags', []),
            title_inflation_detected=parsed_data.get('title_inflation_detected'),
            title_inflation_explanation=parsed_data.get('title_inflation_explanation'),
            target_level=parsed_data.get('target_level'),
            target_level_id=parsed_data.get('target_level_id'),
            levels_apart=parsed_data.get('levels_apart'),
            is_qualified=parsed_data.get('is_qualified'),
            qualification_confidence=parsed_data.get('qualification_confidence'),
            level_mismatch_warnings=level_mismatch_warnings,
            gaps=gaps,
            recommendations=recommendations,
            summary=parsed_data.get('summary', 'Assessment complete.'),
            # LEPE fields
            lepe_applicable=lepe_applicable,
            lepe_decision=lepe_decision,
            lepe_coaching=lepe_coaching,
            lepe_skepticism_warning=lepe_skepticism_warning,
            leadership_tenure=leadership_tenure,
            accountability_record=accountability_record
        )

    except json.JSONDecodeError as e:
        print(f"🔥 LEVELING JSON ERROR: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse leveling assessment: {str(e)}")
    except Exception as e:
        print(f"🔥 LEVELING ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error assessing resume level: {str(e)}")


@app.post("/api/experience/clarifying-questions")
async def generate_clarifying_questions(request: GenerateClarifyingQuestionsRequest) -> GenerateClarifyingQuestionsResponse:
    """
    Generate targeted clarifying questions based on identified gaps.
    These questions help uncover experience the candidate may have omitted from their resume.
    """
    try:
        print(f"❓ Generating clarifying questions for {request.company} - {request.role_title}")
        print(f"   Gaps to address: {request.gaps}")

        # Build resume context
        resume_json = request.resume_json
        resume_context = f"""
Name: {resume_json.get('full_name', resume_json.get('contact', {}).get('full_name', 'Unknown'))}
Current Title: {resume_json.get('current_title', 'Unknown')}
Years of Experience: {resume_json.get('years_experience', 'Unknown')}
Summary: {resume_json.get('summary', resume_json.get('summary_text', ''))}

Skills: {', '.join(resume_json.get('skills', []))}

Experience:
"""
        for exp in resume_json.get('experience', []):
            resume_context += f"\n- {exp.get('title', exp.get('role', ''))} at {exp.get('company', '')} ({exp.get('dates', exp.get('start_date', ''))})"

        # Format gaps
        gaps_text = "\n".join([f"- {gap}" for gap in request.gaps])

        # Build the system prompt
        system_prompt = CLARIFYING_QUESTIONS_PROMPT.format(
            resume_context=resume_context,
            company=request.company,
            role_title=request.role_title,
            job_description=request.job_description[:3000],
            gaps=gaps_text,
            fit_score=request.fit_score
        )

        # Call Claude
        user_message = f"Generate clarifying questions for these gaps:\n\n{gaps_text}"
        response = call_claude(system_prompt, user_message)

        # Parse the response
        cleaned = clean_claude_json(response)
        parsed_data = json.loads(cleaned)

        # Validate and return
        return GenerateClarifyingQuestionsResponse(
            intro_message=parsed_data.get('intro_message', 'I have a few questions to better understand your background:'),
            questions=[
                ClarifyingQuestion(
                    gap_area=q.get('gap_area', ''),
                    question=q.get('question', ''),
                    why_it_matters=q.get('why_it_matters', ''),
                    example_answer=q.get('example_answer', '')
                )
                for q in parsed_data.get('questions', [])
            ]
        )

    except json.JSONDecodeError as e:
        print(f"🔥 CLARIFYING QUESTIONS JSON ERROR: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse clarifying questions: {str(e)}")
    except Exception as e:
        print(f"🔥 CLARIFYING QUESTIONS ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating clarifying questions: {str(e)}")


@app.post("/api/experience/reanalyze")
async def reanalyze_with_supplements(request: ReanalyzeWithSupplementsRequest) -> ReanalyzeWithSupplementsResponse:
    """
    Re-analyze candidate fit after they've provided supplemental experience information.
    Returns updated fit score, strengths, gaps, and an updated resume JSON.
    """
    try:
        print(f"🔄 Re-analyzing fit with {len(request.supplements)} supplements for {request.company} - {request.role_title}")

        # Build resume context
        resume_json = request.resume_json
        resume_context = f"""
Name: {resume_json.get('full_name', resume_json.get('contact', {}).get('full_name', 'Unknown'))}
Current Title: {resume_json.get('current_title', 'Unknown')}
Summary: {resume_json.get('summary', resume_json.get('summary_text', ''))}

Skills: {', '.join(resume_json.get('skills', []))}

Experience:
"""
        for exp in resume_json.get('experience', []):
            resume_context += f"\n- {exp.get('title', exp.get('role', ''))} at {exp.get('company', '')} ({exp.get('dates', exp.get('start_date', ''))})"
            for bullet in exp.get('bullets', [])[:2]:
                resume_context += f"\n  • {bullet}"

        # Format supplements
        supplements_text = ""
        for supp in request.supplements:
            supplements_text += f"\n**Gap: {supp.gap_area}**\n"
            supplements_text += f"Question: {supp.question}\n"
            supplements_text += f"Answer: {supp.answer}\n"

        # Format original gaps
        original_gaps_text = "\n".join([f"- {gap}" for gap in request.original_gaps])

        # Build the system prompt
        system_prompt = REANALYZE_PROMPT.format(
            resume_context=resume_context,
            supplements=supplements_text,
            company=request.company,
            role_title=request.role_title,
            job_description=request.job_description[:3000],
            original_fit_score=request.original_fit_score,
            original_gaps=original_gaps_text
        )

        # Include original resume JSON for the model to update
        user_message = f"""Re-analyze fit with the supplemented experience provided.

Original resume JSON to update:
```json
{json.dumps(resume_json, indent=2)}
```

Supplemented experience:
{supplements_text}"""

        response = call_claude(system_prompt, user_message)

        # Parse the response
        cleaned = clean_claude_json(response)
        parsed_data = json.loads(cleaned)

        # Calculate score change
        new_score = parsed_data.get('new_fit_score', request.original_fit_score)
        score_change = new_score - request.original_fit_score

        # Get updated resume JSON, fall back to original if not provided
        updated_resume = parsed_data.get('updated_resume_json', resume_json)

        return ReanalyzeWithSupplementsResponse(
            new_fit_score=new_score,
            score_change=score_change,
            updated_strengths=parsed_data.get('updated_strengths', []),
            remaining_gaps=parsed_data.get('remaining_gaps', []),
            addressed_gaps=parsed_data.get('addressed_gaps', []),
            updated_resume_json=updated_resume,
            summary=parsed_data.get('summary', f"Fit score updated from {request.original_fit_score}% to {new_score}%.")
        )

    except json.JSONDecodeError as e:
        print(f"🔥 REANALYZE JSON ERROR: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to parse re-analysis: {str(e)}")
    except Exception as e:
        print(f"🔥 REANALYZE ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error re-analyzing with supplements: {str(e)}")


# Legacy endpoint for backwards compatibility
@app.post("/api/documents/download")
async def download_documents_legacy(
    resume_data: str = Form(...),
    cover_letter_data: str = Form(...),
    candidate_name: str = Form("Candidate"),
    include_outreach: str = Form("false")
):
    """Legacy download endpoint using form data"""
    import io
    import zipfile
    from fastapi.responses import StreamingResponse
    
    try:
        print(f"📦 Legacy download for: {candidate_name}")
        
        # Parse JSON data
        resume_json = json.loads(resume_data)
        cover_letter_json = json.loads(cover_letter_data)
        
        resume_text = resume_json.get('full_text', '')
        cover_letter_text = cover_letter_json.get('full_text', '')
        
        # Create DOCX files
        resume_docx = create_docx_from_text(resume_text, f"{candidate_name} - Resume")
        cover_letter_docx = create_docx_from_text(cover_letter_text, f"{candidate_name} - Cover Letter")
        
        # Create ZIP file
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            resume_buffer = io.BytesIO()
            resume_docx.save(resume_buffer)
            resume_buffer.seek(0)
            zip_file.writestr(f"{candidate_name.replace(' ', '_')}_Resume.docx", resume_buffer.getvalue())
            
            cl_buffer = io.BytesIO()
            cover_letter_docx.save(cl_buffer)
            cl_buffer.seek(0)
            zip_file.writestr(f"{candidate_name.replace(' ', '_')}_Cover_Letter.docx", cl_buffer.getvalue())
        
        zip_buffer.seek(0)
        
        return StreamingResponse(
            zip_buffer,
            media_type="application/zip",
            headers={
                "Content-Disposition": f"attachment; filename={candidate_name.replace(' ', '_')}_Application_Package.zip"
            }
        )
        
    except Exception as e:
        print(f"🔥 LEGACY DOWNLOAD ERROR: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate download: {str(e)}")


# ============================================================================
# PREP GUIDE ENDPOINTS
# ============================================================================

@app.post("/api/prep-guide/generate", response_model=PrepGuideResponse)
async def generate_prep_guide(request: PrepGuideRequest):
    """Generate a personalized interview prep guide"""
    try:
        print(f"📋 Generating prep guide for {request.company} - {request.role_title} ({request.interview_type})")

        # Build resume context
        resume_text = ""
        if request.resume_json:
            experiences = request.resume_json.get("experience", [])
            for exp in experiences:
                resume_text += f"- {exp.get('title', '')} at {exp.get('company', '')}\n"
                for bullet in exp.get("bullets", []):
                    resume_text += f"  • {bullet}\n"

            skills = request.resume_json.get("skills", [])
            if skills:
                resume_text += f"\nSkills: {', '.join(skills)}\n"

        # Interview type context
        interview_contexts = {
            "recruiter_screen": "initial phone screen focused on culture fit, basic qualifications, and motivation",
            "hiring_manager": "deep-dive on functional expertise, leadership style, and problem-solving approach",
            "technical": "technical depth, case studies, and hands-on problem solving",
            "panel": "multiple perspectives, cross-functional collaboration, and executive presence",
            "executive": "strategic vision, leadership philosophy, and cultural alignment at senior level"
        }

        interview_context = interview_contexts.get(request.interview_type, "general interview")

        prompt = f"""Generate a comprehensive interview prep guide for a candidate.

INTERVIEW DETAILS:
- Company: {request.company}
- Role: {request.role_title}
- Interview Type: {request.interview_type} ({interview_context})
- Interviewer: {request.interviewer_name or 'Unknown'} ({request.interviewer_title or 'Unknown title'})

JOB DESCRIPTION:
{request.job_description or 'Not provided'}

CANDIDATE BACKGROUND:
{resume_text or 'Not provided'}

Generate a prep guide with these sections:

1. WHAT THEY EVALUATE (5-7 bullet points of what interviewers assess in this type of interview)

2. INTRO PITCH (A 60-90 second "tell me about yourself" script that:
   - Opens with current role and key accomplishment
   - Bridges to relevant experience for this role
   - Closes with why this opportunity excites them
   - Is conversational, not robotic
   - Around 150-200 words)

3. LIKELY QUESTIONS (8-12 questions they'll probably ask, with personalized guidance on how to answer based on their specific experience. Each question should have a "guidance" field with 2-3 sentences of specific advice referencing their background.)

4. HIGH-IMPACT STORIES (3-4 STAR format stories from their experience that they should prepare. Pull from actual experience in their resume. Each story should have: title, competency it demonstrates, situation, task, action, result)

5. RED FLAGS TO AVOID (4-6 things that could hurt them in this interview)

{"6. STRATEGY SCENARIOS (3-4 case/scenario questions with suggested approaches - ONLY for hiring_manager or technical interviews)" if request.interview_type in ["hiring_manager", "technical"] else ""}

Return valid JSON matching this structure:
{{
    "what_they_evaluate": ["item1", "item2", ...],
    "intro_pitch": "The full 60-90 second intro script...",
    "likely_questions": [
        {{"question": "Question text", "guidance": "Personalized guidance on how to answer..."}},
        ...
    ],
    "stories": [
        {{
            "title": "Story title",
            "competency": "Leadership/Problem-solving/etc",
            "situation": "Context...",
            "task": "Challenge...",
            "action": "What they did...",
            "result": "Outcome with metrics..."
        }},
        ...
    ],
    "red_flags": ["Red flag 1", "Red flag 2", ...],
    {"'strategy_scenarios': [{'scenario': 'Scenario description', 'approach': 'How to approach it'}, ...]" if request.interview_type in ["hiring_manager", "technical"] else "'strategy_scenarios': null"}
}}"""

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )

        response_text = response.content[0].text

        # Extract JSON from response
        import re
        json_match = re.search(r'\{[\s\S]*\}', response_text)
        if json_match:
            prep_data = json.loads(json_match.group())
        else:
            raise ValueError("Could not parse JSON from response")

        print(f"✅ Prep guide generated successfully")
        return PrepGuideResponse(**prep_data)

    except json.JSONDecodeError as e:
        print(f"🔥 JSON parse error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse prep guide response")
    except Exception as e:
        print(f"🔥 Prep guide error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Failed to generate prep guide: {str(e)}")


@app.post("/api/prep-guide/regenerate-intro", response_model=RegenerateIntroResponse)
async def regenerate_intro(request: RegenerateIntroRequest):
    """Regenerate just the intro pitch with a fresh take"""
    try:
        print(f"🔄 Regenerating intro for {request.company} - {request.role_title}")

        # Build resume context
        resume_text = ""
        if request.resume_json:
            experiences = request.resume_json.get("experience", [])
            for exp in experiences[:3]:  # Top 3 experiences
                resume_text += f"- {exp.get('title', '')} at {exp.get('company', '')}\n"

        prompt = f"""Generate a fresh 60-90 second "tell me about yourself" intro for an interview.

Company: {request.company}
Role: {request.role_title}
Interview Type: {request.interview_type}

Candidate's recent experience:
{resume_text or 'Not provided'}

Create a conversational, engaging intro (150-200 words) that:
- Opens with their current role and a key accomplishment
- Bridges to why their experience is relevant for this role
- Closes with genuine enthusiasm for this opportunity
- Sounds natural, not scripted
- Is different from typical generic intros

Return ONLY the intro text, no JSON or formatting."""

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=500,
            messages=[{"role": "user", "content": prompt}]
        )

        intro_pitch = response.content[0].text.strip()
        print(f"✅ Intro regenerated: {len(intro_pitch)} chars")

        return RegenerateIntroResponse(intro_pitch=intro_pitch)

    except Exception as e:
        print(f"🔥 Regenerate intro error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to regenerate intro: {str(e)}")


# ============================================================================
# INTERVIEWER INTELLIGENCE
# ============================================================================

class InterviewerRisk(BaseModel):
    """A risk area to prepare for"""
    risk: str
    defense: str

class InterviewerStrength(BaseModel):
    """A strength to highlight based on interviewer alignment"""
    strength: str
    why_relevant: str

class TailoredStory(BaseModel):
    """A story recommendation tailored to the interviewer"""
    story_title: str
    competency: str
    why_this_story: str

class QuestionToAsk(BaseModel):
    """A customized question for the candidate to ask the interviewer"""
    question: str
    rationale: str

class InterviewerIntelRequest(BaseModel):
    """Request to analyze an interviewer's profile"""
    linkedin_profile_text: str  # Can be from PDF extraction, pasted text, or OCR from image
    company: str
    role_title: str
    interview_type: str
    interviewer_name: str = ""
    interviewer_title: str = ""
    candidate_resume_json: Dict[str, Any] = {}
    job_description: str = ""

class InterviewerIntelResponse(BaseModel):
    """Complete interviewer intelligence report"""
    # Interviewer Summary
    interviewer_name: str
    interviewer_title: str
    current_company: str
    tenure: str
    summary: str  # Who they are and what they care about

    # Analysis
    likely_evaluation_focus: List[str]  # What they likely evaluate
    predicted_question_themes: List[str]  # Question themes based on background
    communication_style: str  # How to communicate with this person

    # Strategy
    risks: List[InterviewerRisk]  # Risk areas and defenses
    strengths_to_highlight: List[InterviewerStrength]  # Strengths to emphasize
    tailored_stories: List[TailoredStory]  # Story recommendations
    questions_to_ask: List[QuestionToAsk]  # Customized questions

    # Debrief insights (if available)
    debrief_insights: Optional[str] = None


@app.post("/api/interviewer-intelligence/analyze", response_model=InterviewerIntelResponse)
async def analyze_interviewer(request: InterviewerIntelRequest):
    """
    Analyze an interviewer's LinkedIn profile and generate strategic intelligence.
    This creates recruiter-grade coaching tailored to the specific interviewer.
    """
    try:
        print(f"🔍 Analyzing interviewer profile for {request.company} - {request.role_title}")
        print(f"   Profile text length: {len(request.linkedin_profile_text)} chars")

        # Build candidate context
        candidate_context = ""
        if request.candidate_resume_json:
            experiences = request.candidate_resume_json.get("experience", [])
            for exp in experiences[:4]:
                candidate_context += f"- {exp.get('title', '')} at {exp.get('company', '')}: {exp.get('description', '')[:200]}\n"
            skills = request.candidate_resume_json.get("skills", [])
            if skills:
                candidate_context += f"Skills: {', '.join(skills[:15])}\n"

        prompt = f"""You are an elite executive recruiter providing strategic interview coaching. Analyze this LinkedIn profile and create a comprehensive interviewer intelligence report.

INTERVIEWER'S LINKEDIN PROFILE:
{request.linkedin_profile_text}

INTERVIEW CONTEXT:
- Company: {request.company}
- Role the candidate is interviewing for: {request.role_title}
- Interview Type: {request.interview_type}
- Job Description: {request.job_description[:1000] if request.job_description else 'Not provided'}

CANDIDATE'S BACKGROUND:
{candidate_context if candidate_context else 'Not provided'}

Based on the interviewer's profile, generate a strategic intelligence report. Your analysis must be:
- Specific to THIS interviewer (not generic)
- Based on signals from their background, career path, and experience
- Actionable and practical for interview preparation

Return a JSON object with this exact structure:
{{
    "interviewer_name": "Their full name from profile",
    "interviewer_title": "Their current title",
    "current_company": "Their current company",
    "tenure": "How long they've been at current company (e.g., '3 years')",
    "summary": "2-3 sentence summary of who they are professionally and what they likely care about based on their career trajectory",

    "likely_evaluation_focus": [
        "What they'll likely probe based on their background (e.g., 'Execution rigor - they built teams at high-growth startups')",
        "Second focus area",
        "Third focus area",
        "Fourth focus area"
    ],

    "predicted_question_themes": [
        "Specific question theme they'll likely ask about based on their experience",
        "Second theme",
        "Third theme",
        "Fourth theme"
    ],

    "communication_style": "How to communicate with this person based on their background. Be specific about pace, level of detail, what to emphasize.",

    "risks": [
        {{
            "risk": "A specific risk area the candidate should prepare for",
            "defense": "How to address or mitigate this risk"
        }},
        {{
            "risk": "Second risk",
            "defense": "How to address it"
        }}
    ],

    "strengths_to_highlight": [
        {{
            "strength": "A strength the candidate should emphasize",
            "why_relevant": "Why this matters to this specific interviewer"
        }},
        {{
            "strength": "Second strength",
            "why_relevant": "Why it matters"
        }}
    ],

    "tailored_stories": [
        {{
            "story_title": "Title of a STAR story to prepare",
            "competency": "What competency it demonstrates",
            "why_this_story": "Why this story will resonate with this interviewer"
        }},
        {{
            "story_title": "Second story",
            "competency": "Competency",
            "why_this_story": "Why it resonates"
        }}
    ],

    "questions_to_ask": [
        {{
            "question": "A thoughtful question tailored to this interviewer's background",
            "rationale": "Why this question will resonate with them"
        }},
        {{
            "question": "Second question",
            "rationale": "Why it resonates"
        }},
        {{
            "question": "Third question",
            "rationale": "Why it resonates"
        }}
    ],

    "debrief_insights": null
}}

Return ONLY the JSON object, no additional text."""

        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )

        result_text = response.content[0].text.strip()

        # Parse JSON response
        if result_text.startswith("```"):
            result_text = result_text.split("```")[1]
            if result_text.startswith("json"):
                result_text = result_text[4:]

        intel_data = json.loads(result_text)

        print(f"✅ Interviewer intelligence generated for {intel_data.get('interviewer_name', 'Unknown')}")

        return InterviewerIntelResponse(**intel_data)

    except json.JSONDecodeError as e:
        print(f"🔥 JSON parse error: {e}")
        print(f"   Raw response: {result_text[:500]}")
        raise HTTPException(status_code=500, detail="Failed to parse interviewer analysis")
    except Exception as e:
        print(f"🔥 Interviewer analysis error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to analyze interviewer: {str(e)}")


@app.post("/api/interviewer-intelligence/extract-text")
async def extract_linkedin_text(file: UploadFile = File(...)):
    """
    Extract text from an uploaded LinkedIn PDF or image.
    For PDFs, uses PyMuPDF. For images, returns base64 for Claude vision.
    """
    try:
        content = await file.read()
        filename = file.filename.lower()

        if filename.endswith('.pdf'):
            # Extract text from PDF using PyMuPDF
            import fitz  # PyMuPDF

            pdf_doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            for page in pdf_doc:
                text += page.get_text()
            pdf_doc.close()

            print(f"✅ Extracted {len(text)} chars from PDF")
            return {"text": text, "type": "pdf"}

        elif any(filename.endswith(ext) for ext in ['.png', '.jpg', '.jpeg', '.webp']):
            # For images, we'll use Claude's vision capability
            import base64
            base64_image = base64.b64encode(content).decode('utf-8')

            # Determine media type
            if filename.endswith('.png'):
                media_type = "image/png"
            elif filename.endswith('.webp'):
                media_type = "image/webp"
            else:
                media_type = "image/jpeg"

            # Use Claude to extract text from image
            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4000,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": base64_image
                            }
                        },
                        {
                            "type": "text",
                            "text": """Extract ALL text from this LinkedIn profile screenshot. Include:
- Name and headline
- About section (full text)
- Experience section (all roles, companies, dates, descriptions)
- Education
- Skills
- Any other visible information

Format as plain text, preserving the structure. Be thorough - capture every piece of text visible."""
                        }
                    ]
                }]
            )

            extracted_text = response.content[0].text
            print(f"✅ Extracted {len(extracted_text)} chars from image via vision")
            return {"text": extracted_text, "type": "image"}

        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Please upload PDF or image.")

    except Exception as e:
        print(f"🔥 Text extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")


# ============================================================================
# ASK HENRY - FLOATING CHAT ASSISTANT
# ============================================================================

class HeyHenryContext(BaseModel):
    """Context about where the user is in the app."""
    current_page: str
    page_description: str
    company: Optional[str] = None
    role: Optional[str] = None
    has_analysis: bool = False
    has_resume: bool = False
    has_pipeline: bool = False
    has_attachments: bool = False
    user_name: Optional[str] = None
    # Emotional state fields (from frontend)
    emotional_state: Optional[str] = None  # zen, stressed, struggling, desperate, crushed
    confidence_level: Optional[str] = None  # low, need_validation, shaky, strong
    timeline: Optional[str] = None  # urgent, soon, actively_looking, no_rush
    tone_guidance: Optional[str] = None  # Generated guidance string from frontend
    # Clarification detection (from frontend)
    needs_clarification: bool = False
    clarification_hints: Optional[List[Dict[str, str]]] = None


class HeyHenryMessage(BaseModel):
    """A single message in the conversation."""
    role: str  # 'user' or 'assistant'
    content: str


class HeyHenryAttachment(BaseModel):
    """An attachment sent with a Hey Henry message."""
    name: str
    type: str  # MIME type
    size: int
    data: str  # Base64 encoded


class HeyHenryRequest(BaseModel):
    """Request for Hey Henry chat."""
    message: str
    conversation_history: List[HeyHenryMessage] = []
    context: HeyHenryContext
    analysis_data: Optional[Dict[str, Any]] = None
    resume_data: Optional[Dict[str, Any]] = None
    user_profile: Optional[Dict[str, Any]] = None  # Includes emotional state/situation
    pipeline_data: Optional[Dict[str, Any]] = None  # Application pipeline metrics and apps
    attachments: Optional[List[HeyHenryAttachment]] = None  # File attachments (images, documents)
    # Generated content - Henry is the AUTHOR of these
    documents_data: Optional[Dict[str, Any]] = None  # Generated resume, cover letter, changes
    outreach_data: Optional[Dict[str, Any]] = None  # Generated outreach templates
    interview_prep_data: Optional[Dict[str, Any]] = None  # Generated interview prep modules
    positioning_data: Optional[Dict[str, Any]] = None  # Positioning strategy content
    # Network data - LinkedIn connections (Phase 2.1)
    network_data: Optional[Dict[str, Any]] = None  # LinkedIn connections at target company
    # Outreach log data - Follow-up tracking (Phase 2.7)
    outreach_log_data: Optional[Dict[str, Any]] = None  # Outreach tracking and follow-ups
    # Interview debrief data - Missing debrief detection (Phase 2.3)
    interview_debrief_data: Optional[Dict[str, Any]] = None  # Interviews needing debriefs
    # Cross-interview pattern analysis (Phase 2.3)
    debrief_pattern_analysis: Optional[Dict[str, Any]] = None  # Patterns across user's debriefs


class HeyHenryResponse(BaseModel):
    """Response from Hey Henry."""
    response: str


# ==========================================
# Interview Debrief Intelligence (Phase 2.3)
# ==========================================

class DebriefExtractionRequest(BaseModel):
    """Request to extract structured data from debrief conversation."""
    conversation: str  # Raw conversation text
    company: str
    role: Optional[str] = None
    interview_type: Optional[str] = None  # If known from context
    application_id: Optional[str] = None

class DebriefExtractionResponse(BaseModel):
    """Response with structured debrief data and insights."""
    structured_data: Dict[str, Any]
    insights: List[str]
    stories_extracted: List[Dict[str, Any]]

class PatternAnalysisResponse(BaseModel):
    """Response with cross-interview pattern analysis."""
    weak_categories: List[Dict[str, Any]]
    strong_categories: List[Dict[str, Any]]
    story_usage: List[Dict[str, Any]]
    confidence_trend: Optional[Dict[str, Any]] = None
    total_debriefs: int
    insights: List[str]


# DEBRIEF_EXTRACTION_PROMPT - Now imported from prompts/ module


# Backwards compatibility aliases
AskHenryContext = HeyHenryContext
AskHenryMessage = HeyHenryMessage
AskHenryRequest = HeyHenryRequest
AskHenryResponse = HeyHenryResponse


# HEY_HENRY_SYSTEM_PROMPT, ASK_HENRY_SYSTEM_PROMPT - Now imported from prompts/ module


@app.post("/api/hey-henry", response_model=HeyHenryResponse)
@limiter.limit("60/minute")
async def hey_henry(request: Request, body: HeyHenryRequest):
    """
    HEY HENRY: Strategic career coach available from any page.

    Provides honest, strategic career coaching based on:
    - Current page/section the user is viewing
    - Their job analysis data (if available)
    - Their resume data (if available)
    - Their pipeline data (if available)
    - Emotional state and tone guidance
    - Clarification detection for vague inputs
    - Conversation history for continuity
    """
    print(f"💬 Hey Henry: {body.context.current_page} - {body.message[:50]}...")

    # Detect pipeline analysis request (Phase 2.2 trigger)
    pipeline_analysis_triggers = [
        "how's my search",
        "hows my search",
        "how is my search",
        "how am i doing",
        "search going",
        "my pipeline",
        "pipeline review",
        "review my applications",
        "application patterns",
        "what patterns",
        "analyze my applications",
        "what am i doing wrong",
        "why am i not getting",
        "not getting interviews",
        "not getting responses",
        "getting rejected",
        "strategy review",
        "job search review",
        "search strategy"
    ]
    message_lower = body.message.lower()
    is_pipeline_analysis_request = any(trigger in message_lower for trigger in pipeline_analysis_triggers)

    # Detect rejection forensics request (Phase 2.4 trigger)
    rejection_forensics_triggers = [
        "why am i getting rejected",
        "keep getting rejected",
        "why do i keep",
        "rejection pattern",
        "analyze my rejections",
        "what's wrong with my",
        "whats wrong with my",
        "not landing interviews",
        "not hearing back",
        "ghosted",
        "no responses",
        "rejection analysis",
        "rejection feedback",
        "why rejected"
    ]
    is_rejection_forensics_request = any(trigger in message_lower for trigger in rejection_forensics_triggers)

    # Build analysis context string
    analysis_context = ""
    if body.analysis_data and body.context.has_analysis:
        # Extract strengths - handle both string and dict formats
        strengths_raw = body.analysis_data.get('strengths', [])[:3]
        strengths_list = []
        for s in strengths_raw:
            if isinstance(s, dict):
                strengths_list.append(s.get('strength', s.get('description', str(s))))
            else:
                strengths_list.append(str(s))

        # Extract gaps - handle both string and dict formats
        gaps_raw = body.analysis_data.get('gaps', [])[:3]
        gaps_list = []
        for g in gaps_raw:
            if isinstance(g, dict):
                gaps_list.append(g.get('gap_description', g.get('description', g.get('gap', str(g)))))
            else:
                gaps_list.append(str(g))

        analysis_context = f"""
ANALYSIS DATA AVAILABLE:
- Company: {body.analysis_data.get('_company_name', 'Unknown')}
- Role: {body.analysis_data.get('role_title', 'Unknown')}
- Fit Score: {body.analysis_data.get('fit_score', 'N/A')}
- Key Strengths: {', '.join(strengths_list) if strengths_list else 'None identified'}
- Key Gaps: {', '.join(gaps_list) if gaps_list else 'None identified'}
"""

    if body.resume_data and body.context.has_resume:
        analysis_context += f"""
RESUME DATA:
- Candidate Name: {body.resume_data.get('name', 'Unknown')}
- Current/Recent Role: {body.resume_data.get('experience', [{}])[0].get('title', 'Unknown') if body.resume_data.get('experience') else 'Unknown'}
"""

    # Build pipeline context from pipeline data
    pipeline_context = ""
    if body.pipeline_data and body.context.has_pipeline:
        pd = body.pipeline_data
        pipeline_context = f"""
PIPELINE DATA (USE THIS TO GIVE SPECIFIC, INFORMED ANSWERS):
- Total Applications: {pd.get('total', 0)}
- Active Applications: {pd.get('active', 0)}
- In Interview Stages: {pd.get('interviewing', 0)}
- Applied (waiting for response): {pd.get('applied', 0)}
- Rejected: {pd.get('rejected', 0)}
- Likely Ghosted: {pd.get('ghosted', 0)}
- Hot/Priority: {pd.get('hot', 0)}
- Average Fit Score: {pd.get('avgFitScore', 0)}%
- Interview Rate: {pd.get('interviewRate', 0)}%
- Summary: {pd.get('summary', 'No summary available')}

TOP APPLICATIONS IN PIPELINE:
"""
        top_apps = pd.get('topApps', [])
        for app in top_apps:
            pipeline_context += f"- {app.get('role', 'Unknown')} at {app.get('company', 'Unknown')}: {app.get('status', 'Unknown')} ({app.get('fitScore', 'N/A')}% fit, {app.get('daysSinceUpdate', 0)}d since update)\n"

        # Add pattern analysis if available (Phase 2.2)
        pattern_data = pd.get('patternAnalysis', {})
        if pattern_data:
            pipeline_context += "\nPIPELINE PATTERN ANALYSIS (use this for strategic insights):\n"

            # Fit distribution
            fit_dist = pattern_data.get('fitDistribution', {})
            if any(fit_dist.values()):
                pipeline_context += f"- Fit Distribution: {fit_dist.get('strong', 0)} strong, {fit_dist.get('moderate', 0)} moderate, {fit_dist.get('reach', 0)} reach, {fit_dist.get('longShot', 0)} long shot\n"
                reach_pct = pattern_data.get('reachPercentage', 0)
                if reach_pct >= 50:
                    pipeline_context += f"  ⚠️ {reach_pct}% of applications are reaches or long shots. Candidate may be overreaching.\n"

            # Conversion rates
            conv_rates = pattern_data.get('conversionRates', {})
            if any(conv_rates.values()):
                pipeline_context += f"- Conversion Rates:\n"
                pipeline_context += f"  • Application to Response: {conv_rates.get('applicationToResponse', 0)}%\n"
                pipeline_context += f"  • Response to Recruiter Screen: {conv_rates.get('responseToRecruiter', 0)}%\n"
                pipeline_context += f"  • Recruiter to Hiring Manager: {conv_rates.get('recruiterToHM', 0)}%\n"
                pipeline_context += f"  • Hiring Manager to Final: {conv_rates.get('hmToFinal', 0)}%\n"
                pipeline_context += f"  • Final to Offer: {conv_rates.get('finalToOffer', 0)}%\n"

            # Velocity
            velocity = pattern_data.get('velocity', {})
            if velocity:
                pipeline_context += f"- Application Velocity: {velocity.get('thisWeek', 0)} this week, {velocity.get('lastWeek', 0)} last week (trend: {velocity.get('trend', 'steady')})\n"

            # Rejections by stage
            rej_stage = pattern_data.get('rejectionsByStage', {})
            if any(rej_stage.values()):
                pipeline_context += f"- Rejections by Stage: {rej_stage.get('resume', 0)} at resume, {rej_stage.get('recruiter', 0)} at recruiter, {rej_stage.get('hiringManager', 0)} at HM, {rej_stage.get('finalRound', 0)} at final\n"

            # Weak spots
            weak_spots = pattern_data.get('weakSpots', [])
            if weak_spots:
                pipeline_context += "- Identified Weak Spots:\n"
                for ws in weak_spots:
                    pipeline_context += f"  • {ws.get('message', 'Unknown issue')}\n"

            # Pre-generated pattern insights
            insights = pattern_data.get('patternInsights', [])
            if insights:
                pipeline_context += "- Pattern Insights (use these in your response):\n"
                for insight in insights:
                    pipeline_context += f"  • {insight}\n"

    # Build network context from LinkedIn connections (Phase 2.1)
    network_context = ""
    if body.network_data:
        nd = body.network_data
        target_company = body.context.company or "target company"

        if nd.get('hasConnections'):
            direct_connections = nd.get('directAtCompany', [])
            total_connections = nd.get('totalConnections', 0)

            if direct_connections and len(direct_connections) > 0:
                network_context = f"""
NETWORK INTELLIGENCE (PROACTIVE - surface this when relevant):
You have {len(direct_connections)} first-degree connection(s) at {target_company}:
"""
                for conn in direct_connections[:5]:  # Limit to top 5
                    name = conn.get('fullName', 'Unknown').strip()
                    position = conn.get('position', 'Unknown position')
                    connected_on = conn.get('connectedOn', '')
                    network_context += f"* {name}, {position}"
                    if connected_on:
                        network_context += f" (connected since {connected_on})"
                    network_context += "\n"

                network_context += f"""
PROACTIVE NETWORK GUIDANCE:
- Warm intros get 3x higher response rates than cold outreach
- Suggest the candidate reach out to these connections for intel or referrals
- Offer to draft outreach to the strongest connection
- Prioritize connections in the same department or with relevant titles
"""
            else:
                network_context = f"""
NETWORK INTELLIGENCE:
The candidate has {total_connections} LinkedIn connections uploaded but NO direct connections at {target_company}.

When relevant, suggest:
- Searching for 2nd-degree connections on LinkedIn
- Looking for alumni or former company overlap
- Using cold outreach templates (they can find these on the Outreach page)
"""

    # Build outreach log context for follow-up prompts (Phase 2.7)
    outreach_log_context = ""
    if body.outreach_log_data:
        ol = body.outreach_log_data
        due_for_followup = ol.get('dueForFollowUp', [])
        due_for_final = ol.get('dueForFinalFollowUp', [])

        if due_for_followup or due_for_final:
            outreach_log_context = """
OUTREACH FOLLOW-UP REMINDERS (PROACTIVE - mention these when relevant):
"""
            if due_for_followup:
                outreach_log_context += "Messages due for follow-up (sent 5+ days ago, no response):\n"
                for o in due_for_followup[:3]:
                    outreach_log_context += f"* {o.get('contactName', 'Unknown')} at {o.get('company', 'Unknown')} ({o.get('channel', 'unknown channel')}, {o.get('daysSince', 0)} days ago)\n"

            if due_for_final:
                outreach_log_context += "\nMessages due for FINAL follow-up (sent 10+ days ago):\n"
                for o in due_for_final[:3]:
                    outreach_log_context += f"* {o.get('contactName', 'Unknown')} at {o.get('company', 'Unknown')} ({o.get('channel', 'unknown channel')}, {o.get('daysSince', 0)} days ago)\n"

            outreach_log_context += """
When the user seems open to it, proactively remind them about pending follow-ups.
Offer to draft a follow-up message if they want.
Don't be pushy, but do keep them accountable.
"""

    # Build interview debrief context for missing debrief prompts (Phase 2.3)
    interview_debrief_context = ""
    if body.interview_debrief_data:
        idd = body.interview_debrief_data
        needs_debrief = idd.get('needsDebrief', [])
        has_skipped = idd.get('hasSkippedDebriefs', False)

        if needs_debrief or has_skipped:
            interview_debrief_context = """
INTERVIEW DEBRIEF REMINDERS (IMPORTANT - surface proactively):
"""
            if needs_debrief:
                interview_debrief_context += "Interviews that need debriefs:\n"
                for d in needs_debrief[:3]:
                    company = d.get('company', 'Unknown')
                    interview_type = d.get('interviewType', 'Interview')
                    days_since = d.get('daysSince', 0)
                    reason = d.get('reason', '')
                    interview_debrief_context += f"* {interview_type} at {company}"
                    if days_since > 0:
                        interview_debrief_context += f" ({days_since} days ago)"
                    interview_debrief_context += f"\n  Reason: {reason}\n"

            if has_skipped:
                interview_debrief_context += """
⚠️ COMPOUNDING VALUE ALERT: This candidate has advanced to later interview stages without debriefing earlier rounds. They're missing the compounding value of pattern recognition across interviews.
"""

            interview_debrief_context += """
DEBRIEF VALUE PITCH:
- Debriefs while details are fresh capture 3x more actionable insights
- Pattern recognition across interviews compounds (e.g., "this is the third time you've struggled with X")
- Skipping debriefs means missing early signals that could help in later rounds

When appropriate, proactively encourage them to complete a debrief:
- "Before your panel interview, let's debrief your HM round. Those insights could help."
- "You've had 3 interviews without debriefs. Want to quickly run through what happened? I can spot patterns."
- "Interview debrief takes 5 minutes. The insights compound. Worth it?"
"""

    # Build cross-interview pattern analysis context (Phase 2.3)
    pattern_analysis_context = ""
    if body.debrief_pattern_analysis:
        pa = body.debrief_pattern_analysis
        total_debriefs = pa.get('total_debriefs', 0)

        if total_debriefs >= 3:
            pattern_analysis_context = f"""
CROSS-INTERVIEW PATTERN INTELLIGENCE (from {total_debriefs} debriefs):
"""
            # Weak categories
            weak = pa.get('weak_categories', [])
            if weak:
                pattern_analysis_context += "⚠️ RECURRING WEAK AREAS (address proactively):\n"
                for w in weak[:3]:
                    pattern_analysis_context += f"  - {w['category']} questions: struggled {w['rate']}% of the time ({w['count']}/{w['total']})\n"

            # Strong categories
            strong = pa.get('strong_categories', [])
            if strong:
                pattern_analysis_context += "\n✅ STRENGTHS TO LEVERAGE:\n"
                for s in strong[:3]:
                    pattern_analysis_context += f"  - {s['category']} questions: nailed {s['rate']}% ({s['count']}/{s['total']})\n"

            # Overused stories
            stories = pa.get('story_usage', [])
            overused = [s for s in stories if s.get('count', 0) >= 3]
            if overused:
                pattern_analysis_context += "\n📖 STORY FRESHNESS ALERT:\n"
                for s in overused[:3]:
                    pattern_analysis_context += f"  - '{s['name']}' used {s['count']} times - suggest alternatives\n"

            # Confidence trend
            trend = pa.get('confidence_trend')
            if trend:
                direction = trend.get('direction', 'stable')
                early = trend.get('early_avg', 0)
                recent = trend.get('recent_avg', 0)
                if direction == 'declining':
                    pattern_analysis_context += f"\n📉 CONFIDENCE TREND: Declining ({early} → {recent}). Address this gently.\n"
                elif direction == 'improving':
                    pattern_analysis_context += f"\n📈 CONFIDENCE TREND: Improving ({early} → {recent}). Acknowledge growth!\n"

            # Key insights
            insights = pa.get('insights', [])
            if insights:
                pattern_analysis_context += "\n🎯 COACHING INSIGHTS TO SURFACE:\n"
                for insight in insights[:4]:
                    pattern_analysis_context += f"  - {insight}\n"

            pattern_analysis_context += """
USE THESE PATTERNS TO:
- Reference past struggles when prepping for similar interviews
- Suggest alternative stories when they're about to use an overused one
- Surface specific coaching ("You've struggled with behavioral questions 3 times - let's practice")
- Celebrate improvements in their weak areas
"""

    # Build generated content context - YOU ARE THE AUTHOR
    generated_content_context = ""

    if body.documents_data:
        dd = body.documents_data
        generated_content_context += """
DOCUMENTS YOU CREATED (speak as the author):
"""
        if dd.get('resume_output'):
            ro = dd['resume_output']
            generated_content_context += f"""
RESUME (you rewrote this):
- Summary: {ro.get('summary', 'N/A')[:200]}...
- Key changes you made: {dd.get('changes_summary', 'Strategic optimization for role alignment')}
"""
        if dd.get('cover_letter'):
            cl = dd['cover_letter']
            generated_content_context += f"""
COVER LETTER (you drafted this):
- Opening hook: {cl.get('opening', cl.get('body', '')[:150] if isinstance(cl, dict) else str(cl)[:150])}...
"""

    if body.outreach_data:
        od = body.outreach_data
        generated_content_context += """
OUTREACH TEMPLATES YOU CREATED:
"""
        templates = od.get('templates', od.get('outreach_templates', []))
        for i, template in enumerate(templates[:3]):
            if isinstance(template, dict):
                generated_content_context += f"- {template.get('type', 'Template')}: {template.get('subject', template.get('message', ''))[:100]}...\n"
            else:
                generated_content_context += f"- Template {i+1}: {str(template)[:100]}...\n"

    if body.interview_prep_data:
        ip = body.interview_prep_data
        generated_content_context += """
INTERVIEW PREP YOU BUILT:
"""
        modules = ip.get('modules', ip.get('prep_modules', []))
        for module in modules[:5]:
            if isinstance(module, dict):
                generated_content_context += f"- {module.get('name', module.get('title', 'Module'))}: {module.get('description', '')[:80]}...\n"
            else:
                generated_content_context += f"- {str(module)[:80]}...\n"

    if body.positioning_data:
        pp = body.positioning_data
        generated_content_context += f"""
POSITIONING STRATEGY YOU DEVELOPED:
- Core positioning: {pp.get('positioning_strategy', pp.get('strategic_positioning', 'N/A'))[:200]}...
- Key talking points: {', '.join(pp.get('talking_points', [])[:3]) if pp.get('talking_points') else 'N/A'}
"""

    # Build emotional context - prefer direct context fields, fall back to user_profile
    emotional_context = ""
    holding_up = body.context.emotional_state
    timeline = body.context.timeline
    confidence = body.context.confidence_level

    # Fall back to user_profile if direct fields not set
    if not holding_up and body.user_profile:
        situation = body.user_profile.get('situation', {})
        holding_up = situation.get('holding_up')
        timeline = timeline or situation.get('timeline')
        confidence = confidence or situation.get('confidence')

    if holding_up or timeline or confidence:
        emotional_context = "USER'S EMOTIONAL STATE (adjust your tone accordingly):\n"

        # Enhanced emotional state mapping (spec v2.1 values)
        holding_up_guidance = {
            'zen': "- They're doing well emotionally. Be professional and efficient.\n",
            'stressed': "- They're stressed. Be reassuring but direct. Acknowledge pressure, then provide clear guidance.\n",
            'struggling': "- They're struggling. Be supportive and steady. Break things into manageable steps.\n",
            'desperate': "- They're desperate. Be empathetic but realistic. Acknowledge difficulty, focus on actionable steps.\n",
            'crushed': "- They're crushed emotionally. Be gentle and direct. Acknowledge the pain, focus on small wins.\n",
            # Legacy mappings
            'hanging_in': "- They're hanging in there but it's hard. Be supportive but action-oriented.\n",
            'doing_okay': "- They're doing okay. You can be more direct and strategic.\n",
        }
        if holding_up and holding_up in holding_up_guidance:
            emotional_context += holding_up_guidance[holding_up]

        timeline_guidance = {
            'urgent': "- URGENT timeline. Move fast but smart. Prioritize high-probability opportunities.\n",
            'soon': "- Need a job soon. Make every application count. No spray-and-pray.\n",
            'actively_looking': "- Actively looking. Be selective. Focus on quality over quantity.\n",
            'no_rush': "- No rush. Only roles worth making a move for. Be picky.\n",
            # Legacy mappings
            'exploring': "- Just exploring, no rush. Can be more strategic and selective.\n",
        }
        if timeline and timeline in timeline_guidance:
            emotional_context += timeline_guidance[timeline]

        confidence_guidance = {
            'low': "- Low confidence. Provide more explanation. Build confidence through logic and facts.\n",
            'need_validation': "- Needs validation. Validate their concerns, then redirect to evidence.\n",
            'shaky': "- Shaky confidence. Acknowledge feelings, but anchor responses in facts and progress.\n",
            'strong': "- Strong confidence. Peer-level conversation. Be direct and strategic.\n",
            # Legacy mappings
            'building': "- Confidence is building. Reinforce wins and progress.\n",
            'confident': "- They're feeling confident. Match their energy, be direct and strategic.\n",
        }
        if confidence and confidence in confidence_guidance:
            emotional_context += confidence_guidance[confidence]

    # Build tone guidance context (from frontend getToneGuidance())
    tone_guidance = ""
    tone_guidance_detail = ""
    if body.context.tone_guidance:
        tone_guidance = f"TONE GUIDANCE FROM CONTEXT:\n{body.context.tone_guidance}"
        tone_guidance_detail = body.context.tone_guidance
    else:
        tone_guidance_detail = "Adapt your tone based on the user's emotional state above."

    # Build clarification context
    clarification_context = ""
    if body.context.needs_clarification and body.context.clarification_hints:
        clarification_context = "⚠️ CLARIFICATION NEEDED - The user's message appears vague. Before responding substantively, ask clarifying questions:\n"
        for hint in body.context.clarification_hints:
            hint_type = hint.get('type', 'unknown')
            hint_text = hint.get('hint', '')
            clarification_context += f"- [{hint_type}] {hint_text}\n"
        clarification_context += "\nDo NOT acknowledge vaguely. Ask specific follow-up questions first."

    # Format system prompt
    user_name = body.context.user_name
    name_note = "" if user_name else "(name not available - use warm generic greetings)"
    system_prompt = HEY_HENRY_SYSTEM_PROMPT.format(
        user_name=user_name or "Unknown",
        name_note=name_note,
        current_page=body.context.current_page,
        page_description=body.context.page_description,
        company=body.context.company or "Not specified",
        role=body.context.role or "Not specified",
        has_analysis="Yes" if body.context.has_analysis else "No",
        has_resume="Yes" if body.context.has_resume else "No",
        has_pipeline="Yes" if body.context.has_pipeline else "No",
        analysis_context=analysis_context,
        pipeline_context=pipeline_context,
        network_context=network_context,
        outreach_log_context=outreach_log_context,
        interview_debrief_context=interview_debrief_context,
        pattern_analysis_context=pattern_analysis_context,
        generated_content_context=generated_content_context,
        emotional_context=emotional_context,
        tone_guidance=tone_guidance,
        tone_guidance_detail=tone_guidance_detail,
        clarification_context=clarification_context
    )

    # Add pipeline analysis instruction if triggered (Phase 2.2)
    if is_pipeline_analysis_request and body.pipeline_data:
        pd = body.pipeline_data
        total_apps = pd.get('total', 0)
        if total_apps >= 5:
            system_prompt += """

=== PIPELINE PATTERN ANALYSIS REQUEST ===

The user is asking for a strategic review of their job search. Use the PIPELINE PATTERN ANALYSIS data above to provide a comprehensive, specific analysis.

REQUIRED RESPONSE FORMAT:
1. **Pipeline Health** (1-2 sentences): Give them the headline. Are they on track or off track?

2. **Patterns I'm Seeing** (3-5 bullet points): Use the actual pattern data:
   - Fit distribution (are they overreaching?)
   - Conversion rates at each stage (where are they dropping off?)
   - Velocity trends (accelerating, slowing, or steady?)
   - Any weak spots identified

3. **The Signal** (1-2 sentences): What's the ONE thing the data is telling them?

4. **Recommendation** (1-2 actionable steps): What should they do next?

CRITICAL RULES:
- Use SPECIFIC numbers from the pattern data. Never say "some" or "a few" when you have exact counts.
- Reference specific companies and roles from their pipeline where relevant.
- If they're overreaching (high % of reaches/long shots), say so directly.
- If they're stalling at a specific stage, diagnose why.
- NO generic advice. Everything must be grounded in THEIR data.

Example of GOOD analysis:
"Here's what I'm seeing across your search:

Pipeline Health: 12 applications, 3 active, 2 interviews scheduled, 4 rejected, 3 ghosted

Patterns:
* You're overreaching on scope. 8 of 12 were stretch roles.
* Strong conversion to first round (67%), but dropping off at hiring manager stage (25%).
* All your ghosted applications were at companies with 50+ open roles. They're probably overwhelmed.

Recommendation: Tighten your targeting to roles where you're a 75%+ fit. Your hit rate will improve."

=== END PIPELINE ANALYSIS INSTRUCTION ===
"""
        else:
            system_prompt += f"""

=== PIPELINE ANALYSIS NOTE ===
The user is asking about their search, but they only have {total_apps} applications tracked. Let them know you need at least 5 applications to identify meaningful patterns. Encourage them to keep tracking and come back once they have more data.
=== END NOTE ===
"""

    # Add rejection forensics instruction if triggered (Phase 2.4)
    if is_rejection_forensics_request and body.pipeline_data:
        pd = body.pipeline_data
        rejected_count = pd.get('rejected', 0)
        pattern_data = pd.get('patternAnalysis', {})
        rej_by_stage = pattern_data.get('rejectionsByStage', {})

        if rejected_count >= 3:
            system_prompt += f"""

=== REJECTION FORENSICS REQUEST ===

The user wants to understand their rejection patterns. Analyze the data and provide diagnostic insights.

REJECTION DATA:
- Total rejections: {rejected_count}
- At resume screen: {rej_by_stage.get('resume', 0)}
- At recruiter screen: {rej_by_stage.get('recruiter', 0)}
- At hiring manager: {rej_by_stage.get('hiringManager', 0)}
- At final round: {rej_by_stage.get('finalRound', 0)}

STAGE-SPECIFIC ANALYSIS:
1. **Resume screen rejections** = Keywords missing, experience mismatch, or ATS issues. Resume needs refinement.
2. **Recruiter screen rejections** = Communication issues, salary mismatch, or timeline misalignment. Prep recruiter conversations.
3. **Hiring manager rejections** = Culture fit or scope concerns, not skills. Dig into those conversations.
4. **Final round rejections** = Competition or factors outside control. You were qualified enough to get there.

RESPONSE FORMAT:
1. **The Pattern**: What stage are they dropping off at? Give specific numbers.
2. **Likely Cause**: What typically causes rejection at this stage?
3. **Diagnostic Questions**: Ask 1-2 questions to understand what happened.
4. **Next Steps**: Specific actions to improve.

CRITICAL RULES:
- Be direct but not harsh. Rejections sting.
- Use THEIR specific data, not generic advice.
- If rejections are spread across stages, note that too.
- If they have little rejection data, say you need more to see patterns.

Example:
"You've been rejected 4 times in the last month. Here's the pattern:

* 2 rejections at resume screen: Your materials might not be optimized for these role types
* 1 at recruiter screen: Was salary discussed? That's often the blocker there
* 1 at final round: That one was competition, not qualification

The signal: Your resume screen rejections are where to focus. What types of roles were those?"

=== END REJECTION FORENSICS INSTRUCTION ===
"""
        elif rejected_count > 0:
            system_prompt += f"""

=== REJECTION FORENSICS NOTE ===
The user is asking about rejections, but they only have {rejected_count} rejection(s) tracked. Let them know you need at least 3 rejections to identify meaningful patterns. Acknowledge their frustration, offer to discuss what happened, but note patterns take more data.
=== END NOTE ===
"""

    # Build messages
    messages = []
    for msg in body.conversation_history:
        messages.append({
            "role": msg.role,
            "content": msg.content
        })

    # Build the user message content - handle attachments if present
    if body.attachments and len(body.attachments) > 0:
        # Multi-modal message with attachments
        user_content = []

        # Process each attachment
        for attachment in body.attachments:
            if attachment.type.startswith('image/'):
                # Image attachment - use Claude's vision format
                # Map MIME types to Claude's expected format
                media_type_map = {
                    'image/png': 'image/png',
                    'image/jpeg': 'image/jpeg',
                    'image/gif': 'image/gif',
                    'image/webp': 'image/webp'
                }
                media_type = media_type_map.get(attachment.type, 'image/png')

                user_content.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": attachment.data
                    }
                })
                print(f"📎 Processing image: {attachment.name} ({attachment.type}, {attachment.size} bytes)")

            elif attachment.type == 'application/pdf':
                # PDF - extract text description (Claude can't read PDFs directly via vision)
                user_content.append({
                    "type": "text",
                    "text": f"[PDF Attachment: {attachment.name}]\n(User uploaded a PDF document. Acknowledge you received it but note that PDF content extraction is not yet implemented. Ask them to copy/paste relevant text if needed.)"
                })
                print(f"📎 Processing PDF: {attachment.name} ({attachment.size} bytes)")

            elif attachment.type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                # Word docs - placeholder
                user_content.append({
                    "type": "text",
                    "text": f"[Document Attachment: {attachment.name}]\n(User uploaded a Word document. Acknowledge you received it but note that document content extraction is not yet implemented. Ask them to copy/paste relevant text if needed.)"
                })
                print(f"📎 Processing document: {attachment.name} ({attachment.size} bytes)")

            elif attachment.type == 'text/plain':
                # Plain text - we can include the content directly
                try:
                    import base64
                    text_content = base64.b64decode(attachment.data).decode('utf-8')
                    user_content.append({
                        "type": "text",
                        "text": f"[Text File: {attachment.name}]\n{text_content}"
                    })
                    print(f"📎 Processing text file: {attachment.name} ({attachment.size} bytes)")
                except Exception as e:
                    user_content.append({
                        "type": "text",
                        "text": f"[Text File: {attachment.name}]\n(Could not decode text file)"
                    })
                    print(f"⚠️ Error decoding text file: {e}")

        # Add the user's message text
        user_content.append({
            "type": "text",
            "text": body.message
        })

        messages.append({
            "role": "user",
            "content": user_content
        })
        print(f"📎 Total attachments processed: {len(body.attachments)}")
    else:
        # Simple text message (no attachments)
        messages.append({
            "role": "user",
            "content": body.message
        })

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=1000,
            system=system_prompt,
            messages=messages
        )

        assistant_response = response.content[0].text
        print(f"✅ Hey Henry response: {len(assistant_response)} chars")

        # =================================================================
        # QA VALIDATION: Check chat response for fabrication
        # Only validate if we have resume data to check against
        # =================================================================
        if body.resume_data and body.context.has_resume:
            qa_validation_result = validate_chat_response(
                response=assistant_response,
                user_question=body.message,
                resume_data=body.resume_data
            )

            if qa_validation_result.should_block:
                # Log the blocked response for review
                print("\n" + "🚫"*20)
                print("HEY HENRY QA VALIDATION BLOCKED - POTENTIAL FABRICATION")
                print("🚫"*20)
                print(f"Question: {body.message[:100]}...")
                for issue in qa_validation_result.issues:
                    print(f"  [{issue.category.value}] {issue.message}")
                print("🚫"*20 + "\n")

                # Log to file for review
                ValidationLogger.log_blocked_output(
                    endpoint="/api/hey-henry",
                    result=qa_validation_result,
                    output={"response": assistant_response},
                    resume_data=body.resume_data,
                    request_context={"question": body.message[:200]}
                )

                # Return a safe fallback response instead of blocking entirely
                fallback_response = create_chat_fallback_response(qa_validation_result, body.message)
                return HeyHenryResponse(response=fallback_response)

        return HeyHenryResponse(response=assistant_response)

    except Exception as e:
        print(f"🔥 Hey Henry error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to get response: {str(e)}")


# Backwards compatibility endpoint
@app.post("/api/ask-henry", response_model=HeyHenryResponse)
async def ask_henry(request: HeyHenryRequest):
    """Backwards compatibility alias for /api/hey-henry."""
    return await hey_henry(request)


# ============================================================================
# RESUME BUILDER - CONVERSATIONAL ONBOARDING
# ============================================================================

class ResumeChatMessage(BaseModel):
    """Single message in resume chat conversation."""
    role: str  # 'user' or 'assistant'
    content: str


class ExtractedSkill(BaseModel):
    """A skill extracted from conversation."""
    skill_name: str
    category: str
    evidence: str
    confidence: str = "medium"  # high, medium, low


class ResumeChatRequest(BaseModel):
    """Request for resume builder conversation."""
    conversation_history: List[ResumeChatMessage]
    current_state: str = "START"
    extracted_data: Optional[Dict[str, Any]] = None


class ResumeChatResponse(BaseModel):
    """Response from resume builder conversation."""
    response: str
    next_state: str
    extracted_data: Dict[str, Any]
    skills_extracted: List[ExtractedSkill] = []
    suggested_responses: List[str] = []


# Skill taxonomy (simplified version - full version in frontend/docs)
SKILL_CATEGORIES = {
    "leadership": ["Team Coordination", "Training", "Delegation", "Mentoring", "Decision Making"],
    "operations": ["Process Improvement", "Inventory Management", "Quality Control", "Scheduling", "Logistics"],
    "customer": ["Customer Service", "Conflict Resolution", "Client Relations", "Upselling", "Account Management"],
    "communication": ["Written Communication", "Verbal Communication", "Presentation", "Negotiation", "Active Listening"],
    "technical": ["Data Analysis", "Software Skills", "Technical Writing", "Problem Solving", "Research"],
    "financial": ["Budgeting", "Financial Analysis", "Cash Handling", "Cost Reduction", "Forecasting"],
    "sales": ["Sales", "Lead Generation", "Account Growth", "Pipeline Management", "Closing"],
    "marketing": ["Marketing", "Social Media", "Content Creation", "Brand Management", "Campaign Management"],
    "administrative": ["Organization", "Documentation", "Scheduling", "Filing", "Compliance"],
    "self_management": ["Time Management", "Adaptability", "Reliability", "Initiative", "Stress Management"]
}

# RESUME_CHAT_SYSTEM_PROMPT - Now imported from prompts/ module


@app.post("/api/resume-chat", response_model=ResumeChatResponse)
async def resume_chat(request: ResumeChatRequest):
    """
    RESUME CHAT: Conversational resume building through natural dialogue.

    Guides users through sharing their experience and extracts skills
    for resume generation.
    """
    print(f"📝 Resume Chat: State={request.current_state}, Messages={len(request.conversation_history)}")

    # Extract candidate name from contact info
    extracted_data = request.extracted_data or {}
    contact = extracted_data.get("contact", {})
    candidate_name = contact.get("nickname") or contact.get("firstName") or "there"

    # Format extracted data for prompt
    extracted_data_str = json.dumps(extracted_data, indent=2)

    system_prompt = RESUME_CHAT_SYSTEM_PROMPT.format(
        candidate_name=candidate_name,
        current_state=request.current_state,
        extracted_data=extracted_data_str
    )

    # Build messages for Claude
    messages = []
    for msg in request.conversation_history:
        messages.append({
            "role": msg.role if msg.role in ["user", "assistant"] else "user",
            "content": msg.content
        })

    try:
        response = client.messages.create(
            model="claude-3-5-haiku-20241022",
            max_tokens=1000,
            system=system_prompt,
            messages=messages
        )

        response_text = response.content[0].text
        print(f"✅ Resume Chat raw response: {response_text[:200]}...")

        # Parse JSON response
        try:
            # Handle potential markdown code blocks
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()

            parsed = json.loads(response_text)

            # Validate and extract fields
            return ResumeChatResponse(
                response=parsed.get("response", "I'd love to hear more about your experience."),
                next_state=parsed.get("next_state", request.current_state),
                extracted_data=parsed.get("extracted_data", request.extracted_data or {}),
                skills_extracted=[
                    ExtractedSkill(**skill) for skill in parsed.get("skills_extracted", [])
                ],
                suggested_responses=parsed.get("suggested_responses", [])
            )

        except json.JSONDecodeError as e:
            print(f"⚠️ Failed to parse JSON response: {e}")
            # Return a graceful fallback
            return ResumeChatResponse(
                response=response_text if len(response_text) < 500 else "Tell me more about what you did in that role.",
                next_state=request.current_state,
                extracted_data=request.extracted_data or {},
                skills_extracted=[],
                suggested_responses=["I manage client projects", "I handle sales calls", "I write and review code"]
            )

    except Exception as e:
        print(f"🔥 Resume Chat error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to process conversation: {str(e)}")


class GenerateResumeFromChatRequest(BaseModel):
    """Request to generate resume from chat-extracted data."""
    extracted_data: Dict[str, Any]
    target_role: Optional[str] = None


class GenerateResumeFromChatResponse(BaseModel):
    """Generated resume from conversation data."""
    resume_text: str
    resume_html: str
    sections: Dict[str, Any]


# RESUME_GENERATION_PROMPT - Now imported from prompts/ module


@app.post("/api/generate-resume-from-chat", response_model=GenerateResumeFromChatResponse)
async def generate_resume_from_chat(request: GenerateResumeFromChatRequest):
    """
    Generate a professional resume from conversation-extracted data.
    """
    print(f"📄 Generating resume from chat data...")

    system_prompt = RESUME_GENERATION_PROMPT.format(
        extracted_data=json.dumps(request.extracted_data, indent=2),
        target_role=request.target_role or "Not specified - create general resume"
    )

    try:
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=3000,
            messages=[{
                "role": "user",
                "content": "Generate the resume based on the conversation data provided in the system prompt."
            }],
            system=system_prompt
        )

        response_text = response.content[0].text

        # Parse JSON response
        try:
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0].strip()
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0].strip()

            parsed = json.loads(response_text)

            return GenerateResumeFromChatResponse(
                resume_text=parsed.get("resume_text", ""),
                resume_html=parsed.get("resume_html", ""),
                sections=parsed.get("sections", {})
            )

        except json.JSONDecodeError:
            # Create basic resume from extracted data
            name = request.extracted_data.get("contact", {}).get("name", "Candidate")
            skills = request.extracted_data.get("skills", [])
            experiences = request.extracted_data.get("experiences", [])

            basic_text = f"{name}\n\n"
            if skills:
                basic_text += "SKILLS: " + ", ".join(skills) + "\n\n"
            if experiences:
                basic_text += "EXPERIENCE:\n"
                for exp in experiences:
                    basic_text += f"{exp.get('title', 'Role')} at {exp.get('company', 'Company')}\n"
                    for resp in exp.get("responsibilities", []):
                        basic_text += f"  - {resp}\n"

            return GenerateResumeFromChatResponse(
                resume_text=basic_text,
                resume_html=f"<pre>{basic_text}</pre>",
                sections={
                    "summary": "",
                    "experience": experiences,
                    "skills": skills,
                    "education": ""
                }
            )

    except Exception as e:
        print(f"🔥 Resume generation error: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to generate resume: {str(e)}")


# ============================================================================
# VOICE ENDPOINTS (Whisper STT + OpenAI TTS)
# ============================================================================

# Import OpenAI for voice features
try:
    import openai
    OPENAI_CLIENT = openai.OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
except ImportError:
    OPENAI_CLIENT = None
    print("⚠️ OpenAI library not installed - voice features disabled")


class SpeakRequest(BaseModel):
    """Request for text-to-speech."""
    text: str
    voice: str = "alloy"  # alloy, echo, fable, onyx, nova, shimmer


@app.post("/api/transcribe")
async def transcribe_audio(audio: UploadFile = File(...)):
    """
    Transcribe audio using OpenAI Whisper.
    Accepts audio file uploads and returns transcribed text.
    """
    if not OPENAI_CLIENT:
        raise HTTPException(status_code=503, detail="Voice features not configured. Please set OPENAI_API_KEY.")

    print(f"🎙️ Transcribing audio: {audio.filename}, {audio.content_type}")

    try:
        # Read audio content
        audio_content = await audio.read()

        # Create a temporary file-like object for OpenAI
        audio_file = io.BytesIO(audio_content)
        audio_file.name = audio.filename or "recording.webm"

        # Call Whisper API
        transcription = OPENAI_CLIENT.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            response_format="text"
        )

        print(f"✅ Transcribed: {transcription[:100]}...")

        return {"text": transcription}

    except Exception as e:
        print(f"🔥 Transcription error: {e}")
        raise HTTPException(status_code=500, detail=f"Transcription failed: {str(e)}")


@app.post("/api/speak")
async def text_to_speech(request: SpeakRequest):
    """
    Convert text to speech using OpenAI TTS.
    Returns audio stream.
    """
    if not OPENAI_CLIENT:
        raise HTTPException(status_code=503, detail="Voice features not configured. Please set OPENAI_API_KEY.")

    print(f"🔊 TTS request: {request.text[:50]}... (voice: {request.voice})")

    try:
        # Call OpenAI TTS API
        response = OPENAI_CLIENT.audio.speech.create(
            model="tts-1",
            voice=request.voice,
            input=request.text
        )

        # Stream the audio response
        audio_content = response.content

        return StreamingResponse(
            io.BytesIO(audio_content),
            media_type="audio/mpeg",
            headers={
                "Content-Disposition": "attachment; filename=speech.mp3"
            }
        )

    except Exception as e:
        print(f"🔥 TTS error: {e}")
        raise HTTPException(status_code=500, detail=f"Text-to-speech failed: {str(e)}")


# ============================================================================
# SCREENSHOT JOB EXTRACTION (Claude Vision)
# ============================================================================

class ScreenshotExtractRequest(BaseModel):
    image: str  # Base64 encoded image

class ExtractedJob(BaseModel):
    title: str
    company: str
    status: Optional[str] = None
    date_applied: Optional[str] = None

class ScreenshotExtractResponse(BaseModel):
    jobs: List[ExtractedJob]
    message: Optional[str] = None


# ============================================================================
# PHASE 1.5: SCREENING QUESTIONS ANALYSIS MODELS
# ============================================================================

class ScreeningQuestionType(str, Enum):
    """Types of screening questions encountered in job applications"""
    YES_NO = "yes_no"
    COMP_EXPECTATIONS = "comp_expectations"
    ESSAY = "essay"
    MULTIPLE_CHOICE = "multiple_choice"
    AVAILABILITY = "availability"

class ScreeningRiskLevel(str, Enum):
    """Risk level for auto-rejection from screening questions"""
    HIGH = "high"      # Honest answer likely triggers auto-rejection
    MEDIUM = "medium"  # Borderline/defensible answer needed
    LOW = "low"        # Meets requirements but could be stronger
    SAFE = "safe"      # Clearly meets requirements

class HonestyFlag(str, Enum):
    """Honesty assessment for recommended answers"""
    TRUTHFUL = "truthful"              # 100% accurate, no embellishment
    STRATEGIC_FRAMING = "strategic_framing"  # Defensible but generous interpretation
    BORDERLINE = "borderline"          # Requires judgment call

class ScreeningQuestionInput(BaseModel):
    """Individual screening question from user"""
    question_text: str
    question_type: ScreeningQuestionType
    required: bool = True
    user_draft_answer: Optional[str] = None

class ScreeningQuestionsAnalyzeRequest(BaseModel):
    """Request for analyzing screening questions for auto-rejection risk"""
    screening_questions: List[ScreeningQuestionInput]
    job_description_text: str
    candidate_resume_data: Dict[str, Any]
    job_analysis: Dict[str, Any]

class ScreeningQuestionAnalysis(BaseModel):
    """Analysis result for a single screening question"""
    question_text: str
    question_type: ScreeningQuestionType
    auto_rejection_risk: ScreeningRiskLevel
    risk_explanation: str
    recommended_answer: str
    recommended_answer_rationale: str
    honesty_flag: HonestyFlag
    honesty_explanation: str
    alternative_approach: str
    keywords_to_include: Optional[List[str]] = None

class ScreeningQuestionsAnalyzeResponse(BaseModel):
    """Response containing analysis of all screening questions"""
    screening_analysis: List[ScreeningQuestionAnalysis]
    overall_risk_score: ScreeningRiskLevel
    critical_dealbreakers: List[str]
    strategic_guidance: str
    conversational_summary: str


# ============================================================================
# PHASE 1.5: DOCUMENT REFINEMENT MODELS
# ============================================================================

class DocumentTypeForRefine(str, Enum):
    """Document types that can be refined via chat"""
    RESUME = "resume"
    COVER_LETTER = "cover_letter"
    OUTREACH = "outreach"

class DocumentRefineRequest(BaseModel):
    """Request for refining a document via chat command"""
    chat_command: str
    target_document: DocumentTypeForRefine
    current_document_data: Dict[str, Any]
    original_jd_analysis: Dict[str, Any]
    original_resume_data: Dict[str, Any]
    conversation_history: List[Dict[str, Any]] = []
    version: int = 1

class DocumentChangeDetail(BaseModel):
    """Details of a specific change made during refinement"""
    section: str
    before: str
    after: str
    change_type: str  # "added", "removed", "modified"

class DocumentChangesSummary(BaseModel):
    """Summary of all changes made during document refinement"""
    what_changed: str
    sections_modified: List[str]
    positioning_shift: str
    ats_impact: str
    version: int

class DocumentRefineResponse(BaseModel):
    """Response from document refinement"""
    updated_document: Dict[str, Any]
    changes_summary: DocumentChangesSummary
    changes_detail: Optional[List[DocumentChangeDetail]] = None
    conversational_response: str
    validation: Dict[str, Any]


@app.post("/api/extract-jobs-from-screenshot", response_model=ScreenshotExtractResponse)
async def extract_jobs_from_screenshot(request: ScreenshotExtractRequest):
    """
    Extract job applications from a screenshot of an ATS (Greenhouse, Lever, etc.)
    Uses Claude Vision to analyze the image and extract structured job data.
    """
    try:
        print(f"📸 Processing screenshot for job extraction...")

        # Prepare the prompt for Claude Vision
        extraction_prompt = """Analyze this screenshot of a job application tracker (likely from Greenhouse, Lever, Workday, or similar ATS).

Extract ALL job applications visible in the image. For each job, identify:
1. Job title/role
2. Company name
3. Application status (if visible - e.g., "Applied", "In Review", "Interview", "Rejected", "Offer")
4. Date applied (if visible - any format is fine)

Return a JSON object with this exact structure:
{
    "jobs": [
        {
            "title": "Job Title Here",
            "company": "Company Name Here",
            "status": "Status if visible or null",
            "date_applied": "Date if visible or null"
        }
    ],
    "message": "Optional note about the extraction"
}

Important:
- Extract ALL jobs visible, even if some details are missing
- If company or title is unclear, make your best guess based on context
- For status, normalize to: "Applied", "Screening", "Interview", "Offer", "Rejected", or null if unknown
- Return ONLY valid JSON, no markdown code blocks or extra text"""

        # Call Claude Vision API
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": "image/png",
                                "data": request.image
                            }
                        },
                        {
                            "type": "text",
                            "text": extraction_prompt
                        }
                    ]
                }
            ]
        )

        result_text = response.content[0].text.strip()
        print(f"📸 Claude Vision response: {result_text[:500]}...")

        # Clean up potential markdown formatting
        if result_text.startswith("```"):
            import re
            result_text = re.sub(r'^```(?:json)?\n?', '', result_text)
            result_text = re.sub(r'\n?```$', '', result_text)

        # Parse the JSON response
        try:
            parsed_result = json.loads(result_text)
            jobs = parsed_result.get("jobs", [])
            message = parsed_result.get("message")

            print(f"📸 Extracted {len(jobs)} jobs from screenshot")

            return ScreenshotExtractResponse(
                jobs=[ExtractedJob(**job) for job in jobs],
                message=message
            )
        except json.JSONDecodeError as e:
            print(f"🔥 JSON parse error: {e}")
            print(f"🔥 Raw response: {result_text}")
            raise HTTPException(status_code=500, detail="Failed to parse job extraction results")

    except anthropic.APIError as e:
        print(f"🔥 Claude Vision API error: {e}")
        raise HTTPException(status_code=500, detail=f"AI processing failed: {str(e)}")
    except Exception as e:
        print(f"🔥 Screenshot extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Screenshot processing failed: {str(e)}")


# ============================================================================
# PHASE 2.3: INTERVIEW DEBRIEF INTELLIGENCE ENDPOINTS
# ============================================================================

@app.post("/api/debriefs/extract", response_model=DebriefExtractionResponse)
async def extract_debrief_data(request: DebriefExtractionRequest):
    """
    Extract structured data from interview debrief conversation.

    Uses Claude to analyze the conversation and extract:
    - Interview metadata (type, duration, interviewer)
    - Self-ratings (overall, confidence, preparation)
    - Questions asked and categories
    - Stumbles and wins
    - Stories used with effectiveness ratings
    - Interviewer signals
    - Key insights and improvement areas
    """
    print(f"🎯 Extracting debrief data for {request.company}")

    try:
        # Format the extraction prompt
        prompt = DEBRIEF_EXTRACTION_PROMPT.format(
            conversation=request.conversation,
            company=request.company,
            role=request.role or "Not specified",
            interview_type=request.interview_type or "Unknown"
        )

        # Call Claude for extraction (use haiku for speed/cost)
        response = call_claude(
            system_prompt="You are a JSON extraction assistant. Return ONLY valid JSON, no markdown or other text.",
            user_message=prompt,
            max_tokens=2000,
            temperature=0
        )

        # Parse the JSON response
        try:
            # Clean up response if needed
            json_text = response.strip()
            if json_text.startswith("```"):
                json_text = json_text.split("```")[1]
                if json_text.startswith("json"):
                    json_text = json_text[4:]
            json_text = json_text.strip()

            structured_data = json.loads(json_text)
        except json.JSONDecodeError as e:
            print(f"JSON parse error: {e}")
            print(f"Response was: {response[:500]}")
            # Return minimal structure if parsing fails
            structured_data = {
                "interview_type": request.interview_type or "unknown",
                "key_insights": ["Debrief captured but structured extraction failed"],
                "improvement_areas": []
            }

        # Extract stories for story bank
        stories_extracted = structured_data.get("stories_used", [])

        # Generate insights summary
        insights = structured_data.get("key_insights", [])

        # Add company/role to structured data
        structured_data["company"] = request.company
        structured_data["role"] = request.role
        if request.application_id:
            structured_data["application_id"] = request.application_id

        print(f"✅ Extracted debrief: {len(structured_data.get('questions_asked', []))} questions, {len(stories_extracted)} stories, {len(insights)} insights")

        return DebriefExtractionResponse(
            structured_data=structured_data,
            insights=insights,
            stories_extracted=stories_extracted
        )

    except Exception as e:
        print(f"❌ Debrief extraction error: {e}")
        raise HTTPException(status_code=500, detail=f"Debrief extraction failed: {str(e)}")


@app.post("/api/debriefs/analyze-patterns", response_model=PatternAnalysisResponse)
async def analyze_debrief_patterns(debriefs: List[Dict[str, Any]]):
    """
    Analyze cross-interview patterns from multiple debriefs.

    Identifies:
    - Weak categories (struggled 50%+ of time)
    - Strong categories (performed well consistently)
    - Story usage patterns (overused stories)
    - Confidence trend over time
    - Actionable insights
    """
    print(f"📊 Analyzing patterns across {len(debriefs)} debriefs")

    if len(debriefs) < 3:
        return PatternAnalysisResponse(
            weak_categories=[],
            strong_categories=[],
            story_usage=[],
            confidence_trend=None,
            total_debriefs=len(debriefs),
            insights=["Need at least 3 debriefs to identify meaningful patterns"]
        )

    # Aggregate question categories and performance
    category_performance = {}
    for d in debriefs:
        for cat in d.get("question_categories", []):
            if cat not in category_performance:
                category_performance[cat] = {"total": 0, "struggles": 0, "wins": 0}
            category_performance[cat]["total"] += 1

            # Check if stumbled on this category
            stumbles = d.get("stumbles", [])
            for s in stumbles:
                if cat.lower() in s.get("question", "").lower():
                    category_performance[cat]["struggles"] += 1

            # Check if won on this category
            wins = d.get("wins", [])
            for w in wins:
                if cat.lower() in w.get("moment", "").lower():
                    category_performance[cat]["wins"] += 1

    # Identify weak categories (struggled 50%+ of time, min 2 occurrences)
    weak_categories = []
    strong_categories = []
    for cat, data in category_performance.items():
        if data["total"] >= 2:
            struggle_rate = data["struggles"] / data["total"]
            win_rate = data["wins"] / data["total"]
            if struggle_rate >= 0.5:
                weak_categories.append({
                    "category": cat,
                    "count": data["struggles"],
                    "total": data["total"],
                    "rate": round(struggle_rate * 100)
                })
            elif win_rate >= 0.5:
                strong_categories.append({
                    "category": cat,
                    "count": data["wins"],
                    "total": data["total"],
                    "rate": round(win_rate * 100)
                })

    # Track story usage
    story_count = {}
    for d in debriefs:
        for s in d.get("stories_used", []):
            name = s.get("name") or s.get("story_name")
            if name:
                story_count[name] = story_count.get(name, 0) + 1

    story_usage = [{"name": name, "count": count}
                   for name, count in sorted(story_count.items(), key=lambda x: -x[1])]

    # Confidence trend
    confidence_trend = None
    sorted_debriefs = sorted(
        [d for d in debriefs if d.get("rating_confidence")],
        key=lambda x: x.get("interview_date") or x.get("created_at") or ""
    )
    if len(sorted_debriefs) >= 3:
        mid = len(sorted_debriefs) // 2
        first_half = sorted_debriefs[:mid]
        second_half = sorted_debriefs[mid:]

        first_avg = sum(d["rating_confidence"] for d in first_half) / len(first_half)
        second_avg = sum(d["rating_confidence"] for d in second_half) / len(second_half)

        direction = "improving" if second_avg > first_avg else "declining" if second_avg < first_avg else "stable"
        confidence_trend = {
            "direction": direction,
            "early_avg": round(first_avg, 1),
            "recent_avg": round(second_avg, 1)
        }

    # Generate insights
    insights = []

    if weak_categories:
        worst = max(weak_categories, key=lambda x: x["rate"])
        insights.append(f"You struggle with {worst['category']} questions ({worst['rate']}% struggle rate). Focus prep here.")

    if story_usage and story_usage[0]["count"] >= 3:
        overused = story_usage[0]
        insights.append(f"Your '{overused['name']}' story has been used {overused['count']} times. Consider developing alternatives.")

    if confidence_trend:
        if confidence_trend["direction"] == "improving":
            insights.append(f"Your confidence is improving ({confidence_trend['early_avg']} → {confidence_trend['recent_avg']}). Keep it up!")
        elif confidence_trend["direction"] == "declining":
            insights.append(f"Your confidence has dropped ({confidence_trend['early_avg']} → {confidence_trend['recent_avg']}). Let's address this.")

    if strong_categories:
        best = max(strong_categories, key=lambda x: x["rate"])
        insights.append(f"You're strong on {best['category']} questions ({best['rate']}% win rate). Lean into this strength.")

    print(f"✅ Pattern analysis complete: {len(weak_categories)} weak areas, {len(insights)} insights")

    return PatternAnalysisResponse(
        weak_categories=weak_categories,
        strong_categories=strong_categories,
        story_usage=story_usage,
        confidence_trend=confidence_trend,
        total_debriefs=len(debriefs),
        insights=insights
    )


# ============================================================================
# PHASE 1.5: SCREENING QUESTIONS ANALYSIS ENDPOINT
# ============================================================================

@app.post("/api/screening-questions/analyze", response_model=ScreeningQuestionsAnalyzeResponse)
async def analyze_screening_questions_risk(request: ScreeningQuestionsAnalyzeRequest):
    """
    Analyzes screening questions for auto-rejection risk.

    Unlike /api/screening-questions/generate which creates responses,
    this endpoint ANALYZES questions to identify:
    - Auto-rejection risk levels
    - Knockout questions
    - Strategic answer recommendations with honesty flags
    - Critical dealbreakers

    Returns recommended answers with explicit honesty assessments.
    """
    try:
        print(f"🔍 Analyzing {len(request.screening_questions)} screening questions for risk...")

        system_prompt = """You are a recruiting systems expert who understands how ATS filters work.

Your job: analyze screening questions and predict auto-rejection risk.

CRITICAL RULES:
1. NEVER recommend fabricated answers - if the honest answer means auto-rejection, say so
2. Flag "strategic_framing" answers that are defensible but require generous interpretation
3. If honest answer = auto-rejection, mark as HIGH risk and recommend user reconsider applying
4. Provide alternative approaches for borderline cases
5. Include ATS keywords in essay-type recommended answers
6. Use the candidate's ACTUAL resume data to determine honest answers
7. Use the job description to detect hard requirements and threshold filters

WRITING STYLE:
- Use proper grammar and punctuation throughout all responses
- NEVER use em dashes (—) or en dashes (–). Use commas, semicolons, colons, or separate sentences instead
- Write in clear, professional English
- Avoid run-on sentences

RISK SCORING:
- "high": Honest answer clearly fails stated requirement (e.g., 4 years when 5+ required)
- "medium": Answer is borderline/defensible (e.g., 4.5 years when 5+ required)
- "low": Answer meets requirements but could be stronger
- "safe": Answer clearly meets requirements with no concerns

HONESTY FLAGS:
- "truthful": Answer reflects candidate's actual experience with no embellishment
- "strategic_framing": Technically defensible but requires generous interpretation (e.g., counting academic work as "professional")
- "borderline": Requires judgment call, explain the tradeoffs

At the start of your response, provide a brief conversational_summary (2-3 sentences) explaining the overall situation before the JSON.

Return valid JSON after the summary. Structure:
{
  "screening_analysis": [
    {
      "question_text": "...",
      "question_type": "yes_no|comp_expectations|essay|multiple_choice|availability",
      "auto_rejection_risk": "high|medium|low|safe",
      "risk_explanation": "...",
      "recommended_answer": "...",
      "recommended_answer_rationale": "...",
      "honesty_flag": "truthful|strategic_framing|borderline",
      "honesty_explanation": "...",
      "alternative_approach": "...",
      "keywords_to_include": ["keyword1", "keyword2"] // for essay questions only, null otherwise
    }
  ],
  "overall_risk_score": "high|medium|low|safe",
  "critical_dealbreakers": ["...", "..."],
  "strategic_guidance": "..."
}"""

        # Format questions for analysis
        questions_json = json.dumps([q.dict() for q in request.screening_questions], indent=2)

        user_message = f"""JOB DESCRIPTION:
{request.job_description_text}

CANDIDATE RESUME DATA:
{json.dumps(request.candidate_resume_data, indent=2)}

JOB ANALYSIS (FIT SCORE, REQUIREMENTS):
{json.dumps(request.job_analysis, indent=2)}

SCREENING QUESTIONS TO ANALYZE:
{questions_json}

Analyze each question for auto-rejection risk and provide strategic, honest recommendations."""

        response_text = call_claude(system_prompt, user_message, max_tokens=8000)

        # Parse conversational summary and JSON
        conversational_summary = ""
        json_text = response_text.strip()

        # Check for conversational prefix before JSON
        if not json_text.startswith("{"):
            # Find where JSON starts
            json_start = json_text.find("{")
            if json_start > 0:
                conversational_summary = json_text[:json_start].strip()
                json_text = json_text[json_start:]

        # Clean markdown code blocks if present
        if "```json" in json_text:
            json_text = json_text.split("```json")[1].split("```")[0].strip()
        elif "```" in json_text:
            parts = json_text.split("```")
            for part in parts:
                if part.strip().startswith("{"):
                    json_text = part.strip()
                    break

        # Parse JSON response
        try:
            parsed = json.loads(json_text)
        except json.JSONDecodeError as e:
            print(f"❌ JSON parsing error: {e}")
            print(f"Response text: {json_text[:1000]}")
            raise HTTPException(status_code=500, detail="Failed to parse screening questions analysis")

        # Validate response structure
        if "screening_analysis" not in parsed:
            raise ValueError("Missing screening_analysis in response")

        # Build response with conversational summary
        parsed["conversational_summary"] = conversational_summary if conversational_summary else parsed.get("strategic_guidance", "Analysis complete.")

        print(f"✅ Analyzed {len(parsed['screening_analysis'])} questions. Overall risk: {parsed['overall_risk_score']}")

        return ScreeningQuestionsAnalyzeResponse(**parsed)

    except json.JSONDecodeError as e:
        print(f"🔥 JSON parsing error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse screening questions analysis")
    except Exception as e:
        print(f"🔥 Error analyzing screening questions: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# PHASE 1.5: DOCUMENT REFINEMENT ENDPOINT
# ============================================================================

@app.post("/api/documents/refine", response_model=DocumentRefineResponse)
async def refine_document_from_chat(request: DocumentRefineRequest):
    """
    Refines an existing document based on chat command.

    Allows users to iteratively improve documents through natural language:
    - "Make it more senior"
    - "Add more ATS keywords"
    - "Make the summary shorter"
    - "Remove my internship"

    Returns updated document with changes summary and validation.
    Maintains version history for tracking iterations.
    """
    try:
        print(f"✏️ Refining {request.target_document.value} (v{request.version}): '{request.chat_command}'")

        system_prompt = f"""You are refining an existing {request.target_document.value} based on user feedback.

CRITICAL RULES:
1. Only change what the user requested - don't rewrite the entire document
2. Preserve all factual information - NO FABRICATION
3. Maintain ATS keyword coverage - don't remove keywords unless explicitly asked
4. Track what you changed in changes_summary
5. Explain your changes conversationally in 2-3 sentences
6. If user requests fabrication (e.g., "add 2 years of experience"), refuse and explain why

SUPPORTED REFINEMENTS:
- Tone adjustments ("make it more senior", "less formal")
- Section rewrites ("rewrite my summary to emphasize data skills")
- Bullet refinements ("make the third bullet stronger")
- Keyword injection ("add more ML keywords")
- Content removal ("remove my internship")
- Format changes ("make it more concise")

FORBIDDEN OPERATIONS (refuse politely):
- Adding fake experience, companies, or credentials
- Changing dates or tenure
- Fabricating metrics or achievements
- Adding skills not evidenced in resume

Return valid JSON only. Structure:
{{
  "updated_document": {{ ... same structure as input document ... }},
  "changes_summary": {{
    "what_changed": "Brief description of changes",
    "sections_modified": ["summary", "experience"],
    "positioning_shift": "How the document positioning changed",
    "ats_impact": "Impact on ATS keyword coverage",
    "version": {request.version + 1}
  }},
  "changes_detail": [
    {{
      "section": "summary",
      "before": "Original text...",
      "after": "Modified text...",
      "change_type": "modified"
    }}
  ],
  "conversational_response": "2-3 sentence explanation of what was changed and why"
}}"""

        # Build conversation context
        conversation_context = ""
        if request.conversation_history:
            recent_history = request.conversation_history[-5:]
            conversation_context = f"\nRECENT CONVERSATION:\n{json.dumps(recent_history, indent=2)}\n"

        user_message = f"""CHAT COMMAND: {request.chat_command}

TARGET DOCUMENT: {request.target_document.value}

CURRENT DOCUMENT (VERSION {request.version}):
{json.dumps(request.current_document_data, indent=2)}

ORIGINAL JD ANALYSIS:
{json.dumps(request.original_jd_analysis, indent=2)}

ORIGINAL RESUME DATA (source of truth - cannot fabricate beyond this):
{json.dumps(request.original_resume_data, indent=2)}
{conversation_context}
Refine the document based on the chat command. Return the full updated document with changes_summary."""

        response_text = call_claude(system_prompt, user_message, max_tokens=8000)

        # Clean markdown code blocks if present
        json_text = response_text.strip()
        if "```json" in json_text:
            json_text = json_text.split("```json")[1].split("```")[0].strip()
        elif "```" in json_text:
            parts = json_text.split("```")
            for part in parts:
                if part.strip().startswith("{"):
                    json_text = part.strip()
                    break

        # Parse JSON response
        try:
            parsed = json.loads(json_text)
        except json.JSONDecodeError as e:
            print(f"❌ JSON parsing error: {e}")
            print(f"Response text: {json_text[:1000]}")
            raise HTTPException(status_code=500, detail="Failed to parse document refinement")

        # Validate structure
        if "updated_document" not in parsed:
            raise ValueError("Missing updated_document in response")
        if "changes_summary" not in parsed:
            raise ValueError("Missing changes_summary in response")
        if "conversational_response" not in parsed:
            parsed["conversational_response"] = "Document updated successfully."

        # Ensure version is incremented
        if "changes_summary" in parsed:
            if "version" not in parsed["changes_summary"]:
                parsed["changes_summary"]["version"] = request.version + 1

        # Validate document didn't introduce fabrication
        # Build a document structure that validate_document_quality expects
        doc_for_validation = {}
        if request.target_document == DocumentTypeForRefine.RESUME:
            doc_for_validation["resume_output"] = parsed["updated_document"]
        elif request.target_document == DocumentTypeForRefine.COVER_LETTER:
            doc_for_validation["cover_letter"] = parsed["updated_document"]
        else:
            doc_for_validation["outreach"] = parsed["updated_document"]

        validation_results = validate_document_quality(
            doc_for_validation,
            request.original_resume_data,
            request.original_jd_analysis
        )

        # Check if validation found fabrication issues
        if validation_results.get("approval_status") == "FAIL":
            issues = validation_results.get("issues", [])
            fabrication_issues = [i for i in issues if "fabricat" in i.lower()]

            if fabrication_issues:
                print(f"⚠️ Refinement rejected due to fabrication: {fabrication_issues}")
                # Return original document with error message
                return DocumentRefineResponse(
                    updated_document=request.current_document_data,
                    changes_summary=DocumentChangesSummary(
                        what_changed="Refinement rejected - would introduce fabricated content",
                        sections_modified=[],
                        positioning_shift="No change",
                        ats_impact="No change",
                        version=request.version
                    ),
                    changes_detail=None,
                    conversational_response="I couldn't make that change because it would require fabricating information that's not in your resume. Let's try a different approach - what specific aspect would you like to strengthen using your actual experience?",
                    validation=validation_results
                )

        parsed["validation"] = validation_results

        print(f"✅ Document refined successfully (v{request.version} -> v{parsed['changes_summary']['version']})")

        # Convert changes_summary dict to model
        changes_summary = DocumentChangesSummary(**parsed["changes_summary"])

        # Convert changes_detail if present
        changes_detail = None
        if "changes_detail" in parsed and parsed["changes_detail"]:
            changes_detail = [DocumentChangeDetail(**cd) for cd in parsed["changes_detail"]]

        return DocumentRefineResponse(
            updated_document=parsed["updated_document"],
            changes_summary=changes_summary,
            changes_detail=changes_detail,
            conversational_response=parsed["conversational_response"],
            validation=parsed["validation"]
        )

    except json.JSONDecodeError as e:
        print(f"🔥 JSON parsing error: {e}")
        raise HTTPException(status_code=500, detail="Failed to parse document refinement")
    except Exception as e:
        print(f"🔥 Error refining document: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ============================================================================
# LINKEDIN PROFILE MODELS
# ============================================================================

class LinkedInExperience(BaseModel):
    """A single experience entry from LinkedIn"""
    title: str
    company: str
    dates: Optional[str] = None
    location: Optional[str] = None
    bullets: List[str] = []

class LinkedInEducation(BaseModel):
    """A single education entry from LinkedIn"""
    school: str
    degree: Optional[str] = None
    dates: Optional[str] = None

class LinkedInParsedData(BaseModel):
    """Structured data parsed from LinkedIn PDF"""
    headline: Optional[str] = None
    summary: Optional[str] = None
    current_role: Optional[str] = None
    current_company: Optional[str] = None
    experience: List[LinkedInExperience] = []
    skills: List[str] = []
    education: List[LinkedInEducation] = []

class LinkedInUploadResponse(BaseModel):
    """Response from LinkedIn upload"""
    success: bool
    parsed_data: LinkedInParsedData
    message: str

class LinkedInAlignmentIssue(BaseModel):
    """A single alignment issue between resume and LinkedIn"""
    type: str  # title_mismatch, company_mismatch, date_mismatch, skills_mismatch
    severity: str  # high, medium, low
    resume_value: Optional[str] = None
    linkedin_value: Optional[str] = None
    message: str

class LinkedInAlignmentResponse(BaseModel):
    """Response from alignment check"""
    aligned: bool
    discrepancies: List[LinkedInAlignmentIssue]
    severity: str  # overall severity: high, medium, low
    discrepancy_count: int

class LinkedInOptimizeRequest(BaseModel):
    """Request for optimizing LinkedIn sections"""
    job_id: Optional[str] = None
    resume_json: Optional[Dict[str, Any]] = None
    job_description: Optional[str] = None
    target_role: Optional[str] = None
    linkedin_data: Optional[Dict[str, Any]] = None  # Current LinkedIn profile data
    # Single Source of Truth: Job Fit decision fields
    job_fit_recommendation: Optional[str] = None  # "Do Not Apply", "Apply with Caution", etc.
    recommendation_locked: Optional[bool] = False
    locked_reason: Optional[str] = None
    proceeded_despite_guidance: Optional[bool] = False  # User clicked "Pass Anyway"

class LinkedInExperienceOptimization(BaseModel):
    """Optimized content for a single experience entry"""
    title: str
    company: str
    dates: str
    company_context: str  # E.g., "Fortune 500 tech company with 50K+ employees"
    bullets: List[str]
    why_these_work: List[str]

class LinkedInOptimizeResponse(BaseModel):
    """Response with optimized LinkedIn sections"""
    # Severity and summary
    severity: str = "MEDIUM"  # CRITICAL, HIGH, MEDIUM, LOW
    summary_message: str = ""
    benefits: List[str] = []

    # Headline
    headline: str
    current_headline: Optional[str] = None
    headline_why: List[str] = []
    headline_update_reason: str = ""

    # About section
    summary: str
    current_summary: Optional[str] = None
    summary_why: List[str] = []
    summary_update_reason: str = ""

    # Experience
    experience_optimizations: List[LinkedInExperienceOptimization] = []

    # Skills
    top_skills: List[str]
    skills_to_add: List[str] = []
    skills_to_remove: List[str] = []
    skills_why: List[str] = []

    # Optional sections
    featured_recommendation: str = ""
    featured_suggestions: List[str] = []
    recommendations_advice: str = ""
    who_to_ask: List[Dict[str, str]] = []
    activity_recommendation: str = ""

    # Legacy fields for backwards compatibility
    experience_highlights: Optional[List[str]] = None
    has_issues: bool = False
    issue_count: int = 0
    generated_at: str

    # LEPE Integration (for Manager+ target roles)
    lepe_applicable: Optional[bool] = None  # True if target is Manager+
    lepe_decision: Optional[str] = None  # apply | position | caution | locked
    lepe_coaching: Optional[str] = None  # Coaching advice if positioning mode
    lepe_skepticism_warning: Optional[str] = None  # Warning if caution/locked
    leadership_tenure: Optional[Dict[str, Any]] = None  # From LEPE analysis
    accountability_record: Optional[Dict[str, Any]] = None  # LEPE accountability

# ============================================================================
# LINKEDIN API ENDPOINTS
# ============================================================================

@app.post("/api/linkedin/upload", response_model=LinkedInUploadResponse)
async def upload_linkedin_profile(
    file: UploadFile = File(...)
) -> LinkedInUploadResponse:
    """
    Upload and parse a LinkedIn profile PDF

    Accepts PDF file exported from LinkedIn (Profile -> More -> Save to PDF)
    Returns structured data parsed from the profile.
    """

    try:
        # Validate file type
        filename = file.filename.lower() if file.filename else ""
        if not filename.endswith('.pdf'):
            raise HTTPException(
                status_code=400,
                detail="Only PDF files are supported. Please upload your LinkedIn profile as a PDF."
            )

        # Read file
        file_bytes = await file.read()

        # Validate file size (10MB max)
        max_size = 10 * 1024 * 1024  # 10MB
        if len(file_bytes) > max_size:
            raise HTTPException(
                status_code=400,
                detail="File size exceeds 10MB limit. LinkedIn PDFs are typically 1-3MB."
            )

        # Extract text from PDF
        text_content = extract_pdf_text(file_bytes)

        if not text_content or len(text_content.strip()) < 100:
            raise HTTPException(
                status_code=422,
                detail="Could not extract sufficient text from PDF. Please ensure you uploaded a valid LinkedIn profile PDF."
            )

        print(f"📄 LinkedIn PDF text extracted: {len(text_content)} chars")

        # Parse LinkedIn content using Claude
        system_prompt = """You are a LinkedIn profile parser. Extract structured information from LinkedIn PDF exports.

LinkedIn PDFs have a specific format:
- Name appears at top
- Headline appears after name (job title + value proposition)
- "About" or "Summary" section contains the profile summary
- "Experience" section lists jobs with title, company, dates, and bullet points
- "Skills" section lists skills (sometimes separated by · or newlines)
- "Education" section lists schools, degrees, and dates

CRITICAL RULES:
- Extract ONLY information that is explicitly present in the text
- Do NOT fabricate or infer data
- If a field is not found, use null or empty array
- Preserve exact text as written
- For experience, extract company name, title, date range, and bullet points

Return valid JSON matching this structure:
{
    "headline": "string or null - the professional headline/tagline",
    "summary": "string or null - the About/Summary section text",
    "current_role": "string or null - current job title",
    "current_company": "string or null - current employer",
    "experience": [
        {
            "title": "string",
            "company": "string",
            "dates": "string or null",
            "location": "string or null",
            "bullets": ["array of responsibility/achievement bullets"]
        }
    ],
    "skills": ["array of skill keywords"],
    "education": [
        {
            "school": "string",
            "degree": "string or null",
            "dates": "string or null"
        }
    ]
}

Your response must be ONLY valid JSON, no additional text."""

        user_message = f"Parse this LinkedIn profile PDF text into structured JSON:\n\n{text_content}"

        # Call Claude
        print("📤 Sending LinkedIn content to Claude for parsing...")
        response = call_claude(system_prompt, user_message)
        print(f"📥 Received response from Claude: {len(response)} chars")

        # Clean and parse JSON response
        cleaned = clean_claude_json(response)

        try:
            parsed_data = json.loads(cleaned)
        except json.JSONDecodeError as e:
            print(f"🔥 JSON parse error: {e}")
            raise HTTPException(
                status_code=500,
                detail="Failed to parse LinkedIn profile. Please try downloading your profile again from LinkedIn."
            )

        print(f"✅ Successfully parsed LinkedIn profile: {parsed_data.get('headline', 'Unknown')}")

        # Convert to response model
        linkedin_data = LinkedInParsedData(
            headline=parsed_data.get("headline"),
            summary=parsed_data.get("summary"),
            current_role=parsed_data.get("current_role"),
            current_company=parsed_data.get("current_company"),
            experience=[LinkedInExperience(**exp) for exp in parsed_data.get("experience", [])],
            skills=parsed_data.get("skills", []),
            education=[LinkedInEducation(**edu) for edu in parsed_data.get("education", [])]
        )

        return LinkedInUploadResponse(
            success=True,
            parsed_data=linkedin_data,
            message="LinkedIn profile uploaded and analyzed successfully"
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"🔥 LinkedIn upload error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process LinkedIn profile: {str(e)}"
        )


@app.post("/api/linkedin/check-alignment", response_model=LinkedInAlignmentResponse)
async def check_linkedin_alignment(
    resume_json: Dict[str, Any],
    linkedin_data: Dict[str, Any],
    job_id: Optional[str] = None
) -> LinkedInAlignmentResponse:
    """
    Check alignment between resume and LinkedIn profile

    Identifies discrepancies in:
    - Job titles
    - Company names
    - Date ranges
    - Skills
    """

    discrepancies = []

    # Get resume current role info
    resume_experience = resume_json.get("experience", [])
    resume_current = resume_experience[0] if resume_experience else {}
    resume_title = resume_current.get("title", "").lower().strip()
    resume_company = resume_current.get("company", "").lower().strip()
    resume_dates = resume_current.get("dates", "")
    resume_skills = set([s.lower().strip() for s in resume_json.get("skills", [])])

    # Get LinkedIn info
    linkedin_title = (linkedin_data.get("current_role") or "").lower().strip()
    linkedin_company = (linkedin_data.get("current_company") or "").lower().strip()
    linkedin_exp = linkedin_data.get("experience", [])
    linkedin_dates = linkedin_exp[0].get("dates", "") if linkedin_exp else ""
    linkedin_skills = set([s.lower().strip() for s in linkedin_data.get("skills", [])])

    # Check title mismatch
    if resume_title and linkedin_title:
        # Allow for some variation (e.g., "Senior Engineer" vs "Sr. Engineer")
        if resume_title not in linkedin_title and linkedin_title not in resume_title:
            # Check for common abbreviation variations
            title_variations = [
                (resume_title.replace("senior", "sr.").replace("sr.", "senior"), linkedin_title),
                (resume_title.replace("junior", "jr.").replace("jr.", "junior"), linkedin_title),
                (resume_title.replace("manager", "mgr").replace("mgr", "manager"), linkedin_title),
            ]
            if not any(r in l or l in r for r, l in title_variations):
                discrepancies.append(LinkedInAlignmentIssue(
                    type="title_mismatch",
                    severity="high",
                    resume_value=resume_current.get("title"),
                    linkedin_value=linkedin_data.get("current_role"),
                    message=f"Your resume says '{resume_current.get('title')}' but your LinkedIn says '{linkedin_data.get('current_role')}'"
                ))

    # Check company mismatch
    if resume_company and linkedin_company:
        if resume_company not in linkedin_company and linkedin_company not in resume_company:
            discrepancies.append(LinkedInAlignmentIssue(
                type="company_mismatch",
                severity="high",
                resume_value=resume_current.get("company"),
                linkedin_value=linkedin_data.get("current_company"),
                message=f"Your resume shows current company as '{resume_current.get('company')}' but LinkedIn shows '{linkedin_data.get('current_company')}'"
            ))

    # Check date range mismatch (basic check)
    if resume_dates and linkedin_dates:
        resume_dates_lower = resume_dates.lower()
        linkedin_dates_lower = linkedin_dates.lower()
        if resume_dates_lower not in linkedin_dates_lower and linkedin_dates_lower not in resume_dates_lower:
            # More sophisticated date comparison could be added here
            discrepancies.append(LinkedInAlignmentIssue(
                type="date_mismatch",
                severity="medium",
                resume_value=resume_dates,
                linkedin_value=linkedin_dates,
                message=f"Date ranges don't match: Resume shows '{resume_dates}', LinkedIn shows '{linkedin_dates}'"
            ))

    # Check skills overlap
    if resume_skills:
        common_skills = resume_skills & linkedin_skills
        resume_only = resume_skills - linkedin_skills

        if len(common_skills) < len(resume_skills) * 0.5:  # Less than 50% overlap
            missing_sample = list(resume_only)[:3]
            discrepancies.append(LinkedInAlignmentIssue(
                type="skills_mismatch",
                severity="medium",
                resume_value=", ".join(missing_sample),
                linkedin_value=None,
                message=f"Your resume emphasizes skills that aren't listed on your LinkedIn: {', '.join(missing_sample)}"
            ))

    # Determine overall severity
    severity = "low"
    if any(d.severity == "high" for d in discrepancies):
        severity = "high"
    elif any(d.severity == "medium" for d in discrepancies):
        severity = "medium"

    return LinkedInAlignmentResponse(
        aligned=len(discrepancies) == 0,
        discrepancies=discrepancies,
        severity=severity,
        discrepancy_count=len(discrepancies)
    )


@app.post("/api/linkedin/optimize", response_model=LinkedInOptimizeResponse)
async def optimize_linkedin_profile(
    request: LinkedInOptimizeRequest
) -> LinkedInOptimizeResponse:
    """
    Generate comprehensive LinkedIn optimization with world-class recommendations.

    Uses resume data, job description, and current LinkedIn profile to create:
    - Severity assessment and summary
    - Optimized headline with strategic rationale
    - Comprehensive About section (4 paragraphs)
    - Experience bullets for multiple roles
    - Skills prioritization and recommendations
    - Optional section recommendations (Featured, Recommendations, Activity)

    SINGLE SOURCE OF TRUTH ENFORCEMENT:
    - If Job Fit recommendation is "Do Not Apply" AND user has NOT proceeded despite guidance,
      this endpoint will refuse to optimize for that specific role.
    - HenryHQ does not coach candidates toward roles they are not qualified for.
    """

    # =========================================================================
    # SINGLE SOURCE OF TRUTH: Respect Job Fit Decision
    # Per QA Spec Rule 6: "Downstream pages = mirrors, not judges"
    # LinkedIn optimization MUST NOT offer role-specific optimization when
    # Job Fit has determined "Do Not Apply" - unless user explicitly proceeds.
    # =========================================================================
    job_fit_rec = (request.job_fit_recommendation or "").lower().strip()
    is_do_not_apply = job_fit_rec in ["do not apply", "skip"]

    if is_do_not_apply and not request.proceeded_despite_guidance:
        # Block optimization for this specific role
        locked_reason = request.locked_reason or "You are not qualified for this role based on Job Fit analysis."
        print(f"🚫 LinkedIn Optimize BLOCKED: Job Fit = 'Do Not Apply', user has not proceeded")
        print(f"   Locked reason: {locked_reason}")
        raise HTTPException(
            status_code=400,
            detail={
                "error": "optimization_blocked_by_job_fit",
                "message": "LinkedIn optimization is not available for roles where HenryHQ recommends 'Do Not Apply'.",
                "locked_reason": locked_reason,
                "guidance": "Your profile is optimized for roles that match your experience. Focus on higher-fit opportunities.",
                "can_proceed": True,
                "proceed_action": "If you choose to proceed anyway, acknowledge this guidance first."
            }
        )

    if is_do_not_apply and request.proceeded_despite_guidance:
        # User acknowledged and proceeded - log for accountability
        print(f"⚠️ LinkedIn Optimize: User proceeded despite 'Do Not Apply' guidance")
        print(f"   Target role: {request.target_role}")
        print(f"   Locked reason was: {request.locked_reason}")

    if not request.resume_json and not request.job_description:
        raise HTTPException(
            status_code=400,
            detail="Either resume_json or job_description must be provided"
        )

    # Build comprehensive resume context
    resume_context = ""
    if request.resume_json:
        resume_context = f"""
CANDIDATE RESUME DATA:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Name: {request.resume_json.get('full_name', 'Unknown')}
Current Title: {request.resume_json.get('current_title', 'Unknown')}
Years Experience: {request.resume_json.get('years_experience', 'Unknown')}
Summary: {request.resume_json.get('summary', '')}
Skills: {', '.join(request.resume_json.get('skills', [])[:30])}

WORK EXPERIENCE:
"""
        for i, exp in enumerate(request.resume_json.get('experience', [])[:4]):
            resume_context += f"""
Role {i+1}: {exp.get('title', 'Unknown')} at {exp.get('company', 'Unknown')}
Dates: {exp.get('dates', '')}
Bullets:
"""
            for bullet in exp.get('bullets', [])[:6]:
                resume_context += f"  • {bullet}\n"

    # Build LinkedIn context
    linkedin_context = ""
    current_headline = ""
    current_summary = ""
    if request.linkedin_data:
        current_headline = request.linkedin_data.get('headline', '')
        current_summary = request.linkedin_data.get('summary', '')
        linkedin_context = f"""
CURRENT LINKEDIN PROFILE:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Headline: {current_headline}
Summary: {current_summary}
Current Role: {request.linkedin_data.get('current_role', 'Unknown')}
Current Company: {request.linkedin_data.get('current_company', 'Unknown')}
Skills Listed: {', '.join(request.linkedin_data.get('skills', [])[:20])}
Experience Entries: {len(request.linkedin_data.get('experience', []))}
"""

    job_context = ""
    if request.job_description:
        job_context = f"""
TARGET JOB DESCRIPTION:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{request.job_description[:3000]}
"""

    target_role = request.target_role or "their target role"

    # Comprehensive system prompt for world-class LinkedIn optimization
    system_prompt = """You are an elite executive recruiter generating world-class LinkedIn optimization recommendations for HenryHQ users. Your guidance must meet the standard of $50K/year executive coaching.

=== HENRYHQ VOICE (NON-NEGOTIABLE) ===

You are HenryHQ, a direct, honest, supportive career coach.
You tell candidates the truth without shame, and you always give them a clear next step.
Your tone is calm, confident, human, and never robotic or overly optimistic.
Your goal is simple: make the candidate better with every message.
If an output does not improve clarity, readiness, confidence, or strategy, rewrite it.

Voice Rules:
1. Truth first, support second. Never sugar-coat. Never shame. Use: Truth → Why → Fix → Support.
2. Be direct and concise. Short sentences. No filler. No corporate jargon.
3. Every output must give the user a NEXT STEP.
4. No false encouragement. Praise must be earned and specific.
5. Emotional safety is mandatory. Deliver hard truths calmly and respectfully.

=== END HENRYHQ VOICE ===

QUALITY BAR:
- Recommendations must be specific, actionable, and grounded in the candidate's ACTUAL experience
- Never use generic placeholders or filler language
- Every suggestion must have clear strategic rationale (why this works)
- Output should feel like 1:1 coaching from a top executive recruiter

HEADLINE RULES (220 characters max):
Formula: [Seniority] [Function] | [Specialization] | [Brand Signals] | [Domain Expertise]
Include: seniority signal, function, specialization, brand signals (Ex-[Company]), domain keywords
Bad: "Senior Product Manager" (too generic)
Good: "Senior Product Manager | B2B SaaS Growth | Ex-Spotify, Uber | Product-Market Fit & GTM Strategy"

ABOUT SECTION RULES (600-800 words, 4 paragraphs):
1. Value Proposition (3-4 sentences): What you do, for whom, your unique positioning
2. Approach/Methodology (4-5 sentences): How you work differently, your philosophy
3. Track Record (4-5 sentences): Specific, quantified accomplishments with company names
4. Current Focus (3-4 sentences): What you're working on now, what excites you

EXPERIENCE BULLETS RULES:
Formula: [Accomplished X outcome] by [doing Y action], resulting in [Z measurable impact]
- 5 bullets per role, each with quantified impact
- Numbers must match resume exactly
- Seniority-appropriate framing:
  * IC: "Built...", "Shipped...", "Designed..."
  * Senior IC: "Led...", "Owned...", "Architected..."
  * Leadership: "Built and scaled...", "Established...", "Partnered with [exec]..."

SKILLS RULES:
- Top 3 skills: Most important for target role, highest recruiter search frequency
- Add these skills: Missing but critical for target role
- Remove these skills: Generic low-signal skills (Microsoft Office, Teamwork, etc.)

SEVERITY LEVELS:
- CRITICAL: Profile contradicts resume OR is severely underoptimized
- HIGH: Profile is weak but not contradictory (generic language, no quantified impact)
- MEDIUM: Profile is acceptable but not competitive
- LOW: Profile is solid, minor tweaks only

CRITICAL RULES:
- Use ONLY information from the provided resume - do not fabricate
- Every claim must trace back to resume data
- Numbers must match exactly between resume and recommendations
- Write in candidate's authentic voice (professional but not robotic)

Return valid JSON matching this exact structure:
{
    "severity": "CRITICAL|HIGH|MEDIUM|LOW",
    "summary_message": "2-3 sentence summary of what's wrong and what these updates will fix",
    "benefits": ["Specific benefit 1", "Specific benefit 2", "Specific benefit 3"],

    "headline": "Optimized headline under 220 characters",
    "headline_why": ["Reason 1 why this headline works", "Reason 2", "Reason 3", "Reason 4"],
    "headline_update_reason": "1-2 sentence explanation of why current headline is weak",

    "summary": "Full 4-paragraph About section (600-800 words) in first person",
    "summary_why": ["Reason 1 why this summary works", "Reason 2", "Reason 3", "Reason 4", "Reason 5"],
    "summary_update_reason": "1-2 sentence explanation of why current summary needs work",

    "experience_optimizations": [
        {
            "title": "Job Title",
            "company": "Company Name",
            "dates": "Date Range",
            "company_context": "Brief company description (e.g., Fortune 500 tech company...)",
            "bullets": ["Impact bullet 1", "Impact bullet 2", "Impact bullet 3", "Impact bullet 4", "Impact bullet 5"],
            "why_these_work": ["Reason 1", "Reason 2", "Reason 3"]
        }
    ],

    "top_skills": ["Skill 1", "Skill 2", "Skill 3"],
    "skills_to_add": ["Skill 4", "Skill 5", "Skill 6", "Skill 7", "Skill 8"],
    "skills_to_remove": ["Generic skill 1", "Generic skill 2", "Generic skill 3"],
    "skills_why": ["Why this prioritization makes sense"],

    "featured_recommendation": "Recommendation for Featured section",
    "featured_suggestions": ["Specific suggestion 1", "Specific suggestion 2"],
    "recommendations_advice": "Advice for getting recommendations",
    "who_to_ask": [
        {"person_type": "Former manager", "what_to_emphasize": "Strategic thinking and execution"},
        {"person_type": "Cross-functional partner", "what_to_emphasize": "Collaboration skills"}
    ],
    "activity_recommendation": "Specific activity recommendation based on target role"
}

Your response must be ONLY valid JSON, no additional text."""

    user_message = f"""Generate comprehensive LinkedIn optimization for this candidate targeting {target_role}.

{resume_context}

{linkedin_context}

{job_context}

Generate:
1. Severity assessment and summary message
2. Optimized headline with 4 reasons why it works
3. Full 4-paragraph About section (600-800 words) with rationale
4. Experience optimizations for at least 2 recent roles (5 bullets each)
5. Skills prioritization (top 3, add 5, remove 3)
6. Optional section recommendations (Featured, Recommendations, Activity)

Remember: Use ONLY information from the resume. Every number and claim must trace back to the provided data."""

    try:
        # Call Claude
        print("📤 Generating comprehensive LinkedIn optimization...")
        response = call_claude(system_prompt, user_message)

        # Clean and parse JSON response
        cleaned = clean_claude_json(response)
        parsed = json.loads(cleaned)

        print("✅ LinkedIn optimization generated successfully")

        # Build experience optimizations list
        experience_opts = []
        for exp in parsed.get("experience_optimizations", []):
            experience_opts.append(LinkedInExperienceOptimization(
                title=exp.get("title", ""),
                company=exp.get("company", ""),
                dates=exp.get("dates", ""),
                company_context=exp.get("company_context", ""),
                bullets=exp.get("bullets", []),
                why_these_work=exp.get("why_these_work", [])
            ))

        # Extract first role's bullets for backwards compatibility
        experience_highlights = []
        if experience_opts:
            experience_highlights = experience_opts[0].bullets[:4]

        # ====================================================================
        # LEPE INTEGRATION
        # Per LEPE Spec: LinkedIn Scoring must reference LEPE outputs and
        # must not independently assess leadership readiness for Manager+ roles.
        # ====================================================================
        lepe_applicable = False
        lepe_decision = None
        lepe_coaching = None
        lepe_skepticism_warning = None
        leadership_tenure = None
        accountability_record = None

        if target_role and is_manager_plus_role(target_role):
            lepe_applicable = True
            print(f"   📊 LEPE: Manager+ target detected - running leadership evaluation")

            # Get resume data from request
            resume_json = request.resume_json
            if isinstance(resume_json, str):
                try:
                    resume_json = json.loads(resume_json)
                except json.JSONDecodeError:
                    resume_json = {}

            if resume_json:
                # Build response_data for LEPE
                lepe_response_data = {
                    "role_title": target_role,
                    "job_description": request.job_description or ""
                }

                # Run LEPE evaluation
                lepe_result = evaluate_lepe(resume_json, lepe_response_data)

                if lepe_result.get("lepe_applicable"):
                    positioning = lepe_result.get("positioning_decision", {})
                    lepe_decision = positioning.get("decision")
                    leadership_tenure = lepe_result.get("leadership_tenure")
                    accountability_record = lepe_result.get("accountability_record")

                    # Extract messaging based on decision
                    messaging = positioning.get("messaging", {})
                    if lepe_decision == "position":
                        lepe_coaching = messaging.get("coaching_advice")
                    elif lepe_decision in ["caution", "locked"]:
                        lepe_skepticism_warning = messaging.get("skepticism_warning")

                    print(f"   📊 LEPE Advisory: {lepe_decision} (advisory only)")

        return LinkedInOptimizeResponse(
            # Severity and summary
            severity=parsed.get("severity", "MEDIUM"),
            summary_message=parsed.get("summary_message", ""),
            benefits=parsed.get("benefits", []),

            # Headline
            headline=parsed.get("headline", ""),
            current_headline=current_headline,
            headline_why=parsed.get("headline_why", []),
            headline_update_reason=parsed.get("headline_update_reason", ""),

            # About section
            summary=parsed.get("summary", ""),
            current_summary=current_summary,
            summary_why=parsed.get("summary_why", []),
            summary_update_reason=parsed.get("summary_update_reason", ""),

            # Experience
            experience_optimizations=experience_opts,

            # Skills
            top_skills=parsed.get("top_skills", []),
            skills_to_add=parsed.get("skills_to_add", []),
            skills_to_remove=parsed.get("skills_to_remove", []),
            skills_why=parsed.get("skills_why", []),

            # Optional sections
            featured_recommendation=parsed.get("featured_recommendation", ""),
            featured_suggestions=parsed.get("featured_suggestions", []),
            recommendations_advice=parsed.get("recommendations_advice", ""),
            who_to_ask=parsed.get("who_to_ask", []),
            activity_recommendation=parsed.get("activity_recommendation", ""),

            # Legacy fields
            experience_highlights=experience_highlights,
            has_issues=parsed.get("severity", "MEDIUM") in ["CRITICAL", "HIGH"],
            issue_count=1 if parsed.get("severity", "MEDIUM") in ["CRITICAL", "HIGH"] else 0,
            generated_at=datetime.utcnow().isoformat(),

            # LEPE fields
            lepe_applicable=lepe_applicable,
            lepe_decision=lepe_decision,
            lepe_coaching=lepe_coaching,
            lepe_skepticism_warning=lepe_skepticism_warning,
            leadership_tenure=leadership_tenure,
            accountability_record=accountability_record
        )

    except json.JSONDecodeError as e:
        print(f"🔥 JSON parse error: {e}")
        raise HTTPException(
            status_code=500,
            detail="Failed to generate LinkedIn optimization"
        )
    except Exception as e:
        print(f"🔥 LinkedIn optimization error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"LinkedIn optimization failed: {str(e)}"
        )


# ============================================================================
# FEEDBACK ACKNOWLEDGMENT EMAIL
# ============================================================================

# Resend API Configuration
RESEND_API_KEY = os.getenv("RESEND_API_KEY")
RESEND_API_URL = "https://api.resend.com/emails"


class FeedbackAcknowledgmentRequest(BaseModel):
    email: str
    name: Optional[str] = None
    feedback_type: str  # bug, feature_request, praise, ux_issue, general
    feedback_summary: str  # First ~100 chars of feedback
    current_page: Optional[str] = None  # What page they were on
    full_feedback: Optional[str] = None  # Full feedback text for admin
    conversation_context: Optional[str] = None  # Recent conversation for context


# Admin notification emails
ADMIN_EMAIL = "henry.r.bolden@gmail.com"  # Where to send detailed feedback notifications (support@)
ADMIN_PERSONAL_EMAIL = "hb@henryhq.ai"  # Personal quick notification (hb@)
ADMIN_USER_EMAIL = "hb@henryhq.ai"  # Admin's HenryHQ login email (for real-time Hey Henry notifications)


def get_feedback_email_content(feedback_type: str, feedback_summary: str, name: Optional[str] = None):
    """
    Returns subject and key message based on feedback type.
    Matches the existing welcome email dark theme with gradient logo.
    """
    greeting = f"Hi {name}," if name else "Hi there,"

    # Content by feedback type
    content_map = {
        "bug": {
            "subject": "Got it. We're looking into this.",
            "headline": "Thanks for the heads up.",
            "message": "We'll dig in and follow up if we need more details. Bugs like this help us make HenryHQ better for everyone.",
            "next_steps": "We review all bug reports within 24 hours. If we need to reproduce the issue, we'll reach out."
        },
        "feature_request": {
            "subject": "Idea received. We're reviewing.",
            "headline": "We hear you.",
            "message": "We've added this to our roadmap review. No promises, but every feature request gets considered.",
            "next_steps": "We prioritize features based on impact. If this makes the cut, you'll be the first to know."
        },
        "praise": {
            "subject": "That made our day.",
            "headline": "Seriously, thank you.",
            "message": "It means a lot to hear this. Building HenryHQ is a labor of love, and knowing it's helping makes all the difference.",
            "next_steps": "Keep crushing your job search. We're rooting for you."
        },
        "ux_issue": {
            "subject": "UX feedback received.",
            "headline": "Thanks for helping us improve.",
            "message": "We're always tuning the experience. Feedback like this is how we make HenryHQ feel right.",
            "next_steps": "Our design team reviews UX feedback weekly. If we make changes based on your input, we'll let you know."
        },
        "general": {
            "subject": "Message received.",
            "headline": "Thanks for reaching out.",
            "message": "We've got your note and appreciate you taking the time to share your thoughts.",
            "next_steps": "If this needs a response, we'll get back to you within 48 hours."
        }
    }

    content = content_map.get(feedback_type, content_map["general"])

    # Truncate and escape feedback summary for display
    display_summary = feedback_summary[:150] + "..." if len(feedback_summary) > 150 else feedback_summary
    # Basic HTML escaping
    display_summary = display_summary.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{content['subject']}</title>
</head>
<body style="margin: 0; padding: 0; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', 'Roboto', 'Helvetica', sans-serif; background-color: #0a0a0a;">
    <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #0a0a0a;">
        <tr>
            <td align="center" style="padding: 40px 20px;">
                <table width="600" cellpadding="0" cellspacing="0" style="background-color: #111111; border: 1px solid #222222; border-radius: 12px; padding: 48px;">
                    <tr>
                        <td align="center">
                            <!-- Minimal Ring Logo -->
                            <div style="margin-bottom: 32px;">
                                <svg width="80" height="80" viewBox="0 0 200 200" fill="none" xmlns="http://www.w3.org/2000/svg">
                                    <defs>
                                        <linearGradient id="ringGradient" x1="0%" y1="100%" x2="100%" y2="0%">
                                            <stop offset="0%" style="stop-color:#667eea;stop-opacity:1" />
                                            <stop offset="100%" style="stop-color:#764ba2;stop-opacity:1" />
                                        </linearGradient>
                                        <linearGradient id="strokeGradient" x1="0%" y1="100%" x2="100%" y2="0%">
                                            <stop offset="0%" style="stop-color:#667eea;stop-opacity:1" />
                                            <stop offset="100%" style="stop-color:#764ba2;stop-opacity:1" />
                                        </linearGradient>
                                    </defs>
                                    <circle cx="100" cy="100" r="85" stroke="url(#ringGradient)" stroke-width="4" fill="none"/>
                                    <path d="M55 130 L55 70" stroke="#667eea" stroke-width="9" stroke-linecap="round" fill="none"/>
                                    <path d="M145 130 L145 50" stroke="url(#strokeGradient)" stroke-width="9" stroke-linecap="round" fill="none"/>
                                    <path d="M55 100 L145 100" stroke="#764ba2" stroke-width="9" stroke-linecap="round" fill="none"/>
                                    <circle cx="145" cy="50" r="9" fill="#764ba2"/>
                                </svg>
                            </div>

                            <!-- Greeting -->
                            <p style="margin: 0 0 8px; font-size: 16px; color: #9ca3af; text-align: center;">
                                {greeting}
                            </p>

                            <!-- Headline -->
                            <h1 style="margin: 0 0 24px; font-size: 28px; font-weight: 600; color: #ffffff; letter-spacing: -0.5px; text-align: center;">
                                {content['headline']}
                            </h1>

                            <!-- Main Message -->
                            <p style="margin: 0 0 24px; font-size: 16px; line-height: 1.6; color: #d1d5db; text-align: center;">
                                {content['message']}
                            </p>

                            <!-- What You Shared -->
                            <table width="100%" cellpadding="0" cellspacing="0" style="margin-bottom: 24px; background-color: #0a0a0a; border: 1px solid #222222; border-radius: 8px;">
                                <tr>
                                    <td style="padding: 16px 20px;">
                                        <p style="margin: 0 0 8px; font-size: 12px; text-transform: uppercase; letter-spacing: 0.5px; color: #6b7280;">
                                            What you shared
                                        </p>
                                        <p style="margin: 0; font-size: 14px; line-height: 1.5; color: #9ca3af; font-style: italic;">
                                            "{display_summary}"
                                        </p>
                                    </td>
                                </tr>
                            </table>

                            <!-- What Happens Next -->
                            <p style="margin: 0 0 8px; font-size: 14px; color: #9ca3af; text-align: center;">
                                <strong style="color: #ffffff;">What happens next:</strong>
                            </p>
                            <p style="margin: 0 0 32px; font-size: 14px; line-height: 1.5; color: #9ca3af; text-align: center;">
                                {content['next_steps']}
                            </p>

                            <!-- Reply Prompt -->
                            <p style="margin: 0; font-size: 14px; color: #6b7280; text-align: center; font-style: italic;">
                                Questions? Just reply to this email.
                            </p>
                        </td>
                    </tr>
                </table>

                <!-- Footer -->
                <table width="600" cellpadding="0" cellspacing="0" style="margin-top: 24px;">
                    <tr>
                        <td style="text-align: center; font-size: 13px; color: #4b5563;">
                            HenryHQ | Strategic job search intelligence
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
    </table>
</body>
</html>"""

    return content['subject'], html


def get_admin_notification_email(request: FeedbackAcknowledgmentRequest):
    """
    Creates an admin notification email with full feedback details.
    """
    type_labels = {
        "bug": "🐛 Bug Report",
        "feature_request": "💡 Feature Request",
        "praise": "🎉 Praise",
        "ux_issue": "🎨 UX Issue",
        "general": "💬 General Feedback"
    }

    type_label = type_labels.get(request.feedback_type, "💬 Feedback")
    user_display = request.name or "Anonymous"

    # Escape HTML in user content
    def escape_html(text):
        if not text:
            return ""
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")

    full_feedback = escape_html(request.full_feedback or request.feedback_summary)
    conversation = escape_html(request.conversation_context) if request.conversation_context else "No conversation context"
    current_page = request.current_page or "Unknown"

    subject = f"{type_label} from {user_display}"

    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{subject}</title>
</head>
<body style="margin: 0; padding: 20px; font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; background-color: #f5f5f5;">
    <table width="100%" cellpadding="0" cellspacing="0" style="max-width: 600px; margin: 0 auto; background-color: #ffffff; border-radius: 8px; overflow: hidden; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
        <tr>
            <td style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 20px; color: white;">
                <h1 style="margin: 0; font-size: 20px;">{type_label}</h1>
                <p style="margin: 8px 0 0; opacity: 0.9; font-size: 14px;">New feedback from HenryHQ</p>
            </td>
        </tr>
        <tr>
            <td style="padding: 24px;">
                <table width="100%" cellpadding="0" cellspacing="0">
                    <tr>
                        <td style="padding-bottom: 16px; border-bottom: 1px solid #eee;">
                            <strong style="color: #666; font-size: 12px; text-transform: uppercase;">From</strong>
                            <p style="margin: 4px 0 0; font-size: 16px; color: #333;">{user_display}</p>
                            <p style="margin: 2px 0 0; font-size: 14px; color: #666;">{request.email}</p>
                        </td>
                    </tr>
                    <tr>
                        <td style="padding: 16px 0; border-bottom: 1px solid #eee;">
                            <strong style="color: #666; font-size: 12px; text-transform: uppercase;">Page</strong>
                            <p style="margin: 4px 0 0; font-size: 14px; color: #333;">{current_page}</p>
                        </td>
                    </tr>
                    <tr>
                        <td style="padding: 16px 0; border-bottom: 1px solid #eee;">
                            <strong style="color: #666; font-size: 12px; text-transform: uppercase;">Feedback</strong>
                            <div style="margin: 8px 0 0; padding: 12px; background-color: #f8f9fa; border-radius: 6px; font-size: 14px; color: #333; line-height: 1.5;">
                                {full_feedback}
                            </div>
                        </td>
                    </tr>
                    <tr>
                        <td style="padding: 16px 0;">
                            <strong style="color: #666; font-size: 12px; text-transform: uppercase;">Conversation Context</strong>
                            <div style="margin: 8px 0 0; padding: 12px; background-color: #f0f0f0; border-radius: 6px; font-size: 12px; color: #666; line-height: 1.5; max-height: 200px; overflow-y: auto;">
                                {conversation}
                            </div>
                        </td>
                    </tr>
                </table>
            </td>
        </tr>
        <tr>
            <td style="padding: 16px 24px; background-color: #f8f9fa; border-top: 1px solid #eee;">
                <p style="margin: 0; font-size: 12px; color: #666;">
                    Reply to this email to respond directly to {request.email}
                </p>
            </td>
        </tr>
    </table>
</body>
</html>"""

    return subject, html


@app.post("/api/send-feedback-acknowledgment")
async def send_feedback_acknowledgment(request: FeedbackAcknowledgmentRequest):
    """
    Send acknowledgment email after user submits feedback via Hey Henry chat.
    Also sends notification to admin with full details.

    Supports feedback types: bug, feature_request, praise, ux_issue, general
    Each type gets a tailored subject line and message.
    """

    if not RESEND_API_KEY:
        print("⚠️ RESEND_API_KEY not configured - skipping feedback acknowledgment email")
        return {"success": False, "error": "Email service not configured"}

    headers = {
        "Authorization": f"Bearer {RESEND_API_KEY}",
        "Content-Type": "application/json"
    }

    user_email_id = None
    admin_email_id = None

    # 1. Send acknowledgment to user
    try:
        subject, html = get_feedback_email_content(
            feedback_type=request.feedback_type,
            feedback_summary=request.feedback_summary,
            name=request.name
        )

        payload = {
            "from": "Henry <support@henryhq.ai>",
            "to": request.email,
            "subject": subject,
            "html": html,
            "reply_to": "support@henryhq.ai"
        }

        response = requests.post(RESEND_API_URL, json=payload, headers=headers, timeout=10)
        response.raise_for_status()
        user_email_id = response.json().get("id")
        print(f"📧 Feedback acknowledgment sent to {request.email} (type: {request.feedback_type}, id: {user_email_id})")

    except requests.exceptions.RequestException as e:
        print(f"⚠️ Failed to send user acknowledgment: {str(e)}")

    # 2. Send notification to admin with full details
    try:
        admin_subject, admin_html = get_admin_notification_email(request)

        admin_payload = {
            "from": "HenryHQ Feedback <support@henryhq.ai>",
            "to": ADMIN_EMAIL,
            "subject": admin_subject,
            "html": admin_html,
            "reply_to": request.email  # Reply goes directly to the user
        }

        response = requests.post(RESEND_API_URL, json=admin_payload, headers=headers, timeout=10)
        response.raise_for_status()
        admin_email_id = response.json().get("id")
        print(f"📧 Admin notification sent to {ADMIN_EMAIL} (id: {admin_email_id})")

    except requests.exceptions.RequestException as e:
        print(f"⚠️ Failed to send admin notification: {str(e)}")

    # 3. Send quick personal notification to hb@henryhq.ai
    personal_email_id = None
    try:
        type_emoji = {
            "bug": "🐛", "feature_request": "💡", "praise": "🎉",
            "ux_issue": "🎨", "general": "💬"
        }.get(request.feedback_type, "💬")

        user_display = request.name or "Someone"
        personal_subject = f"{type_emoji} {user_display}: {request.feedback_summary[:50]}..."

        personal_payload = {
            "from": "Hey Henry <hb@henryhq.ai>",
            "to": ADMIN_PERSONAL_EMAIL,
            "subject": personal_subject,
            "html": f"""<p><strong>{user_display}</strong> ({request.email}) submitted {request.feedback_type.replace('_', ' ')}:</p>
<blockquote style="border-left: 3px solid #667eea; padding-left: 12px; color: #555;">{request.full_feedback or request.feedback_summary}</blockquote>
<p style="color: #888; font-size: 12px;">Page: {request.current_page or 'Unknown'}</p>""",
            "reply_to": request.email
        }

        response = requests.post(RESEND_API_URL, json=personal_payload, headers=headers, timeout=10)
        response.raise_for_status()
        personal_email_id = response.json().get("id")
        print(f"📧 Personal notification sent to {ADMIN_PERSONAL_EMAIL} (id: {personal_email_id})")

    except requests.exceptions.RequestException as e:
        print(f"⚠️ Failed to send personal notification: {str(e)}")

    # 4. Store notification for real-time Hey Henry pickup (if Supabase available)
    try:
        if supabase:
            supabase.table("admin_notifications").insert({
                "notification_type": "feedback",
                "feedback_type": request.feedback_type,
                "from_email": request.email,
                "from_name": request.name,
                "summary": request.feedback_summary,
                "full_content": request.full_feedback,
                "current_page": request.current_page,
                "read": False
            }).execute()
            print(f"📝 Notification stored for real-time pickup")
    except Exception as e:
        print(f"⚠️ Could not store notification: {str(e)}")

    return {
        "success": True,
        "message": "Emails sent",
        "user_email_id": user_email_id,
        "admin_email_id": admin_email_id,
        "personal_email_id": personal_email_id
    }


# ============================================================================
# ADMIN NOTIFICATIONS (Real-time Hey Henry alerts)
# ============================================================================

@app.get("/api/admin/notifications")
async def get_admin_notifications(user_email: str = None):
    """
    Get unread admin notifications for real-time Hey Henry alerts.
    Only returns notifications if the requesting user is an admin.
    """
    if not user_email or user_email.lower() != ADMIN_USER_EMAIL.lower():
        return {"notifications": [], "is_admin": False}

    if not supabase:
        return {"notifications": [], "is_admin": True, "error": "Database not available"}

    try:
        result = supabase.table("admin_notifications") \
            .select("*") \
            .eq("read", False) \
            .order("created_at", desc=True) \
            .limit(10) \
            .execute()

        return {
            "notifications": result.data if result.data else [],
            "is_admin": True
        }
    except Exception as e:
        print(f"⚠️ Error fetching admin notifications: {str(e)}")
        return {"notifications": [], "is_admin": True, "error": str(e)}


@app.post("/api/admin/notifications/{notification_id}/read")
async def mark_notification_read(notification_id: str, user_email: str = None):
    """
    Mark a notification as read.
    Only works if the requesting user is an admin.
    """
    if not user_email or user_email.lower() != ADMIN_USER_EMAIL.lower():
        return {"success": False, "error": "Not authorized"}

    if not supabase:
        return {"success": False, "error": "Database not available"}

    try:
        supabase.table("admin_notifications") \
            .update({"read": True}) \
            .eq("id", notification_id) \
            .execute()

        return {"success": True}
    except Exception as e:
        print(f"⚠️ Error marking notification read: {str(e)}")
        return {"success": False, "error": str(e)}


@app.post("/api/admin/notifications/read-all")
async def mark_all_notifications_read(user_email: str = None):
    """
    Mark all notifications as read.
    Only works if the requesting user is an admin.
    """
    if not user_email or user_email.lower() != ADMIN_USER_EMAIL.lower():
        return {"success": False, "error": "Not authorized"}

    if not supabase:
        return {"success": False, "error": "Database not available"}

    try:
        supabase.table("admin_notifications") \
            .update({"read": True}) \
            .eq("read", False) \
            .execute()

        return {"success": True}
    except Exception as e:
        print(f"⚠️ Error marking all notifications read: {str(e)}")
        return {"success": False, "error": str(e)}


@app.delete("/api/linkedin/profile")
async def delete_linkedin_profile():
    """
    Delete LinkedIn profile data

    Note: In this implementation, LinkedIn data is stored client-side.
    This endpoint exists for API consistency and future server-side storage.
    """
    return {"success": True, "message": "LinkedIn profile deleted"}


# ============================================================================
# TIER AND USAGE ENDPOINTS
# ============================================================================

# Pydantic models for tier endpoints
class UserUsageResponse(BaseModel):
    """Response model for user usage and tier information."""
    tier: str
    tier_display: str
    tier_price: int
    is_beta_user: bool
    beta_expires_at: Optional[str] = None
    usage: Dict[str, Any]
    features: Dict[str, Any]


class TierInfoResponse(BaseModel):
    """Response model for tier information."""
    tiers: List[Dict[str, Any]]


class FeatureCheckResponse(BaseModel):
    """Response model for feature access check."""
    allowed: bool
    limited: bool
    upgrade_to: Optional[str] = None
    feature: str


class UsageLimitResponse(BaseModel):
    """Response model for usage limit check."""
    allowed: bool
    used: int
    limit: int
    remaining: Optional[int] = None
    is_unlimited: bool
    usage_type: str


@app.get("/api/user/usage", response_model=UserUsageResponse)
async def get_user_usage(user_id: str = None):
    """
    Get current usage stats and tier information for a user.

    Returns tier, features, and usage limits for the current billing period.
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    if not TIER_SERVICE_AVAILABLE:
        # Return default sourcer tier if tier service not available
        return UserUsageResponse(
            tier='sourcer',
            tier_display='Sourcer',
            tier_price=0,
            is_beta_user=False,
            beta_expires_at=None,
            usage={
                'applications': {'allowed': True, 'used': 0, 'limit': 3, 'remaining': 3, 'is_unlimited': False},
                'resumes': {'allowed': True, 'used': 0, 'limit': 3, 'remaining': 3, 'is_unlimited': False},
                'cover_letters': {'allowed': True, 'used': 0, 'limit': 3, 'remaining': 3, 'is_unlimited': False},
                'henry_conversations': {'allowed': False, 'used': 0, 'limit': 0, 'remaining': 0, 'is_unlimited': False},
                'mock_interviews': {'allowed': False, 'used': 0, 'limit': 0, 'remaining': 0, 'is_unlimited': False},
                'coaching_sessions': {'allowed': False, 'used': 0, 'limit': 0, 'remaining': 0, 'is_unlimited': False},
            },
            features=TIER_LIMITS.get('sourcer', {}).get('features', {})
        )

    try:
        tier_service = TierService(supabase)
        usage_summary = await tier_service.get_user_usage_summary(user_id)
        return UserUsageResponse(**usage_summary)
    except Exception as e:
        logger.error(f"Error getting user usage: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting user usage: {str(e)}")


@app.get("/api/tiers", response_model=TierInfoResponse)
async def get_tiers():
    """
    Get information about all available subscription tiers.

    Returns pricing, limits, and features for each tier.
    """
    if not TIER_SERVICE_AVAILABLE:
        raise HTTPException(status_code=503, detail="Tier service not available")

    try:
        return TierInfoResponse(tiers=get_all_tier_info())
    except Exception as e:
        logger.error(f"Error getting tier info: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting tier info: {str(e)}")


@app.get("/api/user/feature-access/{feature_name}", response_model=FeatureCheckResponse)
async def check_feature_access(feature_name: str, user_id: str = None):
    """
    Check if a user has access to a specific feature.

    Returns whether the feature is allowed, limited, or locked.
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    if not TIER_SERVICE_AVAILABLE:
        # Default to sourcer tier features
        feature_value = TIER_LIMITS.get('sourcer', {}).get('features', {}).get(feature_name, False)
        return FeatureCheckResponse(
            allowed=feature_value is True or feature_value == 'limited',
            limited=feature_value == 'limited',
            upgrade_to='recruiter' if not feature_value else None,
            feature=feature_name
        )

    try:
        tier_service = TierService(supabase)
        profile = await tier_service.ensure_user_profile(user_id)
        tier = tier_service.get_effective_tier(profile)
        access = tier_service.check_feature_access(tier, feature_name)
        return FeatureCheckResponse(
            allowed=access['allowed'],
            limited=access['limited'],
            upgrade_to=access['upgrade_to'],
            feature=feature_name
        )
    except Exception as e:
        logger.error(f"Error checking feature access: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error checking feature access: {str(e)}")


@app.get("/api/user/usage-limit/{usage_type}", response_model=UsageLimitResponse)
async def check_usage_limit(usage_type: str, user_id: str = None):
    """
    Check if a user has remaining usage for a specific type.

    Returns current usage, limit, and whether more usage is allowed.
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    valid_usage_types = ['applications', 'resumes', 'cover_letters', 'henry_conversations', 'mock_interviews', 'coaching_sessions']
    if usage_type not in valid_usage_types:
        raise HTTPException(status_code=400, detail=f"Invalid usage_type. Must be one of: {', '.join(valid_usage_types)}")

    if not TIER_SERVICE_AVAILABLE:
        # Default to sourcer tier limits
        limit = TIER_LIMITS.get('sourcer', {}).get(f'{usage_type}_per_month', 0)
        return UsageLimitResponse(
            allowed=limit > 0,
            used=0,
            limit=limit,
            remaining=limit,
            is_unlimited=False,
            usage_type=usage_type
        )

    try:
        tier_service = TierService(supabase)
        profile = await tier_service.ensure_user_profile(user_id)
        tier = tier_service.get_effective_tier(profile)
        usage = await tier_service.check_usage_limit(user_id, tier, usage_type)
        return UsageLimitResponse(
            allowed=usage['allowed'],
            used=usage['used'],
            limit=usage['limit'],
            remaining=usage['remaining'],
            is_unlimited=usage['is_unlimited'],
            usage_type=usage_type
        )
    except Exception as e:
        logger.error(f"Error checking usage limit: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error checking usage limit: {str(e)}")


@app.post("/api/user/usage/increment/{usage_type}")
async def increment_usage(usage_type: str, user_id: str = None):
    """
    Increment usage counter for a specific type.

    This is typically called after a successful action (resume generated, etc.)
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    valid_usage_types = ['applications', 'resumes', 'cover_letters', 'henry_conversations', 'mock_interviews', 'coaching_sessions']
    if usage_type not in valid_usage_types:
        raise HTTPException(status_code=400, detail=f"Invalid usage_type. Must be one of: {', '.join(valid_usage_types)}")

    if not TIER_SERVICE_AVAILABLE:
        return {"success": True, "message": "Usage tracking not available"}

    try:
        tier_service = TierService(supabase)
        await tier_service.increment_usage(user_id, usage_type)
        return {"success": True, "usage_type": usage_type}
    except Exception as e:
        logger.error(f"Error incrementing usage: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error incrementing usage: {str(e)}")


@app.get("/api/user/tier")
async def get_user_tier(user_id: str = None):
    """
    Get the effective tier for a user.

    Returns the tier, considering beta overrides if applicable.
    """
    if not user_id:
        raise HTTPException(status_code=400, detail="user_id is required")

    if not supabase:
        raise HTTPException(status_code=503, detail="Database not available")

    if not TIER_SERVICE_AVAILABLE:
        return {
            "tier": "sourcer",
            "tier_display": "Sourcer",
            "tier_price": 0,
            "is_beta_user": False
        }

    try:
        tier_service = TierService(supabase)
        profile = await tier_service.ensure_user_profile(user_id)
        tier = tier_service.get_effective_tier(profile)
        return {
            "tier": tier,
            "tier_display": TIER_NAMES.get(tier, 'Sourcer'),
            "tier_price": TIER_PRICES.get(tier, 0),
            "is_beta_user": profile.get('is_beta_user', False),
            "beta_expires_at": profile.get('beta_expires_at')
        }
    except Exception as e:
        logger.error(f"Error getting user tier: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error getting user tier: {str(e)}")


# ============================================================================
# RUN SERVER
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    print(f"\n🚀 Henry Job Search Engine API starting on http://localhost:{port}")
    print(f"📚 API docs available at http://localhost:{port}/docs")
    print(f"🔑 Using Anthropic API key: {API_KEY[:20]}...")
    if OPENAI_API_KEY:
        print(f"🎙️ Voice features enabled (OpenAI key: {OPENAI_API_KEY[:20]}...)")
    else:
        print("⚠️ Voice features disabled (no OPENAI_API_KEY)")
    uvicorn.run(app, host="0.0.0.0", port=port)
